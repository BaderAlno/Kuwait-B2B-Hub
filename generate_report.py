from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_heading(para, text, level=1, color=RGBColor(0x1A, 0x3A, 0x6B)):
    para.clear()
    run = para.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(16)
        para.paragraph_format.space_before = Pt(18)
        para.paragraph_format.space_after  = Pt(6)
    elif level == 2:
        run.font.size = Pt(13)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after  = Pt(4)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x2E, 0x6B, 0xA0)
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(2)

def add_h1(doc, text):
    p = doc.add_paragraph()
    set_heading(p, text, 1)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    set_heading(p, text, 2)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    set_heading(p, text, 3)
    return p

def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def add_table_row(table, cells_data, header=False, bg_color=None):
    row = table.add_row()
    for i, val in enumerate(cells_data):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(val)
        run.font.size = Pt(10)
        if header:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if bg_color:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), bg_color)
            tc_pr.append(shd)
    return row

def shade_row(row, hex_color):
    for cell in row.cells:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tc_pr.append(shd)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1A3A6B')
    pBdr.append(bottom)
    pPr.append(pBdr)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cover_title.add_run("Kuwait B2B Hub")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

cover_sub = doc.add_paragraph()
cover_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = cover_sub.add_run("Project Progress Report")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x2E, 0x6B, 0xA0)

doc.add_paragraph()
add_divider(doc)
doc.add_paragraph()

meta_lines = [
    ("Course:", "Software Engineering — University Project"),
    ("Report Type:", "Progress Report"),
    ("Date:", "April 17, 2026"),
    ("Platform:", "B2B Wholesale Marketplace for Kuwait & GCC"),
]
for lbl, val in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(lbl + "  ")
    r.bold = True
    r.font.size = Pt(12)
    r2 = p.add_run(val)
    r2.font.size = Pt(12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — BRIEF REMINDER ABOUT THE SYSTEM
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "1.  Brief Reminder About the System")
add_divider(doc)

add_h2(doc, "1.1  Purpose")
add_body(doc,
    "Kuwait B2B Hub is a digital wholesale marketplace purpose-built for the Kuwait and "
    "GCC market.  The platform replaces the informal, WhatsApp-driven wholesale ordering "
    "process with a structured, professional environment where verified Kuwaiti brands and "
    "business buyers can connect, negotiate terms, and execute bulk orders reliably.  The "
    "core mission is to professionalise wholesale distribution in Kuwait by providing a "
    "single, trusted platform for catalog management, order processing, brand verification, "
    "pricing transparency, and reputation building."
)

add_h2(doc, "1.2  Target Users")
add_body(doc, "The system serves three distinct actor roles:")
add_bullet(doc,
    "Platform administrators who approve brands, manage users, monitor orders, "
    "handle review moderation, and access revenue analytics.",
    bold_prefix="Admin Users —"
)
add_bullet(doc,
    "Kuwaiti wholesale suppliers / manufacturers who list products, manage catalogs, "
    "set bulk-pricing tiers, approve or reject orders, and track business metrics via a "
    "dedicated dashboard.  Subscription: KD 29 / month with a 14-day free trial.",
    bold_prefix="Brand Owners —"
)
add_bullet(doc,
    "Retail and wholesale buyers who browse verified brands, search and filter "
    "products, place bulk orders, track order status, and leave reviews.  Free to use.",
    bold_prefix="Business Buyers —"
)

add_h2(doc, "1.3  Key Functionalities")
features = [
    ("Authentication & Role-Based Access",
     "Secure registration, login/logout, and protected routes per role (admin / brand_owner / buyer)."),
    ("Marketplace Browsing",
     "Search and filter brands by category, MOQ, rating; view verification badges "
     "(Premium, Verified, New); multi-currency display (KWD, SAR, AED, BHD)."),
    ("Product & Catalog Management",
     "Brand owners create and edit products with bulk-pricing tiers, stock levels, "
     "images, and Excel/CSV bulk-import capability."),
    ("Order Management",
     "Buyers submit order requests; brands approve or reject; full status lifecycle "
     "(pending → approved → completed / rejected) with order-item tracking."),
    ("Brand Verification & Profiles",
     "Admins verify brands via commercial-registration review; tiers: Premium, Verified, New; "
     "WhatsApp click-tracking and business-hours settings."),
    ("Trust & Reputation System",
     "Buyer reviews (1–5 stars), rating aggregation, trust metrics for both brands "
     "(response rate, completion rate, fulfillment days) and buyers (completion rate, "
     "cancellation rate), brand replies to reviews, admin moderation."),
    ("In-App Notifications",
     "Notification panel with read/unread status for order events, new registrations, "
     "and review alerts."),
    ("Admin Dashboard",
     "Platform-wide statistics, pending brand approvals, user management, order analytics, "
     "and review moderation queue."),
    ("Bilingual UI (EN / AR)",
     "Full Arabic right-to-left layout support via next-intl with complete translation files."),
    ("GCC Multi-Market Support",
     "Brands specify served markets (KW, SA, AE, BH); buyers see only relevant listings."),
]
for title, desc in features:
    add_bullet(doc, desc, bold_prefix=title + " —")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — DEVELOPMENT PROCESS MODEL
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "2.  Development Process Model")
add_divider(doc)

add_h2(doc, "2.1  Adopted Model: Agile (Iterative / Sprint-Based)")
add_body(doc,
    "The team adopted an Agile incremental development model, organising work into "
    "iterative sprints rather than committing to a rigid up-front plan.  Each sprint "
    "delivers a working vertical slice of the platform — from authentication, through "
    "the marketplace, to trust features — allowing the team to inspect, adapt, and "
    "reprioritise based on what has been learned."
)

add_h2(doc, "2.2  Rationale")
reasons = [
    ("Evolving Requirements",
     "B2B marketplace requirements for the GCC context are not fully known at the "
     "outset.  Agile allows new needs (e.g., WhatsApp integration, GCC-market targeting) "
     "to be incorporated without rework of earlier phases."),
    ("Early Value Delivery",
     "Core workflows (registration → brand approval → product listing → order) were "
     "delivered first, enabling demo and feedback before peripheral features."),
    ("Risk Reduction",
     "Iterating through authentication, then marketplace, then orders, then trust "
     "means each layer is validated before building on top of it — reducing the risk "
     "of large late-stage defects."),
    ("Team Size",
     "A small university-project team benefits from lightweight Agile practices "
     "(short feedback loops, self-organisation) rather than the formal artefacts "
     "demanded by Waterfall or RUP."),
    ("Technology Uncertainty",
     "Decisions such as which database to finalise (local JSON for MVP vs. Supabase "
     "for production) are intentionally deferred to later sprints, which suits Agile's "
     "philosophy of 'just-enough' up-front design."),
]
for title, desc in reasons:
    add_bullet(doc, desc, bold_prefix=title + " —")

add_h2(doc, "2.3  Sprint Overview")
add_body(doc, "The project has progressed through the following approximate sprints:")

tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0]
for cell, txt in zip(hdr.cells, ["Sprint", "Focus Area", "Key Deliverables"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr, "1A3A6B")

sprints = [
    ("Sprint 1", "Foundation & Auth",
     "Project setup, Next.js App Router scaffold, DB schema design, user registration & login"),
    ("Sprint 2", "Marketplace Core",
     "Brand browsing, product catalog, search/filter, verification badges, multi-currency"),
    ("Sprint 3", "Order Lifecycle",
     "Order creation, approval/rejection flow, order-item tracking, buyer dashboard"),
    ("Sprint 4", "Brand Owner Tools",
     "Brand dashboard, product CRUD, bulk-pricing tiers, Excel import, WhatsApp integration"),
    ("Sprint 5", "Trust & Reputation",
     "Review system, rating aggregation, brand trust metrics, buyer trust metrics, admin moderation"),
    ("Sprint 6", "Admin & Notifications",
     "Admin dashboard, pending approvals, user management, notification panel"),
    ("Sprint 7 (Current)", "Localisation & Polish",
     "Full Arabic/RTL support, translation audit scripts, mobile responsiveness, Supabase auth"),
    ("Sprint 8 (Planned)", "Production Readiness",
     "Supabase DB migration, real-time notifications, payment integration scaffolding, testing"),
]
for s in sprints:
    row = tbl.add_row()
    for cell, val in zip(row.cells, s):
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — PROGRESS OF ALL ACTIVITIES
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "3.  Progress of All Activities")
add_divider(doc)

# ── 3.1 Requirements ──────────────────────────────────────────────────────────
add_h2(doc, "3.1  Requirements & Specifications")
add_h3(doc, "3.1.1  Status of Requirements")
add_body(doc,
    "Core functional and non-functional requirements have been elicited, discussed, and "
    "reflected in the codebase.  The requirements are considered stable for the current "
    "MVP scope; however, the team has not yet produced a standalone, formally reviewed "
    "Software Requirements Specification (SRS) document.  All requirements are currently "
    "traceable through API endpoint definitions, TypeScript interface contracts, and "
    "implemented UI flows."
)

add_h3(doc, "3.1.2  Requirements Traceability Matrix (RTM)")
add_body(doc,
    "The table below maps each functional requirement to its design artefact, "
    "implementation status, and test coverage."
)

rtm_cols = ["Req. ID", "Requirement Description", "Design Artefact", "Implementation Status", "Test Status"]
rtm_data = [
    ("FR-01", "User registration with role selection", "DB schema: users table; /register page", "Complete", "Manual"),
    ("FR-02", "Secure login / logout with session cookies", "/api/auth/login, logout, me", "Complete", "Manual"),
    ("FR-03", "Role-based route protection", "Next.js middleware + server-side auth checks", "Complete", "Manual"),
    ("FR-04", "Browse and search brands (filter by category, MOQ, rating)", "/brands page, search API", "Complete", "Manual"),
    ("FR-05", "Multi-currency price display (KWD/SAR/AED/BHD)", "CurrencyContext, exchange-rate config", "Complete", "Manual"),
    ("FR-06", "Product catalog management (CRUD + bulk import)", "/api/products, brand product pages, XLSX import", "Complete", "Manual"),
    ("FR-07", "Bulk-pricing tiers per product", "Product schema: bulk_pricing_tiers[]", "Complete", "Manual"),
    ("FR-08", "Order creation with MOQ & stock validation", "/api/orders, order-create page", "Complete", "Manual"),
    ("FR-09", "Order approval / rejection by brand owner", "PATCH /api/orders/:id", "Complete", "Manual"),
    ("FR-10", "Order status lifecycle tracking", "Order status enum + buyer dashboard", "Complete", "Manual"),
    ("FR-11", "Brand verification workflow (admin approval)", "/api/admin/brands, admin dashboard", "Complete", "Manual"),
    ("FR-12", "Verification tier badges (Premium/Verified/New)", "Brand schema: verification_tier field", "Complete", "Manual"),
    ("FR-13", "WhatsApp click tracking", "/api/brands/whatsapp-click", "Complete", "Manual"),
    ("FR-14", "Buyer review submission (1–5 stars)", "/api/reviews POST", "Complete", "Manual"),
    ("FR-15", "Brand trust score calculation", "BrandTrust schema, /api/trust", "Partial (70%)", "Not started"),
    ("FR-16", "Admin review moderation queue", "/api/admin/reviews, admin dashboard", "Complete", "Manual"),
    ("FR-17", "In-app notification panel (read/unread)", "Zustand store, /api/notifications", "Complete", "Manual"),
    ("FR-18", "Bilingual UI (English / Arabic RTL)", "next-intl, en.json / ar.json, RTL CSS", "Partial (80%)", "Manual"),
    ("FR-19", "GCC multi-market brand targeting", "Brand schema: areas_served field", "Partial (75%)", "Not started"),
    ("FR-20", "Admin analytics dashboard", "/admin/dashboard, chart components", "Partial (75%)", "Manual"),
    ("NFR-01", "Mobile-first responsive design", "CSS modules, MobileTopBar, BottomNav", "Partial (70%)", "Manual"),
    ("NFR-02", "SEO optimisation", "JSON-LD structured data, Next.js metadata", "Complete", "Manual"),
    ("NFR-03", "Secure auth (HTTP-only cookies)", "Supabase SSR, cookie-based sessions", "Complete", "Manual"),
    ("NFR-04", "Scalable database (Supabase migration)", "Supabase client scaffolding", "Partial (30%)", "Not started"),
    ("NFR-05", "Real-time notifications (WebSocket/SSE)", "Placeholder hooks exist", "Not started (10%)", "Not started"),
    ("NFR-06", "Payment integration", "Not yet scoped", "Not started", "Not started"),
]

rtm_tbl = doc.add_table(rows=1, cols=5)
rtm_tbl.style = 'Table Grid'
rtm_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set column widths
widths = [Cm(1.5), Cm(5.5), Cm(4.0), Cm(3.2), Cm(2.3)]
for i, w in enumerate(widths):
    for cell in rtm_tbl.columns[i].cells:
        cell.width = w

hdr_row = rtm_tbl.rows[0]
for cell, txt in zip(hdr_row.cells, rtm_cols):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr_row, "1A3A6B")

status_colors = {
    "Complete":    "D6EAF8",
    "Partial":     "FEF9E7",
    "Not started": "FDEDEC",
}
for row_data in rtm_data:
    row = rtm_tbl.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(8)
    # colour by status
    status = row_data[3]
    for key, color in status_colors.items():
        if key.lower() in status.lower():
            shade_row(row, color)
            break

doc.add_page_break()

# ── 3.2 Design ────────────────────────────────────────────────────────────────
add_h2(doc, "3.2  Design")
add_h3(doc, "3.2.1  Architectural Overview")
add_body(doc,
    "Kuwait B2B Hub is built on a modern full-stack Next.js 16 (App Router) architecture "
    "deployed on Vercel.  The application follows a layered structure:"
)
arch_layers = [
    ("Presentation Layer", "React/TSX pages and components, CSS Modules, Lucide icons, "
     "Recharts/Chart.js for analytics."),
    ("Application / API Layer", "Next.js Route Handlers providing a RESTful JSON API "
     "under /api/*.  Authentication middleware enforces role-based access."),
    ("Data Access Layer", "A centralised db.ts module abstracts all reads and writes.  "
     "Currently backed by a local db.json file for the MVP; Supabase PostgreSQL is the "
     "target production store."),
    ("External Services", "Supabase (authentication + future DB), WhatsApp Business links, "
     "Vercel hosting."),
]
for title, desc in arch_layers:
    add_bullet(doc, desc, bold_prefix=title + " —")

add_h3(doc, "3.2.2  Database Schema (Entity Descriptions)")
add_body(doc,
    "The following entities have been designed and implemented in the data layer:"
)
entities = [
    ("Users", "id, name, email, password (hashed), role, company_name, verification_status, created_at"),
    ("Brands", "id, owner_id, brand_name, description, logo_url, status, verification_tier, "
               "whatsapp_number, business_hours, auto_reply_message, whatsapp_clicks, created_at"),
    ("Products", "id, brand_id, name, description, price, moq, stock, image_url, bulk_pricing_tiers[], created_at"),
    ("Orders", "id, buyer_id, brand_id, status, total_amount, created_at"),
    ("Order Items", "id, order_id, product_id, quantity, unit_price"),
    ("Reviews", "id, brand_id, order_id, buyer_id, rating, content, anonymous, flagged, brand_reply, status, created_at"),
    ("Notifications", "id, user_id, type, title, body, read, action_url, icon_type, created_at"),
    ("Brand Trust", "brand_id, response_rate, completion_rate, avg_response_hours, total_fulfilled, avg_fulfillment_days, badges[]"),
    ("Buyer Trust", "buyer_id, total_orders, completion_rate, cancellation_rate, badges[]"),
    ("Messages", "id, order_id, sender_id, content, created_at"),
]
for name, fields in entities:
    add_bullet(doc, fields, bold_prefix=name + ":")

add_h3(doc, "3.2.3  Key UML & Design Artefacts")
add_body(doc,
    "The following UML models have been derived from the implemented codebase.  "
    "Formal diagram files will be produced as part of the final report."
)
diagrams = [
    ("Use-Case Diagram",
     "Three actors (Admin, Brand Owner, Buyer) interact with use cases covering "
     "registration, brand management, product management, order management, review "
     "management, and notification handling."),
    ("Class Diagram",
     "Ten entity classes (User, Brand, Product, Order, OrderItem, Review, Notification, "
     "BrandTrust, BuyerTrust, Message) with their attributes and associations "
     "as reflected in the TypeScript interface definitions."),
    ("Sequence Diagram — Order Flow",
     "Buyer browses catalog → selects product → submits order request → Brand Owner "
     "reviews → approves/rejects → system sends notification → Buyer receives update."),
    ("Sequence Diagram — Brand Verification",
     "Brand Owner registers → Admin receives notification → Admin reviews brand → "
     "Approves (sets tier) or Rejects → Brand Owner notified."),
    ("Component Diagram",
     "Identifies Next.js frontend pages, API route handlers, the db.ts data-access "
     "module, Zustand state store, CurrencyContext, and Supabase client as principal components."),
    ("Deployment Diagram",
     "Vercel edge network hosts Next.js app; Supabase hosts auth and future DB; "
     "WhatsApp Business API integrated via external links."),
]
for name, desc in diagrams:
    add_bullet(doc, desc, bold_prefix=name + " —")

add_h3(doc, "3.2.4  Key Architectural Decisions")
decisions = [
    "Next.js App Router was chosen to support SSR for SEO, server-side auth guards, and co-located API routes.",
    "JSON file database for the MVP phase enables rapid iteration without infrastructure overhead; the Supabase migration path is already scaffolded.",
    "Zustand was selected for lightweight client-side notification state to avoid the overhead of Redux in a small-team project.",
    "next-intl provides server-side internationalisation, ensuring Arabic RTL content is rendered correctly on first load.",
    "Multi-currency support is implemented via a React Context with hardcoded exchange rates for the MVP; a live-rates API call is planned.",
]
for d in decisions:
    add_bullet(doc, d)

doc.add_page_break()

# ── 3.3 Implementation ────────────────────────────────────────────────────────
add_h2(doc, "3.3  Implementation")
add_h3(doc, "3.3.1  Overall Completion Summary")

impl_tbl = doc.add_table(rows=1, cols=4)
impl_tbl.style = 'Table Grid'
impl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_row = impl_tbl.rows[0]
for cell, txt in zip(hdr_row.cells, ["Feature Module", "Completion %", "Status", "Remaining Work"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr_row, "1A3A6B")

impl_rows = [
    ("Authentication & Session Management", "95%", "Near-Complete",
     "Full Supabase DB integration, password hashing validation"),
    ("Marketplace Browsing & Search",       "85%", "Mostly Complete",
     "Advanced filters, wishlist, brand comparison"),
    ("Product & Catalog Management",        "90%", "Near-Complete",
     "Inventory alerts, image upload (currently URL-based)"),
    ("Order Lifecycle Management",          "80%", "Mostly Complete",
     "Order messaging/chat, bulk order actions"),
    ("Brand Verification Workflow",         "85%", "Mostly Complete",
     "Document upload, automated CR validation"),
    ("Trust & Reputation System",           "70%", "In Progress",
     "Automated trust-score computation, badge logic"),
    ("Admin Dashboard",                     "75%", "In Progress",
     "Granular permissions, data export, reporting"),
    ("In-App Notifications",                "60%", "In Progress",
     "Real-time WebSocket/SSE, email & SMS notifications"),
    ("Bilingual UI (EN/AR RTL)",            "80%", "In Progress",
     "Complete all missing translation keys"),
    ("GCC Multi-Market Support",            "75%", "In Progress",
     "Market-specific shipping rules"),
    ("Mobile Responsiveness",               "70%", "In Progress",
     "Touch gestures, mobile-specific UX improvements"),
    ("Multi-Currency Support",              "85%", "Mostly Complete",
     "Live exchange-rate API, historical pricing"),
    ("Supabase DB Migration",               "30%", "Early Stage",
     "Full data-layer migration from JSON to PostgreSQL"),
    ("Real-Time Features",                  "10%", "Not Started",
     "WebSocket/Supabase Realtime for notifications & chat"),
    ("Payment Integration",                 "0%",  "Not Started",
     "End-to-end payment gateway integration"),
    ("Shipping Integration",                "0%",  "Not Started",
     "Carrier API, logistics tracking"),
    ("Email / SMS Notifications",           "0%",  "Not Started",
     "Transactional email and SMS via third-party provider"),
]

color_map = {
    "Near-Complete":   "D6EAF8",
    "Mostly Complete": "D5F5E3",
    "In Progress":     "FEF9E7",
    "Early Stage":     "FAD7A0",
    "Not Started":     "FDEDEC",
}

for row_data in impl_rows:
    row = impl_tbl.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
    shade_row(row, color_map.get(row_data[2], "FFFFFF"))

add_body(doc, "")
add_body(doc,
    "Estimated overall platform completion (MVP scope): approximately 72%.  "
    "Core business workflows are functional end-to-end.  Production-readiness features "
    "(Supabase migration, payments, real-time, email) constitute the primary remaining work."
)

add_h3(doc, "3.3.2  Technology Stack")
stack = [
    ("Frontend Framework", "Next.js 16.2.2 with App Router, TypeScript 5"),
    ("Styling",            "CSS Modules (no Tailwind), custom responsive CSS"),
    ("UI Libraries",       "Lucide React (icons), Recharts / Chart.js (analytics)"),
    ("State Management",   "Zustand 5.0.12 (notifications), React Context (currency)"),
    ("Internationalisation","next-intl 4.9.0 — English & Arabic with full RTL support"),
    ("Authentication",     "Supabase Auth with SSR session management"),
    ("Database (MVP)",     "JSON file (db.json) via custom db.ts abstraction layer"),
    ("Database (Target)",  "Supabase PostgreSQL"),
    ("Data Import",        "XLSX 0.18.5 for Excel/CSV bulk product import"),
    ("Hosting",            "Vercel (Next.js native deployment)"),
]
for name, desc in stack:
    add_bullet(doc, desc, bold_prefix=name + ":")

add_h3(doc, "3.3.3  Codebase Metrics")
metrics = [
    "Pages implemented: 15+",
    "API endpoints: 25+",
    "Reusable React components: 20+",
    "Languages supported: 2 (English, Arabic)",
    "Currencies supported: 4 (KWD, SAR, AED, BHD)",
    "External runtime dependencies: 9 principal packages",
    "Translation scripts (Python): 3 (find_missing_keys.py, audit_translations.py, find_all_missing_keys.py)",
    "Lines of TypeScript/TSX: ~3,000+",
]
for m in metrics:
    add_bullet(doc, m)

doc.add_page_break()

# ── 3.4 V&V / Testing ────────────────────────────────────────────────────────
add_h2(doc, "3.4  Verification & Validation / Testing")
add_h3(doc, "3.4.1  Current Testing Approach")
add_body(doc,
    "The project is currently in the MVP phase and formal automated testing infrastructure "
    "has not yet been established.  Verification and validation has been conducted primarily "
    "through manual end-to-end testing against seeded demo data.  ESLint is configured for "
    "static code analysis.  Three Python utility scripts perform translation-key auditing to "
    "validate the completeness of localisation files."
)

add_h3(doc, "3.4.2  Manual Test Coverage Summary")
test_tbl = doc.add_table(rows=1, cols=4)
test_tbl.style = 'Table Grid'
test_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_row = test_tbl.rows[0]
for cell, txt in zip(hdr_row.cells, ["Test Area", "Test Type", "Result", "Notes"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr_row, "1A3A6B")

test_rows = [
    ("User Registration (all roles)",       "Manual / Functional", "Pass", "Tested with 3 roles"),
    ("Login / Logout",                      "Manual / Functional", "Pass", "HTTP-only cookie verified"),
    ("Role-based route protection",         "Manual / Functional", "Pass", "Redirect on unauthorised access"),
    ("Brand creation on owner registration","Manual / Functional", "Pass", "Auto-brand creation confirmed"),
    ("Marketplace brand browsing",          "Manual / Functional", "Pass", "Filters and search work"),
    ("Product CRUD",                        "Manual / Functional", "Pass", "Create, edit, delete confirmed"),
    ("Excel bulk product import",           "Manual / Functional", "Pass", "Tested with sample XLSX"),
    ("Order submission (MOQ validation)",   "Manual / Functional", "Pass", "Below-MOQ orders blocked"),
    ("Order approval / rejection",          "Manual / Functional", "Pass", "Status transitions correct"),
    ("Brand verification (admin)",          "Manual / Functional", "Pass", "Approval and rejection flow tested"),
    ("Review submission",                   "Manual / Functional", "Pass", "Ratings persisted correctly"),
    ("Anonymous review",                    "Manual / Functional", "Pass", "Author hidden in UI"),
    ("Brand reply to review",               "Manual / Functional", "Pass", "Reply persisted and displayed"),
    ("Notification generation",             "Manual / Functional", "Pass", "Notifications appear in panel"),
    ("Currency switching",                  "Manual / Functional", "Pass", "Prices update correctly"),
    ("Arabic RTL layout",                   "Manual / UI",         "Partial", "Some components not yet RTL-adapted"),
    ("Mobile responsiveness",               "Manual / UI",         "Partial", "Main flows work; edge cases pending"),
    ("Admin analytics dashboard",           "Manual / Functional", "Partial", "Charts render; some data gaps"),
    ("Real-time notifications",             "N/A — Not implemented","N/A",   "Planned for Sprint 8"),
    ("Payment flow",                        "N/A — Not implemented","N/A",   "Planned for future sprint"),
    ("Unit tests (automated)",              "Not yet written",     "N/A",   "Planned with Jest + React Testing Library"),
    ("Integration tests (automated)",       "Not yet written",     "N/A",   "Planned with Playwright"),
    ("Performance / load testing",          "Not yet conducted",   "N/A",   "Planned pre-launch"),
]
for row_data in test_rows:
    row = test_tbl.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
    result = row_data[2]
    if result == "Pass":
        shade_row(row, "D5F5E3")
    elif result == "Partial":
        shade_row(row, "FEF9E7")
    elif result == "N/A":
        shade_row(row, "F2F3F4")

add_h3(doc, "3.4.3  Testing Plan — Upcoming Sprints")
testing_plan = [
    "Introduce Jest + React Testing Library for unit tests of utility functions (db.ts helpers, currency converters, trust-score calculators).",
    "Introduce Playwright for end-to-end integration tests covering the critical paths: registration → order placement → approval → review.",
    "Conduct structured manual system testing against a formal test-case document before final submission.",
    "Address the password hashing vulnerability: migrate from plain-text JSON storage to Supabase Auth (bcrypt hashing) before system testing.",
    "Performance baseline test after Supabase migration to confirm acceptable page-load times under simulated concurrent users.",
]
for item in testing_plan:
    add_bullet(doc, item)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — PROGRESS OF THE FINAL REPORT AND PRESENTATION
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "4.  Progress of the Final Report and Presentation")
add_divider(doc)

add_h2(doc, "4.1  Final Report Status")
report_items = [
    ("Title Page & Abstract",       "Complete",    "Written and reviewed"),
    ("System Overview",             "Complete",    "Based on implemented platform"),
    ("Requirements (SRS)",          "In Progress", "Functional requirements extracted from code; formal SRS document being authored"),
    ("Design Chapter (UML)",        "In Progress", "Architecture described; formal UML diagrams being drawn in draw.io / Lucidchart"),
    ("Implementation Chapter",      "In Progress", "Technology stack and structure documented; screenshots pending"),
    ("Testing Chapter",             "In Progress", "Manual test results documented; automated tests section being drafted"),
    ("Progress Report (this doc)",  "Complete",    "Current document"),
    ("Conclusion & Future Work",    "Not Started", "Planned for final sprint"),
    ("References & Bibliography",   "Not Started", "To be compiled from library sources"),
    ("Appendices (code listings)",  "Not Started", "Key code snippets to be included"),
]

rep_tbl = doc.add_table(rows=1, cols=3)
rep_tbl.style = 'Table Grid'
rep_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_row = rep_tbl.rows[0]
for cell, txt in zip(hdr_row.cells, ["Report Section", "Status", "Notes"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr_row, "1A3A6B")

for row_data in report_items:
    row = rep_tbl.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
    status = row_data[1]
    if status == "Complete":
        shade_row(row, "D5F5E3")
    elif status == "In Progress":
        shade_row(row, "FEF9E7")
    else:
        shade_row(row, "FDEDEC")

add_h2(doc, "4.2  Presentation Status")
pres_items = [
    ("Slide deck structure / outline",         "Complete",    "Introduction, demo, architecture, challenges, plan forward"),
    ("Live demo preparation (demo account)",   "In Progress", "Demo data seeded; demo script being rehearsed"),
    ("Architecture & design slides",           "In Progress", "Diagrams being finalised for inclusion"),
    ("Implementation highlights slides",       "In Progress", "Screenshots and code snippets being curated"),
    ("Challenges & lessons-learned slide",     "Not Started", "Content ready; slide design pending"),
    ("Q&A preparation",                        "Not Started", "Anticipated questions being compiled"),
    ("Rehearsal (full run-through)",           "Not Started", "Scheduled for Sprint 8"),
]

pres_tbl = doc.add_table(rows=1, cols=3)
pres_tbl.style = 'Table Grid'
pres_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_row = pres_tbl.rows[0]
for cell, txt in zip(hdr_row.cells, ["Presentation Component", "Status", "Notes"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr_row, "1A3A6B")

for row_data in pres_items:
    row = pres_tbl.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
    status = row_data[1]
    if status == "Complete":
        shade_row(row, "D5F5E3")
    elif status == "In Progress":
        shade_row(row, "FEF9E7")
    else:
        shade_row(row, "FDEDEC")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — CHALLENGES
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "5.  Challenges and How the Team Is Addressing Them")
add_divider(doc)

challenges = [
    (
        "Challenge 1: Database Scalability — JSON File Backing Store",
        "The MVP uses a flat db.json file as its database.  While sufficient for development "
        "and seeded demo data, this approach does not support concurrent writes, lacks transaction "
        "safety, and will not scale to real GCC market traffic.",
        "The team has already scaffolded the Supabase client and authentication.  Sprint 8 is "
        "dedicated to migrating all data-access calls in db.ts to Supabase PostgreSQL queries.  "
        "The abstraction layer in db.ts means the API route handlers will require minimal changes."
    ),
    (
        "Challenge 2: Incomplete Supabase Integration",
        "Authentication currently operates across two systems — some flows use the Supabase Auth "
        "SDK while others fall back to the local JSON store — creating inconsistency and potential "
        "session conflicts.",
        "The team is consolidating authentication entirely through Supabase Auth, removing all "
        "local-credential code paths, and standardising session management via Next.js middleware "
        "before the final submission."
    ),
    (
        "Challenge 3: Internationalisation (i18n) Completeness",
        "Maintaining full parity between the English and Arabic translation files is labour-"
        "intensive.  New UI features introduce English strings that may not immediately have "
        "Arabic equivalents, resulting in fallback rendering.",
        "Three Python utility scripts (find_missing_keys.py, audit_translations.py, "
        "find_all_missing_keys.py) have been written to automatically detect and report "
        "missing translation keys.  The team runs these scripts before each sprint review "
        "and fills gaps iteratively."
    ),
    (
        "Challenge 4: Security — Plain-Text Credentials in JSON Store",
        "Because the MVP database is a JSON file, user passwords are not cryptographically "
        "hashed.  This is a critical security vulnerability that must be resolved before "
        "any public exposure.",
        "Migrating to Supabase Auth (Sprint 8) inherently resolves this: Supabase stores "
        "only bcrypt hashes and never exposes raw passwords.  No user-facing passwords "
        "will be stored in the application layer."
    ),
    (
        "Challenge 5: Absence of Automated Testing",
        "The project currently lacks a formal automated test suite, meaning regressions "
        "can go undetected when new features are added, and the V&V chapter of the final "
        "report cannot reference automated results.",
        "The team has planned dedicated time in Sprint 8 to introduce Jest unit tests for "
        "utility modules and Playwright end-to-end tests for the three critical paths "
        "(registration, order placement, and brand approval).  Manual test-case documents "
        "are being drafted in parallel."
    ),
    (
        "Challenge 6: Real-Time Communication Infrastructure",
        "The notification and messaging systems are designed for real-time delivery but are "
        "currently backed only by polling / static reads.  True WebSocket or Server-Sent "
        "Events infrastructure has not yet been implemented.",
        "The Supabase Realtime service will be used post-migration.  The existing Zustand "
        "notification store and hook placeholders (useRealtimeNotifications) are already "
        "structured to accept a real-time data source with minimal refactoring."
    ),
    (
        "Challenge 7: Payment and Logistics Integration Scope",
        "Payment processing and shipping carrier integration are entirely absent from the "
        "current build.  These are critical features for a production B2B marketplace but "
        "represent significant scope that was intentionally deferred.",
        "The team has agreed to deliver a clearly documented future-work roadmap that scopes "
        "these integrations (Checkout.com or Tap Payments for payments; Aramex / DHL for "
        "logistics), with API selection and a high-level integration architecture included "
        "in the final report's conclusion chapter."
    ),
]

for i, (title, problem, solution) in enumerate(challenges):
    add_h2(doc, title)
    add_body(doc, problem, bold_prefix="Problem:")
    add_body(doc, solution, bold_prefix="Resolution:")
    if i < len(challenges) - 1:
        doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — PLAN FORWARD
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "6.  The Plan Forward")
add_divider(doc)

add_h2(doc, "6.1  Upcoming Milestones")
add_body(doc,
    "The following milestones define the path from the current state to final submission. "
    "Dates are approximate and subject to adjustment based on sprint velocity."
)

milestone_tbl = doc.add_table(rows=1, cols=4)
milestone_tbl.style = 'Table Grid'
milestone_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_row = milestone_tbl.rows[0]
for cell, txt in zip(hdr_row.cells, ["Milestone", "Target Date", "Owner", "Success Criteria"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr_row, "1A3A6B")

milestones = [
    ("Supabase full auth consolidation",
     "Late April 2026", "Full Team",
     "All auth flows use Supabase; no local-JSON credential logic remains"),
    ("Missing translation keys resolved",
     "Late April 2026", "Frontend Lead",
     "Audit scripts report 0 missing keys in both EN and AR"),
    ("Supabase database migration (all entities)",
     "Early May 2026", "Backend Lead",
     "All CRUD operations use Supabase PostgreSQL; db.json deprecated"),
    ("Real-time notifications via Supabase Realtime",
     "Early May 2026", "Backend Lead",
     "Notifications delivered within 2 seconds of trigger event"),
    ("Automated unit tests (utility modules)",
     "Mid-May 2026", "Full Team",
     "Jest tests cover db helpers, currency converters, trust-score logic"),
    ("Automated E2E tests (3 critical paths)",
     "Mid-May 2026", "Full Team",
     "Playwright tests pass for registration, order, brand-approval flows"),
    ("Mobile responsiveness finalised",
     "Mid-May 2026", "Frontend Lead",
     "All pages pass manual mobile review on 375 px and 768 px viewports"),
    ("Final report first full draft",
     "Late May 2026", "Full Team",
     "All chapters drafted and reviewed internally"),
    ("Presentation slides complete",
     "Early June 2026", "Full Team",
     "Slides reviewed; live-demo rehearsal conducted"),
    ("Final report submission",
     "Mid-June 2026", "Full Team",
     "Report submitted to university portal"),
    ("Final presentation / viva",
     "Late June 2026", "Full Team",
     "Presentation delivered; Q&A handled"),
]
for row_data in milestones:
    row = milestone_tbl.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)

add_h2(doc, "6.2  Sprint 8 Detailed Work Breakdown")
add_body(doc,
    "Sprint 8 is the final active development sprint and focuses on production-readiness "
    "and academic deliverable completion.  Key tasks:"
)
sprint8 = [
    "Migrate db.ts data-access layer from JSON to Supabase PostgreSQL (all 10 entity types).",
    "Consolidate Supabase Auth — remove all local JSON authentication code paths.",
    "Implement Supabase Realtime subscription for notifications.",
    "Write Jest unit tests for: currency converters, trust-score calculations, order-validation helpers.",
    "Write Playwright E2E tests for: new-user registration, buyer order placement, admin brand approval.",
    "Resolve all missing i18n translation keys (guided by audit scripts).",
    "Complete mobile-responsiveness review and fix remaining layout issues.",
    "Author the Conclusion & Future Work chapter for the final report.",
    "Curate code screenshots and UML diagrams for report appendices.",
    "Full presentation rehearsal with timed run-through.",
]
for task in sprint8:
    add_bullet(doc, task)

add_h2(doc, "6.3  Risk Register")
risk_tbl = doc.add_table(rows=1, cols=4)
risk_tbl.style = 'Table Grid'
risk_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr_row = risk_tbl.rows[0]
for cell, txt in zip(hdr_row.cells, ["Risk", "Likelihood", "Impact", "Mitigation"]):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
shade_row(hdr_row, "1A3A6B")

risks = [
    ("Supabase migration takes longer than estimated",
     "Medium", "High",
     "Start migration in parallel sub-branches; JSON fallback remains available"),
    ("Automated tests surface unexpected bugs late",
     "Medium", "Medium",
     "Prioritise critical-path tests first; accept partial coverage for submission"),
    ("Scope creep — new features requested during final sprint",
     "Low", "Medium",
     "Feature freeze applied; new items logged in backlog only"),
    ("Team member unavailability near submission",
     "Low", "High",
     "Shared knowledge across all modules; documentation kept up to date"),
    ("Real-time implementation complexity exceeds estimate",
     "Medium", "Low",
     "Supabase Realtime is well-documented; fall back to polling if blocked"),
]
for row_data in risks:
    row = risk_tbl.add_row()
    for i, val in enumerate(row_data):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
    likelihood = row_data[1]
    if likelihood == "High":
        shade_row(row, "FDEDEC")
    elif likelihood == "Medium":
        shade_row(row, "FEF9E7")
    else:
        shade_row(row, "D5F5E3")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CLOSING
# ══════════════════════════════════════════════════════════════════════════════
add_h1(doc, "Summary")
add_divider(doc)
add_body(doc,
    "Kuwait B2B Hub is a purpose-built wholesale marketplace for the Kuwait and GCC market, "
    "currently approximately 72% complete at the MVP feature level.  All core business "
    "workflows — authentication, marketplace browsing, product management, order lifecycle, "
    "brand verification, trust and reputation, and in-app notifications — are implemented "
    "and manually tested.  The primary outstanding work centres on database migration to "
    "Supabase, real-time infrastructure, automated testing, and the academic deliverables "
    "(final report and presentation)."
)
add_body(doc,
    "The team is following an Agile sprint model, is on track for final submission, and "
    "has clear mitigation strategies for all identified risks.  The project demonstrates "
    "a solid understanding of full-stack web development, bilingual and multi-currency design, "
    "trust-centric marketplace architecture, and GCC-market-specific requirements."
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— End of Progress Report —")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# ── Save ─────────────────────────────────────────────────────────────────────
out_path = "/Users/baderalnoumas/Desktop/B2B/Progress_Report_Final.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
