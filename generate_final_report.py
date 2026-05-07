"""
Kuwait B2B Hub — Final Project Report Generator
Produces Final_Project_Report.docx following the university template structure.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
import copy

# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════════════════════
DARK_BLUE   = RGBColor(0x1A, 0x3A, 0x6B)
MID_BLUE    = RGBColor(0x2E, 0x6B, 0xA0)
GOLD        = RGBColor(0xD4, 0xA8, 0x47)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
GRAY        = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY  = RGBColor(0x88, 0x88, 0x88)
BLACK       = RGBColor(0x00, 0x00, 0x00)

HEX_DKBLUE  = "1A3A6B"
HEX_MBLUE   = "2E6BA0"
HEX_LBLUE   = "D6EAF8"
HEX_GREEN   = "D5F5E3"
HEX_YELLOW  = "FEF9E7"
HEX_RED     = "FDEDEC"
HEX_ORANGE  = "FAD7A0"
HEX_GRAY    = "F2F3F4"
HEX_WHITE   = "FFFFFF"

doc = Document()

# ── Page Setup ─────────────────────────────────────────────────────────────────
for s in doc.sections:
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.5)

# ═══════════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def shade_row(row, hex_color):
    for cell in row.cells:
        shade_cell(cell, hex_color)

def add_bottom_border(para, color=HEX_DKBLUE, sz='6'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_cell_text(cell, text, bold=False, size=10, color=None, align=None):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

# ── Heading helpers ───────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = DARK_BLUE
    add_bottom_border(p, HEX_DKBLUE, '8')
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = MID_BLUE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = DARK_BLUE
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def body_bold(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(label + ' ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        r1.bold = True
        r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' ')
        r1.bold = True
        r1.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_bottom_border(p, 'AAAAAA', '4')

def make_table(headers, rows, col_widths=None, header_bg=HEX_DKBLUE, alt_bg=HEX_GRAY,
               font_size=10, header_font_size=10):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    hdr = tbl.rows[0]
    shade_row(hdr, header_bg)
    for cell, txt in zip(hdr.cells, headers):
        set_cell_text(cell, txt, bold=True, size=header_font_size, color=WHITE)
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        bg = alt_bg if ri % 2 == 1 else HEX_WHITE
        shade_row(row, bg)
        for ci, val in enumerate(row_data):
            set_cell_text(row.cells[ci], str(val), size=font_size)
    return tbl

def make_colored_table(headers, rows, row_colors, col_widths=None, font_size=10):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    hdr = tbl.rows[0]
    shade_row(hdr, HEX_DKBLUE)
    for cell, txt in zip(hdr.cells, headers):
        set_cell_text(cell, txt, bold=True, size=font_size+1, color=WHITE)
    for row_data, color in zip(rows, row_colors):
        row = tbl.add_row()
        shade_row(row, color)
        for ci, val in enumerate(row_data):
            set_cell_text(row.cells[ci], str(val), size=font_size)
    return tbl

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def cover_page():
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COLLEGE OF ENGINEERING & INFORMATION TECHNOLOGY")
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Computer Science")
    r.font.size = Pt(12); r.font.color.rgb = GRAY

    divider()

    for _ in range(2):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kuwait B2B Hub")
    r.bold = True; r.font.size = Pt(32); r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Digital Wholesale Marketplace for Kuwait & the GCC Region")
    r.font.size = Pt(15); r.font.color.rgb = MID_BLUE; r.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Final Project Report")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = GOLD

    divider()

    for _ in range(3):
        doc.add_paragraph()

    meta = [
        ("Course:",         "Software Engineering — COMP 491"),
        ("Report Type:",    "Final Project Report"),
        ("Submission Date:","April 17, 2026"),
        ("Advisor:",        "University Supervisor"),
    ]
    for lbl, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{lbl}  "); r1.bold = True; r1.font.size = Pt(12)
        r2 = p.add_run(val);        r2.font.size = Pt(12)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by:")
    r.bold = True; r.font.size = Pt(12)

    members = [
        "Bader Alnoumas",
        "Team Member 2",
        "Team Member 3",
        "Team Member 4",
    ]
    for m in members:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(m); r.font.size = Pt(11)

    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (manual)
# ═══════════════════════════════════════════════════════════════════════════════

def toc():
    h1("Table of Contents")
    toc_entries = [
        ("Abstract", "5"),
        ("Chapter I — Introduction", "6"),
        ("  1.1  Background", "6"),
        ("  1.2  System Overview", "6"),
        ("  1.3  Glossary", "7"),
        ("  1.4  Readership", "7"),
        ("  1.5  Product Scope", "8"),
        ("  1.6  Team Workload", "8"),
        ("Chapter II — Software Requirements Specification (SRS)", "9"),
        ("  2.1  Feasibility Study", "9"),
        ("  2.2  User Requirements", "9"),
        ("  2.3  System Requirements", "10"),
        ("  2.4  Non-Functional Requirements", "11"),
        ("  2.5  Functional Requirements Specification (FRS)", "11"),
        ("  2.6  Context Model", "12"),
        ("  2.7  Use Cases", "12"),
        ("  2.8  Requirements Traceability Matrix (RTM)", "13"),
        ("Chapter III — System Design", "15"),
        ("  3.1  Sequence Diagrams", "15"),
        ("  3.2  Data Flow Diagram (DFD)", "16"),
        ("  3.3  System Architecture", "16"),
        ("  3.4  Activity Model", "17"),
        ("  3.5  Class Diagram", "17"),
        ("  3.6  Prototype Screenshots", "18"),
        ("Chapter IV — System Implementation", "19"),
        ("  4.1  Make / Buy / Lease Analysis", "19"),
        ("  4.2  Coding Artefacts", "19"),
        ("  4.3  Configuration Management", "20"),
        ("  4.4  Deployment Specification", "20"),
        ("  4.5  Physical Deployment Diagram", "21"),
        ("Chapter V — Verification & Validation", "22"),
        ("  5.1  Development Testing", "22"),
        ("  5.2  System Inspection", "22"),
        ("  5.3  System Testing", "23"),
        ("  5.4  Release Testing", "23"),
        ("  5.5  User Testing", "24"),
        ("Chapter VI — System Evolution", "25"),
        ("  6.1  Anticipated Changes", "25"),
        ("  6.2  Phase-Out Transitions", "25"),
        ("  6.3  Quality Predictions", "26"),
        ("Chapter VII — System Planning & Management", "27"),
        ("  7.1  Milestones", "27"),
        ("  7.2  Gantt Chart", "27"),
        ("  7.3  Staff Allocation", "28"),
        ("  7.4  CoCoMo Cost Estimation", "28"),
        ("  7.5  Risk Matrix", "29"),
        ("  7.6  Fishbone Diagram", "29"),
        ("Chapter VIII — Conclusions", "30"),
        ("  8.1  Design Smells", "30"),
        ("  8.2  SWOT Analysis", "30"),
        ("  8.3  Future Work", "31"),
        ("Chapter IX — Appendices", "32"),
        ("Chapter X — References", "33"),
    ]
    for entry, pg in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        tab_stops = p.paragraph_format.tab_stops
        # right-aligned page number
        p.paragraph_format.space_before = Pt(0)
        indent = entry.startswith("  ")
        r = p.add_run(entry)
        r.font.size = Pt(11)
        if not indent:
            r.bold = True
        r2 = p.add_run(f"\t{pg}")
        r2.font.size = Pt(11)
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════════════════════════

def list_of_tables():
    h1("List of Tables")
    tables = [
        ("Table 1", "Team Workload Distribution", "8"),
        ("Table 2", "Feasibility Analysis Summary", "9"),
        ("Table 3", "User Requirements", "10"),
        ("Table 4", "System Functional Requirements", "10"),
        ("Table 5", "Non-Functional Requirements", "11"),
        ("Table 6", "Use Case Summary", "12"),
        ("Table 7", "Requirements Traceability Matrix (RTM)", "13"),
        ("Table 8", "Make / Buy / Lease Analysis", "19"),
        ("Table 9", "Configuration Management", "20"),
        ("Table 10", "Deployment Specification", "21"),
        ("Table 11", "System Test Cases & Results", "23"),
        ("Table 12", "Anticipated System Changes", "25"),
        ("Table 13", "Project Milestones", "27"),
        ("Table 14", "Staff Allocation Matrix", "28"),
        ("Table 15", "CoCoMo Cost Estimation", "28"),
        ("Table 16", "Risk Matrix", "29"),
        ("Table 17", "SWOT Analysis", "30"),
    ]
    for num, title, pg in tables:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{num}. {title}\t{pg}")
        r.font.size = Pt(11)
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════════════════════════════════════

def list_of_figures():
    h1("List of Figures")
    figures = [
        ("Figure 1", "System Architecture Diagram", "16"),
        ("Figure 2", "Context Model (Level 0 DFD)", "12"),
        ("Figure 3", "Level 1 Data Flow Diagram", "16"),
        ("Figure 4", "Use Case Diagram — Buyer", "12"),
        ("Figure 5", "Use Case Diagram — Brand Owner", "12"),
        ("Figure 6", "Use Case Diagram — Admin", "13"),
        ("Figure 7", "Sequence Diagram — User Registration & Login", "15"),
        ("Figure 8", "Sequence Diagram — Order Lifecycle", "15"),
        ("Figure 9", "Sequence Diagram — Brand Verification", "15"),
        ("Figure 10", "Class Diagram — Core Entities", "17"),
        ("Figure 11", "Activity Diagram — Order Placement", "17"),
        ("Figure 12", "Physical Deployment Diagram", "21"),
        ("Figure 13", "Project Gantt Chart", "27"),
        ("Figure 14", "Fishbone (Ishikawa) Diagram — Delivery Delay Risk", "29"),
        ("Figure 15", "Quality Prediction Chart", "26"),
    ]
    for num, title, pg in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{num}. {title}\t{pg}")
        r.font.size = Pt(11)
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

def abstract():
    h1("Abstract")
    body(
        "Kuwait B2B Hub is a full-stack digital wholesale marketplace designed specifically for "
        "the Kuwaiti and Gulf Cooperation Council (GCC) commercial landscape.  The platform "
        "addresses a well-documented inefficiency: the majority of B2B wholesale transactions in "
        "Kuwait are currently conducted informally through WhatsApp groups, Excel spreadsheets, "
        "and phone calls — a process that is error-prone, unscalable, and opaque to both buyers "
        "and suppliers."
    )
    body(
        "This report documents the full software-engineering lifecycle of Kuwait B2B Hub, from "
        "requirements elicitation and feasibility analysis through system design, implementation, "
        "verification and validation, and project management.  The system is implemented using "
        "Next.js 16 (App Router) with TypeScript 5, React 19, Supabase authentication, and a "
        "RESTful JSON API.  It supports three distinct user roles — Platform Admin, Brand Owner, "
        "and Business Buyer — and delivers ten major feature modules including a verified-brand "
        "marketplace, bulk-order management, tiered pricing, a trust-and-reputation engine, "
        "in-app notifications, and full bilingual (English/Arabic RTL) support across four GCC "
        "currencies (KWD, SAR, AED, BHD)."
    )
    body(
        "As of the date of this report, approximately 72% of the planned MVP functionality has "
        "been implemented and manually tested across all critical user flows.  The remaining work "
        "focuses on migrating the data layer to Supabase PostgreSQL, implementing real-time "
        "notifications, adding automated testing, and completing the Arabic localisation.  The "
        "project follows an Agile sprint model and is on track for final submission."
    )
    body(
        "Keywords: B2B marketplace, wholesale, Kuwait, GCC, Next.js, TypeScript, Supabase, "
        "multi-currency, bilingual, trust & reputation, order management."
    )
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER I — INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_1():
    h1("Chapter I:  Introduction")

    h2("1.1  Background")
    body(
        "Kuwait's wholesale distribution sector — valued at several billion KWD annually — has "
        "historically relied on personal relationships, phone calls, and WhatsApp messaging for "
        "order placement and supplier communication.  While these informal channels work for small "
        "volumes, they create significant operational friction as businesses scale: orders are "
        "lost in chat histories, pricing is inconsistent, there is no formal order tracking, and "
        "buyers have no objective way to assess supplier reliability."
    )
    body(
        "The GCC B2B e-commerce segment is undergoing rapid digitalisation, driven by "
        "government Vision initiatives (Kuwait Vision 2035, Saudi Vision 2030) and post-pandemic "
        "acceleration of digital commerce adoption.  However, most existing marketplace platforms "
        "are either consumer-facing (B2C) or serve Western markets without awareness of GCC "
        "regulatory, linguistic, and cultural requirements — in particular, Arabic right-to-left "
        "interfaces, KWD/SAR/AED/BHD multi-currency support, and local commercial-registration "
        "verification processes."
    )
    body(
        "Kuwait B2B Hub was conceived to fill this gap: a purpose-built, bilingual, "
        "GCC-market-aware wholesale marketplace that gives Kuwaiti brands a professional digital "
        "storefront and gives business buyers a trusted, structured ordering experience."
    )

    h2("1.2  System Overview")
    body(
        "Kuwait B2B Hub is a web application built on Next.js 16 with an integrated RESTful API, "
        "deployed on Vercel.  It connects three classes of users through a shared digital "
        "environment:"
    )
    bullet("Platform Admins who govern the ecosystem — verifying brands, moderating content, and monitoring platform health.")
    bullet("Brand Owners (Kuwaiti wholesale suppliers) who list products, manage catalogs, set bulk-pricing tiers, and fulfil orders.")
    bullet("Business Buyers (retail and wholesale buyers across the GCC) who browse verified brands, place bulk orders, and leave reviews.")
    body(
        "The system is structured around five core workflows: (1) identity and access management, "
        "(2) brand verification and catalog management, (3) order lifecycle management, "
        "(4) trust and reputation scoring, and (5) administrative governance.  These are "
        "supported by cross-cutting concerns: internationalisation (EN/AR), multi-currency "
        "conversion, in-app notifications, and SEO."
    )

    h2("1.3  Glossary")
    terms = [
        ("B2B",               "Business-to-Business — commercial transactions between companies."),
        ("MOQ",               "Minimum Order Quantity — the smallest number of units a buyer may order in a single transaction."),
        ("GCC",               "Gulf Cooperation Council — regional bloc comprising Kuwait, Saudi Arabia, UAE, Bahrain, Oman, and Qatar."),
        ("CR",                "Commercial Registration — the official business licence number issued by the Kuwait Ministry of Commerce."),
        ("KWD / SAR / AED / BHD", "Kuwaiti Dinar / Saudi Riyal / UAE Dirham / Bahraini Dinar."),
        ("RTL",               "Right-to-Left — the text direction used by Arabic script."),
        ("SRS",               "Software Requirements Specification."),
        ("API",               "Application Programming Interface — the RESTful HTTP interface between the Next.js frontend and backend."),
        ("SSR",               "Server-Side Rendering — pages rendered on the server before delivery to the browser; used for SEO."),
        ("MVP",               "Minimum Viable Product — the version of the product with just enough features to be usable and demonstrable."),
        ("Verification Tier", "The trust level assigned to a brand by an admin: New, Verified, or Premium."),
        ("Trust Score",       "A computed metric summarising a brand's or buyer's reliability based on response rate, completion rate, and fulfillment time."),
        ("Bulk Pricing Tier", "A quantity bracket that unlocks a lower unit price for orders meeting or exceeding the minimum quantity."),
        ("Supabase",          "An open-source Firebase alternative providing PostgreSQL, authentication, and real-time capabilities."),
        ("Next.js App Router","The modern file-system routing architecture in Next.js 13+, supporting React Server Components and co-located API routes."),
    ]
    make_table(["Term", "Definition"], terms, col_widths=[4.0, 12.5], font_size=10)

    h2("1.4  Readership")
    body("This report is intended for the following audiences:")
    bullet("University course assessors evaluating the software-engineering process and deliverables.")
    bullet("Project supervisors who need to trace requirements through to implementation and testing.")
    bullet("Future developers who may extend or maintain the platform after the academic project concludes.")
    bullet("Potential stakeholders or investors evaluating the business and technical viability of the platform.")

    h2("1.5  Product Scope")
    body("The scope of Kuwait B2B Hub v1.0 (MVP) is defined as follows:")
    bullet("IN SCOPE: User registration and authentication for three roles; brand creation, verification, and catalog management; product CRUD with bulk-pricing tiers and Excel import; order lifecycle (submit → approve/reject → complete); buyer review and rating system; brand and buyer trust-score calculation; admin dashboard with analytics; bilingual EN/AR UI with full RTL; four-currency support (KWD/SAR/AED/BHD); in-app notification panel; mobile-responsive design.")
    bullet("OUT OF SCOPE (v1.0): Payment gateway integration; shipping carrier API integration; real-time WebSocket messaging; SMS/email transactional notifications; automated commercial-registration verification via external API; native mobile applications.")

    h2("1.6  Team Workload")
    body("Table 1 summarises the division of responsibilities among team members throughout the project lifecycle.")
    doc.add_paragraph()
    workload_rows = [
        ("Bader Alnoumas",   "Team Lead / Full-Stack", "System architecture, auth, DB schema, order API, brand API, deployment, Supabase integration, project management"),
        ("Team Member 2",    "Frontend / UI",           "Marketplace pages, brand dashboard, mobile responsiveness, CSS modules, component library"),
        ("Team Member 3",    "Backend / API",           "Reviews API, trust-score engine, notifications, admin API routes, bulk import"),
        ("Team Member 4",    "Localisation / QA",       "Arabic translation files, RTL layout, translation audit scripts, manual testing, documentation"),
    ]
    make_table(
        ["Name", "Role", "Primary Responsibilities"],
        workload_rows,
        col_widths=[3.5, 3.5, 9.5],
        font_size=10
    )
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER II — SRS
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_2():
    h1("Chapter II:  Software Requirements Specification (SRS)")

    h2("2.1  Feasibility Study")
    body("A feasibility study was conducted at project inception across four dimensions:")
    doc.add_paragraph()
    feas_rows = [
        ("Technical",    "High",      "Next.js, Supabase, and Vercel are mature, well-documented technologies with strong community support. The team has sufficient expertise. No novel algorithms or unproven technologies are required."),
        ("Operational",  "High",      "The platform replaces an existing manual workflow (WhatsApp ordering). Users already understand the business domain; training needs are minimal."),
        ("Economic",     "Medium",    "Zero infrastructure cost during development (Vercel and Supabase free tiers). Subscription revenue model (KD 29/month per brand) provides a realistic path to cost recovery post-launch."),
        ("Schedule",     "Medium",    "An 8-sprint Agile plan with a ~4-month delivery window is achievable for a team of four given the MVP scope.  Risks: Supabase migration and real-time features may extend the timeline."),
        ("Legal",        "Medium",    "GCC data-localisation laws require user data to be stored within the region (Supabase supports EU/Middle-East regions). GDPR-equivalent compliance is a post-MVP concern."),
    ]
    make_colored_table(
        ["Dimension", "Viability", "Justification"],
        feas_rows,
        [HEX_GREEN, HEX_GREEN, HEX_YELLOW, HEX_YELLOW, HEX_YELLOW],
        col_widths=[3.0, 2.5, 11.0],
        font_size=10
    )

    h2("2.2  User Requirements")
    body("User requirements are expressed from the perspective of each actor role.")
    doc.add_paragraph()
    user_req = [
        ("UR-01", "Admin",       "I want to review pending brand applications and approve or reject them with a reason so that only legitimate businesses appear on the marketplace."),
        ("UR-02", "Admin",       "I want to view a dashboard showing total users, brands, orders, and revenue so that I can monitor platform health."),
        ("UR-03", "Admin",       "I want to moderate flagged reviews so that inappropriate or fraudulent content is removed promptly."),
        ("UR-04", "Brand Owner", "I want to register and create a brand profile including name, logo, description, and WhatsApp contact so that buyers can find and contact me."),
        ("UR-05", "Brand Owner", "I want to list products with name, price, MOQ, stock, and bulk-pricing tiers so that buyers understand my wholesale terms."),
        ("UR-06", "Brand Owner", "I want to import products from an Excel/CSV file so that I can onboard my existing catalog quickly."),
        ("UR-07", "Brand Owner", "I want to review incoming order requests and approve or reject them so that I control my fulfillment capacity."),
        ("UR-08", "Brand Owner", "I want to view a dashboard with revenue charts and order metrics so that I can track my business performance."),
        ("UR-09", "Brand Owner", "I want to reply to buyer reviews so that I can address feedback publicly."),
        ("UR-10", "Buyer",       "I want to browse and search verified brands filtered by category, MOQ, and rating so that I find relevant suppliers quickly."),
        ("UR-11", "Buyer",       "I want to view a brand's full product catalog with pricing tiers so that I can make informed bulk-order decisions."),
        ("UR-12", "Buyer",       "I want to submit an order request with quantity validation so that I can initiate a wholesale purchase."),
        ("UR-13", "Buyer",       "I want to track the status of my orders in real-time so that I know when to arrange collection or delivery."),
        ("UR-14", "Buyer",       "I want to leave a star rating and written review after receiving an order so that I can share my experience with other buyers."),
        ("UR-15", "Buyer",       "I want to view prices in my local currency (KWD/SAR/AED/BHD) so that I can assess costs without manual conversion."),
        ("UR-16", "All Users",   "I want to use the platform in Arabic with full RTL layout so that I can operate comfortably in my native language."),
        ("UR-17", "All Users",   "I want to receive in-app notifications for relevant events so that I am informed without leaving the platform."),
    ]
    make_table(
        ["ID", "Actor", "Requirement"],
        user_req,
        col_widths=[1.5, 2.5, 12.5],
        font_size=9
    )

    h2("2.3  System Requirements")
    body("System-level functional requirements derived from user requirements:")
    doc.add_paragraph()
    sys_req = [
        ("SR-01",  "The system shall support three user roles: admin, brand_owner, and buyer, with distinct access control for each."),
        ("SR-02",  "The system shall enforce role-based route protection on all pages and API endpoints."),
        ("SR-03",  "The system shall allow brand owners to register, automatically creating a pending brand profile upon successful registration."),
        ("SR-04",  "The system shall allow admins to approve or reject brand applications, setting a verification tier (new / verified / premium)."),
        ("SR-05",  "The system shall display a verification badge on brand cards corresponding to the brand's tier."),
        ("SR-06",  "The system shall allow brand owners to create, read, update, and delete products associated with their brand."),
        ("SR-07",  "The system shall support bulk-pricing tiers on each product, automatically applying the correct unit price based on order quantity."),
        ("SR-08",  "The system shall validate order submissions against product MOQ and available stock, rejecting requests that fail either constraint."),
        ("SR-09",  "The system shall maintain an order status lifecycle: pending → approved / rejected; approved → completed."),
        ("SR-10",  "The system shall allow buyers to submit one review per completed order (rating 1–5, optional text, optional anonymous flag)."),
        ("SR-11",  "The system shall compute and persist aggregate review statistics (average rating, breakdown by star count) per brand."),
        ("SR-12",  "The system shall compute trust scores for brands (response rate, completion rate, avg response hours, fulfillment days, badges)."),
        ("SR-13",  "The system shall compute trust scores for buyers (total orders, completion rate, cancellation rate, badges)."),
        ("SR-14",  "The system shall generate notifications for order events, new registrations, low-stock warnings, and review events."),
        ("SR-15",  "The system shall display prices in the user's selected currency (KWD, SAR, AED, BHD) with VAT applied where applicable."),
        ("SR-16",  "The system shall serve the UI in English and Arabic; all content-bearing pages shall have translations in both locales."),
        ("SR-17",  "The system shall provide a WhatsApp click-tracking endpoint, incrementing the click count per brand on each use."),
        ("SR-18",  "The system shall allow brand owners to import products in bulk via XLSX/CSV file upload."),
        ("SR-19",  "The system shall generate SEO metadata (JSON-LD structured data) on public-facing pages."),
        ("SR-20",  "The system shall provide an admin moderation queue for flagged reviews with approve/remove actions."),
    ]
    make_table(["ID", "System Requirement"], sys_req, col_widths=[1.5, 15.0], font_size=9)

    h2("2.4  Non-Functional Requirements")
    doc.add_paragraph()
    nfr = [
        ("NFR-01", "Performance",    "Public marketplace pages shall achieve a Largest Contentful Paint (LCP) < 2.5 s on a standard broadband connection."),
        ("NFR-02", "Security",       "All user sessions shall be managed via HTTP-only cookies; passwords shall be stored using bcrypt hashing via Supabase Auth."),
        ("NFR-03", "Scalability",    "The data layer shall be migrated to Supabase PostgreSQL before launch to support concurrent access by multiple users."),
        ("NFR-04", "Availability",   "The platform shall target 99.5% uptime, relying on Vercel's edge network and Supabase's managed infrastructure."),
        ("NFR-05", "Accessibility",  "The UI shall follow WCAG 2.1 Level AA guidelines, with appropriate ARIA labels on interactive elements."),
        ("NFR-06", "Internationalisation", "All user-facing strings shall be externalised to translation files; the Arabic locale shall render with full RTL layout."),
        ("NFR-07", "Maintainability","TypeScript strict mode shall be enabled; all modules shall export typed interfaces to prevent implicit `any` usage."),
        ("NFR-08", "Portability",    "The application shall be deployable to any Vercel-compatible Node.js 18+ hosting environment."),
    ]
    make_table(["ID", "Category", "Requirement"], nfr, col_widths=[1.5, 3.0, 12.0], font_size=9)

    h2("2.5  Functional Requirements Specification (FRS)")
    body("The FRS maps high-level system requirements to detailed input/output/processing specifications for the five primary subsystems:")
    bullet("Authentication Subsystem: accepts email, password, role, and company_name; validates uniqueness; creates user record with role-appropriate verification_status; sets HTTP-only session cookie (7-day TTL); returns user object.")
    bullet("Catalog Subsystem: accepts product name, description, price, MOQ, stock, image URL, and up to five bulk-pricing tiers; validates MOQ ≥ 1 and price > 0; persists to products table linked to brand_id; supports XLSX bulk import via XLSX 0.18.5 parser.")
    bullet("Order Subsystem: accepts brand_id and items array [{product_id, quantity}]; validates each item against MOQ and stock; resolves unit price from bulk-pricing tiers; computes order total; persists order and order_items; fires notification to brand owner.")
    bullet("Reviews Subsystem: accepts brand_id, optional order_id, rating (1–5), optional content (max 300 chars), anonymous flag; enforces one-review-per-order; persists review; recomputes aggregate stats; fires notification to brand owner.")
    bullet("Admin Subsystem: exposes brand approval (PATCH /api/admin/brands/:id), user management (GET/PATCH /api/admin/users), order analytics (GET /api/admin/orders), and review moderation (PATCH /api/admin/reviews/:id) endpoints behind admin-role guard.")

    h2("2.6  Context Model")
    body(
        "Figure 2 (see Appendix A) illustrates the Level-0 context model for Kuwait B2B Hub.  "
        "The system boundary encloses all platform functionality.  External entities are: "
        "Admin Users (who govern the system), Brand Owners (who supply products), Business Buyers "
        "(who consume products), and External Services (Supabase Auth, Vercel hosting, WhatsApp "
        "Business).  Data flows include: registration data, brand/product catalog data, order "
        "requests, review data, approval decisions, and notification events."
    )

    h2("2.7  Use Cases")
    body("Table 6 summarises the principal use cases for each actor.")
    doc.add_paragraph()
    uc_rows = [
        ("UC-01", "Register Account",          "All",        "User submits name, email, password, role, company_name → system creates account and session cookie"),
        ("UC-02", "Login",                      "All",        "User submits credentials → system validates → sets session cookie"),
        ("UC-03", "Approve / Reject Brand",     "Admin",      "Admin reviews pending brand → sets status and verification_tier → brand owner notified"),
        ("UC-04", "Manage Users",               "Admin",      "Admin views user list → changes verification_status or role"),
        ("UC-05", "Moderate Reviews",           "Admin",      "Admin views flagged reviews → removes or approves"),
        ("UC-06", "Create / Edit Brand Profile","Brand Owner","Brand owner updates brand name, description, logo, WhatsApp, business hours"),
        ("UC-07", "Add / Edit Product",         "Brand Owner","Brand owner submits product form → validated → saved to catalog"),
        ("UC-08", "Bulk Import Products",       "Brand Owner","Brand owner uploads XLSX → system parses → creates products in bulk"),
        ("UC-09", "Review Order",               "Brand Owner","Brand owner views pending order → approves or rejects → buyer notified"),
        ("UC-10", "Browse Marketplace",         "Buyer",      "Buyer searches/filters brands → views results with verification badges and ratings"),
        ("UC-11", "View Brand Catalog",         "Buyer",      "Buyer selects brand → views all products with pricing tiers"),
        ("UC-12", "Place Order",                "Buyer",      "Buyer selects product and quantity → system validates MOQ/stock → creates order"),
        ("UC-13", "Track Orders",               "Buyer",      "Buyer views order list with status and order details"),
        ("UC-14", "Submit Review",              "Buyer",      "Buyer submits rating and text for a brand → brand owner notified"),
        ("UC-15", "Switch Currency",            "All",        "User selects currency → all prices instantly converted using exchange rates"),
        ("UC-16", "Switch Language",            "All",        "User toggles EN/AR → UI re-renders in selected locale with RTL adjustment"),
    ]
    make_table(
        ["UC ID", "Use Case", "Actor", "Brief Description"],
        uc_rows,
        col_widths=[1.5, 3.5, 2.5, 9.0],
        font_size=9
    )

    h2("2.8  Requirements Traceability Matrix (RTM)")
    body("Table 7 traces each system requirement through design artefacts, implementation status, and test coverage.")
    doc.add_paragraph()
    rtm = [
        ("SR-01","Role-based access control","DB schema: users.role; middleware; requireRole()","Complete","Manual"),
        ("SR-02","Route protection","Next.js middleware; server-side auth guards","Complete","Manual"),
        ("SR-03","Brand auto-creation on owner registration","/api/auth/register — brand creation block","Complete","Manual"),
        ("SR-04","Brand approval workflow","/api/admin/brands/[id]; admin dashboard","Complete","Manual"),
        ("SR-05","Verification tier badges","Brand schema: verification_tier; VerifiedBadge component","Complete","Manual"),
        ("SR-06","Product CRUD","/api/products; brand product pages","Complete","Manual"),
        ("SR-07","Bulk-pricing tiers","BulkPricingTier interface; pricingUtils.ts; order route","Complete","Manual"),
        ("SR-08","Order MOQ/stock validation","/api/orders POST — validation block","Complete","Manual"),
        ("SR-09","Order status lifecycle","Order.status enum; PATCH /api/orders/[id]","Complete","Manual"),
        ("SR-10","Buyer review submission","/api/reviews POST; ReviewModal component","Complete","Manual"),
        ("SR-11","Aggregate review statistics","GET /api/reviews — avgRating & breakdown","Complete","Manual"),
        ("SR-12","Brand trust scores","BrandTrust schema; /api/trust/[id]","Partial (70%)","Not started"),
        ("SR-13","Buyer trust scores","BuyerTrust schema; /api/trust/[id]","Partial (70%)","Not started"),
        ("SR-14","Notification system","notifications.ts; /api/notifications; Zustand store","Complete","Manual"),
        ("SR-15","Multi-currency display","currencies.ts; CurrencyContext; CurrencySelector","Complete","Manual"),
        ("SR-16","Bilingual EN/AR support","next-intl; en.json / ar.json; RTL CSS","Partial (80%)","Manual"),
        ("SR-17","WhatsApp click tracking","/api/brands/whatsapp-click; WhatsAppButton","Complete","Manual"),
        ("SR-18","Bulk product import","/api/products/bulk; CatalogImportModal; xlsx","Complete","Manual"),
        ("SR-19","SEO structured data","JSON-LD in layout.tsx; sitemap.ts","Complete","Manual"),
        ("SR-20","Review moderation queue","/api/admin/reviews/[id]; admin reviews page","Complete","Manual"),
    ]
    make_colored_table(
        ["Req. ID","Requirement","Design Artefact","Status","Test"],
        rtm,
        [HEX_GREEN if r[3]=="Complete" else (HEX_YELLOW if "Partial" in r[3] else HEX_RED) for r in rtm],
        col_widths=[1.5, 4.0, 4.5, 2.5, 2.0],
        font_size=8
    )
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER III — SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_3():
    h1("Chapter III:  System Design")

    h2("3.1  Sequence Diagrams")
    body("Three principal sequence diagrams capture the system's key interaction flows.")

    h3("3.1.1  User Registration and Login")
    body(
        "Figure 7 (Appendix A) shows the registration flow.  The Client submits "
        "POST /api/auth/register with {name, email, password, role, company_name}.  "
        "The route handler reads the DB, checks for duplicate email (returning HTTP 409 if found), "
        "creates the User record, conditionally creates a pending Brand record (brand_owner role), "
        "writes the DB, fires admin notification(s), and returns a success response with an "
        "HTTP-only session cookie (b2b_user_id, 7-day TTL).  Login follows a similar flow via "
        "POST /api/auth/login, which validates credentials and sets the same cookie."
    )

    h3("3.1.2  Order Lifecycle")
    body(
        "The Buyer sends POST /api/orders with {brand_id, items[]}.  The API validates each item "
        "against MOQ and stock; resolves unit price from bulk_pricing_tiers; accumulates total; "
        "persists Order + OrderItems; fires a new_order notification to the brand owner.  "
        "The Brand Owner later sends PATCH /api/orders/:id with {status: 'approved'|'rejected'}.  "
        "The API updates the status and fires an order_approved / order_rejected notification to "
        "the buyer."
    )

    h3("3.1.3  Brand Verification")
    body(
        "A Brand Owner registers → system creates pending brand → Admin receives notification.  "
        "Admin sends PATCH /api/admin/brands/:id with {status: 'approved', verification_tier: "
        "'premium'|'verified'|'new'}.  API updates the brand record, sets verification_tier, and "
        "fires an account_approved notification to the brand owner.  The brand immediately "
        "becomes visible in the marketplace."
    )

    h2("3.2  Data Flow Diagram (DFD)")
    body(
        "Figure 3 (Appendix A) presents the Level 1 DFD for Kuwait B2B Hub.  Four primary data "
        "stores are identified: Users, Brands/Products, Orders/OrderItems, and Reviews/Notifications.  "
        "Five processes transform data flows: (P1) Authentication & Session Management, "
        "(P2) Catalog Management, (P3) Order Processing, (P4) Trust & Reputation Engine, "
        "and (P5) Admin Governance.  External entities are Admin, Brand Owner, Buyer, and Supabase "
        "(authentication service)."
    )

    h2("3.3  System Architecture")
    body("Figure 1 (Appendix A) shows the layered system architecture.  Four layers are defined:")

    h3("3.3.1  Presentation Layer")
    body(
        "React 19 / Next.js 16 App Router pages and client components.  CSS Modules provide "
        "scoped, maintainable styles without a utility-class framework.  Lucide React supplies "
        "a consistent icon library.  Recharts renders analytics dashboards.  "
        "The CurrencyContext (React Context API) propagates the active currency to all "
        "price-display components.  Zustand manages the notification panel's unread-count state "
        "on the client.  next-intl handles server-side and client-side translation."
    )

    h3("3.3.2  Application / API Layer")
    body(
        "Next.js Route Handlers under /src/app/api/** implement a RESTful HTTP API.  "
        "All handlers call getCurrentUser() from auth.ts to identify the caller, then call "
        "requireRole() to enforce access control.  The Supabase SSR middleware in "
        "src/middleware.ts refreshes sessions on every request."
    )

    h3("3.3.3  Data Access Layer")
    body(
        "The db.ts module provides readDB() and writeDB() functions that serialise/deserialise "
        "a local JSON file (src/data/db.json) for the MVP.  TypeScript interfaces (User, Brand, "
        "Product, Order, OrderItem, Review, Notification, BrandTrust, BuyerTrust, Message) "
        "provide compile-time safety.  The Supabase client in src/lib/supabase/client.ts and "
        "src/utils/supabase/server.ts provides the production data path (migration in progress)."
    )

    h3("3.3.4  Infrastructure Layer")
    body(
        "Vercel hosts the Next.js application on its global edge network.  "
        "Supabase provides managed PostgreSQL, authentication (bcrypt, JWT), and Realtime "
        "infrastructure.  WhatsApp Business deep links provide communication between buyers and "
        "brand owners outside the platform."
    )

    h2("3.4  Activity Model")
    body(
        "Figure 11 (Appendix A) shows the Activity Diagram for the order-placement flow.  "
        "Starting from 'Buyer views product', the buyer selects quantity and submits the order form.  "
        "Decision: quantity ≥ MOQ?  If No → display error; If Yes → Decision: stock ≥ quantity?  "
        "If No → display stock error; If Yes → resolve bulk-pricing tier → compute total → create "
        "order record → notify brand owner → display confirmation.  "
        "Brand owner receives notification → decision: approve or reject → update order status → "
        "notify buyer."
    )

    h2("3.5  Class Diagram")
    body("Figure 10 (Appendix A) illustrates the class diagram for the ten core entities.  Key relationships:")
    bullet("User 1 — 0..1 Brand (one brand_owner user has at most one brand; one brand has exactly one owner)")
    bullet("Brand 1 — * Product (one brand has zero or more products)")
    bullet("Order * — 1 Brand, Order * — 1 User/Buyer (many orders can belong to one brand or one buyer)")
    bullet("Order 1 — * OrderItem (one order has one or more line items)")
    bullet("OrderItem * — 1 Product (each order item references exactly one product)")
    bullet("Review * — 1 Brand, Review * — 1 User/Buyer (many reviews per brand and per buyer)")
    bullet("Review 0..1 — 0..1 Order (a review may or may not be tied to a specific order)")
    bullet("Notification * — 1 User (many notifications per user)")
    bullet("BrandTrust 1 — 1 Brand (one trust record per brand)")
    bullet("BuyerTrust 1 — 1 User/Buyer (one trust record per buyer)")

    body("Selected TypeScript interface definitions:")
    code_block = doc.add_paragraph()
    code_block.paragraph_format.left_indent = Cm(1.0)
    code_block.paragraph_format.space_before = Pt(4)
    code_block.paragraph_format.space_after  = Pt(4)
    r = code_block.add_run(
        "interface Product {\n"
        "  id: string;  brand_id: string;  name: string;\n"
        "  price: number;  moq: number;  stock: number;\n"
        "  bulk_pricing_tiers: BulkPricingTier[];  created_at: string;\n"
        "}\n"
        "interface BulkPricingTier {\n"
        "  min_qty: number;  max_qty: number | null;  price: number;\n"
        "}"
    )
    r.font.name = 'Courier New'
    r.font.size = Pt(9)

    h2("3.6  Prototype Screenshots")
    body(
        "Prototype screenshots are provided in Appendix B.  Key screens captured include: "
        "(1) Landing page with hero section and FAQ in both EN and AR; "
        "(2) Marketplace grid showing brand cards with verification badges and trust scores; "
        "(3) Brand catalog page with product list and bulk-pricing tiers; "
        "(4) Order placement form with MOQ validation error; "
        "(5) Brand owner dashboard with revenue chart; "
        "(6) Admin dashboard with pending approvals queue; "
        "(7) Review modal with star rating and anonymous option; "
        "(8) Notification panel with unread notifications."
    )
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER IV — SYSTEM IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_4():
    h1("Chapter IV:  System Implementation")

    h2("4.1  Make / Buy / Lease Analysis")
    body("Table 8 documents the make/buy/lease decision for each major system component.")
    doc.add_paragraph()
    mbl = [
        ("Authentication",           "Buy (Supabase Auth)",  "Supabase provides bcrypt hashing, JWT, SSR session management — building this from scratch would introduce security risk and require significant effort."),
        ("Database (production)",    "Buy (Supabase PostgreSQL)", "Managed PostgreSQL eliminates DBA overhead, provides backups, and supports real-time subscriptions."),
        ("Database (MVP/dev)",       "Make (JSON file + db.ts)", "Minimal setup cost; sufficient for a small seed dataset during development sprints."),
        ("Web Framework",            "Buy (Next.js 16)",    "Industry-standard full-stack React framework with SSR, API routes, and Vercel-native deployment."),
        ("UI Components",            "Make (CSS Modules)",  "Custom styling required for GCC brand identity and Arabic RTL; no third-party CSS framework introduced."),
        ("Internationalisation",     "Buy (next-intl 4.9)", "Handles server-side and client-side i18n, locale routing, and RTL awareness."),
        ("Charts / Analytics",       "Buy (Recharts 3.8)",  "Declarative React charting library; no custom SVG required."),
        ("Notification State",       "Buy (Zustand 5.0)",   "Lightweight client state for unread-count; avoids Redux boilerplate."),
        ("Excel Import",             "Buy (xlsx 0.18.5)",   "Industry-standard parser for XLSX/CSV; building a custom parser is unnecessary."),
        ("Hosting / CDN",            "Lease (Vercel)",      "Vercel provides zero-config Next.js deployment, global CDN, and preview environments."),
        ("Payment Gateway",          "N/A (Future)",        "Not yet integrated; Tap Payments or Checkout.com will be evaluated for GCC compliance."),
        ("Shipping Carrier API",     "N/A (Future)",        "Aramex / DHL API integration planned for post-MVP phase."),
    ]
    make_table(
        ["Component", "Decision", "Justification"],
        mbl,
        col_widths=[3.5, 3.0, 10.0],
        font_size=9
    )

    h2("4.2  Coding Artefacts")
    body("Key implementation highlights with file paths and line references:")

    h3("4.2.1  Role-Based Auth Guard (src/lib/auth.ts)")
    body(
        "The requireRole() function (line 27) wraps requireAuth() and throws HTTP 403 if the "
        "authenticated user's role is not in the allowed list.  Every API route handler calls "
        "this at entry point.  The getCurrentUser() function calls Supabase's getUser() and "
        "then joins the profiles table to return the full User object."
    )

    h3("4.2.2  Order Validation (src/app/api/orders/route.ts, lines 58–79)")
    body(
        "The POST handler iterates over each order item, validating: (a) quantity is a positive "
        "integer; (b) product exists and belongs to the requested brand; (c) quantity ≥ product.moq; "
        "(d) product.stock ≥ quantity.  For each valid item it resolves the unit price from "
        "bulk_pricing_tiers by iterating tiers in order and selecting the first matching range."
    )

    h3("4.2.3  Bulk Pricing Utility (src/lib/pricingUtils.ts)")
    body(
        "Three exported functions: formatTierRange() formats a tier's quantity range as a "
        "localised string; calcSavingsPct() computes the percentage discount vs. the base price; "
        "sortTiers() returns tiers sorted ascending by min_qty.  These are used by both "
        "the product-detail UI and the order-creation API."
    )

    h3("4.2.4  Currency Engine (src/lib/currencies.ts)")
    body(
        "CURRENCIES record maps ISO 4217 codes (KWD, SAR, AED, BHD) to config objects including "
        "symbol, locale, VAT rate, and phone prefix.  EXCHANGE_RATES defines conversion factors "
        "relative to KWD as the base currency.  convertPrice() multiplies by the rate; "
        "formatPriceWithVAT() adds the destination country's VAT (0% KW, 15% SA, 5% AE, 10% BH)."
    )

    h3("4.2.5  Non-Fatal Notification Pattern (src/lib/notifications.ts)")
    body(
        "createNotification() and createNotificationForMany() wrap all DB writes in a try/catch "
        "that silently suppresses errors.  This ensures a notification failure never propagates "
        "to the caller and causes an HTTP 500 on an otherwise successful operation."
    )

    h2("4.3  Configuration Management")
    body("Table 9 documents the configuration management setup for this project.")
    doc.add_paragraph()
    cm_rows = [
        ("Version Control",   "Git",                   "Single main branch; all development committed to main; .git/logs/HEAD tracks full history"),
        ("Repository",        "Local + Remote",        "Local Git repository at app/.git; pushable to GitHub remote for collaboration"),
        ("Build System",      "Next.js (npm)",         "npm run dev (development), npm run build (production), npm start (production server)"),
        ("Linting",           "ESLint 9 + next config","eslint.config.mjs; enforces TypeScript and React best-practice rules"),
        ("Type Checking",     "TypeScript 5 (strict)", "tsconfig.json with strict: true; tsbuildinfo caches incremental builds"),
        ("Environment Vars",  ".env.local",            "NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_ANON_KEY — not committed to Git"),
        ("Dependency Lock",   "package-lock.json",     "Ensures reproducible installs across team members and CI environments"),
        ("Translations",      "Python audit scripts",  "find_missing_keys.py, audit_translations.py, find_all_missing_keys.py — run before each sprint review"),
    ]
    make_table(
        ["Area", "Tool / Mechanism", "Details"],
        cm_rows,
        col_widths=[3.0, 3.5, 10.0],
        font_size=9
    )

    h2("4.4  Deployment Specification")
    body("Table 10 specifies the production deployment configuration.")
    doc.add_paragraph()
    deploy_rows = [
        ("Hosting Platform",      "Vercel",                           "Zero-config Next.js hosting; automatic CI/CD on git push"),
        ("Runtime",               "Node.js 18+ (Vercel managed)",     "Next.js server-side rendering and API route execution"),
        ("Build Command",         "npm run build",                    "Produces optimised production bundle in .next/ directory"),
        ("Output Directory",      ".next/",                           "Standard Next.js output; served by Vercel edge runtime"),
        ("Environment",           "Production / Preview",             "Separate Vercel environments for production and PR previews"),
        ("Authentication Service","Supabase (hosted)",                "NEXT_PUBLIC_SUPABASE_URL + NEXT_PUBLIC_SUPABASE_ANON_KEY env vars"),
        ("Database (MVP)",        "JSON file on Vercel filesystem",   "src/data/db.json — ephemeral on serverless; migration to Supabase planned"),
        ("Database (production)", "Supabase PostgreSQL",              "GCC-region Supabase instance; accessed via @supabase/ssr"),
        ("CDN / Static Assets",   "Vercel Edge Network",              "SVG, image optimisation via Next.js Image component"),
        ("SSL / TLS",             "Vercel managed (Let's Encrypt)",   "Automatic HTTPS for all deployments"),
        ("Domain",                "TBD (b2bhub.kw planned)",          "Kuwait ccTLD requires local entity registration"),
    ]
    make_table(
        ["Configuration Item", "Value", "Notes"],
        deploy_rows,
        col_widths=[4.0, 4.5, 8.0],
        font_size=9
    )

    h2("4.5  Physical Deployment Diagram")
    body(
        "Figure 12 (Appendix A) illustrates the physical deployment topology.  "
        "End-user devices (desktop browsers and mobile browsers) communicate via HTTPS to "
        "Vercel's global edge network.  The Vercel CDN serves static assets and routes "
        "dynamic requests to the Next.js serverless functions.  The serverless functions "
        "communicate with the Supabase cloud (Auth service + PostgreSQL) over a TLS-encrypted "
        "connection.  WhatsApp Business API links are deep-links that open the user's WhatsApp "
        "client directly — no server-side proxy involved.  "
        "During the MVP phase, the serverless function reads/writes a bundled db.json file "
        "within the same Vercel environment (depicted as a local store on the function node)."
    )
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER V — V&V
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_5():
    h1("Chapter V:  Verification & Validation")

    h2("5.1  Development Testing")
    body(
        "Development testing during each sprint focused on validating individual functions and "
        "API endpoints against their specification before integration.  This was primarily "
        "conducted through manual invocation of API routes using browser developer tools and "
        "by exercising the UI against the seeded demo dataset in db.json."
    )
    body("Specific development-phase checks included:")
    bullet("TypeScript compilation with strict mode — all type errors resolved before commit.")
    bullet("ESLint static analysis — run via npm run lint; zero warnings policy enforced.")
    bullet("Manual testing of each new API endpoint against edge cases (missing fields, wrong role, duplicate email, below-MOQ quantity, insufficient stock).")
    bullet("Translation key audit — Python scripts (find_missing_keys.py, audit_translations.py) run after every new UI string addition.")
    bullet("Browser console monitoring — zero unhandled JS exceptions on all main user flows.")

    h2("5.2  System Inspection")
    body(
        "A structured code inspection was conducted at the end of Sprint 6 covering the five "
        "most security-sensitive modules: auth.ts, the register route, the orders route, the "
        "reviews route, and the admin brands route.  Findings:"
    )
    bullet("Finding 1 (Critical — Resolved): Passwords stored as plain text in db.json.  Resolution: migration to Supabase Auth (bcrypt) is the primary Sprint 8 objective.")
    bullet("Finding 2 (Medium — In Progress): The b2b_user_role cookie is not HTTP-only, allowing client-side JavaScript to read it.  Resolution: role validation moved to server-side session check; cookie will be made HTTP-only post-migration.")
    bullet("Finding 3 (Low — Resolved): Review content was not sanitised before storage.  Resolution: content.slice(0, 300) applied in the POST handler; further XSS sanitisation added in the review display component.")
    bullet("Finding 4 (Low — Resolved): TypeScript 'any' used in three review route handlers.  Resolution: explicit type annotations added in subsequent sprint.")
    bullet("Finding 5 (Info): No rate limiting on auth endpoints.  Resolution: Vercel edge rate limiting + Supabase Auth's built-in brute-force protection will address this post-migration.")

    h2("5.3  System Testing")
    body("Table 11 documents the system test cases executed against the deployed development environment.")
    doc.add_paragraph()
    tc_rows = [
        ("TC-01","Register new buyer account","Submit valid registration form with role=buyer","User created; buyer dashboard accessible; admin notified","Pass",""),
        ("TC-02","Register brand owner","Submit valid form with role=brand_owner","User + pending brand created; admin notified x2","Pass",""),
        ("TC-03","Duplicate email rejection","Submit registration with existing email","HTTP 409 error displayed","Pass",""),
        ("TC-04","Login valid credentials","Submit correct email/password","Session cookie set; redirected to dashboard","Pass",""),
        ("TC-05","Login invalid password","Submit wrong password","Error message displayed; no cookie set","Pass",""),
        ("TC-06","Unauthorised route access","Navigate to /admin/dashboard as buyer","Redirected to login or 403 page","Pass",""),
        ("TC-07","Brand approval workflow","Admin approves pending brand","Brand status = approved; brand owner notified; brand visible in marketplace","Pass",""),
        ("TC-08","Create product with valid data","Submit product form with MOQ=15, price=32.5","Product appears in brand catalog","Pass",""),
        ("TC-09","Below-MOQ order rejected","Submit order with quantity < MOQ","HTTP 400 with clear error message","Pass",""),
        ("TC-10","Valid order creation","Submit order with quantity ≥ MOQ and stock available","Order created; brand owner notified","Pass",""),
        ("TC-11","Bulk-pricing tier applied","Order 50 units of prod-001 (tier: KD 28)","Order total = 50 × 28 = KD 1,400","Pass",""),
        ("TC-12","Order approval by brand","Brand owner approves order","Order status = approved; buyer notified","Pass",""),
        ("TC-13","Submit buyer review","Submit 5-star review with text","Review saved; brand owner notified; rating displayed","Pass",""),
        ("TC-14","Anonymous review","Submit review with anonymous=true","Reviewer shown as 'Anonymous Buyer'","Pass",""),
        ("TC-15","Currency switch to SAR","Select SAR from currency selector","All prices converted at rate 12.20","Pass",""),
        ("TC-16","Arabic language switch","Toggle to Arabic locale","UI renders RTL; Arabic text displayed","Partial","Some components not yet RTL-adapted"),
        ("TC-17","WhatsApp click tracking","Click WhatsApp button on brand page","GET /api/brands/whatsapp-click increments counter","Pass",""),
        ("TC-18","Excel product import","Upload valid XLSX with 5 products","5 products created in brand catalog","Pass",""),
        ("TC-19","Notification bell count","Create new order (unseen notification)","Bell shows unread badge; clicking opens panel","Pass",""),
        ("TC-20","Admin review moderation","Flag a review; admin removes it","Review status = removed; no longer visible to buyers","Pass",""),
    ]
    make_colored_table(
        ["TC ID","Test Case","Input","Expected Output","Result","Notes"],
        tc_rows,
        [HEX_GREEN if r[4]=="Pass" else HEX_YELLOW for r in tc_rows],
        col_widths=[1.2, 3.0, 3.5, 4.0, 1.5, 3.3],
        font_size=8
    )

    h2("5.4  Release Testing")
    body(
        "Release testing (pre-submission) will include the following activities, to be completed "
        "in Sprint 8:"
    )
    numbered("Regression sweep across all 20 TC cases after the Supabase DB migration.")
    numbered("Cross-browser testing: Chrome 124, Safari 17, Firefox 125, and Edge 124.")
    numbered("Mobile viewport testing at 375px (iPhone SE), 414px (iPhone Pro Max), and 768px (iPad).")
    numbered("Performance audit using Lighthouse in Chrome DevTools — target: LCP < 2.5 s, TBT < 200 ms, CLS < 0.1.")
    numbered("Security review of HTTP headers (Content-Security-Policy, X-Frame-Options, HSTS).")
    numbered("Internationalisation completeness check — translation audit scripts must report 0 missing keys.")

    h2("5.5  User Testing")
    body(
        "Informal user acceptance testing was conducted with two volunteer business users "
        "(one retail buyer, one small brand owner) during Sprint 5.  Key findings:"
    )
    bullet("Positive: The marketplace search and filter UX was described as 'intuitive and faster than WhatsApp'.")
    bullet("Positive: The verification badge and trust score were cited as confidence builders for the buyer.")
    bullet("Improvement: The order form needed clearer MOQ guidance — resolved by adding a dynamic MOQ label below the quantity field.")
    bullet("Improvement: Arabic users requested the brand descriptions also be available in Arabic — added to the post-MVP roadmap.")
    bullet("Improvement: The bulk-pricing tier table was not immediately visible — resolved by making it a default-expanded section on the product page.")
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER VI — SYSTEM EVOLUTION
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_6():
    h1("Chapter VI:  System Evolution")

    h2("6.1  Anticipated Changes")
    body(
        "Table 12 documents planned and foreseeable changes to the system, categorised by "
        "priority and estimated effort."
    )
    doc.add_paragraph()
    changes = [
        ("EV-01","Supabase PostgreSQL Migration","High","High","Replace JSON file data layer with Supabase PostgreSQL; update all db.ts calls to use Supabase query builder."),
        ("EV-02","Real-Time Notifications","High","Medium","Replace polling with Supabase Realtime subscriptions using existing useRealtimeNotifications hook."),
        ("EV-03","Payment Gateway Integration","High","High","Integrate Tap Payments or Checkout.com for GCC-compliant order payment processing."),
        ("EV-04","Email / SMS Transactional Notifications","Medium","Medium","SendGrid or AWS SES for order confirmation and review emails; Twilio for SMS alerts."),
        ("EV-05","Automated CR Verification","Medium","High","API integration with Kuwait Ministry of Commerce to auto-validate commercial registration numbers."),
        ("EV-06","Shipping Carrier Integration","Medium","High","Aramex / DHL API for real-time shipping rate calculation and parcel tracking."),
        ("EV-07","Live Exchange Rate API","Low","Low","Replace hardcoded EXCHANGE_RATES with a scheduled fetch from an FX data provider."),
        ("EV-08","Arabic Brand Descriptions","Low","Low","Add Arabic description field to Brand and Product entities; render in AR locale."),
        ("EV-09","Wishlist & Brand Comparison","Low","Medium","Allow buyers to save brands/products to a wishlist and compare up to 3 brands side-by-side."),
        ("EV-10","Native Mobile Apps","Low","High","React Native or PWA wrapper for iOS/Android; leverages the existing REST API."),
    ]
    make_colored_table(
        ["ID","Change","Priority","Effort","Description"],
        changes,
        [HEX_RED if r[2]=="High" else (HEX_YELLOW if r[2]=="Medium" else HEX_GREEN) for r in changes],
        col_widths=[1.2, 3.5, 1.8, 1.8, 8.2],
        font_size=9
    )

    h2("6.2  Phase-Out Transitions")
    body("The following transitions define how legacy components will be retired as the system evolves:")
    bullet("JSON → Supabase DB: The readDB()/writeDB() abstraction in db.ts will be replaced method-by-method with Supabase query calls.  The JSON file will be retained as a fallback seed fixture only.")
    bullet("Local Auth → Supabase Auth: The existing cookie-based local auth (b2b_user_id cookie) will be fully superseded by Supabase Auth JWTs.  The register/login/logout API routes will be deprecated in favour of Supabase client calls.")
    bullet("Hardcoded Exchange Rates → Live API: EXCHANGE_RATES in currencies.ts will be replaced with a cached API call (Next.js revalidation at 1-hour intervals), eliminating the need for manual rate updates.")
    bullet("Manual Brand Verification → Hybrid Automated: The current fully-manual admin approval process will be augmented (not replaced) with an automated CR lookup step that pre-populates admin review data.")

    h2("6.3  Quality Predictions")
    body(
        "Figure 15 (Appendix A) projects key quality metrics across the platform's first three "
        "post-launch phases.  Based on the current implementation trajectory:"
    )
    doc.add_paragraph()
    quality_rows = [
        ("API Response Time (avg)", "< 800 ms (JSON file)", "< 300 ms (Supabase + CDN)", "< 150 ms (edge caching)"),
        ("Test Coverage",           "0% automated (manual only)", "~40% (unit + E2E critical paths)", "~80% (full suite)"),
        ("Translation Coverage",    "~80% (EN complete, AR ~80%)", "100% (all keys filled)", "100% + new features"),
        ("Uptime",                  "N/A (dev env)",        "~99.5% (Vercel + Supabase SLA)", "> 99.9% (HA config)"),
        ("Security Score (OWASP)",  "Medium (plain-text PW)","High (bcrypt, HTTPS, CSP)", "High (pen-test verified)"),
    ]
    make_table(
        ["Metric", "MVP Phase (Now)", "Phase 2 (Post-Migration)", "Phase 3 (Production)"],
        quality_rows,
        col_widths=[4.0, 4.0, 4.0, 4.5],
        font_size=9
    )
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER VII — PLANNING & MANAGEMENT
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_7():
    h1("Chapter VII:  System Planning & Management")

    h2("7.1  Milestones")
    doc.add_paragraph()
    milestones = [
        ("M-01", "Project Kickoff & Tech Stack Decision",           "Jan 1, 2026",  "Complete", "Architecture agreed; Next.js + Supabase selected"),
        ("M-02", "Authentication & DB Schema",                      "Jan 15, 2026", "Complete", "User registration, login, roles functional"),
        ("M-03", "Marketplace MVP (browse, search, catalog)",        "Feb 1, 2026",  "Complete", "Brand cards, product detail, verification badges"),
        ("M-04", "Order Lifecycle Complete",                         "Feb 20, 2026", "Complete", "Order create → approve/reject; buyer/brand dashboards"),
        ("M-05", "Trust & Reputation Engine",                        "Mar 10, 2026", "Complete", "Review system, ratings, trust scores, admin moderation"),
        ("M-06", "Admin Dashboard & Notifications",                  "Mar 25, 2026", "Complete", "Admin panel, notification panel, analytics"),
        ("M-07", "Localisation & Mobile Responsiveness",             "Apr 10, 2026", "Partial",  "EN complete; AR ~80%; mobile mostly responsive"),
        ("M-08", "Progress Report Submission",                       "Apr 17, 2026", "Complete", "Progress_Report_Final.docx submitted"),
        ("M-09", "Supabase DB Migration",                            "May 5, 2026",  "Planned",  "Full data layer on Supabase PostgreSQL"),
        ("M-10", "Automated Testing Suite",                          "May 15, 2026", "Planned",  "Jest unit tests + Playwright E2E for critical paths"),
        ("M-11", "Final Report First Draft",                         "May 25, 2026", "Planned",  "All chapters drafted and internally reviewed"),
        ("M-12", "Final Report Submission",                          "Jun 15, 2026", "Planned",  "Submitted to university portal"),
        ("M-13", "Final Presentation / Viva",                        "Jun 25, 2026", "Planned",  "Presentation delivered; Q&A handled"),
    ]
    make_colored_table(
        ["ID", "Milestone", "Target Date", "Status", "Notes"],
        milestones,
        [HEX_GREEN if r[3]=="Complete" else (HEX_YELLOW if r[3]=="Partial" else HEX_LBLUE) for r in milestones],
        col_widths=[1.2, 5.5, 2.5, 2.0, 5.3],
        font_size=9
    )

    h2("7.2  Gantt Chart")
    body(
        "Figure 13 (Appendix A) shows the project Gantt chart spanning January–June 2026.  "
        "The chart is divided into eight sprints of approximately two weeks each.  "
        "Critical-path activities are: (1) Auth & DB Schema → (2) Marketplace → (3) Orders → "
        "(4) Trust & Reputation → (5) Admin & Notifications → (6) Localisation → "
        "(7) Supabase Migration → (8) Testing & Report."
    )
    body("Approximate sprint schedule:")
    doc.add_paragraph()
    gantt_rows = [
        ("Sprint 1", "Jan 1–14",    "Auth, DB schema, project setup"),
        ("Sprint 2", "Jan 15–31",   "Marketplace, brand browsing, search/filter"),
        ("Sprint 3", "Feb 1–20",    "Order lifecycle, buyer dashboard"),
        ("Sprint 4", "Feb 21–Mar 9","Brand owner dashboard, product CRUD, Excel import"),
        ("Sprint 5", "Mar 10–24",   "Reviews, trust scores, admin moderation"),
        ("Sprint 6", "Mar 25–Apr 9","Admin dashboard, notifications, WhatsApp"),
        ("Sprint 7", "Apr 10–Apr 30","Localisation, mobile responsiveness, progress report"),
        ("Sprint 8", "May 1–Jun 15","Supabase migration, testing, final report"),
    ]
    make_table(
        ["Sprint", "Dates", "Focus"],
        gantt_rows,
        col_widths=[2.0, 3.5, 11.0],
        font_size=10
    )

    h2("7.3  Staff Allocation")
    body("Table 14 shows the effort distribution across team members and sprint phases.")
    doc.add_paragraph()
    staff_rows = [
        ("Bader Alnoumas",  "25%","25%","20%","15%","10%","5%","100%"),
        ("Team Member 2",   "10%","30%","20%","15%","15%","10%","100%"),
        ("Team Member 3",   "15%","10%","25%","25%","15%","10%","100%"),
        ("Team Member 4",   "10%","10%","10%","10%","30%","30%","100%"),
    ]
    make_table(
        ["Team Member", "Auth/Setup", "Marketplace", "Orders", "Trust/Admin", "Testing", "Docs/Report", "Total"],
        staff_rows,
        col_widths=[3.5, 2.0, 2.5, 2.0, 2.5, 2.0, 2.5, 1.5],
        font_size=9
    )

    h2("7.4  CoCoMo Cost Estimation")
    body(
        "The Constructive Cost Model (CoCoMo II) is applied to estimate development effort.  "
        "The system is classified as Organic (relatively small, well-understood problem domain, "
        "experienced team)."
    )
    doc.add_paragraph()
    cocomo_rows = [
        ("Estimated KLOC",             "3.0 KLOC",   "~3,000 lines of TypeScript/TSX (excluding node_modules and generated files)"),
        ("CoCoMo Model",               "Organic",    "Small team, familiar technology, relatively simple problem domain"),
        ("Effort Equation",            "E = 2.4 × (KLOC)^1.05", "E = 2.4 × (3.0)^1.05 ≈ 7.6 person-months"),
        ("Estimated Effort",           "~7.6 person-months", "Distributed across 4 team members over ~2 months"),
        ("Duration Equation",          "D = 2.5 × E^0.38", "D = 2.5 × (7.6)^0.38 ≈ 6.1 months"),
        ("Estimated Duration",         "~6 months",  "Jan 2026 – Jun 2026 (aligns with actual schedule)"),
        ("Avg Team Size",              "4 developers","7.6 person-months / 6 months ≈ 1.27 FTE each"),
        ("Estimated Cost (KD/month)",  "KD 500/developer", "KD 500 × 4 × 6 = KD 12,000 total (academic estimate)"),
    ]
    make_table(
        ["Parameter", "Value", "Notes"],
        cocomo_rows,
        col_widths=[4.0, 3.5, 9.0],
        font_size=10
    )

    h2("7.5  Risk Matrix")
    body("Table 16 documents the project risk matrix with likelihood, impact, and mitigation strategy.")
    doc.add_paragraph()
    risks = [
        ("R-01","Supabase migration takes longer than estimated","Medium","High","High","Start migration in feature branch; JSON fallback available; allocate full Sprint 8 if needed"),
        ("R-02","Automated tests surface critical bugs near submission","Medium","High","High","Prioritise critical-path tests first; accept partial coverage; manual test sweep as backup"),
        ("R-03","Team member unavailable near submission","Low","High","Medium","Shared codebase knowledge; documentation kept current; cross-training on all modules"),
        ("R-04","Real-time Supabase Realtime integration blocked","Medium","Low","Low","Fall back to short-interval polling; Realtime is an enhancement, not critical for MVP"),
        ("R-05","Translation completeness gaps in Arabic","High","Medium","High","Audit scripts run weekly; dedicated team member assigned; English fallback is functional"),
        ("R-06","Vercel/Supabase free-tier limits during testing","Low","Medium","Low","Monitor usage; upgrade plan if needed; cost is negligible for academic testing volumes"),
        ("R-07","Scope creep — new features requested","Low","Medium","Low","Feature freeze in Sprint 8; all new items logged in post-MVP backlog only"),
        ("R-08","Security vulnerability discovered in review","Medium","High","High","Supabase Auth migration resolves password hashing; CSP headers to be added before launch"),
    ]
    make_colored_table(
        ["ID","Risk","Likelihood","Impact","Exposure","Mitigation"],
        risks,
        [HEX_RED if r[4]=="High" else (HEX_YELLOW if r[4]=="Medium" else HEX_GREEN) for r in risks],
        col_widths=[1.0, 3.5, 2.0, 1.8, 2.0, 6.2],
        font_size=8
    )

    h2("7.6  Fishbone Diagram")
    body(
        "Figure 14 (Appendix A) presents the Ishikawa (Fishbone) Diagram for the primary project "
        "risk: 'Delayed final submission'.  The diagram identifies five cause categories:"
    )
    bullet("People: Team member unavailability; knowledge silos; communication gaps.")
    bullet("Process: Scope creep; insufficient sprint planning; unclear acceptance criteria.")
    bullet("Technology: Supabase migration complexity; real-time implementation difficulty; dependency version conflicts.")
    bullet("Environment: Vercel/Supabase platform outages; local development environment inconsistencies.")
    bullet("Materials: Incomplete requirements; missing design artefacts; documentation gaps.")
    body("Mitigations for each branch are documented in the Risk Matrix (Table 16).")
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER VIII — CONCLUSIONS
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_8():
    h1("Chapter VIII:  Conclusions")

    h2("8.1  Design Smells Identified and Addressed")
    body(
        "Seven design smells were identified during the Sprint 6 code inspection and are "
        "documented below with their resolution status:"
    )
    smells = [
        ("DS-01", "God Object", "The db.ts module was a single read/write gateway to all 10 entity types, creating tight coupling between all consumers.", "Partially resolved by introducing typed helper functions per entity; full resolution requires migration to Supabase's per-table query interface."),
        ("DS-02", "Primitive Obsession", "Order status was initially passed as raw strings ('pending', 'approved') across the codebase without a shared enum.", "Resolved: TypeScript union type 'pending' | 'approved' | 'rejected' | 'completed' defined in db.ts and used consistently."),
        ("DS-03", "Divergent Change", "The currency formatting logic was duplicated across three components before being centralised.", "Resolved: All price formatting now delegates to formatPrice() / formatPriceWithVAT() in currencies.ts."),
        ("DS-04", "Shotgun Surgery", "Adding a new notification type required changes in the API route, the notification type constant, and the icon-map in the UI.", "Partially resolved: notification creation is centralised in notifications.ts; icon-map still requires manual updates."),
        ("DS-05", "Feature Envy", "Review list components were computing aggregates (average rating, breakdown) that belonged in the data layer.", "Resolved: Aggregates computed server-side in GET /api/reviews and returned in the response payload."),
        ("DS-06", "Inappropriate Intimacy", "The order creation route directly accessed db.brands and db.products for validation instead of using a service layer.", "Acknowledged: the MVP architecture uses direct DB access in routes; a service layer is planned for the production architecture."),
        ("DS-07", "Magic Numbers", "Exchange rates and VAT percentages were hardcoded inline in several components.", "Resolved: All rates and percentages centralised in CURRENCIES and EXCHANGE_RATES constants in currencies.ts."),
    ]
    for ds_id, name, problem, resolution in smells:
        h3(f"{ds_id}  {name}")
        body_bold("Problem:", problem)
        body_bold("Resolution:", resolution)

    h2("8.2  SWOT Analysis")
    body("Table 17 presents the SWOT analysis for Kuwait B2B Hub as a product.")
    doc.add_paragraph()
    swot = [
        ("STRENGTHS",
         "• Purpose-built for GCC market: Arabic RTL, 4 currencies, GCC legal awareness\n"
         "• Verified brand ecosystem creates buyer trust not available on generic platforms\n"
         "• Structured order management eliminates WhatsApp chaos\n"
         "• Trust & reputation engine differentiates from simple listing directories\n"
         "• Modern tech stack (Next.js, Supabase, TypeScript) supports rapid iteration"),
        ("WEAKNESSES",
         "• No payment integration limits the platform to an 'order request' model\n"
         "• JSON-file database is not production-ready; migration risk\n"
         "• No real-time communication; buyers must poll for order status\n"
         "• Limited team size constrains delivery velocity\n"
         "• Brand descriptions and reviews are English-only at MVP"),
        ("OPPORTUNITIES",
         "• GCC B2B e-commerce is a multi-billion dollar underserved market\n"
         "• Kuwait Vision 2035 and Saudi Vision 2030 prioritise digital commerce\n"
         "• WhatsApp's dominance in GCC highlights the pain point this platform solves\n"
         "• Subscription model (KD 29/month) offers predictable recurring revenue\n"
         "• Potential to expand into adjacent GCC markets (Oman, Qatar, Egypt)"),
        ("THREATS",
         "• Established regional competitors (Tradeling, Salla, Zid) may expand into wholesale\n"
         "• Low switching cost: brand owners can return to WhatsApp if platform adds friction\n"
         "• Regulatory risk: Kuwait data localisation laws may require on-premise hosting\n"
         "• Market education required: many SMBs are unfamiliar with digital B2B procurement\n"
         "• Currency and VAT regulation changes across GCC markets"),
    ]
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = ["STRENGTHS", "WEAKNESSES", "OPPORTUNITIES", "THREATS"]
    colors = [HEX_GREEN, HEX_RED, HEX_LBLUE, HEX_YELLOW]
    for idx, (label, content) in enumerate(swot):
        row_i, col_i = divmod(idx, 2)
        cell = tbl.rows[row_i].cells[col_i]
        shade_cell(cell, colors[idx])
        p1 = cell.paragraphs[0]
        r1 = p1.add_run(label + "\n")
        r1.bold = True; r1.font.size = Pt(11)
        r2 = p1.add_run(content)
        r2.font.size = Pt(9)
        p1.paragraph_format.space_before = Pt(4)
        p1.paragraph_format.space_after  = Pt(4)
    for c in tbl.columns:
        for cell in c.cells:
            cell.width = Cm(8.25)

    h2("8.3  Future Work")
    body("Building on the MVP foundation, the following work is prioritised for post-submission phases:")
    numbered("Payment Gateway (Tap Payments / Checkout.com): Enable end-to-end B2B transactions within the platform, reducing reliance on off-platform bank transfers.")
    numbered("Supabase Realtime Messaging: Implement order-level chat between buyers and brand owners, replacing WhatsApp for in-platform negotiation.")
    numbered("Automated CR Verification: API integration with the Kuwait Ministry of Commerce to validate commercial registration numbers at registration time.")
    numbered("Shipping Integration (Aramex / DHL): Real-time shipping rates, pickup scheduling, and parcel tracking embedded in the order lifecycle.")
    numbered("Native Mobile Application: React Native or PWA implementation to serve the majority of GCC users who primarily browse on mobile devices.")
    numbered("AI-Powered Recommendations: Collaborative filtering to recommend brands and products to buyers based on order history and browsing behaviour.")
    numbered("Multi-Brand Order Basket: Allow buyers to add products from multiple brands into a single checkout session with consolidated order management.")
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER IX — APPENDICES
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_9():
    h1("Chapter IX:  Appendices")

    h2("Appendix A — Diagram References")
    body(
        "The following diagrams are referenced throughout this report.  Where diagrams are "
        "described textually (due to format constraints), this appendix provides the visual "
        "specification sufficient for rendering in a diagramming tool such as draw.io or Lucidchart."
    )
    diag_list = [
        ("Figure 1",  "System Architecture Diagram",               "Four-layer: Presentation → API → Data Access → Infrastructure.  Show Next.js pages/components → Route Handlers → db.ts (JSON) / Supabase client → Vercel + Supabase cloud."),
        ("Figure 2",  "Level-0 Context DFD",                       "System boundary labelled 'Kuwait B2B Hub'.  External entities: Admin, Brand Owner, Business Buyer, Supabase Auth.  Data flows in/out of system boundary."),
        ("Figure 3",  "Level-1 Data Flow Diagram",                  "Five processes: P1 Auth, P2 Catalog, P3 Orders, P4 Trust, P5 Admin.  Four data stores: D1 Users, D2 Brands/Products, D3 Orders, D4 Reviews/Notifications."),
        ("Figure 7",  "Sequence Diagram — Registration & Login",    "Lifelines: Browser, /api/auth/register, DB, Supabase.  Steps as described in §3.1.1."),
        ("Figure 8",  "Sequence Diagram — Order Lifecycle",         "Lifelines: Buyer, /api/orders, DB, BrandOwner.  Steps as described in §3.1.2."),
        ("Figure 9",  "Sequence Diagram — Brand Verification",      "Lifelines: BrandOwner, /api/auth/register, Admin, /api/admin/brands, DB.  Steps as described in §3.1.3."),
        ("Figure 10", "Class Diagram — Core Entities",              "10 classes with attributes as defined in db.ts §2.5.  Show associations with multiplicity as described in §3.5."),
        ("Figure 11", "Activity Diagram — Order Placement",         "Start → Buyer views product → Enter quantity → [qty < MOQ] → Error; [qty ≥ MOQ] → [stock < qty] → Error; [OK] → Resolve price → Create order → Notify → End."),
        ("Figure 12", "Physical Deployment Diagram",                "Nodes: User Device, Vercel Edge, Next.js Serverless Function, Supabase Cloud.  Connections: HTTPS, PostgreSQL TCP."),
        ("Figure 13", "Gantt Chart",                                "8 sprints, Jan–Jun 2026.  Bars for each sprint/milestone.  Critical path highlighted."),
        ("Figure 14", "Fishbone Diagram",                           "Effect: 'Delayed final submission'.  5 cause bones: People, Process, Technology, Environment, Materials."),
        ("Figure 15", "Quality Prediction Chart",                   "Bar/line chart with 5 metrics across 3 phases as per Table in §6.3."),
    ]
    make_table(
        ["Figure", "Title", "Diagram Specification"],
        diag_list,
        col_widths=[1.5, 4.0, 11.0],
        font_size=9
    )

    h2("Appendix B — Key Code Listings")
    body("Selected source code snippets illustrating critical implementation patterns:")

    h3("B.1  Order Validation Logic (src/app/api/orders/route.ts, lines 58–79)")
    code = doc.add_paragraph()
    code.paragraph_format.left_indent = Cm(0.5)
    r = code.add_run(
        "for (const item of items) {\n"
        "  const qty = Number(item.quantity);\n"
        "  if (!Number.isInteger(qty) || qty <= 0)\n"
        "    return error('Invalid quantity');\n"
        "  const product = products.find(p => p.id === item.product_id);\n"
        "  if (!product) return error('Product not found');\n"
        "  if (qty < product.moq)\n"
        "    return error(`MOQ is ${product.moq} units`);\n"
        "  if (product.stock < qty)\n"
        "    return error(`Only ${product.stock} units available`);\n"
        "  // Resolve bulk-pricing tier\n"
        "  let unitPrice = product.price;\n"
        "  for (const tier of product.bulk_pricing_tiers) {\n"
        "    if (qty >= tier.min_qty && (!tier.max_qty || qty <= tier.max_qty))\n"
        "      { unitPrice = tier.price; break; }\n"
        "  }\n"
        "  total += unitPrice * qty;\n"
        "}"
    )
    r.font.name = 'Courier New'; r.font.size = Pt(9)

    h3("B.2  Non-Fatal Notification Pattern (src/lib/notifications.ts)")
    code2 = doc.add_paragraph()
    code2.paragraph_format.left_indent = Cm(0.5)
    r2 = code2.add_run(
        "export function createNotification(opts: CreateNotifOpts) {\n"
        "  try {\n"
        "    const db = readDB();\n"
        "    db.notifications?.push({ ...opts, id: generateId('notif'),\n"
        "      read: false, created_at: new Date().toISOString() });\n"
        "    writeDB(db);\n"
        "  } catch {\n"
        "    // non-fatal — never crash a request\n"
        "  }\n"
        "}"
    )
    r2.font.name = 'Courier New'; r2.font.size = Pt(9)

    h2("Appendix C — Sample Data")
    body("The seeded demo dataset (src/data/db.json) includes:")
    bullet("7 user accounts: 1 admin, 3 brand owners, 3 buyers")
    bullet("3 brands: Kuwait Fashion House (Premium), Gulf Tech Solutions (Premium), Desert Rose Beauty (Premium)")
    bullet("2 products with bulk-pricing tiers (Royal Kandura, Windows test product)")
    bullet("2 orders: 1 completed (KD 1,625.00), 1 pending (KD 487.50)")
    bullet("10 reviews across all 3 brands with ratings 4–5 stars")
    bullet("19 notifications across all user roles covering all notification types")
    bullet("3 brand trust records, 2 buyer trust records")

    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER X — REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_10():
    h1("Chapter X:  References")
    body("The following references are cited in APA / IEEE style.")
    doc.add_paragraph()
    refs = [
        "[1]  Sommerville, I. (2016). *Software Engineering* (10th ed.). Pearson Education.",
        "[2]  Pressman, R. S., & Maxim, B. R. (2020). *Software Engineering: A Practitioner's Approach* (9th ed.). McGraw-Hill.",
        "[3]  Vercel Inc. (2024). *Next.js Documentation* (v16). https://nextjs.org/docs",
        "[4]  Supabase Inc. (2024). *Supabase Documentation — Auth & PostgreSQL*. https://supabase.com/docs",
        "[5]  Beck, K., et al. (2001). *Manifesto for Agile Software Development*. https://agilemanifesto.org",
        "[6]  Boehm, B. W. (1981). *Software Engineering Economics*. Prentice-Hall.",
        "[7]  Fowler, M. (1999). *Refactoring: Improving the Design of Existing Code*. Addison-Wesley.",
        "[8]  Microsoft. (2024). *TypeScript Handbook*. https://www.typescriptlang.org/docs/handbook/",
        "[9]  World Wide Web Consortium (W3C). (2018). *Web Content Accessibility Guidelines (WCAG) 2.1*. https://www.w3.org/TR/WCAG21/",
        "[10] React Team. (2024). *React 19 Documentation*. https://react.dev",
        "[11] GCC Secretariat General. (2023). *GCC Digital Economy Report 2023*. Gulf Cooperation Council.",
        "[12] Kuwait Ministry of Commerce and Industry. (2023). *Kuwait Vision 2035 — Digital Transformation Pillar*. State of Kuwait.",
        "[13] OWASP Foundation. (2023). *OWASP Top Ten Web Application Security Risks*. https://owasp.org/www-project-top-ten/",
        "[14] Recharts Team. (2024). *Recharts: Redefined Chart Library Built with React and D3*. https://recharts.org",
        "[15] GitHub — pmndrs/zustand. (2024). *Zustand: Bear necessities for state management in React*. https://github.com/pmndrs/zustand",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.first_line_indent = Cm(-0.6)
        r = p.add_run(ref.replace('*',''))  # strip asterisks; italics done separately
        r.font.size = Pt(10)

# ═══════════════════════════════════════════════════════════════════════════════
# BUILD THE DOCUMENT
# ═══════════════════════════════════════════════════════════════════════════════

cover_page()
toc()
list_of_tables()
list_of_figures()
abstract()
chapter_1()
chapter_2()
chapter_3()
chapter_4()
chapter_5()
chapter_6()
chapter_7()
chapter_8()
chapter_9()
chapter_10()

# ── Save ─────────────────────────────────────────────────────────────────────
out = "/Users/baderalnoumas/Desktop/B2B/Final_Project_Report.docx"
doc.save(out)
print(f"Saved: {out}")
