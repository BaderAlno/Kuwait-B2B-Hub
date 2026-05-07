#!/usr/bin/env python3
"""
Kuwait B2B Hub — Final Report Generator (Phase 2)
Opens Final_Report_v1.docx, appends Chapters III and IV,
saves as Final_Report_v2.docx.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ───────────────────────────────────────────────────────────────────

BUILD_DIR   = "/Users/baderalnoumas/Desktop/B2B/Final_Report_Build"
INPUT_FILE  = os.path.join(BUILD_DIR, "Final_Report_v1.docx")
OUTPUT_FILE = os.path.join(BUILD_DIR, "Final_Report_v2.docx")

doc = Document(INPUT_FILE)

# ─── Helpers (same contract as Phase 1) ──────────────────────────────────────

def page_break(doc):
    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def body(doc, text):
    doc.add_paragraph(text, style="Normal")


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def make_table(doc, headers, rows, col_widths=None):
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return table


def placeholder(doc, fig_num, caption):
    """Grey centred placeholder box for a diagram image."""
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ FIGURE {fig_num} — {caption} ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("(Insert diagram image here)")
    r2.italic = True
    r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r2.font.size = Pt(9)
    doc.add_paragraph()


def monospace_block(doc, text, font_size=8):
    """Render a preformatted text block in Courier New."""
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size  = Pt(font_size)
    doc.add_paragraph()


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER III — SYSTEM DESIGN
# ════════════════════════════════════════════════════════════════════════════

page_break(doc)
h1(doc, "CHAPTER III — SYSTEM DESIGN")

body(doc,
    "This chapter presents the complete system design of the Kuwait B2B Hub. It covers the "
    "behavioural models (sequence diagrams and activity model), the structural models (class "
    "diagram and component design), the data flow models, the architectural pattern, finite state "
    "machine specifications for the two core workflows, and prototype screenshots illustrating "
    "the delivered user interface."
)

# ── §3.1 Sequence Diagrams ───────────────────────────────────────────────────

h2(doc, "3.1  Sequence Diagrams")

body(doc,
    "Three sequence diagrams model the time-ordered message exchanges between system participants "
    "for the platform's most critical workflows. Each diagram is described narratively below; "
    "the corresponding rendered diagram is inserted as a figure placeholder for final production."
)

# 3.1.1
h3(doc, "3.1.1  User Registration & Login (Figure 1)")

body(doc,
    "This diagram models two consecutive interactions: a new user registering an account (POST "
    "/api/auth/register) and an existing user logging in (POST /api/auth/login). Four participants "
    "are shown: Browser, API Server, Database, and Notification Service."
)

body(doc,
    "Registration flow: The browser submits the registration form payload (name, email, password, "
    "role, company_name). The API validates all fields (400 if missing) then reads the database "
    "snapshot. If the email is already registered it returns 409 Conflict. Otherwise the API "
    "generates a user ID, creates the user record, and — if the role is brand_owner — also creates "
    "a pending brand record and a second admin notification. Verification status is set to "
    "'approved' for buyers and 'pending' for brand owners. The API then writes the updated database "
    "and sets two HTTP-only session cookies (b2b_user_id with httpOnly=true, b2b_user_role with "
    "httpOnly=false) carrying a 7-day maxAge on the 200 response."
)

body(doc,
    "Login flow: The browser submits email and password. The API reads the database, finds the "
    "matching user record, and compares credentials. On a match it attaches the same two session "
    "cookies to the 200 response. On a mismatch it returns 401 Invalid email or password. No "
    "database write occurs during a successful login — only cookies are set."
)

body(doc,
    "Design note: the authentication for these routes is entirely custom — no Supabase Auth SDK "
    "calls are made in register/route.ts or login/route.ts. Sessions are established via HTTP-only "
    "cookies set directly on the NextResponse object."
)

placeholder(doc, 1, "Sequence Diagram — User Registration & Login")

# 3.1.2
h3(doc, "3.1.2  Order Placement & Review Lifecycle (Figure 2)")

body(doc,
    "This diagram models the full buyer-to-brand order lifecycle across two phases: order "
    "submission and order status management. Four participants: Buyer, API Server, Database, "
    "Brand Owner."
)

body(doc,
    "Order placement phase: The buyer sends POST /api/orders with brand_id and an items array. "
    "The API calls getCurrentUser() to resolve the session; if the role is not buyer it returns "
    "403. The API reads the database to locate the brand (404 if not found or not approved). "
    "For each item in the order the API validates product membership in the brand, enforces MOQ "
    "(400 if quantity below minimum), enforces stock level (400 if insufficient), and resolves "
    "the unit price by iterating bulk_pricing_tiers to match the quantity range — falling back to "
    "the base product price if no tier matches. The API then inserts the order record and order "
    "items atomically, creates a 'new_order' notification for the brand owner, and returns 201 "
    "Created with the order object."
)

body(doc,
    "Status update phase: The brand owner sends PATCH /api/orders/:id with a new status "
    "(approved, rejected, or completed). The API validates the role (403 if not brand_owner or "
    "admin), verifies the order belongs to the brand, updates the status, and dispatches a "
    "buyer notification whose type matches the new status: order_approved, order_rejected, or "
    "order_completed. The 'order_completed' notification includes a call-to-action prompting "
    "the buyer to leave a review."
)

placeholder(doc, 2, "Sequence Diagram — Order Placement & Review Lifecycle")

# 3.1.3
h3(doc, "3.1.3  Brand Verification (Figure 3)")

body(doc,
    "This diagram models the administrative brand verification workflow. Five participants: "
    "Admin, API Server, Database, Brand Owner, Notification Service."
)

body(doc,
    "The admin sends PATCH /api/admin/brands/:id with a status field and an optional "
    "verification_tier. The API calls getCurrentUser() and returns 403 if the role is not admin. "
    "The API reads the database to locate the brand (404 if not found) and validates the status "
    "value (400 if not in {approved, rejected}). On approval: the API atomically updates "
    "brand.status = 'approved' and the owner's verification_status = 'approved', creates an "
    "'account_approved' notification for the brand owner, and dispatches a 'new_brand' "
    "notification to every active buyer on the platform. The brand becomes visible on the "
    "marketplace immediately. On rejection: the status is set to 'rejected' on both records and "
    "only the brand owner receives an 'account_rejected' notification. If a verification_tier "
    "is provided in the request body, the API performs a second database read to obtain a fresh "
    "snapshot, updates the tier field, and persists — preventing stale-state overwrites."
)

placeholder(doc, 3, "Sequence Diagram — Brand Verification")

# ── §3.2 Data Flow Model ─────────────────────────────────────────────────────

h2(doc, "3.2  Data Flow Model")

body(doc,
    "Data Flow Diagrams (DFDs) decompose the system into processes, data stores, and external "
    "entities, showing how data moves through the platform. Two levels are presented: Level 0 "
    "(Context Diagram) and Level 1 (sub-process decomposition)."
)

h3(doc, "3.2.1  DFD Level 0 — Context Diagram (Figure 4)")

body(doc,
    "At the context level the Kuwait B2B Hub is treated as a single process 'B2B Platform' "
    "sitting at the centre of four external entities:"
)

make_table(doc,
    ["External Entity", "Data Flows INTO Platform", "Data Flows OUT OF Platform"],
    [
        ("Business Buyers",
         "Registration data; login credentials; search/filter queries; order requests (brand_id, "
         "product_id, quantity); review submissions (rating, content, anonymous flag); "
         "notification read events.",
         "Session cookies; brand/product catalogue listings; order confirmations and status "
         "updates; notifications; trust scores; multi-currency prices."),
        ("Brand Owners",
         "Brand profile data; product listings (name, price, MOQ, stock, tiers); order status "
         "updates (approved / rejected / completed); review replies; Excel catalogue files.",
         "Session cookies; order request details; review notifications; brand analytics "
         "dashboard data; trust score and badge assignments."),
        ("Admin Users",
         "Brand verification decisions (approve/reject, tier); user management actions; "
         "review moderation commands (flag, remove).",
         "Pending brand queue; platform-wide analytics (user count, order volume, revenue); "
         "review moderation queue."),
        ("External Services",
         "JWT session validation responses (Supabase Auth); PostgreSQL query results "
         "(Supabase DB); live GCC currency exchange rates (Exchange Rate API).",
         "JWT authentication requests; database read/write queries; WhatsApp contact deep-links; "
         "Git push triggers for Vercel CI/CD pipeline."),
    ],
    col_widths=[1.4, 2.7, 2.4]
)

placeholder(doc, 4, "DFD Level 0 — Context Diagram")

h3(doc, "3.2.2  DFD Level 1 — Sub-Process Decomposition (Figure 5)")

body(doc,
    "At Level 1 the single 'B2B Platform' process is decomposed into five sub-processes, each "
    "corresponding to a platform subsystem. Three data stores are shared across processes:"
)

make_table(doc,
    ["Data Store", "Contents"],
    [
        ("DS-1 Users & Sessions",
         "User profiles (id, name, email, password_hash, role, company_name, "
         "verification_status, created_at); active session cookie references."),
        ("DS-2 Brands & Products",
         "Brand records (id, owner_id, brand_name, status, verification_tier, "
         "whatsapp_clicks, ...); Product records (id, brand_id, name, price, moq, stock, "
         "bulk_pricing_tiers[], ...)."),
        ("DS-3 Orders, Reviews & Notifications",
         "Order records + Order Items; Review records (with anonymous flag and brand_reply); "
         "Notification records (user_id, type, read, action_url); Message records."),
    ],
    col_widths=[1.8, 4.7]
)

body(doc, "Sub-processes and their data flows:")

make_table(doc,
    ["Process", "Inputs", "Outputs", "Reads / Writes"],
    [
        ("P1 Authentication & User Mgmt",
         "Registration form; login credentials; session cookie",
         "Session cookies; user profile; admin notifications",
         "Reads + Writes DS-1"),
        ("P2 Brand Management",
         "Brand profile updates; verification decisions; WhatsApp click events",
         "Updated brand records; brand owner / buyer notifications; marketplace listings",
         "Reads + Writes DS-1, DS-2"),
        ("P3 Product & Order Management",
         "Product CRUD payloads; Excel import files; order requests; status updates",
         "Product records; order confirmations; status notifications; analytics data",
         "Reads + Writes DS-2, DS-3"),
        ("P4 Trust, Reviews & Reputation",
         "Review submissions; review replies; moderation commands; order history",
         "Active reviews; brand replies; trust scores; buyer trust profiles; badges",
         "Reads DS-3; Writes DS-3"),
        ("P5 Notifications & Real-Time",
         "System events from P1–P4; notification read events",
         "Notification panel content; unread badge count; real-time delivery (planned)",
         "Reads + Writes DS-3"),
    ],
    col_widths=[1.5, 1.9, 1.9, 1.2]
)

placeholder(doc, 5, "DFD Level 1 — Sub-Process Decomposition")

# ── §3.3 Architectural Pattern ───────────────────────────────────────────────

h2(doc, "3.3  Architectural Pattern")

body(doc,
    "The Kuwait B2B Hub adopts a four-layer MVC-inspired architecture implemented on the "
    "Next.js 16 App Router platform. The design draws on the Model-View-Controller pattern "
    "while adapting it to the constraints and capabilities of a server-rendered, serverless "
    "deployment environment."
)

h3(doc, "3.3.1  Layer Mapping to MVC Concepts")

make_table(doc,
    ["MVC Concept", "Kuwait B2B Hub Realisation", "Key Files"],
    [
        ("Model",
         "TypeScript interfaces defined in src/lib/db.ts constitute the data model. Each "
         "interface (User, Brand, Product, Order, OrderItem, Review, Notification, BrandTrust, "
         "BuyerTrust, Message) serves as a compile-time contract enforced across all layers. "
         "The readDB() / writeDB() abstraction acts as the repository layer between "
         "the model and the underlying storage.",
         "src/lib/db.ts"),
        ("View",
         "React/TSX components and Next.js page routes render all UI. Server Components "
         "generate SSR HTML on first load for SEO and performance; Client Components add "
         "interactivity via browser-side hydration. CSS Modules provide scoped styling "
         "without a utility-class framework dependency.",
         "src/app/**/page.tsx; src/components/*.tsx"),
        ("Controller",
         "Next.js App Router Route Handlers serve as controllers, receiving HTTP requests, "
         "applying RBAC guards, delegating to the data layer, and returning JSON responses. "
         "src/middleware.ts acts as a front-controller intercepting all requests for "
         "session validation before routing proceeds.",
         "src/app/api/**/route.ts; src/middleware.ts"),
        ("Service / Infrastructure",
         "Cross-cutting concerns (auth resolution, notification creation, currency formatting, "
         "ID generation) are encapsulated in src/lib/* utility modules consumed by "
         "controllers. Supabase and Vercel constitute the infrastructure layer beyond "
         "the application boundary.",
         "src/lib/auth.ts; src/lib/notifications.ts; src/lib/currencies.ts"),
    ],
    col_widths=[1.3, 3.4, 1.8]
)

h3(doc, "3.3.2  Next.js App Router Design Decisions")

for decision in [
    "Server Components are used for all data-fetching pages (marketplace, brand profiles, "
    "order listings), eliminating client-side loading spinners and enabling immediate "
    "content-available rendering for search engine crawlers.",
    "Client Components are restricted to interactive UI (notification bell, currency selector, "
    "review modal, order quantity input) where browser event listeners are required. "
    "The use='client' directive is applied only where necessary.",
    "API Route Handlers replace a separate Express or Fastify server, enabling co-location "
    "of frontend and backend code in a single repository and a single Vercel deployment. "
    "All 25 handlers are stateless serverless functions.",
    "Middleware (src/middleware.ts) runs at the edge before any route handler, validating "
    "the b2b_user_id cookie and redirecting unauthenticated requests to /login — providing "
    "a uniform auth gate without duplicating guard logic in every handler.",
    "The @/ path alias (configured in tsconfig.json paths) enables absolute imports across "
    "all 93 source files, eliminating fragile relative path chains and simplifying "
    "refactoring.",
]:
    bullet(doc, decision)

# ── §3.4 Activity/Process Model ──────────────────────────────────────────────

h2(doc, "3.4  Activity / Process Model — Order Placement")

body(doc,
    "The activity diagram below models the decision-heavy order placement workflow, capturing "
    "all guards, error paths, and database writes that occur between a buyer initiating an "
    "order and the system returning a 201 Created response. This workflow is implemented in "
    "the POST handler of src/app/api/orders/route.ts."
)

body(doc,
    "Process narrative: The buyer navigates to a brand's product page and enters an order "
    "quantity. The system first checks whether the entered quantity meets the product's "
    "minimum order quantity (MOQ). If the quantity is below the MOQ, the system displays an "
    "error (e.g., 'Minimum order quantity is N units') and the workflow terminates with no "
    "database writes. If the MOQ is satisfied, the system checks whether the requested quantity "
    "is available in stock. If the stock is insufficient the system displays a stock error and "
    "terminates. When both guards pass, the system iterates the product's bulk_pricing_tiers[] "
    "array to find the tier whose min_qty ≤ quantity ≤ max_qty and assigns that tier's price "
    "as the unit price. If no tier matches, the base product.price is used as the fallback. "
    "The system calculates the total order amount as the sum of (unit_price × quantity) across "
    "all order items. The system then performs two sequential database writes: (1) INSERT an "
    "order record with status='pending' and the computed total_amount; (2) INSERT the "
    "corresponding order_items rows. Finally, the system dispatches a 'new_order' notification "
    "to the brand owner and returns a 201 Created response with the full order object."
)

placeholder(doc, 6, "Activity Diagram — Order Placement")

# ── §3.5 Finite State Machine ────────────────────────────────────────────────

h2(doc, "3.5  Finite State Machine (FSM) Specifications")

body(doc,
    "Two finite state machines formally specify the permissible state transitions for the "
    "Kuwait B2B Hub's two principal lifecycle entities: orders and brand verification records. "
    "Both are modelled as Mealy machines — outputs (notifications and database writes) are "
    "associated with transitions rather than states."
)

h3(doc, "3.5.1  Order Lifecycle FSM")

body(doc,
    "States: pending (initial), approved, rejected, completed (terminal), trust_updated (silent "
    "terminal). An order is created in the pending state by the POST /api/orders handler and "
    "can only be advanced forward — no backward transitions are permitted."
)

make_table(doc,
    ["Transition #", "From State", "Event / Trigger", "Actor", "To State", "Output / Side-Effect"],
    [
        ("T-01", "—  (none)",   "POST /api/orders succeeds — buyer submits order",
         "Buyer",       "pending",
         "INSERT order + order_items; notify brand owner: 'New Order Received!'"),
        ("T-02", "pending",     "PATCH /api/orders/:id {status: 'approved'}",
         "Brand Owner", "approved",
         "UPDATE order.status; notify buyer: 'Order Approved!'"),
        ("T-03", "pending",     "PATCH /api/orders/:id {status: 'rejected'}",
         "Brand Owner", "rejected",
         "UPDATE order.status; notify buyer: 'Order Request Declined'"),
        ("T-04", "approved",    "PATCH /api/orders/:id {status: 'completed'}",
         "Brand Owner", "completed",
         "UPDATE order.status; notify buyer: 'Order Fulfilled! — leave a review?'"),
        ("T-05", "approved",    "PATCH /api/orders/:id {status: 'rejected'} by admin",
         "Admin",       "rejected",
         "UPDATE order.status; notify buyer: 'Order Request Declined' (admin override)"),
        ("T-06", "pending",     "GET /api/orders/:id — admin reviews order",
         "Admin",       "pending",
         "Read-only; no state change; no notification"),
        ("T-07", "approved",    "GET /api/orders/:id — brand owner views item detail",
         "Brand Owner", "approved",
         "Read-only; no state change; no notification"),
        ("T-08", "rejected",    "GET /api/orders/:id — buyer views outcome",
         "Buyer",       "rejected",
         "Read-only; no state change; no notification"),
        ("T-09", "completed",   "POST /api/reviews — buyer submits a review",
         "Buyer",       "completed",
         "Review record created; notify brand owner: 'New Review Posted'"),
        ("T-10", "completed",   "GET /api/trust/:id — trust recalculation triggered",
         "System",      "trust_updated",
         "BrandTrust.completion_rate and BuyerTrust metrics recomputed and persisted"),
        ("T-11", "rejected",    "GET /api/trust/:id — trust recalculation triggered",
         "System",      "trust_updated",
         "BuyerTrust.cancellation_rate incremented; BrandTrust response metrics updated"),
    ],
    col_widths=[0.6, 0.8, 1.8, 0.9, 0.8, 1.6]
)

h3(doc, "3.5.2  Brand Verification FSM")

body(doc,
    "States: pending (initial), approved, rejected, tier_updated (sub-state of approved). "
    "The brand FSM governs the administrative approval lifecycle from registration through to "
    "marketplace visibility."
)

make_table(doc,
    ["Transition #", "From State", "Event / Trigger", "Actor", "To State", "Output / Side-Effect"],
    [
        ("BV-01", "—  (none)",    "Brand owner registers → F-006 auto-creates brand",
         "Brand Owner", "pending",
         "INSERT brand {status: 'pending'}; notify all admins: 'New Brand Awaiting Approval'"),
        ("BV-02", "pending",      "PATCH /api/admin/brands/:id {status: 'approved'}",
         "Admin",       "approved",
         "UPDATE brand.status + owner.verification_status = 'approved'; notify brand owner: "
         "'Your Brand is Now Live!'; notify all buyers: 'New Brand on the Platform'"),
        ("BV-03", "pending",      "PATCH /api/admin/brands/:id {status: 'rejected'}",
         "Admin",       "rejected",
         "UPDATE brand.status + owner.verification_status = 'rejected'; notify brand owner: "
         "'Brand Verification Update — not approved'; buyers NOT notified"),
        ("BV-04", "approved",     "PATCH /api/admin/brands/:id {verification_tier: 'premium' | 'verified' | 'new'}",
         "Admin",       "tier_updated",
         "Fresh DB read → UPDATE brand.verification_tier; badge UI updated on all brand cards "
         "and profile pages; no buyer notification"),
        ("BV-05", "approved",     "Brand owner updates profile via PATCH /api/brands/:id",
         "Brand Owner", "approved",
         "UPDATE brand fields (name, description, logo_url, whatsapp_number, business_hours, "
         "auto_reply_message); no status change; no notification"),
    ],
    col_widths=[0.65, 0.9, 1.85, 0.9, 0.9, 1.3]
)

# ── §3.6 Component/Unit Design ───────────────────────────────────────────────

h2(doc, "3.6  Component / Unit Design")

body(doc,
    "The following six components constitute the core non-UI technical units of the Kuwait B2B "
    "Hub. Each is a discrete module with a defined public interface consumed by multiple "
    "callers. Their design directly shapes the maintainability and testability of the system."
)

make_table(doc,
    ["Component", "File", "Public Interface", "Callers", "Responsibility"],
    [
        ("Auth Resolver",
         "src/lib/auth.ts",
         "getCurrentUser(): Promise<User | null>",
         "All 25 API route handlers (first call in every handler body)",
         "Resolves Supabase JWT session to a full User profile from the database. Returns null "
         "if session is missing, expired, or the user record cannot be found. All protected "
         "routes treat a null return as HTTP 401."),
        ("Data Access Module",
         "src/lib/db.ts",
         "readDB(): DB; writeDB(db: DB): void; generateId(prefix: string): string",
         "All API route handlers; Trust Engine",
         "Single source of truth for all database reads and writes in the MVP. The readDB / "
         "writeDB abstraction decouples API handlers from the storage implementation, enabling "
         "the PostgreSQL migration without handler changes."),
        ("Notification Service",
         "src/lib/notifications.ts",
         "createNotification(userId, type, title, body, actionUrl): void; "
         "createNotificationForMany(userIds[], ...): void",
         "All state-changing API route handlers (orders, brands, reviews, auth)",
         "Constructs and persists notification records for targeted users or user groups. "
         "Abstracts the notification creation pattern used in all 10 system event types."),
        ("Currency Module",
         "src/lib/currencies.ts",
         "CURRENCIES: Record<string, CurrencyConfig>; formatCurrency(amount, code): string",
         "CurrencyContext.tsx; CurrencySelector.tsx; Exchange-rate API handler; "
         "all price-rendering components",
         "Single source of truth for all GCC currency configuration: ISO codes, display symbols, "
         "decimal precision, VAT rates, locale strings, and phone prefixes. Consumed by "
         "every component that formats or converts monetary values."),
        ("Notification Store",
         "src/store/notificationStore.ts",
         "useNotificationStore(): { notifications, unreadCount, setNotifications, markRead }",
         "NotificationBell.tsx; RealtimeProvider.tsx",
         "Zustand-powered client-side store for notification state. Maintains the unread "
         "count badge, the notification panel list, and the markRead action. Designed to "
         "accept data from either a REST poll or a future Supabase Realtime subscription."),
        ("Supabase Server Client Factory",
         "src/utils/supabase/server.ts",
         "createClient(): Promise<SupabaseClient>",
         "src/lib/auth.ts; any server-side code requiring an authenticated DB connection",
         "Creates a server-side Supabase client using the @supabase/ssr package, reading and "
         "writing session tokens exclusively through Next.js cookies(). Ensures auth state is "
         "scoped to the server request context and never exposed to client-side JavaScript. "
         "Foundational security primitive of the entire backend layer."),
    ],
    col_widths=[1.2, 1.3, 1.5, 1.3, 2.2]
)

# ── §3.7 Class Diagram ───────────────────────────────────────────────────────

h2(doc, "3.7  Class Diagram — Core Entity Model")

body(doc,
    "The class diagram is derived directly from the TypeScript interface definitions in "
    "src/lib/db.ts — the authoritative data contract for all ten entity types in the Kuwait "
    "B2B Hub. The table below enumerates each class, its attributes, and its key associations. "
    "Optional fields (marked with ? in TypeScript) are noted in the Attributes column. "
    "Figure 7 shows the rendered UML class diagram."
)

make_table(doc,
    ["Class", "Attributes", "Key Associations", "Notes"],
    [
        ("User",
         "id: String; name: String; email: String; password: String; role: String; "
         "company_name: String; verification_status: String; whatsapp_number?: String; "
         "created_at: String",
         "1 User owns 0..1 Brand (via owner_id); 1 User places 0..* Orders (via buyer_id); "
         "1 User writes 0..* Reviews (via buyer_id); 1 User receives 0..* Notifications (via user_id)",
         "role ∈ {admin, brand_owner, buyer}; "
         "verification_status ∈ {pending, approved, rejected}"),
        ("Brand",
         "id: String; owner_id: String; brand_name: String; description: String; "
         "logo_url: String; status: String; verification_tier: String; "
         "whatsapp_number?: String; business_hours?: String; auto_reply_message?: String; "
         "whatsapp_clicks?: Number; created_at: String",
         "1 Brand lists 0..* Products (via brand_id); 1 Brand receives 0..* Orders; "
         "1 Brand receives 0..* Reviews; 1 Brand has 0..1 BrandTrust (composition)",
         "status ∈ {pending, approved, rejected}; "
         "verification_tier ∈ {premium, verified, new}"),
        ("Product",
         "id: String; brand_id: String; name: String; description: String; price: Number; "
         "moq: Number; stock: Number; image_url: String; "
         "bulk_pricing_tiers: BulkPricingTier[]; created_at: String",
         "1 Product belongs to 1 Brand; 1 Product embeds 0..* BulkPricingTier (composition); "
         "1 Product is referenced by 0..* OrderItems",
         "bulk_pricing_tiers[] is a value-object array embedded inside Product — "
         "not a separate table"),
        ("BulkPricingTier",
         "min_qty: Number; max_qty: Number | null; price: Number",
         "Embedded in Product.bulk_pricing_tiers[] (no independent identity)",
         "max_qty = null means no upper bound (open-ended tier); "
         "tiers are iterated in order to match the requested quantity"),
        ("Order",
         "id: String; buyer_id: String; brand_id: String; status: String; "
         "total_amount: Number; created_at: String",
         "1 Order contains 1..* OrderItems (composition); 1 Order is placed by 1 User (buyer); "
         "1 Order belongs to 1 Brand; 0..1 Order is linked to 0..1 Review",
         "status ∈ {pending, approved, rejected, completed}"),
        ("OrderItem",
         "id: String; order_id: String; product_id: String; quantity: Number; "
         "unit_price: Number",
         "1 OrderItem belongs to 1 Order; 1 OrderItem references 1 Product",
         "unit_price is the bulk-pricing-resolved price at the time of order creation — "
         "immutable after insert"),
        ("Review",
         "id: String; brand_id: String; order_id: String | null; buyer_id: String; "
         "rating: Number; content: String; anonymous: Boolean; flagged: Boolean; "
         "brand_reply: String | null; status?: String; created_at: String",
         "1 Review is written by 1 User (buyer); 1 Review is received by 1 Brand; "
         "0..1 Review is linked to 0..1 Order (nullable)",
         "order_id is nullable — review can exist independently of a specific order; "
         "status ∈ {active, removed}"),
        ("Notification",
         "id: String; user_id: String; type: String; title: String; body: String; "
         "read: Boolean; action_url: String; icon_type: String; created_at: String",
         "1 Notification is addressed to 1 User (via user_id)",
         "Single flat reference — same schema serves buyers, brand owners, and admins; "
         "type field distinguishes the event category"),
        ("BrandTrust",
         "brand_id: String; response_rate: Number; completion_rate: Number; "
         "avg_response_hours: Number; total_fulfilled: Number; orders_this_month: Number; "
         "avg_fulfillment_days: Number; badges: String[]",
         "1 BrandTrust is composed by 1 Brand (via brand_id) — composition, not association",
         "Read-only trust profile; no back-reference to Order — "
         "keeps trust calculation decoupled from live order mutations"),
        ("BuyerTrust",
         "buyer_id: String; total_orders: Number; completion_rate: Number; "
         "cancellation_rate: Number; badges: String[]",
         "1 BuyerTrust is composed by 1 User-buyer (via buyer_id)",
         "Symmetric counterpart to BrandTrust; provides brand owners with buyer "
         "reliability context when evaluating incoming orders"),
    ],
    col_widths=[1.2, 2.5, 1.8, 1.0]
)

placeholder(doc, 7, "Class Diagram — Core Entities (extracted from src/lib/db.ts)")

# ── §3.8 Prototype Screenshots ───────────────────────────────────────────────

h2(doc, "3.8  Prototype Screenshots")

body(doc,
    "The following eight screenshots document the Kuwait B2B Hub interface as delivered at the "
    "time of final submission. All screenshots are taken from the live Vercel deployment against "
    "the seeded demo dataset."
)

screenshot_list = [
    ("SS-01", "Landing Page — Hero section, platform value proposition, bilingual toggle, call-to-action buttons for buyer and brand owner registration."),
    ("SS-02", "Marketplace — Brand catalogue grid with verification tier badges (Premium / Verified / New), star ratings, MOQ indicators, and search/filter controls."),
    ("SS-03", "Brand Profile Page — Brand header, verification badge, trust score panel (response rate, completion rate, fulfilment days), product grid, and review section with brand replies."),
    ("SS-04", "Brand Owner Dashboard — Order management table with status lifecycle controls, revenue summary card, product inventory overview, and notification bell."),
    ("SS-05", "Order Submission — Product detail page with MOQ warning, bulk-pricing tier visualisation, quantity input, and order total calculation in selected GCC currency."),
    ("SS-06", "Admin Dashboard — Platform analytics (total users, brands, orders, revenue), pending brand approval queue, and review moderation panel."),
    ("SS-07", "Arabic RTL Layout — Marketplace page with full right-to-left layout: mirrored navigation, Arabic product names and descriptions, RTL-aligned trust metrics."),
    ("SS-08", "Notification Panel — Unread notification list with type-specific icons, timestamps, and mark-as-read interaction; unread count badge on bell icon."),
]

make_table(doc,
    ["Ref", "Description"],
    [(ref, desc) for ref, desc in screenshot_list],
    col_widths=[0.7, 5.8]
)

for ref, desc in screenshot_list:
    placeholder(doc, ref, desc)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER IV — SYSTEM IMPLEMENTATION
# ════════════════════════════════════════════════════════════════════════════

page_break(doc)
h1(doc, "CHAPTER IV — SYSTEM IMPLEMENTATION")

body(doc,
    "This chapter documents the complete implementation of the Kuwait B2B Hub. It presents the "
    "Make-Buy-Lease acquisition decisions for all platform components, a software reuse analysis "
    "of the twelve externally-sourced libraries and services, the full codebase folder structure "
    "and artefact inventory, the configuration management record including the change request log "
    "and version baseline, and the complete development and deployment specification covering "
    "hardware, software, tooling, local setup, and production architecture."
)

# ── §4.1 Make, Buy or Lease Analysis ────────────────────────────────────────

h2(doc, "4.1  Make, Buy or Lease (MBL) Analysis")

body(doc,
    "A Make-Buy-Lease evaluation was conducted for each major technical component of the Kuwait "
    "B2B Hub to justify acquisition decisions, clarify licensing obligations, and classify the "
    "level at which external software assets are reused. Decisions were guided by three criteria: "
    "development effort saved, suitability to the Kuwait B2B domain, and long-term "
    "maintainability."
)

body(doc, "Reuse Level Definitions (Sommerville, Software Engineering, 10th ed.):")

make_table(doc,
    ["Level", "Definition"],
    [
        ("System",    "An entire off-the-shelf system or managed platform is adopted with no internal modification."),
        ("Component", "A library, framework, or SDK is integrated as a composable subsystem."),
        ("Object",    "Individual modules, classes, or hooks are consumed directly."),
        ("Abstract",  "Algorithms, patterns, or domain logic are designed and built from first principles."),
    ],
    col_widths=[1.2, 5.3]
)

make_table(doc,
    ["#", "Component", "Version", "Decision", "License", "Justification", "Reuse Level"],
    [
        ("1",  "Next.js Framework",          "16.2.2",   "Buy",
         "MIT",
         "Provides App Router, SSR, API route handlers, middleware pipeline, and TypeScript "
         "support. Eliminates building a custom server, router, or bundler — saves ~80+ hours.",
         "System"),
        ("2",  "React",                       "19.2.4",   "Buy",
         "MIT",
         "Underpins all UI rendering. React 19 Server Components eliminate redundant "
         "client-side JavaScript bundles on data-heavy pages.",
         "Component"),
        ("3",  "TypeScript",                  "^5.x",     "Buy",
         "Apache 2.0",
         "Compile-time type safety across all 93 source files. Interface definitions in db.ts "
         "serve as canonical data contracts enforced at every layer.",
         "System"),
        ("4",  "Supabase Auth",               "^2.101.1", "Lease",
         "Apache 2.0 / Supabase Cloud ToS",
         "Managed JWT session validation via supabase.auth.getUser(). Freemium tier sufficient "
         "for prototype scale.",
         "System"),
        ("5",  "@supabase/ssr SDK",           "^0.10.0",  "Buy",
         "Apache 2.0",
         "Bridges Supabase Auth with Next.js server-side cookie handling — required for "
         "session tokens in Server Components where localStorage is unavailable.",
         "Component"),
        ("6",  "PostgreSQL Database",         "Managed",  "Lease",
         "PostgreSQL Licence / Supabase Cloud ToS",
         "Supabase-managed PostgreSQL with PostgREST, real-time subscriptions, and automated "
         "backups. JSON prototype migration to PostgreSQL is the primary production step.",
         "System"),
        ("7",  "Vercel Hosting Platform",     "N/A",      "Lease",
         "Proprietary (Vercel ToS)",
         "First-party Next.js deployment. Provides automatic CI/CD, serverless execution for "
         "all 25 API routes, global CDN, and preview environments at zero cost (Hobby tier).",
         "System"),
        ("8",  "next-intl",                   "^4.9.0",   "Buy",
         "MIT",
         "Full Arabic/English bilingual support with RTL/LTR switching via messages/ar.json "
         "and messages/en.json. Handles locale routing via next.config.ts plugin.",
         "Component"),
        ("9",  "Zustand",                     "^5.0.12",  "Buy",
         "MIT",
         "Lightweight client-side state management for notificationStore.ts. Zero-boilerplate "
         "hook API chosen over Redux for a small-team project.",
         "Object"),
        ("10", "Recharts",                    "^3.8.1",   "Buy",
         "MIT",
         "Composable React charting library for order volume and performance visualisations "
         "on admin and brand dashboards.",
         "Component"),
        ("11", "lucide-react",                "^1.7.0",   "Buy",
         "ISC",
         "Icon component library providing all UI iconography across 15+ components.",
         "Object"),
        ("12", "xlsx",                         "^0.18.5",  "Buy",
         "Apache 2.0",
         "Enables Excel catalogue import in CatalogImportModal.tsx — brand owners bulk-upload "
         "product listings via .xlsx files.",
         "Component"),
        ("13", "Custom Trust Engine",          "N/A",      "Make",
         "Proprietary",
         "No off-the-shelf solution for Kuwait-specific B2B trust scoring. Calculates "
         "response_rate, completion_rate, avg_fulfillment_days, and badge thresholds for "
         "BrandTrust and BuyerTrust. Core differentiating IP.",
         "Abstract"),
        ("14", "Custom RBAC Middleware",       "N/A",      "Make",
         "Proprietary",
         "Three-tier role model (admin, brand_owner, buyer) does not map to any generic RBAC "
         "library. Enforced via getCurrentUser() guard in auth.ts and middleware.ts.",
         "Abstract"),
        ("15", "Custom REST API Layer",        "N/A",      "Make",
         "Proprietary",
         "All 25 API handlers encode domain-specific rules: MOQ enforcement, bulk pricing "
         "resolution, WhatsApp click tracking, multi-role order lifecycle management.",
         "Component"),
    ],
    col_widths=[0.3, 1.5, 0.7, 0.6, 1.2, 2.1, 0.8]
)

body(doc, "Summary by Decision:")

make_table(doc,
    ["Decision", "Count", "Share", "Commentary"],
    [
        ("Buy (open-source libraries)",  "9",  "60%",
         "Deliberate maximisation of open-source adoption to reduce development effort."),
        ("Lease (managed cloud services)", "3", "20%",
         "Infrastructure delegated to Supabase and Vercel; no servers to manage."),
        ("Make (custom-built in-house)",  "3",  "20%",
         "Concentrated exclusively on domain logic constituting the platform's "
         "competitive differentiation — Trust Engine, RBAC, and B2B API ruleset."),
    ],
    col_widths=[1.8, 0.7, 0.7, 3.3]
)

body(doc,
    "No proprietary closed-source software was purchased, keeping total licensing cost at "
    "zero for the prototype phase."
)

# ── §4.2 Software Reuse Analysis ─────────────────────────────────────────────

h2(doc, "4.2  Software Reuse Analysis")

body(doc,
    "The Software Reuse Analysis focuses on the twelve externally-sourced components (nine "
    "Buy + three Lease decisions) and catalogues the reuse level, integration depth, and "
    "the specific codebase location at which each component is consumed. Reuse classification "
    "follows Sommerville's four-level taxonomy (System, Component, Object, Abstract)."
)

make_table(doc,
    ["#", "Component", "Version", "Reuse Level", "License",
     "Primary Integration Point", "Dependency Type"],
    [
        ("1",  "Next.js Framework",     "16.2.2",   "System",
         "MIT",
         "next.config.ts; entire src/app/ directory (App Router convention)",
         "Framework — entire application is scaffolded inside Next.js"),
        ("2",  "React",                 "19.2.4",   "Component",
         "MIT",
         "All .tsx files; JSX transform via tsconfig.json jsx: 'react-jsx'",
         "Runtime library — all UI components are React components"),
        ("3",  "TypeScript",            "^5.x",     "System",
         "Apache 2.0",
         "tsconfig.json (strict: true); all .ts and .tsx source files",
         "Compiler — type-checks all 93 source files at build time"),
        ("4",  "Supabase Auth (Lease)", "^2.101.1", "System",
         "Apache 2.0 / Supabase ToS",
         "src/utils/supabase/server.ts → consumed by src/lib/auth.ts",
         "Managed service — JWT session management delegated entirely to Supabase"),
        ("5",  "@supabase/ssr SDK",     "^0.10.0",  "Component",
         "Apache 2.0",
         "src/utils/supabase/server.ts (createServerClient); "
         "src/utils/supabase/client.ts (createBrowserClient)",
         "SDK — adapts Supabase Auth to Next.js cookie-based SSR sessions"),
        ("6",  "PostgreSQL (Lease)",    "Managed",  "System",
         "PostgreSQL Lic. / Supabase ToS",
         "Target production store; currently accessed via Supabase dashboard "
         "for auth; db.ts will point to PostgreSQL post-migration",
         "Managed service — no local PostgreSQL installation required"),
        ("7",  "Vercel (Lease)",        "N/A",      "System",
         "Proprietary ToS",
         "vercel.json (implicit); Git remote → Vercel CI/CD webhook",
         "Managed service — hosts all Next.js SSR pages and serverless API routes"),
        ("8",  "next-intl",             "^4.9.0",   "Component",
         "MIT",
         "src/i18n/request.ts; src/i18n.ts; next.config.ts (withNextIntl plugin); "
         "messages/ar.json + messages/en.json",
         "Plugin + library — wraps the Next.js router for locale-aware routing"),
        ("9",  "Zustand",               "^5.0.12",  "Object",
         "MIT",
         "src/store/notificationStore.ts (create()); "
         "consumed by NotificationBell.tsx and RealtimeProvider.tsx",
         "Library — single store slice; zero-boilerplate hook API"),
        ("10", "Recharts",              "^3.8.1",   "Component",
         "MIT",
         "src/app/admin/dashboard/page.tsx; src/app/brand/dashboard/page.tsx",
         "Library — composable chart components rendered inside React pages"),
        ("11", "lucide-react",          "^1.7.0",   "Object",
         "ISC",
         "15+ component files (Bell, ShoppingCart, Star, Check, X, ...)",
         "Library — individual SVG icon components imported by name"),
        ("12", "xlsx",                  "^0.18.5",  "Component",
         "Apache 2.0",
         "src/components/CatalogImportModal.tsx; src/app/api/products/bulk/route.ts",
         "Library — XLSX.read() and XLSX.utils.sheet_to_json() called directly"),
    ],
    col_widths=[0.3, 1.3, 0.7, 0.9, 1.2, 2.0, 1.1]
)

# ── §4.3 Coding Artefacts & Project Structure ────────────────────────────────

h2(doc, "4.3  Coding Artefacts & Project Structure")

body(doc,
    "The Kuwait B2B Hub was developed using a modular, component-driven architecture built on "
    "the Next.js 16 App Router paradigm. All core application logic, routing, and UI components "
    "are encapsulated within the src/ directory, enforcing a clear separation between "
    "presentation, business logic, and data access layers. The directory tree below represents "
    "the final physical structure of the workspace repository. API routes are shown in flattened "
    "path notation for brevity."
)

body(doc, "Figure 8 — Kuwait B2B Hub: Final Project Folder Structure")

folder_tree = """\
Kuwait B2B Hub
/app
├── messages/
│   ├── ar.json
│   └── en.json
├── public/
│   ├── file.svg  globe.svg  next.svg  vercel.svg  window.svg
│   └── robots.txt
├── src/
│   ├── app/
│   │   ├── admin/
│   │   │   ├── brands/page.tsx
│   │   │   ├── dashboard/page.tsx
│   │   │   ├── orders/page.tsx
│   │   │   └── users/page.tsx
│   │   ├── api/
│   │   │   ├── admin/brands/[id]/route.ts    admin/brands/route.ts
│   │   │   ├── admin/orders/route.ts         admin/reviews/[id]/route.ts
│   │   │   ├── admin/users/route.ts
│   │   │   ├── auth/login/route.ts           auth/logout/route.ts
│   │   │   ├── auth/me/route.ts              auth/register/route.ts
│   │   │   ├── brands/[id]/route.ts          brands/route.ts
│   │   │   ├── brands/whatsapp-click/route.ts
│   │   │   ├── exchange-rates/route.ts       messages/route.ts
│   │   │   ├── notifications/[id]/route.ts   notifications/route.ts
│   │   │   ├── orders/[id]/route.ts          orders/route.ts
│   │   │   ├── products/[id]/route.ts        products/bulk/route.ts
│   │   │   ├── products/route.ts
│   │   │   ├── reviews/[id]/reply/route.ts   reviews/[id]/route.ts
│   │   │   ├── reviews/route.ts              trust/[id]/route.ts
│   │   ├── brand/
│   │   │   ├── dashboard/page.tsx
│   │   │   ├── orders/[id]/page.tsx          orders/page.tsx
│   │   │   ├── products/[id]/edit/page.tsx   products/new/page.tsx
│   │   │   ├── products/page.tsx             profile/page.tsx
│   │   ├── brands/[id]/page.tsx
│   │   ├── dashboard/page.tsx                login/page.tsx
│   │   ├── marketplace/page.tsx              notifications/page.tsx
│   │   ├── orders/[id]/page.tsx              orders/page.tsx
│   │   ├── pending/page.tsx                  register/page.tsx
│   │   ├── favicon.ico  globals.css  layout.tsx
│   │   ├── page.module.css  page.tsx  sitemap.ts
│   │   ├── components/
│   │   │   ├── providers/RealtimeProvider.tsx
│   │   │   ├── BottomNav.tsx      BrandAvatar.tsx     BrandCard.tsx
│   │   │   ├── BrandSidebar.tsx   BuyerTrustCard.tsx  CatalogImportModal.tsx
│   │   │   ├── CurrencySelector.tsx  ImageUrlInput.tsx  LandingPage.tsx
│   │   │   ├── LanguageToggle.tsx    MarketModal.tsx    MobileTopBar.tsx
│   │   │   ├── Navbar.tsx         NotificationBell.tsx  Providers.tsx
│   │   │   ├── RatingBreakdown.tsx  ReviewCard.tsx     ReviewModal.tsx
│   │   │   ├── StarRating.tsx     StatusBadge.tsx     TrustScore.tsx
│   │   │   └── VerifiedBadge.tsx  WhatsAppButton.tsx
│   │   ├── contexts/CurrencyContext.tsx
│   │   ├── data/db.json
│   │   ├── hooks/
│   │   │   ├── useInView.ts              useRealtimeBrandOrders.ts
│   │   │   ├── useRealtimeNotifications.ts  useRealtimeOrder.ts
│   │   ├── i18n/request.ts
│   │   ├── lib/
│   │   │   ├── supabase/client.ts
│   │   │   ├── auth.ts          currencies.ts    db.ts
│   │   │   ├── formatters.ts    i18n.ts          notifications.ts
│   │   │   └── pricingUtils.ts
│   │   ├── store/notificationStore.ts
│   │   └── utils/supabase/
│   │       ├── client.ts  middleware.ts  server.ts
│   ├── i18n.ts  middleware.ts  navigation.ts
├── next.config.ts
├── package.json
└── tsconfig.json"""

monospace_block(doc, folder_tree, font_size=7)

body(doc,
    "The project enforces strict architectural boundaries: server-side API logic is isolated "
    "in /src/app/api; reusable UI elements, custom React hooks, and global notification state "
    "(managed via Zustand) reside in /components, /hooks, and /store respectively. Supabase "
    "client initialisation and authentication middleware are grouped under /utils/supabase, "
    "providing a secure, globally accessible database connection across server and client components."
)

body(doc, "SLOC Summary by Category:")

make_table(doc,
    ["Category", "Files", "LOC", "KLOC", "Share"],
    [
        ("Pages — Buyer / Public",          "11", "2,344", "2.34", "20.6%"),
        ("Components",                       "24", "3,802", "3.80", "33.5%"),
        ("Pages — Brand Portal",             "8",  "2,213", "2.21", "19.5%"),
        ("API Routes",                       "25", "1,289", "1.29", "11.4%"),
        ("Pages — Admin",                    "4",  "790",   "0.79", "7.0%"),
        ("Lib / Utilities",                  "11", "485",   "0.48", "4.3%"),
        ("Hooks",                            "4",  "228",   "0.23", "2.0%"),
        ("Store / Context",                  "2",  "138",   "0.14", "1.2%"),
        ("Config / i18n / Middleware",       "4",  "64",    "0.06", "0.6%"),
        ("TOTAL",                            "93", "11,353","11.35","100%"),
    ],
    col_widths=[2.6, 0.6, 0.9, 0.8, 0.8]
)

body(doc,
    "The project totals 11.35 KLOC across 93 source files. The largest single category is "
    "Components (3.80 KLOC, 33.5%), driven primarily by LandingPage.tsx (913 lines), "
    "CatalogImportModal.tsx (560 lines), and NotificationBell.tsx (337 lines). The 25 API "
    "route handlers account for 1.29 KLOC, reflecting the lightweight, function-per-route "
    "server design of the Next.js App Router."
)

# ── §4.4 Configuration Management ───────────────────────────────────────────

h2(doc, "4.4  Configuration Management")

body(doc,
    "Configuration Management (CM) encompasses the practices used to systematically control, "
    "track, and audit all changes to the Kuwait B2B Hub codebase throughout its development "
    "lifecycle. Procedures follow ISO/IEC 12207 guidelines, ensuring every version of the "
    "system is reproducible, traceable, and consistent across all development environments."
)

h3(doc, "4.4.1  Versioning Scheme")

body(doc,
    "The Kuwait B2B Hub adopts a four-part versioning scheme:  "
    "PHASE . MAJOR . MINOR . PATCH"
)

make_table(doc,
    ["Field", "Position", "Trigger for Increment", "Example"],
    [
        ("Phase", "1st", "Advancement to a new development phase",                   "1 → 2"),
        ("Major", "2nd", "Architectural change, new subsystem, or breaking API change", "1.0 → 1.1"),
        ("Minor", "3rd", "Non-breaking feature addition or workflow enhancement",    "1.1.0 → 1.1.1"),
        ("Patch", "4th", "Bug fix, security patch, or configuration adjustment",     "1.1.1.0 → 1.1.1.1"),
    ],
    col_widths=[0.8, 0.8, 3.5, 1.4]
)

make_table(doc,
    ["Phase", "Code", "Scope"],
    [
        ("Planning & Requirements", "1.x.x.x", "Project scaffolding, schema design, environment setup"),
        ("Design & Architecture",   "2.x.x.x", "Component structure, routing hierarchy, database interface design"),
        ("Implementation",          "3.x.x.x", "Feature development, API routes, UI components, integrations"),
        ("Testing & Stabilisation", "4.x.x.x", "Defect resolution, security hardening, performance tuning"),
    ],
    col_widths=[1.8, 1.0, 3.7]
)

body(doc,
    "The project entered version 3.0.0.0 at the commencement of full implementation and "
    "reached 3.5.2.1 at the time of final submission."
)

h3(doc, "4.4.2  Configuration Items")

make_table(doc,
    ["CI ID", "Item", "Type", "Storage Location"],
    [
        ("CI-01", "Application Source Code (src/)",              "Software",      "Git repository"),
        ("CI-02", "Database Schema Interfaces (src/lib/db.ts)",  "Software",      "Git repository"),
        ("CI-03", "Environment Configuration (.env.local)",      "Configuration", "Secured — NOT version-controlled"),
        ("CI-04", "Dependency Manifest (package.json)",          "Configuration", "Git repository"),
        ("CI-05", "Internationalisation Messages (messages/)",   "Data",          "Git repository"),
        ("CI-06", "Final Project Report (.docx)",                "Document",      "Team shared drive"),
        ("CI-07", "Prototype Database (src/data/db.json)",       "Data",          "Git repository"),
        ("CI-08", "Deployment Configuration (next.config.ts)",  "Configuration", "Git repository"),
    ],
    col_widths=[0.8, 2.5, 1.2, 2.0]
)

h3(doc, "4.4.3  Change Request Log")

body(doc,
    "Changes are typed using the IEEE 1219 taxonomy — Corrective (defect/security fix), "
    "Adaptive (environment change), Perfective (new capability), Preventive (proactive "
    "restructure)."
)

make_table(doc,
    ["CR ID", "Change Description", "Type", "Priority", "Old Ver.", "New Ver.", "Status"],
    [
        ("CR-001",
         "Initial project scaffold — Next.js 16 App Router setup, TypeScript configuration, "
         "package.json baseline.",
         "Perfective", "High", "3.0.0.0", "3.1.0.0", "Closed"),
        ("CR-002",
         "Database schema & interface design — Defined all 10 TypeScript interfaces in db.ts. "
         "Established readDB() / writeDB() / generateId() utilities.",
         "Perfective", "High", "3.1.0.0", "3.2.0.0", "Closed"),
        ("CR-003",
         "Arabic language & RTL support — Integrated next-intl v4.9; created ar.json and "
         "en.json; added LanguageToggle component and locale-aware routing middleware. "
         "Refactored all hard-coded strings across 24 UI components.",
         "Adaptive", "High", "3.2.0.0", "3.3.0.0", "Closed"),
        ("CR-004",
         "Supabase Auth integration — Migrated session management to Supabase JWT. Introduced "
         "src/utils/supabase/server.ts and client.ts using @supabase/ssr. Updated "
         "getCurrentUser() to delegate to supabase.auth.getUser().",
         "Adaptive", "High", "3.3.0.0", "3.4.0.0", "Closed"),
        ("CR-005",
         "XSS vulnerability remediation — Unsanitised output identified in BrandCard.tsx, "
         "ReviewCard.tsx, and LandingPage.tsx. Applied React JSX escaping; hardened cookies: "
         "b2b_user_id set to httpOnly: true, SameSite: 'lax'.",
         "Corrective", "Critical", "3.4.0.0", "3.4.0.1", "Closed"),
        ("CR-006",
         "Bulk pricing tier system — Extended Product interface with bulk_pricing_tiers[]. "
         "Updated POST /api/orders to resolve unit_price from matching tier range. Added UI "
         "controls in product creation and edit pages.",
         "Perfective", "High", "3.4.0.1", "3.4.1.0", "Closed"),
        ("CR-007",
         "Trust Score Engine — Designed BrandTrust and BuyerTrust scoring with "
         "response_rate, completion_rate, avg_fulfillment_days, and badge metrics.",
         "Perfective", "High", "3.4.1.0", "3.5.0.0", "Closed"),
        ("CR-008",
         "Real-time notification system — Introduced Zustand v5 store, RealtimeProvider.tsx, "
         "NotificationBell.tsx, and createNotification() / createNotificationForMany() "
         "integrated into all key workflows.",
         "Perfective", "Medium", "3.5.0.0", "3.5.1.0", "Closed"),
        ("CR-009",
         "Multi-currency GCC support — Added CurrencyConfig interface and CURRENCIES registry "
         "covering KWD, SAR, AED, QAR, BHD, OMR with ISO 4217 codes, VAT rates, locale strings.",
         "Perfective", "Medium", "3.5.1.0", "3.5.2.0", "Closed"),
        ("CR-010",
         "Excel catalogue bulk import — Added CatalogImportModal.tsx (560 LOC) enabling brand "
         "owners to upload .xlsx product catalogues parsed via the xlsx library and submitted "
         "to POST /api/products/bulk.",
         "Perfective", "Low", "3.5.2.0", "3.5.2.1", "Closed"),
    ],
    col_widths=[0.7, 3.0, 0.9, 0.8, 0.7, 0.7, 0.7]
)

h3(doc, "4.4.4  Version Baseline Summary")

make_table(doc,
    ["Milestone", "Version", "Key Deliverable"],
    [
        ("Project Scaffold",       "3.1.0.0", "Next.js 16 setup, TypeScript baseline"),
        ("Schema Freeze",          "3.2.0.0", "All 10 db.ts interfaces defined"),
        ("Bilingual Release",      "3.3.0.0", "Arabic RTL support live"),
        ("Auth Stabilisation",     "3.4.0.0", "Supabase JWT sessions active"),
        ("Security Patch",         "3.4.0.1", "XSS vulnerability closed"),
        ("B2B Pricing Complete",   "3.4.1.0", "Bulk pricing tiers implemented"),
        ("Trust Engine Release",   "3.5.0.0", "BrandTrust / BuyerTrust scoring live"),
        ("Notifications Release",  "3.5.1.0", "Real-time notification system active"),
        ("GCC Currency Support",   "3.5.2.0", "Six-currency selector live"),
        ("Final Submission",       "3.5.2.1", "Excel import; full feature freeze"),
    ],
    col_widths=[2.0, 1.2, 3.3]
)

# ── §4.5 Development & Deployment Specification ──────────────────────────────

h2(doc, "4.5  Development & Deployment Specification")

h3(doc, "4.5.1  Hardware Requirements")

make_table(doc,
    ["Specification", "Minimum (Dev)", "Recommended (Dev)", "Production (Vercel — Managed)"],
    [
        ("CPU",
         "Dual-core 2.0 GHz (x86-64 or ARM64)",
         "Quad-core 2.5 GHz+ (Apple M-series or Intel i5/i7)",
         "Managed serverless — no configuration required"),
        ("RAM",
         "8 GB",
         "16 GB",
         "Per-function memory: 1024 MB (Vercel default)"),
        ("Storage",
         "10 GB free (SSD preferred)",
         "50 GB free SSD",
         "Ephemeral — Supabase handles data persistence"),
        ("Display",
         "1280 × 720",
         "1920 × 1080 or higher",
         "N/A"),
        ("Network",
         "Broadband (10 Mbps+) — required for Supabase cloud",
         "Broadband (50 Mbps+)",
         "Vercel global edge CDN"),
        ("Architecture",
         "x86-64 or Apple Silicon (ARM64)",
         "Apple Silicon M1/M2/M3 or x86-64",
         "Serverless (x86-64, Linux)"),
    ],
    col_widths=[1.3, 1.5, 1.8, 1.9]
)

body(doc,
    "The development team worked primarily on macOS Ventura (Darwin 22.5.0) on Apple Silicon. "
    "No hardware-level incompatibilities were identified on Windows 11 or Ubuntu 22.04 LTS "
    "during peer reviews."
)

h3(doc, "4.5.2  Software Requirements")

body(doc, "Operating System:")
make_table(doc,
    ["Platform", "Version", "Support Status"],
    [
        ("macOS",            "Ventura 13.x or later",          "Primary — used by development team"),
        ("Windows",          "11 (Build 22000+)",               "Compatible"),
        ("Ubuntu / Debian",  "22.04 LTS or later",             "Compatible"),
    ],
    col_widths=[1.5, 2.0, 3.0]
)

body(doc, "Runtime & Package Manager:")
make_table(doc,
    ["Software", "Version Used", "Minimum Required", "Notes"],
    [
        ("Node.js", "v24.14.0", "v20.0.0 LTS",
         "Next.js 16 requires Node ≥ 20. Developed and tested on Node 24. "
         "@types/node: ^20 confirms Node 20 as minimum target."),
        ("npm",     "v11.9.0",  "v10.0.0",
         "Ships with Node 24. Used for all dependency installation and scripting."),
        ("Git",     "2.40+",    "2.30+",
         "Version control and change tracking."),
    ],
    col_widths=[1.2, 1.2, 1.3, 2.8]
)

body(doc, "Core Framework & Language:")
make_table(doc,
    ["Software", "Version", "License", "Role"],
    [
        ("Next.js",     "16.2.2", "MIT",        "Full-stack React framework — App Router, SSR, API routes, middleware"),
        ("React",       "19.2.4", "MIT",        "UI component rendering (Server + Client Components)"),
        ("TypeScript",  "^5.x",   "Apache 2.0", "Statically typed — strict mode enabled across all 93 source files"),
        ("ESLint",      "^9.x",   "MIT",        "Static analysis (eslint-config-next: 16.2.2)"),
    ],
    col_widths=[1.2, 0.9, 1.0, 3.4]
)

body(doc, "TypeScript compiler options (tsconfig.json):")
make_table(doc,
    ["Option", "Value", "Effect"],
    [
        ("target",           "ES2017",     "Broad browser and Node compatibility"),
        ("strict",           "true",       "All strict type-checking rules enabled"),
        ("moduleResolution", "bundler",    "Optimised for Next.js 16 bundler-based resolution"),
        ("jsx",              "react-jsx",  "Automatic JSX transform — no import React needed"),
        ("incremental",      "true",       "Caches compilations to speed up rebuilds"),
        ("paths",            "@/* → ./src/*", "Absolute imports via @/ alias across codebase"),
    ],
    col_widths=[1.5, 1.3, 3.7]
)

body(doc, "Cloud Services & Backend Infrastructure:")
make_table(doc,
    ["Service", "Version / Tier", "Role", "Configuration"],
    [
        ("Supabase Auth",      "@supabase/supabase-js ^2.101.1",
         "JWT session management",
         "NEXT_PUBLIC_SUPABASE_URL + NEXT_PUBLIC_SUPABASE_ANON_KEY"),
        ("Supabase SSR",       "@supabase/ssr ^0.10.0",
         "Server-side cookie session handling",
         "Initialised in src/utils/supabase/server.ts and client.ts"),
        ("Supabase PostgreSQL","Managed (Free Tier)",
         "Auth session storage; production data layer post-migration",
         "Hosted on Supabase cloud — no local PostgreSQL required"),
        ("Vercel",             "Hobby Tier",
         "Production deployment, CDN, serverless execution",
         "Linked to Git repository for automatic CI/CD"),
    ],
    col_widths=[1.3, 1.6, 1.7, 1.9]
)

body(doc, "Key Runtime Dependencies:")
make_table(doc,
    ["Package", "Version", "Purpose"],
    [
        ("next-intl",    "^4.9.0",   "Arabic/English bilingual support via next.config.ts plugin"),
        ("zustand",      "^5.0.12",  "Client-side notification state management"),
        ("recharts",     "^3.8.1",   "Data visualisation for dashboards"),
        ("lucide-react", "^1.7.0",   "Icon component library"),
        ("xlsx",         "^0.18.5",  "Excel catalogue import parsing"),
        ("uuid",         "^13.0.0",  "UUID generation for entity IDs"),
    ],
    col_widths=[1.5, 1.0, 4.0]
)

h3(doc, "4.5.3  IDE & Developer Tooling")

make_table(doc,
    ["Tool", "Version", "Role"],
    [
        ("Visual Studio Code", "Latest stable",
         "Primary IDE — TypeScript IntelliSense, integrated terminal, ESLint extension"),
        ("Claude Code CLI",    "Latest",
         "AI-assisted development and refactoring"),
        ("Supabase Dashboard", "Web (cloud)",
         "Database inspection, auth user management, real-time event monitoring"),
        ("Vercel Dashboard",   "Web (cloud)",
         "Deployment management, environment variables, build logs, preview URLs"),
        ("Git",                "2.40+",
         "Source control — branching, change tracking, pull request workflow"),
        ("Browser DevTools",   "Chrome / Safari",
         "Network inspection, cookie debugging, RTL layout verification, manual E2E testing"),
    ],
    col_widths=[1.6, 1.3, 3.6]
)

h3(doc, "4.5.4  Testing & Quality Assurance")

make_table(doc,
    ["Tool / Method", "Type", "Coverage", "Notes"],
    [
        ("ESLint (eslint-config-next)",     "Static analysis",        "All .ts / .tsx files",
         "Run via npm run lint"),
        ("TypeScript Compiler (tsc)",       "Type checking",          "All 93 source files",
         "Catches type violations before runtime; strict mode enforced"),
        ("Manual browser testing",          "Functional / E2E",       "All user flows",
         "Chrome and Safari; all role-based paths tested manually against seeded demo data"),
        ("Supabase dashboard",              "Integration",            "Auth + database layer",
         "Sessions and real-time events verified via table editor"),
        ("Automated unit testing",          "—",                      "0% (none configured)",
         "No test framework present. Vitest + React Testing Library recommended for Sprint 8+"),
    ],
    col_widths=[1.9, 1.2, 1.5, 2.0]
)

h3(doc, "4.5.5  Local Development Setup")

body(doc, "To run the Kuwait B2B Hub locally from source:")

monospace_block(doc,
"""\
# 1. Clone the repository
git clone <repository-url> && cd app

# 2. Install dependencies (Node.js v20+ required)
npm install

# 3. Configure environment variables
cp .env.example .env.local
#    Set: NEXT_PUBLIC_SUPABASE_URL
#         NEXT_PUBLIC_SUPABASE_ANON_KEY

# 4. Start development server
npm run dev        # → http://localhost:3000

# 5. Static analysis
npm run lint

# 6. Production build validation (optional)
npm run build && npm run start""",
    font_size=8
)

body(doc,
    "No local database installation, no Docker, and no paid tooling are required. "
    "The entire stack is either open-source or available on a free tier — fully "
    "reproducible from a single npm install."
)

h3(doc, "4.5.6  Deployment Architecture")

body(doc, "The production deployment pipeline follows a fully automated Git-to-CDN workflow:")

monospace_block(doc,
"""\
Developer Machine
      │  git push
      ▼
 Git Repository
      │  Webhook trigger
      ▼
 Vercel CI/CD Pipeline
  ├── npm install
  ├── tsc --noEmit  (type check)
  ├── next build
  └── Deploy to Vercel Edge Network
            ├── Static assets   →  CDN (global edge)
            ├── Page routes     →  Serverless functions (SSR)
            └── API routes      →  Serverless functions (/api/*)
                                        │
                                        └── Supabase Cloud
                                             ├── Auth (JWT sessions)
                                             └── PostgreSQL (data layer)""",
    font_size=8
)

make_table(doc,
    ["Environment", "Trigger", "Purpose"],
    [
        ("Development (localhost:3000)", "npm run dev",                   "Local hot-reload development"),
        ("Preview (*.vercel.app)",       "Git push to non-main branch",   "Stakeholder review and QA testing"),
        ("Production (custom domain)",   "Merge to main branch",          "Live system — all traffic"),
    ],
    col_widths=[2.1, 2.1, 2.3]
)

# ── §4.6 Physical Deployment Diagram ────────────────────────────────────────

h2(doc, "4.6  Physical Deployment Diagram")

body(doc,
    "The physical deployment diagram maps the Kuwait B2B Hub's runtime artefacts to their "
    "execution nodes. Three primary nodes are in scope: the Developer Machine (local), the "
    "Vercel Edge Network (global serverless), and Supabase Cloud (GCC-region managed services). "
    "A fourth external node — the End-User Browser — communicates exclusively with the Vercel "
    "Edge Network over HTTPS."
)

body(doc,
    "Node descriptions: (1) Developer Machine — runs Node.js 24, npm 11, and the Next.js "
    "development server (localhost:3000). Connects to Supabase Cloud for authentication and "
    "database access during local development. (2) Vercel Edge Network — hosts all compiled "
    "Next.js page routes as serverless functions, serves static assets (JavaScript bundles, "
    "CSS, SVGs) from a global CDN, and routes all /api/* requests to individual serverless "
    "function instances. Each function has 1024 MB memory and a 10-second execution timeout. "
    "(3) Supabase Cloud — provides the managed PostgreSQL database instance, the Supabase Auth "
    "JWT session service, and (planned) real-time WebSocket subscriptions for notification "
    "delivery. Deployed in a GCC-proximate region for latency optimisation. "
    "(4) End-User Browser — all clients (buyers, brand owners, admins) access the platform "
    "exclusively via HTTPS. No platform software is installed on end-user devices beyond "
    "a standard web browser."
)

placeholder(doc, 9, "Physical Deployment Diagram — Kuwait B2B Hub")

# ─── SAVE ─────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_FILE)
print("PHASE 2 COMPLETE")
print(f"Output saved to: {OUTPUT_FILE}")
