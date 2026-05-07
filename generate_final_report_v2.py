"""
Kuwait B2B Hub — Final Project Report v2 Generator
Fixes applied:
  1. Cover page: CEAS header, student IDs, advisor name, team size [4]
  2. Team Workload table: Ch I–X with X marks and % weights
  3. Real names for all team members
  4. Chapter IV §4.2: VS Code-style file-tree image embedded
  5. Chapter VII: Activity Network diagram (table-based)
  6. Appendix A: note about draw.io / Lucidchart manual insertion
"""

# ── Dependencies ──────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image, ImageDraw, ImageFont
import io, os

# ═══════════════════════════════════════════════════════════════════════════════
# TEAM & PROJECT CONSTANTS  (edit here if names change)
# ═══════════════════════════════════════════════════════════════════════════════
TEAM = [
    {"name": "Bader Alnoumas",    "id": "220043570", "role": "Team Lead / Full-Stack Engineer"},
    {"name": "Ahmad Al-Rashidi",  "id": "220038812", "role": "Frontend / UI Developer"},
    {"name": "Omar Al-Mutairi",   "id": "220051234", "role": "Backend / API Developer"},
    {"name": "Noor Al-Dosari",    "id": "220047891", "role": "Localisation & QA Engineer"},
]
ADVISOR   = "Dr. Mohammed Al-Enezi"
COLLEGE   = "College of Engineering and Applied Sciences (CEAS)"
DEPT      = "Department of Computer Science"
COURSE    = "Software Engineering — COMP 491"
DATE      = "April 17, 2026"

# ═══════════════════════════════════════════════════════════════════════════════
# COLOURS
# ═══════════════════════════════════════════════════════════════════════════════
DARK_BLUE  = RGBColor(0x1A, 0x3A, 0x6B)
MID_BLUE   = RGBColor(0x2E, 0x6B, 0xA0)
GOLD       = RGBColor(0xD4, 0xA8, 0x47)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
GRAY       = RGBColor(0x55, 0x55, 0x55)
LIGHT_GRAY = RGBColor(0x88, 0x88, 0x88)
BLACK      = RGBColor(0x00, 0x00, 0x00)

HEX_DKBLUE = "1A3A6B"
HEX_MBLUE  = "2E6BA0"
HEX_LBLUE  = "D6EAF8"
HEX_GREEN  = "D5F5E3"
HEX_YELLOW = "FEF9E7"
HEX_RED    = "FDEDEC"
HEX_ORANGE = "FAD7A0"
HEX_GRAY   = "F2F3F4"
HEX_WHITE  = "FFFFFF"
HEX_DGRAY  = "CCCCCC"

# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT SETUP
# ═══════════════════════════════════════════════════════════════════════════════
doc = Document()
for s in doc.sections:
    s.top_margin    = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin   = Cm(3.0)
    s.right_margin  = Cm(2.5)

# ═══════════════════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def shade_row(row, hex_color):
    for cell in row.cells:
        shade_cell(cell, hex_color)

def add_bottom_border(para, color=HEX_DKBLUE, sz='8'):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    sz)
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pBdr.append(bot)
    pPr.append(pBdr)

def set_cell_text(cell, text, bold=False, size=10, color=None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = DARK_BLUE
    add_bottom_border(p, HEX_DKBLUE, '8')
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = MID_BLUE
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = DARK_BLUE
    return p

def body(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def body_bold(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    r1 = p.add_run(label + ' '); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text);        r2.font.size = Pt(11)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' '); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)
    return p

def numbered(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(2)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + ' '); r1.bold = True; r1.font.size = Pt(11)
    r2 = p.add_run(text); r2.font.size = Pt(11)
    return p

def divider(color='AAAAAA'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_bottom_border(p, color, '4')

def code_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Courier New'; r.font.size = Pt(9)
    return p

def make_table(headers, rows, col_widths=None, header_bg=HEX_DKBLUE,
               alt_bg=HEX_GRAY, font_size=10, header_font_size=10):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    hdr = tbl.rows[0]
    shade_row(hdr, header_bg)
    for cell, txt in zip(hdr.cells, headers):
        set_cell_text(cell, txt, bold=True, size=header_font_size, color=WHITE)
    for ri, row_data in enumerate(rows):
        row = tbl.add_row()
        shade_row(row, alt_bg if ri % 2 == 1 else HEX_WHITE)
        for ci, val in enumerate(row_data):
            set_cell_text(row.cells[ci], str(val), size=font_size)
    return tbl

def make_colored_table(headers, rows, row_colors, col_widths=None, font_size=10):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in tbl.columns[i].cells:
                cell.width = Cm(w)
    hdr = tbl.rows[0]
    shade_row(hdr, HEX_DKBLUE)
    for cell, txt in zip(hdr.cells, headers):
        set_cell_text(cell, txt, bold=True, size=font_size + 1, color=WHITE)
    for row_data, color in zip(rows, row_colors):
        row = tbl.add_row()
        shade_row(row, color)
        for ci, val in enumerate(row_data):
            set_cell_text(row.cells[ci], str(val), size=font_size)
    return tbl

# ═══════════════════════════════════════════════════════════════════════════════
# SECTION 4.2 IMAGE — VS Code-style file explorer
# ═══════════════════════════════════════════════════════════════════════════════

FILE_TREE = [
    ("Kuwait B2B Hub",          0, "root"),
    ("app/",                    1, "folder"),
    ("  public/",               2, "folder"),
    ("    file.svg",            3, "file"),
    ("    robots.txt",          3, "file"),
    ("  src/",                  2, "folder"),
    ("    app/",                3, "folder"),
    ("      api/",              4, "folder"),
    ("        auth/",           5, "folder"),
    ("          login/route.ts",6, "ts"),
    ("          logout/route.ts",6,"ts"),
    ("          me/route.ts",   6, "ts"),
    ("          register/route.ts",6,"ts"),
    ("        brands/route.ts", 5, "ts"),
    ("        orders/route.ts", 5, "ts"),
    ("        products/route.ts",5,"ts"),
    ("        reviews/route.ts", 5,"ts"),
    ("        notifications/route.ts",5,"ts"),
    ("        admin/",          5, "folder"),
    ("          brands/route.ts",6,"ts"),
    ("          users/route.ts", 6,"ts"),
    ("          reviews/route.ts",6,"ts"),
    ("      marketplace/page.tsx",4,"tsx"),
    ("      brands/[id]/page.tsx",4,"tsx"),
    ("      dashboard/page.tsx",4,"tsx"),
    ("      orders/page.tsx",   4, "tsx"),
    ("      login/page.tsx",    4, "tsx"),
    ("      register/page.tsx", 4, "tsx"),
    ("      layout.tsx",        4, "tsx"),
    ("      page.tsx",          4, "tsx"),
    ("    components/",         3, "folder"),
    ("      BrandCard.tsx",     4, "tsx"),
    ("      Navbar.tsx",        4, "tsx"),
    ("      ReviewModal.tsx",   4, "tsx"),
    ("      TrustScore.tsx",    4, "tsx"),
    ("      NotificationBell.tsx",4,"tsx"),
    ("      CurrencySelector.tsx",4,"tsx"),
    ("      VerifiedBadge.tsx", 4, "tsx"),
    ("    lib/",                3, "folder"),
    ("      auth.ts",           4, "ts"),
    ("      db.ts",             4, "ts"),
    ("      currencies.ts",     4, "ts"),
    ("      notifications.ts",  4, "ts"),
    ("      pricingUtils.ts",   4, "ts"),
    ("    data/db.json",        3, "json"),
    ("    middleware.ts",       3, "ts"),
    ("  messages/",             2, "folder"),
    ("    en.json",             3, "json"),
    ("    ar.json",             3, "json"),
    ("  package.json",          2, "json"),
    ("  next.config.ts",        2, "ts"),
    ("  tsconfig.json",         2, "json"),
]

def build_file_tree_image(save_path: str, width=780, scale=2):
    """Render a VS Code-style file explorer panel as a PNG."""
    BG       = (30, 30, 30)
    TITLE_BG = (37, 37, 38)
    ACTBAR   = (51, 51, 51)
    TXT_CLR  = (204, 204, 204)
    DIM_CLR  = (150, 150, 150)
    FOLD_CLR = (229, 192, 123)
    TS_CLR   = (66, 153, 225)
    TSX_CLR  = (97, 175, 239)
    JSON_CLR = (229, 192, 123)
    FILE_CLR = (171, 178, 191)
    ROOT_CLR = (229, 229, 229)
    SEL_BG   = (37, 56, 80)

    LINE_H   = 22 * scale
    INDENT   = 14 * scale
    PAD_LEFT = 12 * scale
    ICON_W   = 18 * scale
    FONT_SZ  = int(13 * scale)
    TITLE_H  = 30 * scale

    height = TITLE_H + LINE_H * len(FILE_TREE) + 16 * scale
    W = width * scale

    img = Image.new("RGB", (W, height), BG)
    draw = ImageDraw.Draw(img)

    # Try to load a monospace font, fall back gracefully
    font = None
    bold_font = None
    for fp in [
        "/System/Library/Fonts/Menlo.ttc",
        "/System/Library/Fonts/Monaco.ttf",
        "/Library/Fonts/Courier New.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
    ]:
        if os.path.exists(fp):
            try:
                font      = ImageFont.truetype(fp, FONT_SZ)
                bold_font = ImageFont.truetype(fp, FONT_SZ)
                break
            except Exception:
                pass
    if font is None:
        font = bold_font = ImageFont.load_default()

    # Title bar
    draw.rectangle([0, 0, W, TITLE_H], fill=TITLE_BG)
    draw.text((PAD_LEFT, int(TITLE_H * 0.18)), "EXPLORER", fill=(189, 189, 189), font=bold_font)
    draw.text((int(W * 0.45), int(TITLE_H * 0.18)), "Kuwait B2B Hub — VS Code", fill=DIM_CLR, font=font)

    # File rows
    for i, (label, depth, ftype) in enumerate(FILE_TREE):
        y = TITLE_H + i * LINE_H
        x = PAD_LEFT + depth * INDENT

        # Highlight selected row (first folder)
        if i == 1:
            draw.rectangle([0, y, W, y + LINE_H - 1], fill=SEL_BG)

        # Icon
        icon_x = x
        icon_y = y + int(LINE_H * 0.1)
        if ftype == "folder":
            draw.rectangle([icon_x, icon_y + int(LINE_H*0.2),
                            icon_x + ICON_W, icon_y + int(LINE_H*0.75)], fill=FOLD_CLR)
            draw.rectangle([icon_x, icon_y + int(LINE_H*0.1),
                            icon_x + int(ICON_W*0.55), icon_y + int(LINE_H*0.3)], fill=FOLD_CLR)
        elif ftype == "ts":
            draw.rectangle([icon_x + 2, icon_y + int(LINE_H*0.15),
                            icon_x + ICON_W - 2, icon_y + int(LINE_H*0.8)], fill=TS_CLR)
            draw.text((icon_x + int(ICON_W*0.1), icon_y + int(LINE_H*0.15)),
                      "ts", fill=(255,255,255), font=ImageFont.load_default())
        elif ftype == "tsx":
            draw.rectangle([icon_x + 2, icon_y + int(LINE_H*0.15),
                            icon_x + ICON_W - 2, icon_y + int(LINE_H*0.8)], fill=TSX_CLR)
            draw.text((icon_x, icon_y + int(LINE_H*0.15)),
                      "tsx", fill=(255,255,255), font=ImageFont.load_default())
        elif ftype == "json":
            draw.rectangle([icon_x + 2, icon_y + int(LINE_H*0.15),
                            icon_x + ICON_W - 2, icon_y + int(LINE_H*0.8)], fill=JSON_CLR)
        elif ftype == "root":
            draw.rectangle([icon_x, icon_y + int(LINE_H*0.15),
                            icon_x + ICON_W, icon_y + int(LINE_H*0.8)], fill=(50, 100, 160))
        else:
            draw.rectangle([icon_x + 3, icon_y + int(LINE_H*0.2),
                            icon_x + ICON_W - 3, icon_y + int(LINE_H*0.75)], fill=FILE_CLR)

        # Label
        text_x = icon_x + ICON_W + int(6 * scale)
        text_y = y + int(LINE_H * 0.12)
        clr = ROOT_CLR if ftype == "root" else (FOLD_CLR if ftype == "folder" else TXT_CLR)
        draw.text((text_x, text_y), label.lstrip(), fill=clr, font=font)

    # Bottom status bar
    draw.rectangle([0, height - 8 * scale, W, height], fill=(0, 122, 204))

    img.save(save_path, "PNG", dpi=(144, 144))
    return save_path


# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE  (Fix #1 + #3)
# ═══════════════════════════════════════════════════════════════════════════════

def cover_page():
    for _ in range(2):
        doc.add_paragraph()

    # College header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COLLEGE)
    r.bold = True; r.font.size = Pt(13); r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DEPT)
    r.font.size = Pt(12); r.font.color.rgb = GRAY

    divider()
    doc.add_paragraph()
    doc.add_paragraph()

    # Project title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kuwait B2B Hub")
    r.bold = True; r.font.size = Pt(32); r.font.color.rgb = DARK_BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Digital Wholesale Marketplace for Kuwait & the GCC Region")
    r.font.size = Pt(15); r.font.color.rgb = MID_BLUE; r.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Final Project Report")
    r.bold = True; r.font.size = Pt(18); r.font.color.rgb = GOLD

    divider()
    doc.add_paragraph()
    doc.add_paragraph()

    # Metadata
    meta = [
        ("Course:",         COURSE),
        ("Report Type:",    "Final Project Report"),
        ("Submission Date:", DATE),
        ("Project Advisor:", ADVISOR),
        ("Team Size:",       "[4]"),
    ]
    for lbl, val in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{lbl}  "); r1.bold = True; r1.font.size = Pt(12)
        r2 = p.add_run(val);        r2.font.size = Pt(12)

    doc.add_paragraph()

    # Team members with IDs
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared by:"); r.bold = True; r.font.size = Pt(12)

    doc.add_paragraph()

    # Team table on cover
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    shade_row(hdr, HEX_DKBLUE)
    for cell, txt in zip(hdr.cells, ["Name", "Student ID", "Role"]):
        set_cell_text(cell, txt, bold=True, size=11, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for m in TEAM:
        row = tbl.add_row()
        shade_row(row, HEX_WHITE)
        set_cell_text(row.cells[0], m["name"], size=11, bold=True)
        set_cell_text(row.cells[1], m["id"],   size=11,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], m["role"], size=10)
    for i, w in enumerate([5.5, 3.0, 8.0]):
        for cell in tbl.columns[i].cells:
            cell.width = Cm(w)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

def toc():
    h1("Table of Contents")
    entries = [
        ("Abstract", "5"),
        ("Chapter I — Introduction", "6"),
        ("  1.1  Background", "6"),
        ("  1.2  System Overview", "6"),
        ("  1.3  Glossary", "7"),
        ("  1.4  Readership", "7"),
        ("  1.5  Product Scope", "8"),
        ("  1.6  Team Workload", "8"),
        ("Chapter II — Software Requirements Specification (SRS)", "9"),
        ("  2.1  Feasibility Study", "9"),
        ("  2.2  User Requirements", "9"),
        ("  2.3  System Requirements", "10"),
        ("  2.4  Non-Functional Requirements", "11"),
        ("  2.5  Functional Requirements Specification (FRS)", "11"),
        ("  2.6  Context Model", "12"),
        ("  2.7  Use Cases", "12"),
        ("  2.8  Requirements Traceability Matrix (RTM)", "13"),
        ("Chapter III — System Design", "15"),
        ("  3.1  Sequence Diagrams", "15"),
        ("  3.2  Data Flow Diagram (DFD)", "16"),
        ("  3.3  System Architecture", "16"),
        ("  3.4  Activity Model", "17"),
        ("  3.5  Class Diagram", "17"),
        ("  3.6  Prototype Screenshots", "18"),
        ("Chapter IV — System Implementation", "19"),
        ("  4.1  Make / Buy / Lease Analysis", "19"),
        ("  4.2  Coding Artefacts & Project Structure", "19"),
        ("  4.3  Configuration Management", "21"),
        ("  4.4  Deployment Specification", "21"),
        ("  4.5  Physical Deployment Diagram", "22"),
        ("Chapter V — Verification & Validation", "23"),
        ("  5.1  Development Testing", "23"),
        ("  5.2  System Inspection", "23"),
        ("  5.3  System Testing", "24"),
        ("  5.4  Release Testing", "25"),
        ("  5.5  User Testing", "25"),
        ("Chapter VI — System Evolution", "26"),
        ("  6.1  Anticipated Changes", "26"),
        ("  6.2  Phase-Out Transitions", "26"),
        ("  6.3  Quality Predictions", "27"),
        ("Chapter VII — System Planning & Management", "28"),
        ("  7.1  Milestones", "28"),
        ("  7.2  Gantt Chart", "29"),
        ("  7.3  Activity Network Diagram", "29"),
        ("  7.4  Staff Allocation", "30"),
        ("  7.5  CoCoMo Cost Estimation", "30"),
        ("  7.6  Risk Matrix", "31"),
        ("  7.7  Fishbone Diagram", "31"),
        ("Chapter VIII — Conclusions", "32"),
        ("  8.1  Design Smells", "32"),
        ("  8.2  SWOT Analysis", "33"),
        ("  8.3  Future Work", "33"),
        ("Chapter IX — Appendices", "34"),
        ("Chapter X — References", "36"),
    ]
    for entry, pg in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(entry)
        r1.font.size = Pt(11)
        if not entry.startswith("  "):
            r1.bold = True
        r2 = p.add_run(f"\t{pg}")
        r2.font.size = Pt(11)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF TABLES
# ═══════════════════════════════════════════════════════════════════════════════

def list_of_tables():
    h1("List of Tables")
    tables = [
        ("Table 1",  "Team Workload Distribution by Chapter", "8"),
        ("Table 2",  "Feasibility Analysis Summary", "9"),
        ("Table 3",  "User Requirements", "10"),
        ("Table 4",  "System Functional Requirements", "10"),
        ("Table 5",  "Non-Functional Requirements", "11"),
        ("Table 6",  "Use Case Summary", "12"),
        ("Table 7",  "Requirements Traceability Matrix (RTM)", "13"),
        ("Table 8",  "Make / Buy / Lease Analysis", "19"),
        ("Table 9",  "Configuration Management", "21"),
        ("Table 10", "Deployment Specification", "21"),
        ("Table 11", "System Test Cases & Results", "24"),
        ("Table 12", "Anticipated System Changes", "26"),
        ("Table 13", "Project Milestones", "28"),
        ("Table 14", "Activity Network — Sprint Dependencies", "29"),
        ("Table 15", "Staff Allocation Matrix", "30"),
        ("Table 16", "CoCoMo Cost Estimation", "30"),
        ("Table 17", "Risk Matrix", "31"),
        ("Table 18", "SWOT Analysis", "33"),
    ]
    for num, title, pg in tables:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{num}. {title}\t{pg}")
        r.font.size = Pt(11)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF FIGURES
# ═══════════════════════════════════════════════════════════════════════════════

def list_of_figures():
    h1("List of Figures")
    figures = [
        ("Figure 1",  "System Architecture Diagram", "17"),
        ("Figure 2",  "Context Model (Level 0 DFD)", "12"),
        ("Figure 3",  "Level 1 Data Flow Diagram", "16"),
        ("Figure 4",  "Use Case Diagram — All Actors", "13"),
        ("Figure 5",  "Sequence Diagram — User Registration & Login", "15"),
        ("Figure 6",  "Sequence Diagram — Order Lifecycle", "15"),
        ("Figure 7",  "Sequence Diagram — Brand Verification", "16"),
        ("Figure 8",  "Class Diagram — Core Entities", "18"),
        ("Figure 9",  "Activity Diagram — Order Placement", "17"),
        ("Figure 10", "VS Code Project Explorer — B2B Folder Structure", "20"),
        ("Figure 11", "Physical Deployment Diagram", "22"),
        ("Figure 12", "Project Gantt Chart", "29"),
        ("Figure 13", "Activity Network Diagram — Sprint Dependencies", "30"),
        ("Figure 14", "Fishbone (Ishikawa) Diagram — Delayed Submission Risk", "32"),
        ("Figure 15", "Quality Prediction Chart", "27"),
    ]
    for num, title, pg in figures:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"{num}. {title}\t{pg}")
        r.font.size = Pt(11)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

def abstract():
    h1("Abstract")
    body(
        "Kuwait B2B Hub is a full-stack digital wholesale marketplace designed specifically for "
        "the Kuwaiti and Gulf Cooperation Council (GCC) commercial landscape.  The platform "
        "addresses a well-documented inefficiency: the majority of B2B wholesale transactions in "
        "Kuwait are currently conducted informally through WhatsApp groups, Excel spreadsheets, "
        "and phone calls — a process that is error-prone, unscalable, and opaque to both buyers "
        "and suppliers."
    )
    body(
        "This report documents the full software-engineering lifecycle of Kuwait B2B Hub, from "
        "requirements elicitation and feasibility analysis through system design, implementation, "
        "verification and validation, and project management.  The system is built on Next.js 16 "
        "(App Router) with TypeScript 5, React 19, Supabase authentication, and a RESTful JSON "
        "API.  It supports three user roles — Platform Admin, Brand Owner, and Business Buyer — "
        "and delivers ten major feature modules including a verified-brand marketplace, bulk-order "
        "management, tiered pricing, a trust-and-reputation engine, in-app notifications, and "
        "full bilingual (English/Arabic RTL) support across four GCC currencies."
    )
    body(
        "As of submission, approximately 72% of planned MVP functionality has been implemented "
        "and manually tested.  The project follows an Agile sprint model with eight sprints "
        "spanning January–June 2026 and is on track for final delivery."
    )
    body(
        "Keywords: B2B marketplace, wholesale, Kuwait, GCC, Next.js, TypeScript, Supabase, "
        "multi-currency, bilingual, trust & reputation, order management."
    )
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER I — INTRODUCTION  (Fix #2: rebuilt team workload table)
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_1():
    h1("Chapter I:  Introduction")

    h2("1.1  Background")
    body(
        "Kuwait's wholesale distribution sector has historically relied on personal relationships, "
        "phone calls, and WhatsApp messaging for order placement and supplier communication.  "
        "While these informal channels work for small volumes, they create significant operational "
        "friction at scale: orders are lost in chat histories, pricing is inconsistent, there is "
        "no formal order tracking, and buyers have no objective way to assess supplier reliability."
    )
    body(
        "The GCC B2B e-commerce segment is undergoing rapid digitalisation, driven by government "
        "Vision initiatives (Kuwait Vision 2035, Saudi Vision 2030) and post-pandemic acceleration "
        "of digital commerce adoption.  Most existing marketplace platforms are either consumer-"
        "facing (B2C) or serve Western markets without GCC linguistic and regulatory awareness."
    )
    body(
        "Kuwait B2B Hub fills this gap: a purpose-built, bilingual, GCC-market-aware wholesale "
        "marketplace giving Kuwaiti brands a professional digital storefront and giving buyers a "
        "trusted, structured ordering experience."
    )

    h2("1.2  System Overview")
    body(
        "Kuwait B2B Hub is a web application built on Next.js 16 with an integrated RESTful API "
        "deployed on Vercel.  It connects three classes of users through a shared digital environment:"
    )
    bullet("Platform Admins — govern the ecosystem: verify brands, moderate reviews, monitor platform health.")
    bullet("Brand Owners — Kuwaiti wholesale suppliers who list products, manage catalogs, set bulk-pricing tiers, and fulfil orders.")
    bullet("Business Buyers — retail and wholesale buyers across the GCC who browse, order, and review brands.")

    h2("1.3  Glossary")
    terms = [
        ("B2B",               "Business-to-Business — commercial transactions between companies."),
        ("MOQ",               "Minimum Order Quantity — the smallest number of units allowed per transaction."),
        ("GCC",               "Gulf Cooperation Council — Kuwait, Saudi Arabia, UAE, Bahrain, Oman, Qatar."),
        ("CR",                "Commercial Registration — Kuwait Ministry of Commerce business licence number."),
        ("KWD / SAR / AED / BHD","Kuwaiti Dinar / Saudi Riyal / UAE Dirham / Bahraini Dinar."),
        ("RTL",               "Right-to-Left — text direction used for Arabic script."),
        ("SRS",               "Software Requirements Specification."),
        ("API",               "Application Programming Interface — RESTful HTTP interface."),
        ("SSR",               "Server-Side Rendering — pages rendered server-side for SEO."),
        ("MVP",               "Minimum Viable Product — first releasable version."),
        ("Verification Tier", "Trust level assigned to a brand: New, Verified, or Premium."),
        ("Trust Score",       "Computed reliability metric for brands and buyers."),
        ("Bulk Pricing Tier", "Quantity bracket that unlocks a lower unit price."),
        ("Supabase",          "Open-source Firebase alternative: PostgreSQL + Auth + Realtime."),
        ("App Router",        "Next.js 13+ routing architecture with React Server Components."),
    ]
    make_table(["Term", "Definition"], terms, col_widths=[4.0, 12.5], font_size=10)

    h2("1.4  Readership")
    body("This report is intended for:")
    bullet("University course assessors evaluating the software-engineering process.")
    bullet("Project supervisors tracing requirements through to implementation.")
    bullet("Future developers extending or maintaining the platform.")
    bullet("Potential stakeholders evaluating technical and business viability.")

    h2("1.5  Product Scope")
    body("Scope of Kuwait B2B Hub v1.0 (MVP):")
    bullet("IN SCOPE: User authentication (3 roles); brand verification & catalog management; product CRUD with tiered pricing & Excel import; order lifecycle; buyer reviews & trust scores; admin dashboard; bilingual EN/AR UI; four-currency support; in-app notifications; mobile-responsive design.")
    bullet("OUT OF SCOPE (v1.0): Payment gateway; shipping carrier API; real-time WebSocket messaging; SMS/email notifications; automated CR verification; native mobile apps.")

    h2("1.6  Team Workload")
    body(
        "Table 1 shows the contribution of each team member across the ten report chapters.  "
        "'X' denotes the primary lead; 'x' denotes a secondary contributor.  "
        "The Weight column shows the percentage each chapter contributes to the final report."
    )
    doc.add_paragraph()

    # ── FIX #2 — Chapter-by-chapter workload table ────────────────────────────
    members = [m["name"].split()[0] for m in TEAM]   # first names for columns
    headers = ["Chapter", "Topic", "Weight"] + members

    # (chapter, topic, weight%, Bader, Ahmad, Omar, Noor)
    wl_rows = [
        ("Ch. I",   "Introduction",                "8%",  "X", "x", "",  "x"),
        ("Ch. II",  "SRS / Requirements",          "12%", "x", "",  "X", "x"),
        ("Ch. III", "System Design",               "12%", "X", "x", "x", ""),
        ("Ch. IV",  "Implementation",              "12%", "X", "x", "x", ""),
        ("Ch. V",   "Verification & Validation",   "10%", "x", "",  "x", "X"),
        ("Ch. VI",  "System Evolution",            "8%",  "x", "",  "X", "x"),
        ("Ch. VII", "Planning & Management",       "10%", "X", "x", "",  "x"),
        ("Ch. VIII","Conclusions",                 "8%",  "x", "x", "",  "X"),
        ("Ch. IX",  "Appendices",                  "10%", "X", "x", "x", "x"),
        ("Ch. X",   "References",                  "10%", "x", "X", "x", "x"),
    ]

    n_cols = len(headers)
    tbl = doc.add_table(rows=1, cols=n_cols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Column widths: Chapter | Topic | Weight | 4 × name
    col_w = [1.8, 4.2, 1.5] + [2.2] * 4
    for i, w in enumerate(col_w):
        for cell in tbl.columns[i].cells:
            cell.width = Cm(w)

    # Header row
    hdr = tbl.rows[0]
    shade_row(hdr, HEX_DKBLUE)
    for cell, txt in zip(hdr.cells, headers):
        set_cell_text(cell, txt, bold=True, size=10, color=WHITE,
                      align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, row_data in enumerate(wl_rows):
        row = tbl.add_row()
        shade_row(row, HEX_GRAY if ri % 2 == 1 else HEX_WHITE)
        for ci, val in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if ci >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            bold  = (val == "X")
            color = DARK_BLUE if val == "X" else (MID_BLUE if val == "x" else None)
            set_cell_text(row.cells[ci], val, size=10, bold=bold,
                          color=color, align=align)

    # Legend row
    leg_row = tbl.add_row()
    shade_row(leg_row, "EBF5FB")
    set_cell_text(leg_row.cells[0], "Legend", bold=True, size=9,
                  color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    # merge cells [1..n_cols-1]
    for ci in range(1, n_cols):
        leg_row.cells[0].merge(leg_row.cells[1])
    set_cell_text(leg_row.cells[0], "Legend:   X = Primary Lead     x = Secondary Contributor",
                  size=9, italic=True)

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER II — SRS
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_2():
    h1("Chapter II:  Software Requirements Specification (SRS)")

    h2("2.1  Feasibility Study")
    body("A feasibility study was conducted at project inception across five dimensions:")
    doc.add_paragraph()
    feas = [
        ("Technical",   "High",   "Next.js, Supabase, and Vercel are mature, well-documented technologies. No novel algorithms or unproven tech required."),
        ("Operational", "High",   "Replaces existing manual workflow (WhatsApp ordering). Users already understand the domain; training needs are minimal."),
        ("Economic",    "Medium", "Zero infrastructure cost during development (free tiers). KD 29/month subscription provides a realistic revenue path post-launch."),
        ("Schedule",    "Medium", "8-sprint Agile plan over ~4 months is achievable for a team of four given the MVP scope. Supabase migration is the main schedule risk."),
        ("Legal",       "Medium", "GCC data-localisation laws require regional data storage. Supabase supports GCC-adjacent regions. GDPR-equivalent compliance is a post-MVP concern."),
    ]
    make_colored_table(
        ["Dimension", "Viability", "Justification"],
        feas,
        [HEX_GREEN, HEX_GREEN, HEX_YELLOW, HEX_YELLOW, HEX_YELLOW],
        col_widths=[3.0, 2.5, 11.0], font_size=10
    )

    h2("2.2  User Requirements")
    user_req = [
        ("UR-01","Admin",       "Review pending brand applications and approve or reject them with a reason."),
        ("UR-02","Admin",       "View dashboard showing total users, brands, orders, and revenue."),
        ("UR-03","Admin",       "Moderate flagged reviews — remove or approve."),
        ("UR-04","Brand Owner", "Register and create a brand profile (name, logo, description, WhatsApp)."),
        ("UR-05","Brand Owner", "List products with name, price, MOQ, stock, and bulk-pricing tiers."),
        ("UR-06","Brand Owner", "Import products from an Excel/CSV file."),
        ("UR-07","Brand Owner", "Review and approve or reject incoming order requests."),
        ("UR-08","Brand Owner", "View a revenue and order-metrics dashboard."),
        ("UR-09","Brand Owner", "Reply to buyer reviews publicly."),
        ("UR-10","Buyer",       "Browse and search verified brands filtered by category, MOQ, and rating."),
        ("UR-11","Buyer",       "View a brand's full product catalog with pricing tiers."),
        ("UR-12","Buyer",       "Submit an order request with quantity validation."),
        ("UR-13","Buyer",       "Track order status in real time."),
        ("UR-14","Buyer",       "Leave a star rating and written review after receiving an order."),
        ("UR-15","Buyer",       "View prices in local currency (KWD/SAR/AED/BHD)."),
        ("UR-16","All",         "Use the platform in Arabic with full RTL layout."),
        ("UR-17","All",         "Receive in-app notifications for relevant events."),
    ]
    make_table(["ID", "Actor", "Requirement"], user_req,
               col_widths=[1.5, 2.5, 12.5], font_size=9)

    h2("2.3  System Requirements")
    sys_req = [
        ("SR-01","Support three user roles (admin, brand_owner, buyer) with distinct access control."),
        ("SR-02","Enforce role-based route protection on all pages and API endpoints."),
        ("SR-03","Create a pending brand profile automatically upon brand_owner registration."),
        ("SR-04","Allow admins to approve/reject brands and set a verification tier."),
        ("SR-05","Display a verification badge (Premium/Verified/New) on brand cards."),
        ("SR-06","Support full CRUD for products associated with a brand."),
        ("SR-07","Apply bulk-pricing tiers automatically based on order quantity."),
        ("SR-08","Validate orders against product MOQ and available stock."),
        ("SR-09","Maintain order status lifecycle: pending → approved/rejected → completed."),
        ("SR-10","Allow buyers to submit one review per completed order (1–5 stars)."),
        ("SR-11","Compute and return aggregate review statistics (avg rating, breakdown)."),
        ("SR-12","Compute trust scores for brands (response rate, completion rate, etc.)."),
        ("SR-13","Compute trust scores for buyers (total orders, completion rate, etc.)."),
        ("SR-14","Generate in-app notifications for order events, registrations, and reviews."),
        ("SR-15","Display prices in the user's selected currency with market-specific VAT."),
        ("SR-16","Serve the UI in English and Arabic; support full RTL for Arabic locale."),
        ("SR-17","Track WhatsApp click counts per brand via a dedicated API endpoint."),
        ("SR-18","Support bulk product import via XLSX/CSV file upload."),
        ("SR-19","Generate SEO structured data (JSON-LD) on public-facing pages."),
        ("SR-20","Provide an admin moderation queue for flagged reviews."),
    ]
    make_table(["ID", "System Requirement"], sys_req,
               col_widths=[1.5, 15.0], font_size=9)

    h2("2.4  Non-Functional Requirements")
    nfr = [
        ("NFR-01","Performance",        "Public pages: LCP < 2.5 s on standard broadband."),
        ("NFR-02","Security",           "HTTP-only session cookies; bcrypt password hashing via Supabase Auth."),
        ("NFR-03","Scalability",        "Data layer migrated to Supabase PostgreSQL before launch."),
        ("NFR-04","Availability",       "99.5% uptime target via Vercel edge + Supabase managed infra."),
        ("NFR-05","Accessibility",      "WCAG 2.1 Level AA; ARIA labels on interactive elements."),
        ("NFR-06","Internationalisation","All strings externalised to i18n files; Arabic renders RTL."),
        ("NFR-07","Maintainability",    "TypeScript strict mode; typed interfaces for all entities."),
        ("NFR-08","Portability",        "Deployable to any Vercel-compatible Node.js 18+ environment."),
    ]
    make_table(["ID", "Category", "Requirement"], nfr,
               col_widths=[1.5, 3.0, 12.0], font_size=9)

    h2("2.5  Functional Requirements Specification (FRS)")
    body("Detailed input/process/output specifications per subsystem:")
    bullet("Authentication: accepts {email, password, role, company_name}; validates uniqueness; creates user + optional brand record; sets HTTP-only cookie (7-day TTL).")
    bullet("Catalog: accepts {name, description, price, MOQ, stock, image_url, bulk_pricing_tiers[]}; validates MOQ ≥ 1 and price > 0; supports XLSX bulk import.")
    bullet("Orders: accepts {brand_id, items[{product_id, quantity}]}; validates MOQ and stock; resolves tier price; persists Order + OrderItems; fires brand notification.")
    bullet("Reviews: accepts {brand_id, order_id?, rating 1–5, content ≤ 300 chars, anonymous}; enforces one-review-per-order; recomputes aggregates; fires brand notification.")
    bullet("Admin: exposes brand approval, user management, order analytics, and review moderation — all behind admin-role guard.")

    h2("2.6  Context Model")
    body(
        "Figure 2 (Appendix A) shows the Level-0 context model.  The system boundary encloses all "
        "platform functionality.  External entities: Admin Users, Brand Owners, Business Buyers, "
        "and External Services (Supabase Auth, Vercel, WhatsApp Business).  Data flows include "
        "registration data, catalog data, order requests, review data, approval decisions, and "
        "notification events."
    )

    h2("2.7  Use Cases")
    uc = [
        ("UC-01","Register Account",           "All",        "Submit registration form → system creates account and session."),
        ("UC-02","Login",                       "All",        "Submit credentials → validate → set session cookie."),
        ("UC-03","Approve / Reject Brand",      "Admin",      "Review pending brand → set status + tier → notify brand owner."),
        ("UC-04","Manage Users",                "Admin",      "View user list → change verification status or role."),
        ("UC-05","Moderate Reviews",            "Admin",      "View flagged reviews → remove or approve."),
        ("UC-06","Create / Edit Brand Profile", "Brand Owner","Update brand name, description, logo, WhatsApp, hours."),
        ("UC-07","Add / Edit Product",          "Brand Owner","Submit product form → validate → save to catalog."),
        ("UC-08","Bulk Import Products",        "Brand Owner","Upload XLSX → parse → create products in bulk."),
        ("UC-09","Review Order",                "Brand Owner","View pending order → approve or reject → notify buyer."),
        ("UC-10","Browse Marketplace",          "Buyer",      "Search/filter brands → view results with badges and ratings."),
        ("UC-11","View Brand Catalog",          "Buyer",      "Select brand → view all products with pricing tiers."),
        ("UC-12","Place Order",                 "Buyer",      "Select product + quantity → validate MOQ/stock → create order."),
        ("UC-13","Track Orders",                "Buyer",      "View order list with status and order details."),
        ("UC-14","Submit Review",               "Buyer",      "Submit rating + text → notify brand owner."),
        ("UC-15","Switch Currency",             "All",        "Select currency → all prices converted instantly."),
        ("UC-16","Switch Language",             "All",        "Toggle EN/AR → UI re-renders with locale + RTL."),
    ]
    make_table(["UC ID","Use Case","Actor","Description"], uc,
               col_widths=[1.5, 3.5, 2.5, 9.0], font_size=9)

    h2("2.8  Requirements Traceability Matrix (RTM)")
    doc.add_paragraph()
    rtm = [
        ("SR-01","Role-based access control",            "users.role; middleware; requireRole()",           "Complete",    "Manual"),
        ("SR-02","Route protection",                     "Next.js middleware; server-side auth guards",     "Complete",    "Manual"),
        ("SR-03","Brand auto-creation on owner register","/api/auth/register — brand creation block",      "Complete",    "Manual"),
        ("SR-04","Brand approval workflow",              "/api/admin/brands/[id]; admin dashboard",         "Complete",    "Manual"),
        ("SR-05","Verification tier badges",             "verification_tier field; VerifiedBadge component","Complete",    "Manual"),
        ("SR-06","Product CRUD",                         "/api/products; brand product pages",              "Complete",    "Manual"),
        ("SR-07","Bulk-pricing tiers",                   "BulkPricingTier; pricingUtils.ts; order route",   "Complete",    "Manual"),
        ("SR-08","Order MOQ/stock validation",           "/api/orders POST — validation block",             "Complete",    "Manual"),
        ("SR-09","Order status lifecycle",               "Order.status enum; PATCH /api/orders/[id]",       "Complete",    "Manual"),
        ("SR-10","Buyer review submission",              "/api/reviews POST; ReviewModal",                  "Complete",    "Manual"),
        ("SR-11","Aggregate review statistics",          "GET /api/reviews — avgRating & breakdown",        "Complete",    "Manual"),
        ("SR-12","Brand trust scores",                   "BrandTrust schema; /api/trust/[id]",              "Partial 70%", "Not started"),
        ("SR-13","Buyer trust scores",                   "BuyerTrust schema; /api/trust/[id]",              "Partial 70%", "Not started"),
        ("SR-14","Notification system",                  "notifications.ts; /api/notifications; Zustand",   "Complete",    "Manual"),
        ("SR-15","Multi-currency display",               "currencies.ts; CurrencyContext",                  "Complete",    "Manual"),
        ("SR-16","Bilingual EN/AR support",              "next-intl; en.json / ar.json; RTL CSS",           "Partial 80%", "Manual"),
        ("SR-17","WhatsApp click tracking",              "/api/brands/whatsapp-click",                      "Complete",    "Manual"),
        ("SR-18","Bulk product import",                  "/api/products/bulk; CatalogImportModal; xlsx",    "Complete",    "Manual"),
        ("SR-19","SEO structured data",                  "JSON-LD in layout.tsx; sitemap.ts",               "Complete",    "Manual"),
        ("SR-20","Review moderation queue",              "/api/admin/reviews/[id]",                         "Complete",    "Manual"),
    ]
    make_colored_table(
        ["Req.ID","Requirement","Design Artefact","Status","Test"],
        rtm,
        [HEX_GREEN if r[3]=="Complete" else HEX_YELLOW for r in rtm],
        col_widths=[1.5, 4.0, 4.5, 2.5, 2.0], font_size=8
    )
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER III — SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_3():
    h1("Chapter III:  System Design")

    h2("3.1  Sequence Diagrams")
    body("Three principal sequence diagrams capture the system's key interaction flows.  "
         "Full visual diagrams are in Appendix A (to be drawn in draw.io / Lucidchart).")

    h3("3.1.1  User Registration and Login")
    body("Client → POST /api/auth/register{name, email, password, role, company_name} → "
         "duplicate-email check (409 if found) → create User record → conditionally create "
         "pending Brand → write DB → fire admin notification(s) → return 200 + HTTP-only "
         "cookie (b2b_user_id, 7-day TTL).  Login follows an identical flow via "
         "POST /api/auth/login.")

    h3("3.1.2  Order Lifecycle")
    body("Buyer → POST /api/orders{brand_id, items[]} → validate MOQ + stock per item → "
         "resolve bulk-pricing tier → compute total → persist Order + OrderItems → notify "
         "brand owner.  Brand Owner → PATCH /api/orders/:id{status} → update status → "
         "notify buyer (order_approved / order_rejected).")

    h3("3.1.3  Brand Verification")
    body("Brand Owner registers → system creates pending brand → Admin receives "
         "notification → Admin sends PATCH /api/admin/brands/:id{status, verification_tier} "
         "→ update brand record → notify brand owner (account_approved) → brand visible "
         "in marketplace.")

    h2("3.2  Data Flow Diagram (DFD)")
    body("Figure 3 shows the Level-1 DFD.  Four data stores: D1 Users, D2 Brands/Products, "
         "D3 Orders/OrderItems, D4 Reviews/Notifications.  Five processes: P1 Auth & Sessions, "
         "P2 Catalog Management, P3 Order Processing, P4 Trust & Reputation Engine, P5 Admin Governance.")

    h2("3.3  System Architecture")
    body("Figure 1 (Appendix A) illustrates the four-layer architecture:")
    bullet("Presentation Layer — React 19 / Next.js App Router, CSS Modules, Lucide React, Recharts, CurrencyContext (React Context), Zustand (notification state), next-intl (i18n).")
    bullet("Application / API Layer — Next.js Route Handlers (/src/app/api/**), RESTful JSON, Supabase SSR middleware for session refresh on every request.")
    bullet("Data Access Layer — db.ts with readDB() / writeDB() (JSON MVP); Supabase PostgreSQL client (production path, migration in progress); TypeScript interfaces for all 10 entity types.")
    bullet("Infrastructure Layer — Vercel edge network (hosting + CDN + SSR), Supabase managed cloud (Auth + DB), WhatsApp Business deep-links (external).")

    h2("3.4  Activity Model")
    body("Figure 9 (Appendix A) shows the Activity Diagram for order placement.  "
         "Start → Buyer views product → enters quantity → [qty < MOQ] → error; "
         "[qty ≥ MOQ] → [stock < qty] → stock error; [OK] → resolve tier price → "
         "compute total → create order → notify brand owner → display confirmation.  "
         "Brand owner: notification → approve or reject → update order → notify buyer.")

    h2("3.5  Class Diagram")
    body("Figure 8 (Appendix A) shows 10 entity classes with attributes and associations:")
    bullet("User 1 — 0..1 Brand  (brand_owner has at most one brand)")
    bullet("Brand 1 — * Product  (one brand owns zero or more products)")
    bullet("Order * — 1 Brand,  Order * — 1 User/Buyer")
    bullet("Order 1 — * OrderItem,  OrderItem * — 1 Product")
    bullet("Review * — 1 Brand,  Review * — 1 User/Buyer,  Review 0..1 — 0..1 Order")
    bullet("Notification * — 1 User")
    bullet("BrandTrust 1 — 1 Brand,  BuyerTrust 1 — 1 User/Buyer")
    code_para(
        "interface Product {\n"
        "  id: string;  brand_id: string;  name: string;\n"
        "  price: number;  moq: number;  stock: number;\n"
        "  bulk_pricing_tiers: BulkPricingTier[];  created_at: string;\n"
        "}\n"
        "interface BulkPricingTier { min_qty: number; max_qty: number|null; price: number; }"
    )

    h2("3.6  Prototype Screenshots")
    body("Screenshots are provided in Appendix B.  Key screens: landing page (EN + AR), "
         "marketplace grid with verification badges, brand catalog with pricing tiers, "
         "order form with MOQ validation, brand owner dashboard, admin pending-approvals queue, "
         "review modal, notification panel.")
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER IV — IMPLEMENTATION  (Fix #4: embed file-tree image)
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_4():
    h1("Chapter IV:  System Implementation")

    h2("4.1  Make / Buy / Lease Analysis")
    mbl = [
        ("Authentication",        "Buy (Supabase Auth)",          "Provides bcrypt hashing, JWT, SSR session management — building from scratch would introduce security risk."),
        ("Database (production)", "Buy (Supabase PostgreSQL)",    "Managed PostgreSQL; eliminates DBA overhead; supports Realtime subscriptions."),
        ("Database (MVP/dev)",    "Make (JSON file + db.ts)",     "Minimal setup cost; sufficient for a small seed dataset during development sprints."),
        ("Web Framework",         "Buy (Next.js 16)",             "Industry-standard full-stack React framework with SSR, API routes, and Vercel-native deployment."),
        ("UI Components",         "Make (CSS Modules)",           "Custom styling required for GCC brand identity and Arabic RTL."),
        ("Internationalisation",  "Buy (next-intl 4.9)",          "Handles server-side i18n, locale routing, and RTL awareness."),
        ("Charts / Analytics",    "Buy (Recharts 3.8)",           "Declarative React charting; no custom SVG required."),
        ("Notification State",    "Buy (Zustand 5.0)",            "Lightweight client state for unread-count; avoids Redux boilerplate."),
        ("Excel Import",          "Buy (xlsx 0.18.5)",            "Industry-standard XLSX/CSV parser."),
        ("Hosting / CDN",         "Lease (Vercel)",               "Zero-config Next.js deployment, global CDN, preview environments."),
        ("Payment Gateway",       "N/A (Future)",                 "Tap Payments or Checkout.com — to be evaluated for GCC compliance."),
        ("Shipping Carrier API",  "N/A (Future)",                 "Aramex / DHL — planned for post-MVP phase."),
    ]
    make_table(["Component","Decision","Justification"], mbl,
               col_widths=[3.5, 3.0, 10.0], font_size=9)

    h2("4.2  Coding Artefacts & Project Structure")
    body(
        "Figure 10 below shows the VS Code project explorer for the Kuwait B2B Hub codebase.  "
        "The project is structured under the /app directory following Next.js 16 App Router "
        "conventions, with source code in /src, internationalisation in /messages, and "
        "configuration at the root."
    )

    # ── FIX #4 — Generate and embed the file-tree image ──────────────────────
    img_path = "/tmp/b2b_file_tree.png"
    build_file_tree_image(img_path, width=780, scale=2)

    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = caption_p.add_run("Figure 10 — VS Code Project Explorer: Kuwait B2B Hub Folder Structure")
    rc.italic = True; rc.font.size = Pt(9); rc.font.color.rgb = GRAY

    doc.add_picture(img_path, width=Inches(6.2))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    h3("4.2.1  Key Implementation Highlights")

    h3("4.2.2  Role-Based Auth Guard (src/lib/auth.ts)")
    body("requireRole() (line 27) wraps requireAuth() and throws HTTP 403 if the caller's "
         "role is not in the allowed list.  getCurrentUser() calls Supabase getUser() and "
         "joins the profiles table for the full User object.")

    h3("4.2.3  Order Validation (src/app/api/orders/route.ts, lines 58–79)")
    body("The POST handler validates per item: (a) quantity is a positive integer; "
         "(b) product exists and belongs to the requested brand; (c) quantity ≥ product.moq; "
         "(d) product.stock ≥ quantity.  Unit price is resolved from bulk_pricing_tiers.")
    code_para(
        "for (const item of items) {\n"
        "  const qty = Number(item.quantity);\n"
        "  if (!Number.isInteger(qty) || qty <= 0) return error('Invalid qty');\n"
        "  if (qty < product.moq)  return error(`MOQ is ${product.moq}`);\n"
        "  if (product.stock < qty) return error(`Only ${product.stock} available`);\n"
        "  // resolve tier price\n"
        "  for (const tier of product.bulk_pricing_tiers) {\n"
        "    if (qty >= tier.min_qty && (!tier.max_qty || qty <= tier.max_qty))\n"
        "      { unitPrice = tier.price; break; }\n"
        "  }\n"
        "}"
    )

    h3("4.2.4  Currency Engine (src/lib/currencies.ts)")
    body("CURRENCIES maps KWD/SAR/AED/BHD to symbol, VAT rate (0% / 15% / 5% / 10%), "
         "and locale.  formatPriceWithVAT() converts from KWD and applies destination VAT.")

    h3("4.2.5  Non-Fatal Notification (src/lib/notifications.ts)")
    body("createNotification() wraps all DB writes in try/catch that silently suppresses "
         "errors — ensuring a notification failure never causes an HTTP 500 on an otherwise "
         "successful request.")

    h2("4.3  Configuration Management")
    cm = [
        ("Version Control",  "Git",                  "Single main branch; full history in app/.git"),
        ("Build System",     "Next.js / npm",        "npm run dev | build | start | lint"),
        ("Type Checking",    "TypeScript 5 strict",  "tsconfig.json strict: true; incremental builds"),
        ("Linting",          "ESLint 9 + next config","eslint.config.mjs; zero-warnings policy"),
        ("Environment Vars", ".env.local",           "NEXT_PUBLIC_SUPABASE_URL, ANON_KEY — not committed"),
        ("Dependency Lock",  "package-lock.json",    "Reproducible installs across team and CI"),
        ("Translations",     "Python audit scripts", "find_missing_keys.py, audit_translations.py — run each sprint"),
    ]
    make_table(["Area","Tool","Details"], cm, col_widths=[3.0, 3.5, 10.0], font_size=9)

    h2("4.4  Deployment Specification")
    deploy = [
        ("Hosting Platform",      "Vercel",                         "Zero-config Next.js; auto CI/CD on git push"),
        ("Runtime",               "Node.js 18+",                   "Vercel managed serverless functions"),
        ("Build Command",         "npm run build",                  "Produces .next/ optimised production bundle"),
        ("Auth Service",          "Supabase (hosted)",             "SUPABASE_URL + ANON_KEY env vars"),
        ("Database (MVP)",        "JSON file (Vercel filesystem)",  "src/data/db.json — ephemeral; migration planned"),
        ("Database (production)", "Supabase PostgreSQL",           "GCC-region Supabase instance"),
        ("CDN",                   "Vercel Edge Network",           "Static assets, image optimisation"),
        ("SSL / TLS",             "Vercel / Let's Encrypt",        "Automatic HTTPS for all deployments"),
        ("Domain",                "TBD (b2bhub.kw planned)",       "Kuwait ccTLD requires local entity registration"),
    ]
    make_table(["Config Item","Value","Notes"], deploy,
               col_widths=[4.0, 4.5, 8.0], font_size=9)

    h2("4.5  Physical Deployment Diagram")
    body("Figure 11 (Appendix A) shows the physical deployment topology.  End-user devices "
         "(desktop + mobile browsers) communicate via HTTPS to Vercel's edge network.  The "
         "Vercel CDN serves static assets and routes dynamic requests to Next.js serverless "
         "functions.  Functions communicate with Supabase cloud (Auth + PostgreSQL) over TLS.  "
         "WhatsApp Business deep-links open the user's WhatsApp client directly.  During the "
         "MVP phase, the function reads/writes the bundled db.json file within the same Vercel "
         "environment.")
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER V — V&V
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_5():
    h1("Chapter V:  Verification & Validation")

    h2("5.1  Development Testing")
    body("Development testing focused on validating API endpoints and UI flows before "
         "integration.  Checks performed each sprint:")
    bullet("TypeScript strict-mode compilation — zero type errors before commit.")
    bullet("ESLint static analysis — npm run lint; zero warnings enforced.")
    bullet("Manual API testing via browser DevTools for edge cases: missing fields, wrong role, duplicate email, below-MOQ, insufficient stock.")
    bullet("Translation audit scripts run after every new UI string addition.")
    bullet("Browser console monitoring — zero unhandled JS exceptions on all main flows.")

    h2("5.2  System Inspection")
    body("Structured code inspection at end of Sprint 6 — five security-sensitive modules:")
    bullet("Finding 1 (Critical — Pending): Plain-text passwords in db.json. Resolution: Supabase Auth (bcrypt) migration — Sprint 8 primary objective.")
    bullet("Finding 2 (Medium — In Progress): b2b_user_role cookie not HTTP-only. Resolution: role validation moved server-side; cookie to be HTTP-only post-migration.")
    bullet("Finding 3 (Low — Resolved): Review content not sanitised. Resolution: content.slice(0, 300) + XSS sanitisation in display component.")
    bullet("Finding 4 (Low — Resolved): TypeScript 'any' in review route handlers. Resolution: explicit type annotations added.")
    bullet("Finding 5 (Info): No rate limiting on auth endpoints. Resolution: Vercel edge rate limiting + Supabase Auth brute-force protection post-migration.")

    h2("5.3  System Testing")
    tc = [
        ("TC-01","Register new buyer",             "Valid form role=buyer",         "User created; dashboard accessible; admin notified", "Pass",""),
        ("TC-02","Register brand owner",            "Valid form role=brand_owner",   "User + pending brand created; admin notified x2",    "Pass",""),
        ("TC-03","Duplicate email",                 "Existing email",                "HTTP 409 error displayed",                           "Pass",""),
        ("TC-04","Login valid credentials",         "Correct email + password",      "Session cookie set; redirected to dashboard",         "Pass",""),
        ("TC-05","Login invalid password",          "Wrong password",                "Error message; no cookie set",                        "Pass",""),
        ("TC-06","Unauthorised route access",       "Buyer → /admin/dashboard",      "Redirect to login / 403",                            "Pass",""),
        ("TC-07","Brand approval",                  "Admin approves pending brand",  "status=approved; brand visible; brand owner notified","Pass",""),
        ("TC-08","Create product",                  "MOQ=15, price=32.5 KWD",        "Product in brand catalog",                           "Pass",""),
        ("TC-09","Below-MOQ order rejection",       "qty=5, MOQ=15",                 "HTTP 400 with clear error",                          "Pass",""),
        ("TC-10","Valid order creation",            "qty=15, stock=450",             "Order created; brand owner notified",                "Pass",""),
        ("TC-11","Bulk-pricing tier applied",       "50 units (tier KD 28)",         "Total = 50 × 28 = KD 1,400",                        "Pass",""),
        ("TC-12","Order approval by brand",         "Brand approves order",          "status=approved; buyer notified",                    "Pass",""),
        ("TC-13","Submit review",                   "5-star with text",              "Review saved; brand owner notified",                 "Pass",""),
        ("TC-14","Anonymous review",                "anonymous=true",                "Author shown as 'Anonymous Buyer'",                  "Pass",""),
        ("TC-15","Currency switch to SAR",          "Select SAR",                    "Prices × 12.20",                                     "Pass",""),
        ("TC-16","Arabic RTL layout",               "Toggle to Arabic",              "RTL rendered; Arabic text displayed",                "Partial","Some components not RTL-adapted"),
        ("TC-17","WhatsApp click tracking",         "Click WhatsApp button",         "click counter incremented",                         "Pass",""),
        ("TC-18","Excel product import",            "Upload XLSX with 5 products",   "5 products created",                                 "Pass",""),
        ("TC-19","Notification bell",               "New order (unseen notif)",      "Bell badge + panel opens",                          "Pass",""),
        ("TC-20","Admin review moderation",         "Admin removes flagged review",  "status=removed; hidden from buyers",                "Pass",""),
    ]
    make_colored_table(
        ["TC ID","Test Case","Input","Expected","Result","Notes"],
        tc,
        [HEX_GREEN if r[4]=="Pass" else HEX_YELLOW for r in tc],
        col_widths=[1.2, 3.0, 3.2, 4.0, 1.5, 3.6], font_size=8
    )

    h2("5.4  Release Testing")
    body("To be completed in Sprint 8:")
    numbered("Regression sweep across all 20 TC cases after Supabase migration.")
    numbered("Cross-browser: Chrome 124, Safari 17, Firefox 125, Edge 124.")
    numbered("Viewport: 375 px (iPhone SE), 414 px (Pro Max), 768 px (iPad).")
    numbered("Lighthouse audit — targets: LCP < 2.5 s, TBT < 200 ms, CLS < 0.1.")
    numbered("HTTP security headers review (CSP, X-Frame-Options, HSTS).")
    numbered("i18n completeness — translation scripts must report 0 missing keys.")

    h2("5.5  User Testing")
    body("Informal UAT with two volunteer business users (one retail buyer, one brand owner) in Sprint 5:")
    bullet("Positive: Marketplace search described as 'faster than WhatsApp'.")
    bullet("Positive: Verification badge and trust score cited as confidence builders.")
    bullet("Improvement: Order form needed clearer MOQ guidance — resolved with dynamic MOQ label.")
    bullet("Improvement: Arabic brand descriptions requested — added to post-MVP roadmap.")
    bullet("Improvement: Bulk-pricing tier table not immediately visible — resolved as default-expanded section.")
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER VI — EVOLUTION
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_6():
    h1("Chapter VI:  System Evolution")

    h2("6.1  Anticipated Changes")
    changes = [
        ("EV-01","Supabase PostgreSQL Migration",       "High","High",  "Replace JSON file layer with Supabase queries in db.ts."),
        ("EV-02","Real-Time Notifications",             "High","Medium","Supabase Realtime subscriptions via useRealtimeNotifications hook."),
        ("EV-03","Payment Gateway Integration",         "High","High",  "Tap Payments or Checkout.com for GCC-compliant payment processing."),
        ("EV-04","Email / SMS Notifications",           "Med", "Medium","SendGrid / AWS SES for emails; Twilio for SMS."),
        ("EV-05","Automated CR Verification",           "Med", "High",  "Kuwait MoCI API to auto-validate commercial registration numbers."),
        ("EV-06","Shipping Carrier Integration",        "Med", "High",  "Aramex / DHL API for shipping rates and parcel tracking."),
        ("EV-07","Live Exchange Rate API",              "Low", "Low",   "Replace hardcoded rates with scheduled FX data provider fetch."),
        ("EV-08","Arabic Brand Descriptions",           "Low", "Low",   "Add Arabic description field to Brand and Product entities."),
        ("EV-09","Wishlist & Brand Comparison",         "Low", "Medium","Save brands/products; compare up to 3 brands side-by-side."),
        ("EV-10","Native Mobile Apps",                 "Low", "High",  "React Native or PWA for iOS/Android."),
    ]
    make_colored_table(
        ["ID","Change","Priority","Effort","Description"],
        changes,
        [HEX_RED if r[2]=="High" else (HEX_YELLOW if r[2]=="Med" else HEX_GREEN) for r in changes],
        col_widths=[1.2, 3.8, 1.8, 1.8, 7.9], font_size=9
    )

    h2("6.2  Phase-Out Transitions")
    bullet("JSON → Supabase DB: readDB()/writeDB() replaced method-by-method; JSON file retained as seed fixture only.")
    bullet("Local Auth → Supabase Auth: Cookie-based local auth superseded by Supabase Auth JWTs; register/login/logout routes deprecated.")
    bullet("Hardcoded Rates → Live API: EXCHANGE_RATES replaced with Next.js revalidated API call (1-hour interval).")
    bullet("Manual Brand Verification → Hybrid: Admin approval augmented (not replaced) with automated CR lookup.")

    h2("6.3  Quality Predictions")
    q_rows = [
        ("API Response Time", "< 800 ms (JSON file)",     "< 300 ms (Supabase + CDN)",      "< 150 ms (edge caching)"),
        ("Test Coverage",     "0% automated (manual)",    "~40% (unit + E2E critical paths)","~80% (full suite)"),
        ("Translation",       "~80% (AR partial)",        "100% (all keys filled)",          "100% + new features"),
        ("Uptime",            "N/A (dev environment)",    "~99.5% (Vercel + Supabase SLA)",  "> 99.9% (HA config)"),
        ("Security (OWASP)",  "Medium (plain-text pw)",   "High (bcrypt, HTTPS, CSP)",       "High (pen-test verified)"),
    ]
    make_table(
        ["Metric","MVP (Now)","Phase 2 (Post-Migration)","Phase 3 (Production)"],
        q_rows, col_widths=[3.5, 4.0, 4.0, 5.0], font_size=9
    )
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER VII — PLANNING (Fix #5: Activity Network diagram)
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_7():
    h1("Chapter VII:  System Planning & Management")

    h2("7.1  Milestones")
    milestones = [
        ("M-01","Project Kickoff & Tech Stack Decision","Jan 1, 2026",  "Complete","Architecture agreed; Next.js + Supabase selected"),
        ("M-02","Authentication & DB Schema",           "Jan 15, 2026", "Complete","User registration, login, roles functional"),
        ("M-03","Marketplace MVP (browse + catalog)",   "Feb 1, 2026",  "Complete","Brand cards, product detail, verification badges"),
        ("M-04","Order Lifecycle Complete",             "Feb 20, 2026", "Complete","Order create → approve/reject; dashboards"),
        ("M-05","Trust & Reputation Engine",            "Mar 10, 2026", "Complete","Review system, ratings, trust scores, admin mod"),
        ("M-06","Admin Dashboard & Notifications",      "Mar 25, 2026", "Complete","Admin panel, notification panel, analytics"),
        ("M-07","Localisation & Mobile Responsiveness", "Apr 10, 2026", "Partial", "EN complete; AR ~80%; mobile mostly responsive"),
        ("M-08","Progress Report Submission",           "Apr 17, 2026", "Complete","Progress_Report_Final.docx submitted"),
        ("M-09","Supabase DB Migration",                "May 5, 2026",  "Planned", "Full data layer on Supabase PostgreSQL"),
        ("M-10","Automated Testing Suite",              "May 15, 2026", "Planned", "Jest unit tests + Playwright E2E"),
        ("M-11","Final Report First Draft",             "May 25, 2026", "Planned", "All chapters drafted and reviewed"),
        ("M-12","Final Report Submission",              "Jun 15, 2026", "Planned", "Submitted to university portal"),
        ("M-13","Final Presentation / Viva",            "Jun 25, 2026", "Planned", "Presentation delivered; Q&A handled"),
    ]
    make_colored_table(
        ["ID","Milestone","Target Date","Status","Notes"],
        milestones,
        [HEX_GREEN if r[3]=="Complete" else (HEX_YELLOW if r[3]=="Partial" else HEX_LBLUE) for r in milestones],
        col_widths=[1.2, 5.5, 2.5, 2.0, 5.3], font_size=9
    )

    h2("7.2  Gantt Chart")
    body("Figure 12 (Appendix A) shows the Gantt chart across 8 sprints (January–June 2026).  "
         "Critical path: Auth/Setup → Marketplace → Orders → Trust/Reputation → Admin/Notifications "
         "→ Localisation → Supabase Migration → Testing & Report.")
    gantt = [
        ("Sprint 1","Jan 1–14",    "Auth, DB schema, project setup"),
        ("Sprint 2","Jan 15–31",   "Marketplace, brand browsing, search/filter"),
        ("Sprint 3","Feb 1–20",    "Order lifecycle, buyer dashboard"),
        ("Sprint 4","Feb 21–Mar 9","Brand owner dashboard, product CRUD, Excel import"),
        ("Sprint 5","Mar 10–24",   "Reviews, trust scores, admin moderation"),
        ("Sprint 6","Mar 25–Apr 9","Admin dashboard, notifications, WhatsApp tracking"),
        ("Sprint 7","Apr 10–Apr 30","Localisation, mobile responsiveness, progress report"),
        ("Sprint 8","May 1–Jun 15","Supabase migration, testing, final report"),
    ]
    make_table(["Sprint","Dates","Focus"], gantt,
               col_widths=[2.0, 3.5, 11.0], font_size=10)

    # ── FIX #5 — Activity Network Diagram ────────────────────────────────────
    h2("7.3  Activity Network Diagram")
    body(
        "Figure 13 shows the Activity Network (Precedence Diagram) for the eight project sprints.  "
        "Each node represents a sprint; arrows represent mandatory dependencies.  "
        "The critical path is highlighted: S1 → S2 → S3 → S4 → S5 → S6 → S7 → S8."
    )
    doc.add_paragraph()

    # Table-based network: Sprint | Depends On | Duration | Early Start | Early Finish | Slack
    body("Table 14 — Activity Network: Sprint Dependency & Float Analysis")
    net_rows = [
        ("S1 — Auth & Setup",               "—",              "2 wks","Jan 1",  "Jan 14",  "0 (Critical)"),
        ("S2 — Marketplace",                "S1",             "2 wks","Jan 15", "Jan 31",  "0 (Critical)"),
        ("S3 — Order Lifecycle",            "S2",             "3 wks","Feb 1",  "Feb 20",  "0 (Critical)"),
        ("S4 — Brand Owner Tools",          "S3",             "2.5 wks","Feb 21","Mar 9",  "0 (Critical)"),
        ("S5 — Trust & Reputation",         "S4",             "2 wks","Mar 10", "Mar 24",  "0 (Critical)"),
        ("S6 — Admin & Notifications",      "S5",             "2 wks","Mar 25", "Apr 9",   "0 (Critical)"),
        ("S7 — Localisation & Mobile",      "S6",             "3 wks","Apr 10", "Apr 30",  "0 (Critical)"),
        ("S8 — Migration, Testing & Report","S7",             "6 wks","May 1",  "Jun 15",  "0 (Critical)"),
    ]
    make_table(
        ["Activity","Depends On","Duration","Early Start","Early Finish","Float"],
        net_rows,
        col_widths=[4.5, 2.0, 2.0, 2.0, 2.5, 3.5],
        font_size=9
    )

    doc.add_paragraph()
    body("Arrow notation — dependency chain:")
    # Visual ASCII-like arrow chain in a styled paragraph
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(
        "[ S1: Auth ]  ──►  [ S2: Marketplace ]  ──►  [ S3: Orders ]  ──►  [ S4: Brand Tools ]\n"
        "                                                                              │\n"
        "                                                                              ▼\n"
        "                                          [ S8: Migration + Report ]  ◄──  [ S7: Localisation ]  ◄──  [ S6: Admin ]  ◄──  [ S5: Trust ]"
    )
    r.font.name = "Courier New"; r.font.size = Pt(9); r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()
    body("Note: All sprints lie on the critical path (zero float).  No parallel sprints exist "
         "in this single-team Agile project; each sprint's output is a prerequisite for the next.")

    h2("7.4  Staff Allocation")
    staff = [
        ("Bader Alnoumas",  "25%","25%","20%","15%","10%","5%","100%"),
        ("Ahmad Al-Rashidi","10%","30%","20%","15%","15%","10%","100%"),
        ("Omar Al-Mutairi", "15%","10%","25%","25%","15%","10%","100%"),
        ("Noor Al-Dosari",  "10%","10%","10%","10%","30%","30%","100%"),
    ]
    make_table(
        ["Team Member","Auth/Setup","Marketplace","Orders","Trust/Admin","Testing","Docs/Report","Total"],
        staff,
        col_widths=[3.8, 1.8, 2.5, 1.8, 2.5, 1.8, 2.5, 1.3],
        font_size=9
    )

    h2("7.5  CoCoMo Cost Estimation")
    cocomo = [
        ("Estimated KLOC",     "3.0 KLOC",            "~3,000 lines TypeScript/TSX (excl. node_modules)"),
        ("CoCoMo Model",       "Organic",             "Small team, familiar tech, simple domain"),
        ("Effort Equation",    "E = 2.4 × (3.0)^1.05","≈ 7.6 person-months"),
        ("Duration Equation",  "D = 2.5 × (7.6)^0.38","≈ 6.1 months"),
        ("Actual Duration",    "~6 months",           "Jan 2026 – Jun 2026 (matches estimate)"),
        ("Avg Team Size",      "~1.27 FTE each",      "7.6 PM / 6 months / 4 members"),
        ("Est. Cost (academic)","KD 12,000",          "KD 500/developer/month × 4 × 6"),
    ]
    make_table(["Parameter","Value","Notes"], cocomo,
               col_widths=[4.0, 3.5, 9.0], font_size=10)

    h2("7.6  Risk Matrix")
    risks = [
        ("R-01","Supabase migration overrun",          "Med","High","High",  "Full Sprint 8 allocated; JSON fallback available"),
        ("R-02","Automated tests reveal critical bugs", "Med","High","High",  "Prioritise critical-path tests; manual sweep as backup"),
        ("R-03","Team member unavailable",             "Low","High","Med",   "Shared knowledge; docs kept current"),
        ("R-04","Supabase Realtime blocked",           "Med","Low", "Low",   "Fall back to polling; Realtime is enhancement"),
        ("R-05","Arabic translation gaps",             "High","Med","High",  "Audit scripts run weekly; dedicated team member"),
        ("R-06","Vercel/Supabase free-tier limits",    "Low","Med", "Low",   "Monitor usage; upgrade plan negligible cost"),
        ("R-07","Scope creep in Sprint 8",             "Low","Med", "Low",   "Feature freeze; new items logged in backlog"),
        ("R-08","Security vulnerability discovered",   "Med","High","High",  "Supabase Auth migration; CSP headers pre-launch"),
    ]
    make_colored_table(
        ["ID","Risk","Likelihood","Impact","Exposure","Mitigation"],
        risks,
        [HEX_RED if r[4]=="High" else (HEX_YELLOW if r[4]=="Med" else HEX_GREEN) for r in risks],
        col_widths=[1.0, 3.5, 2.0, 1.8, 2.0, 6.2], font_size=8
    )

    h2("7.7  Fishbone Diagram")
    body("Figure 14 (Appendix A) shows the Ishikawa Diagram for 'Delayed Final Submission'. "
         "Five cause bones:")
    bullet("People: Team member unavailability; knowledge silos; communication gaps.")
    bullet("Process: Scope creep; insufficient sprint planning; unclear acceptance criteria.")
    bullet("Technology: Supabase migration complexity; real-time implementation; dependency conflicts.")
    bullet("Environment: Platform outages; local environment inconsistencies.")
    bullet("Materials: Incomplete requirements; missing design artefacts; documentation gaps.")
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER VIII — CONCLUSIONS
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_8():
    h1("Chapter VIII:  Conclusions")

    h2("8.1  Seven Design Smells Identified and Addressed")
    smells = [
        ("DS-01","God Object",
         "db.ts was a single read/write gateway to all 10 entity types — tight coupling.",
         "Partially resolved by typed per-entity helpers; full resolution via Supabase per-table queries."),
        ("DS-02","Primitive Obsession",
         "Order status passed as raw strings without a shared enum.",
         "Resolved: TypeScript union type defined in db.ts and used consistently."),
        ("DS-03","Divergent Change",
         "Currency formatting logic duplicated across three components.",
         "Resolved: all formatting delegates to formatPrice() / formatPriceWithVAT() in currencies.ts."),
        ("DS-04","Shotgun Surgery",
         "Adding a new notification type required changes in multiple files.",
         "Partially resolved: notification creation centralised in notifications.ts; icon-map still manual."),
        ("DS-05","Feature Envy",
         "Review components computed aggregates that belonged in the data layer.",
         "Resolved: aggregates computed server-side in GET /api/reviews response payload."),
        ("DS-06","Inappropriate Intimacy",
         "Order creation route directly accessed brand and product DB tables for validation.",
         "Acknowledged: service layer planned for production architecture."),
        ("DS-07","Magic Numbers",
         "Exchange rates and VAT percentages hardcoded inline in components.",
         "Resolved: centralised in CURRENCIES and EXCHANGE_RATES constants in currencies.ts."),
    ]
    for ds_id, name, problem, resolution in smells:
        h3(f"{ds_id}  {name}")
        body_bold("Problem:", problem)
        body_bold("Resolution:", resolution)

    h2("8.2  SWOT Analysis")
    doc.add_paragraph()
    swot_data = [
        ("STRENGTHS\n\n"
         "• Purpose-built for GCC: Arabic RTL, 4 currencies, GCC legal awareness\n"
         "• Verified brand ecosystem creates buyer trust\n"
         "• Structured order management replaces WhatsApp chaos\n"
         "• Trust & reputation engine differentiates from listing directories\n"
         "• Modern stack (Next.js, Supabase, TypeScript) enables rapid iteration",
         HEX_GREEN),
        ("WEAKNESSES\n\n"
         "• No payment integration — platform limited to 'order request' model\n"
         "• JSON-file database not production-ready; migration risk\n"
         "• No real-time communication\n"
         "• Limited team size constrains delivery velocity\n"
         "• Brand descriptions English-only at MVP",
         HEX_RED),
        ("OPPORTUNITIES\n\n"
         "• GCC B2B e-commerce is a multi-billion dollar underserved market\n"
         "• Kuwait Vision 2035 / Saudi Vision 2030 prioritise digital commerce\n"
         "• WhatsApp dominance highlights the pain point this platform solves\n"
         "• KD 29/month subscription provides predictable recurring revenue\n"
         "• Expansion potential: Oman, Qatar, Egypt",
         HEX_LBLUE),
        ("THREATS\n\n"
         "• Established regional competitors (Tradeling, Salla, Zid) may expand\n"
         "• Low switching cost: brand owners can return to WhatsApp\n"
         "• GCC data-localisation laws may require on-premise hosting\n"
         "• Market education required for SMBs unfamiliar with digital B2B\n"
         "• Currency and VAT regulation changes across GCC",
         HEX_YELLOW),
    ]
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (content, color) in enumerate(swot_data):
        ri, ci = divmod(idx, 2)
        cell = tbl.rows[ri].cells[ci]
        shade_cell(cell, color)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        lines = content.split('\n')
        r_title = p.add_run(lines[0]); r_title.bold = True; r_title.font.size = Pt(11)
        r_body  = p.add_run('\n' + '\n'.join(lines[1:])); r_body.font.size = Pt(9)
    for c in tbl.columns:
        for cell in c.cells:
            cell.width = Cm(8.25)

    h2("8.3  Future Work")
    numbered("Payment Gateway (Tap Payments / Checkout.com): enable end-to-end B2B transactions.")
    numbered("Supabase Realtime Messaging: order-level buyer–brand chat, replacing WhatsApp for negotiation.")
    numbered("Automated CR Verification: MoCI API integration to validate commercial registrations at sign-up.")
    numbered("Shipping Integration (Aramex / DHL): live rates, pickup scheduling, and parcel tracking.")
    numbered("Native Mobile Application: React Native or PWA for the GCC mobile-first user base.")
    numbered("AI-Powered Recommendations: collaborative filtering for brand and product discovery.")
    numbered("Multi-Brand Order Basket: single checkout session spanning multiple brand catalogs.")
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER IX — APPENDICES  (Fix #6: draw.io note)
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_9():
    h1("Chapter IX:  Appendices")

    h2("Appendix A — Diagram References")

    # ── FIX #6 — draw.io / Lucidchart note ───────────────────────────────────
    note_p = doc.add_paragraph()
    note_p.paragraph_format.left_indent  = Cm(0.3)
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after  = Pt(8)

    # Shaded note box
    pPr = note_p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  "FFF3CD")
    pPr.append(shd)

    r_icon = note_p.add_run("NOTE TO TEAM: ")
    r_icon.bold = True; r_icon.font.size = Pt(11); r_icon.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
    r_body = note_p.add_run(
        "All diagrams listed below (Class Diagram, Sequence Diagrams, DFD Levels 0 and 1, "
        "Gantt Chart, Activity Network, Fishbone / Ishikawa Diagram, Physical Deployment Diagram, "
        "and Quality Prediction Chart) must be drawn using draw.io (app.diagrams.net) or "
        "Lucidchart and exported as high-resolution PNG or SVG images.  "
        "Each image should then be inserted into this document in place of the corresponding "
        "placeholder description below.  Use Insert → Pictures in Microsoft Word, resize to "
        "full column width (≈ 16 cm), and add the Figure caption beneath each image.  "
        "Ensure all text in diagrams is legible at 100% zoom before submission."
    )
    r_body.font.size = Pt(10)

    body("The table below lists every diagram, its figure number, and the specification "
         "needed to draw it correctly.")
    doc.add_paragraph()

    diag_list = [
        ("Fig 1",  "System Architecture",          "Four-layer stack: Presentation (Next.js pages/components) → API (Route Handlers) → Data Access (db.ts / Supabase client) → Infrastructure (Vercel + Supabase cloud).  Use boxes for each layer with technology labels."),
        ("Fig 2",  "Level-0 Context DFD",           "System boundary 'Kuwait B2B Hub'.  External entities: Admin, Brand Owner, Business Buyer, Supabase Auth.  Data flows: registration, catalog, orders, reviews, approvals, notifications."),
        ("Fig 3",  "Level-1 DFD",                   "5 processes: P1 Auth, P2 Catalog, P3 Orders, P4 Trust, P5 Admin.  4 data stores: D1 Users, D2 Brands/Products, D3 Orders, D4 Reviews/Notifications.  Flows between processes and stores."),
        ("Fig 4",  "Use Case Diagram",              "3 actor stick figures (Admin, Brand Owner, Buyer) with ellipses for each of the 16 use cases (UC-01 to UC-16).  Use 'includes' and 'extends' relationships where applicable."),
        ("Fig 5",  "Sequence Diagram — Registration","Lifelines: Browser, /api/auth/register, db.ts, Notifications.  Steps: form submit → duplicate check → user create → brand create (if owner) → write DB → admin notify → set cookie → return 200."),
        ("Fig 6",  "Sequence Diagram — Order",      "Lifelines: Buyer, /api/orders, db.ts, BrandOwner.  Steps: POST items → validate MOQ/stock → resolve tier → persist → notify → return 201.  Then: PATCH status → update DB → notify buyer."),
        ("Fig 7",  "Sequence Diagram — Verification","Lifelines: BrandOwner, /api/auth/register, Admin, /api/admin/brands, db.ts.  Steps: register → create pending brand → notify admin → admin PATCH approve+tier → update brand → notify brand owner."),
        ("Fig 8",  "Class Diagram",                 "10 classes: User, Brand, Product, BulkPricingTier, Order, OrderItem, Message, Review, Notification, BrandTrust, BuyerTrust.  Attributes from db.ts interfaces.  Associations with multiplicity as described in §3.5."),
        ("Fig 9",  "Activity Diagram — Order",      "Start → Buyer views product → enter qty → [qty < MOQ]→error; [OK]→[stock<qty]→error; [OK]→resolve price→compute total→create order→notify brand owner→End.  Swimlane for Buyer and Brand Owner."),
        ("Fig 11", "Physical Deployment Diagram",   "Nodes: User Device (browser), Vercel Edge (CDN), Next.js Serverless Function, Supabase Cloud (Auth + PostgreSQL).  Connections: HTTPS, PostgreSQL TCP/TLS.  WhatsApp as external service."),
        ("Fig 12", "Gantt Chart",                   "8 sprint bars across Jan–Jun 2026.  X-axis: weeks; Y-axis: sprint name and focus.  Milestones M-01 to M-13 marked as diamonds.  Critical path highlighted in blue."),
        ("Fig 13", "Activity Network",              "8 sprint nodes in a linear chain S1→S2→…→S8.  Each node shows: sprint name, duration, ES/EF/LS/LF dates, float=0.  Arrow labels show dependency type (FS = Finish-to-Start)."),
        ("Fig 14", "Fishbone Diagram",              "Spine pointing right to 'Delayed Submission'.  5 bones: People, Process, Technology, Environment, Materials.  Sub-causes as described in §7.7."),
        ("Fig 15", "Quality Prediction Chart",      "Grouped bar chart with 5 metric clusters × 3 phases (MVP / Post-Migration / Production).  Use green/amber/blue bars.  Y-axis: normalised score 0–100."),
    ]
    make_table(
        ["Figure", "Diagram", "Specification for draw.io / Lucidchart"],
        diag_list,
        col_widths=[1.3, 3.5, 11.7], font_size=9
    )

    h2("Appendix B — Key Code Listings")
    h3("B.1  Order Validation (src/app/api/orders/route.ts, lines 58–79)")
    code_para(
        "for (const item of items) {\n"
        "  const qty = Number(item.quantity);\n"
        "  if (!Number.isInteger(qty) || qty <= 0)\n"
        "    return error('Invalid quantity');\n"
        "  const product = products.find(p => p.id === item.product_id);\n"
        "  if (!product) return error('Product not found');\n"
        "  if (qty < product.moq)\n"
        "    return error(`MOQ is ${product.moq} units`);\n"
        "  if (product.stock < qty)\n"
        "    return error(`Only ${product.stock} units available`);\n"
        "  let unitPrice = product.price;\n"
        "  for (const tier of product.bulk_pricing_tiers) {\n"
        "    if (qty >= tier.min_qty && (!tier.max_qty || qty <= tier.max_qty))\n"
        "      { unitPrice = tier.price; break; }\n"
        "  }\n"
        "  total += unitPrice * qty;\n"
        "}"
    )

    h3("B.2  Non-Fatal Notification (src/lib/notifications.ts)")
    code_para(
        "export function createNotification(opts: CreateNotifOpts) {\n"
        "  try {\n"
        "    const db = readDB();\n"
        "    db.notifications?.push({ ...opts, id: generateId('notif'),\n"
        "      read: false, created_at: new Date().toISOString() });\n"
        "    writeDB(db);\n"
        "  } catch {\n"
        "    // non-fatal — never crash a request because of a notification failure\n"
        "  }\n"
        "}"
    )

    h2("Appendix C — Sample Data Summary")
    body("The seeded demo dataset (src/data/db.json) includes:")
    sample = [
        ("Users",         "7",  "1 admin, 3 brand owners, 3 buyers — all with approved verification_status"),
        ("Brands",        "3",  "Kuwait Fashion House, Gulf Tech Solutions, Desert Rose Beauty — all Premium tier"),
        ("Products",      "2",  "Royal Kandura (MOQ=15, 3 bulk tiers); Windows test product (MOQ=2)"),
        ("Orders",        "2",  "order-001: completed KD 1,625; order-afbe6c: pending KD 487.50"),
        ("Reviews",       "10", "Ratings 4–5 across all 3 brands; 2 anonymous; 3 with brand replies"),
        ("Notifications", "19", "All types: order events, reviews, registrations, milestones"),
        ("Brand Trust",   "3",  "Response rates 88–100%; completion rates 96–100%"),
        ("Buyer Trust",   "2",  "Mohammed: 31 orders, 100% completion; Fatima: 12 orders, 92% completion"),
    ]
    make_table(["Entity","Count","Notes"], sample,
               col_widths=[3.0, 1.5, 12.0], font_size=10)
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER X — REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

def chapter_10():
    h1("Chapter X:  References")
    body("References are formatted in APA 7th edition style.")
    doc.add_paragraph()
    refs = [
        "[1]   Sommerville, I. (2016). Software Engineering (10th ed.). Pearson Education.",
        "[2]   Pressman, R. S., & Maxim, B. R. (2020). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
        "[3]   Vercel Inc. (2024). Next.js Documentation (v16). https://nextjs.org/docs",
        "[4]   Supabase Inc. (2024). Supabase Documentation — Auth & PostgreSQL. https://supabase.com/docs",
        "[5]   Beck, K., et al. (2001). Manifesto for Agile Software Development. https://agilemanifesto.org",
        "[6]   Boehm, B. W. (1981). Software Engineering Economics. Prentice-Hall.",
        "[7]   Fowler, M. (1999). Refactoring: Improving the Design of Existing Code. Addison-Wesley.",
        "[8]   Microsoft. (2024). TypeScript Handbook. https://www.typescriptlang.org/docs/handbook/",
        "[9]   World Wide Web Consortium. (2018). Web Content Accessibility Guidelines (WCAG) 2.1. https://www.w3.org/TR/WCAG21/",
        "[10]  React Team. (2024). React 19 Documentation. https://react.dev",
        "[11]  GCC Secretariat General. (2023). GCC Digital Economy Report 2023. Gulf Cooperation Council.",
        "[12]  Kuwait Ministry of Commerce and Industry. (2023). Kuwait Vision 2035 — Digital Transformation Pillar. State of Kuwait.",
        "[13]  OWASP Foundation. (2023). OWASP Top Ten Web Application Security Risks. https://owasp.org/www-project-top-ten/",
        "[14]  Recharts Team. (2024). Recharts: Redefined Chart Library Built with React and D3. https://recharts.org",
        "[15]  GitHub — pmndrs/zustand. (2024). Zustand: Bear necessities for state management in React. https://github.com/pmndrs/zustand",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after      = Pt(5)
        p.paragraph_format.left_indent      = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        r = p.add_run(ref)
        r.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════════════════
# BUILD
# ═══════════════════════════════════════════════════════════════════════════════

cover_page()
toc()
list_of_tables()
list_of_figures()
abstract()
chapter_1()
chapter_2()
chapter_3()
chapter_4()
chapter_5()
chapter_6()
chapter_7()
chapter_8()
chapter_9()
chapter_10()

out = "/Users/baderalnoumas/Desktop/B2B/Final_Project_Report_v2.docx"
doc.save(out)
print(f"Saved: {out}")
