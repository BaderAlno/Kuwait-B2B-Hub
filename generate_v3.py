#!/usr/bin/env python3
"""Generate Final_Project_Report_v3.docx from v2 with all requested fixes."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = '/Users/baderalnoumas/Desktop/B2B/Final_Project_Report_v2.docx'
OUTPUT = '/Users/baderalnoumas/Desktop/B2B/Final_Project_Report_v3.docx'

doc = Document(INPUT)

# ─────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────

def replace_in_runs(runs, old, new):
    for run in runs:
        if old in run.text:
            run.text = run.text.replace(old, new)

def clear_para_and_set(para, text):
    """Replace all runs with a single run containing text."""
    bold  = para.runs[0].bold  if para.runs else False
    size  = para.runs[0].font.size if para.runs else None
    align = para.alignment
    for r in list(para.runs):
        r._r.getparent().remove(r._r)
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size

def set_cell(cell, text, bold=False, size_pt=10, font='Calibri'):
    """Clear a table cell and set text."""
    while len(cell.paragraphs) > 1:
        cell.paragraphs[-1]._p.getparent().remove(cell.paragraphs[-1]._p)
    p = cell.paragraphs[0]
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size_pt)

def make_mono_p_elem(text, size_pt=8):
    """Return a w:p XML element with monospace Courier New text."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Courier New')
    rFonts.set(qn('w:hAnsi'), 'Courier New')
    rFonts.set(qn('w:cs'), 'Courier New')
    rPr.append(rFonts)
    for tag in ('w:sz', 'w:szCs'):
        el = OxmlElement(tag)
        el.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(el)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def replace_table_in_body(old_table, new_table):
    """Swap old table XML element for new table XML element in doc body."""
    old_tbl = old_table._tbl
    new_tbl = new_table._tbl
    parent  = old_tbl.getparent()
    idx     = list(parent).index(old_tbl)
    # Detach new table from end of body
    new_tbl.getparent().remove(new_tbl)
    # Replace old table
    parent.remove(old_tbl)
    parent.insert(idx, new_tbl)

# ─────────────────────────────────────────────────────────────
# 1. COVER PAGE FIXES
# ─────────────────────────────────────────────────────────────

for para in doc.paragraphs:
    t = para.text.strip()

    if t == 'College of Engineering and Applied Sciences (CEAS)':
        clear_para_and_set(
            para,
            'College of Engineering and Applied Sciences (CEAS)'
            ' | CSIS-330 Software Engineering | Dr. Aaron Rasheed Rababaah'
        )

    elif 'Course:' in t and ('COMP 491' in t or 'Software Engineering' in t):
        replace_in_runs(para.runs, 'Software Engineering \u2014 COMP 491', 'CSIS-330 Software Engineering')
        replace_in_runs(para.runs, 'COMP 491', 'CSIS-330 Software Engineering')

    elif 'Project Advisor:' in t and 'Dr. Mohammed Al-Enezi' in t:
        replace_in_runs(para.runs, 'Dr. Mohammed Al-Enezi', 'Dr. Aaron Rasheed Rababaah')

    elif 'Team Size:' in t and '[4]' in t:
        replace_in_runs(para.runs, '[4]', '4')

# ─────────────────────────────────────────────────────────────
# 2. GLOBAL NAME REPLACEMENT (paragraphs + table cells)
# ─────────────────────────────────────────────────────────────

NAME_MAP = {
    'Ahmad Al-Rashidi': 'Abdullah Abduljaleel',
    'Omar Al-Mutairi':  'Abdullah Subhi',
    'Noor Al-Dosari':   'Salah Abdulfattah',
}

for para in doc.paragraphs:
    for run in para.runs:
        for old, new in NAME_MAP.items():
            if old in run.text:
                run.text = run.text.replace(old, new)

for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    for old, new in NAME_MAP.items():
                        if old in run.text:
                            run.text = run.text.replace(old, new)

# ─────────────────────────────────────────────────────────────
# 3. UPDATE COVER TABLE (Table 0) — student IDs for new members
# ─────────────────────────────────────────────────────────────

team_tbl = doc.tables[0]
NEW_MEMBERS = [
    ('Abdullah Abduljaleel', '2021-00002', 'Frontend / UI Developer'),
    ('Abdullah Subhi',       '2021-00003', 'Backend / API Developer'),
    ('Salah Abdulfattah',    '2021-00004', 'Localisation & QA Engineer'),
]
for i, (name, sid, role) in enumerate(NEW_MEMBERS, start=2):
    if i < len(team_tbl.rows):
        row = team_tbl.rows[i]
        set_cell(row.cells[0], name)
        set_cell(row.cells[1], sid)
        set_cell(row.cells[2], role)

# ─────────────────────────────────────────────────────────────
# 4. REPLACE WORKLOAD TABLE (Table 2)
# ─────────────────────────────────────────────────────────────
# New format: No. | Lead Member | Ch1(7%) … Ch10(2%)
# Each member leads chapters summing to ~25%

WL_HEADERS = [
    'No.', 'Lead Member',
    'Ch1\n(7%)', 'Ch2\n(7%)', 'Ch3\n(20%)', 'Ch4\n(25%)',
    'Ch5\n(15%)', 'Ch6\n(10%)', 'Ch7\n(7%)', 'Ch8\n(5%)',
    'Ch9\n(2%)', 'Ch10\n(2%)'
]
# marks: index 0→Ch1 … 9→Ch10
WL_DATA = [
    ('1', 'Bader Alnoumas',       ['', '', '',  'X', '',  '',  '',  '',  '',  ''  ]),  # Ch4=25%
    ('2', 'Abdullah Abduljaleel', ['', '', 'X', '',  '',  '',  '',  '', 'X', 'X' ]),  # Ch3+Ch9+Ch10=24%
    ('3', 'Abdullah Subhi',       ['', '', '',  '',  'X', 'X', '',  '',  '',  ''  ]),  # Ch5+Ch6=25%
    ('4', 'Salah Abdulfattah',    ['X', 'X', '', '', '',  '',  'X', 'X', '',  ''  ]),  # Ch1+Ch2+Ch7+Ch8=26%
]

new_wl = doc.add_table(rows=len(WL_DATA) + 2, cols=12)
new_wl.style = 'Table Grid'
# Header row
for i, h in enumerate(WL_HEADERS):
    set_cell(new_wl.rows[0].cells[i], h, bold=True, size_pt=9)
# Data rows
for r, (no, name, marks) in enumerate(WL_DATA, start=1):
    set_cell(new_wl.rows[r].cells[0], no,   size_pt=9)
    set_cell(new_wl.rows[r].cells[1], name, size_pt=9)
    for c, m in enumerate(marks):
        set_cell(new_wl.rows[r].cells[2 + c], m, size_pt=9)
# Legend row
legend_row = new_wl.rows[-1]
legend_text = 'X = Chapter lead.  Each member leads chapters totalling ~25% of report weight.'
set_cell(legend_row.cells[0], legend_text, size_pt=8)
legend_row.cells[0].merge(legend_row.cells[-1])

replace_table_in_body(doc.tables[2], new_wl)

# ─────────────────────────────────────────────────────────────
# 5. ADD FILE TREE UNDER SECTION 4.2 (after "Figure 10" caption)
# ─────────────────────────────────────────────────────────────

FILE_TREE_LINES = [
    'B2B/app/',
    '\u251c\u2500\u2500 messages/',
    '\u2502   \u251c\u2500\u2500 en.json                         (English i18n strings)',
    '\u2502   \u2514\u2500\u2500 ar.json                         (Arabic i18n strings)',
    '\u251c\u2500\u2500 public/                             (Static assets)',
    '\u251c\u2500\u2500 src/',
    '\u2502   \u251c\u2500\u2500 app/                            (Next.js 16 App Router pages)',
    '\u2502   \u2502   \u251c\u2500\u2500 admin/',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 brands/page.tsx',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 dashboard/page.tsx',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 orders/page.tsx',
    '\u2502   \u2502   \u2502   \u2514\u2500\u2500 users/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 api/',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 admin/brands/[id]/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 admin/brands/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 admin/orders/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 admin/reviews/[id]/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 admin/users/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 auth/login/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 auth/logout/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 auth/me/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 auth/register/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 brands/[id]/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 brands/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 brands/whatsapp-click/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 exchange-rates/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 messages/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 notifications/[id]/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 notifications/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 orders/[id]/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 orders/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 products/[id]/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 products/bulk/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 products/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 reviews/[id]/reply/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 reviews/[id]/route.ts',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 reviews/route.ts',
    '\u2502   \u2502   \u2502   \u2514\u2500\u2500 trust/[id]/route.ts',
    '\u2502   \u2502   \u251c\u2500\u2500 brand/',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 dashboard/page.tsx',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 orders/[id]/page.tsx',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 orders/page.tsx',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 products/[id]/edit/page.tsx',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 products/new/page.tsx',
    '\u2502   \u2502   \u2502   \u251c\u2500\u2500 products/page.tsx',
    '\u2502   \u2502   \u2502   \u2514\u2500\u2500 profile/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 brands/[id]/page.tsx      (Public brand detail)',
    '\u2502   \u2502   \u251c\u2500\u2500 dashboard/page.tsx         (Buyer dashboard)',
    '\u2502   \u2502   \u251c\u2500\u2500 login/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 marketplace/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 notifications/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 orders/[id]/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 orders/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 pending/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 register/page.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 globals.css',
    '\u2502   \u2502   \u251c\u2500\u2500 layout.tsx                 (Root layout, metadata, JSON-LD)',
    '\u2502   \u2502   \u251c\u2500\u2500 page.tsx                   (Landing page)',
    '\u2502   \u2502   \u2514\u2500\u2500 sitemap.ts',
    '\u2502   \u251c\u2500\u2500 components/                 (Reusable UI components)',
    '\u2502   \u2502   \u251c\u2500\u2500 providers/RealtimeProvider.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 BottomNav.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 BrandAvatar.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 BrandCard.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 BrandSidebar.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 BuyerTrustCard.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 CatalogImportModal.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 CurrencySelector.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 ImageUrlInput.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 LandingPage.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 LanguageToggle.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 MarketModal.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 MobileTopBar.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 Navbar.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 NotificationBell.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 Providers.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 RatingBreakdown.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 ReviewCard.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 ReviewModal.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 StarRating.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 StatusBadge.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 TrustScore.tsx',
    '\u2502   \u2502   \u251c\u2500\u2500 VerifiedBadge.tsx',
    '\u2502   \u2502   \u2514\u2500\u2500 WhatsAppButton.tsx',
    '\u2502   \u251c\u2500\u2500 contexts/',
    '\u2502   \u2502   \u2514\u2500\u2500 CurrencyContext.tsx        (Multi-currency React context)',
    '\u2502   \u251c\u2500\u2500 data/',
    '\u2502   \u2502   \u2514\u2500\u2500 db.json                    (Seed data \u2014 MVP JSON store)',
    '\u2502   \u251c\u2500\u2500 hooks/',
    '\u2502   \u2502   \u251c\u2500\u2500 useInView.ts',
    '\u2502   \u2502   \u251c\u2500\u2500 useRealtimeBrandOrders.ts',
    '\u2502   \u2502   \u251c\u2500\u2500 useRealtimeNotifications.ts',
    '\u2502   \u2502   \u2514\u2500\u2500 useRealtimeOrder.ts',
    '\u2502   \u251c\u2500\u2500 i18n/request.ts',
    '\u2502   \u251c\u2500\u2500 lib/                        (Core business logic)',
    '\u2502   \u2502   \u251c\u2500\u2500 supabase/client.ts',
    '\u2502   \u2502   \u251c\u2500\u2500 auth.ts                    (requireRole, getCurrentUser)',
    '\u2502   \u2502   \u251c\u2500\u2500 currencies.ts              (FX rates, VAT, formatPriceWithVAT)',
    '\u2502   \u2502   \u251c\u2500\u2500 db.ts                      (readDB / writeDB \u2014 JSON layer)',
    '\u2502   \u2502   \u251c\u2500\u2500 formatters.ts',
    '\u2502   \u2502   \u251c\u2500\u2500 i18n.ts',
    '\u2502   \u2502   \u251c\u2500\u2500 notifications.ts           (createNotification \u2014 non-fatal)',
    '\u2502   \u2502   \u2514\u2500\u2500 pricingUtils.ts            (Bulk pricing tier resolution)',
    '\u2502   \u251c\u2500\u2500 store/',
    '\u2502   \u2502   \u2514\u2500\u2500 notificationStore.ts       (Zustand unread-count store)',
    '\u2502   \u251c\u2500\u2500 utils/supabase/',
    '\u2502   \u2502   \u251c\u2500\u2500 client.ts',
    '\u2502   \u2502   \u251c\u2500\u2500 middleware.ts',
    '\u2502   \u2502   \u2514\u2500\u2500 server.ts',
    '\u2502   \u251c\u2500\u2500 i18n.ts',
    '\u2502   \u251c\u2500\u2500 middleware.ts               (SSR session refresh + route guards)',
    '\u2502   \u2514\u2500\u2500 navigation.ts',
    '\u251c\u2500\u2500 next.config.ts',
    '\u251c\u2500\u2500 package.json',
    '\u2514\u2500\u2500 tsconfig.json',
]

# Find "Figure 10" caption paragraph
fig10_p = None
for para in doc.paragraphs:
    if para.text.strip().startswith('Figure 10'):
        fig10_p = para
        break

if fig10_p:
    anchor = fig10_p._p
    # Insert lines in reverse so they appear in correct order after anchor
    for line in reversed(FILE_TREE_LINES):
        anchor.addnext(make_mono_p_elem(line, size_pt=7))
    print(f"File tree inserted after: {fig10_p.text[:60]}")
else:
    print("WARNING: Figure 10 paragraph not found — file tree not inserted.")

# ─────────────────────────────────────────────────────────────
# 6. REPLACE ACTIVITY NETWORK TABLE (Table 17) with Predecessor+Successor
# ─────────────────────────────────────────────────────────────
# After the workload table was replaced, table indices shift.
# The Activity Network was originally Table 17. We need to find it by content.

AN_HEADERS = ['Activity', 'Predecessor', 'Successor', 'Duration',
              'Early Start', 'Early Finish', 'Float']

AN_DATA = [
    ('S1 \u2014 Auth & Setup',             '\u2014',  'S2',        '2 wks',   'Jan 1',  'Jan 14', '0 (Critical)'),
    ('S2 \u2014 Marketplace',              'S1',     'S3',        '2 wks',   'Jan 15', 'Jan 31', '0 (Critical)'),
    ('S3 \u2014 Order Lifecycle',          'S2',     'S4',        '3 wks',   'Feb 1',  'Feb 20', '0 (Critical)'),
    ('S4 \u2014 Brand Owner Tools',        'S3',     'S5',        '2.5 wks', 'Feb 21', 'Mar 9',  '0 (Critical)'),
    ('S5 \u2014 Trust & Reputation',       'S4',     'S6',        '2 wks',   'Mar 10', 'Mar 24', '0 (Critical)'),
    ('S6 \u2014 Admin & Notifications',    'S5',     'S7',        '2 wks',   'Mar 25', 'Apr 9',  '0 (Critical)'),
    ('S7 \u2014 Localisation & Mobile',    'S6',     'S8',        '3 wks',   'Apr 10', 'Apr 30', '0 (Critical)'),
    ('S8 \u2014 Migration, Testing & Report', 'S7',  '\u2014 (End)', '6 wks', 'May 1', 'Jun 15', '0 (Critical)'),
]

# Find the activity network table by looking for the old header content
an_old_table = None
for tbl in doc.tables:
    if tbl.rows and tbl.rows[0].cells:
        header_text = tbl.rows[0].cells[0].text.strip()
        if 'Activity' in header_text and len(tbl.rows) == 9:
            an_old_table = tbl
            break
    # Also check if first row has "S1" or "Auth"
    if tbl.rows and len(tbl.rows) > 1:
        row1_text = tbl.rows[1].cells[0].text.strip() if tbl.rows[1].cells else ''
        if 'Auth' in row1_text and 'S1' in row1_text:
            an_old_table = tbl
            break

if an_old_table is None:
    # Fall back: try known index (was 17, but table 2 was replaced; overall count unchanged)
    # Workload table replaced doesn't change count; so Activity Network is still index 17
    if len(doc.tables) > 17:
        an_old_table = doc.tables[17]
        print("Using fallback table index 17 for Activity Network.")

if an_old_table:
    new_an = doc.add_table(rows=1 + len(AN_DATA) + 1, cols=7)
    new_an.style = 'Table Grid'
    for i, h in enumerate(AN_HEADERS):
        set_cell(new_an.rows[0].cells[i], h, bold=True, size_pt=9)
    for r, row_data in enumerate(AN_DATA, start=1):
        for c, val in enumerate(row_data):
            set_cell(new_an.rows[r].cells[c], val, size_pt=9)
    # Note row
    note_row = new_an.rows[-1]
    set_cell(note_row.cells[0],
             'Note: All 8 sprints lie on the critical path (zero float). '
             'Each sprint output is a prerequisite for the next.', size_pt=8)
    note_row.cells[0].merge(note_row.cells[-1])
    replace_table_in_body(an_old_table, new_an)
    print("Activity Network table replaced.")
else:
    print("WARNING: Activity Network table not found.")

# ─────────────────────────────────────────────────────────────
# 7. ENHANCE APPENDIX A NOTE
# ─────────────────────────────────────────────────────────────

APP_A_NOTE = (
    '\u26a0\ufe0f  IMPORTANT \u2014 MANUAL DIAGRAM INSERTION REQUIRED\n\n'
    'The following UML and project diagrams MUST be created in draw.io or Lucidchart '
    'and inserted manually as high-resolution PNG or SVG images into this document '
    'at the figure placeholders indicated in each chapter:\n\n'
    '  \u2022  Figure 8   \u2014 Class Diagram (10 entity classes with relationships)\n'
    '  \u2022  Figure 5   \u2014 Sequence Diagram: User Registration & Login\n'
    '  \u2022  Figure 6   \u2014 Sequence Diagram: Order Lifecycle\n'
    '  \u2022  Figure 7   \u2014 Sequence Diagram: Brand Verification\n'
    '  \u2022  Figure 2   \u2014 Context Model / Level-0 DFD\n'
    '  \u2022  Figure 3   \u2014 Level-1 Data Flow Diagram (DFD)\n'
    '  \u2022  Figure 12  \u2014 Gantt Chart (8 sprints, Jan\u2013Jun 2026)\n'
    '  \u2022  Figure 9   \u2014 Activity Diagram: Order Placement workflow\n'
    '  \u2022  Figure 11  \u2014 Physical Deployment Diagram (Vercel + Supabase)\n'
    '  \u2022  Figure 13  \u2014 Activity Network Diagram: Sprint dependency graph\n'
    '  \u2022  Figure 14  \u2014 Fishbone / Ishikawa Diagram: Delayed Submission Risk\n'
    '  \u2022  Figure 15  \u2014 Quality Prediction Chart\n\n'
    'Steps: (1) Open draw.io or Lucidchart. '
    '(2) Draw each diagram using specifications in the table below. '
    '(3) Export as PNG at 150 dpi minimum or SVG. '
    '(4) In Microsoft Word, place cursor at the figure placeholder, '
    'then Insert \u2192 Pictures \u2192 select the exported image. '
    '(5) Set image width to 14 cm and align centre.'
)

for para in doc.paragraphs:
    if 'NOTE TO TEAM' in para.text and 'draw.io' in para.text:
        clear_para_and_set(para, APP_A_NOTE)
        print("Appendix A note updated.")
        break

# ─────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f"\nSaved: {OUTPUT}")
