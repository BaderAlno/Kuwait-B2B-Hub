#!/usr/bin/env python3
"""Phase 4: Chapters VIII, IX, X + front matter (Cover, ToC, LoT, LoF).
   Reads Final_Report_v3.docx, appends new chapters, prepends front matter,
   saves as Final_Report_FINAL.docx.
"""

import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT_FILE  = "/Users/baderalnoumas/Desktop/B2B/Final_Report_Build/Final_Report_v3.docx"
OUTPUT_FILE = "/Users/baderalnoumas/Desktop/B2B/Final_Report_FINAL.docx"

# ── helpers ────────────────────────────────────────────────────────────────────

def page_break(doc):
    doc.add_page_break()

def h1(doc, t):
    doc.add_heading(t, level=1)

def h2(doc, t):
    doc.add_heading(t, level=2)

def h3(doc, t):
    doc.add_heading(t, level=3)

def body(doc, t):
    doc.add_paragraph(t, style="Normal")

def bullet(doc, t):
    doc.add_paragraph(t, style="List Bullet")

def make_table(doc, headers, rows, col_widths=None):
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return table

def placeholder(doc, fig_id, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[FIGURE {fig_id} — {caption}]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("(Insert diagram here)")
    r2.italic = True
    r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    doc.add_paragraph()

def mono(doc, text, fs=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(fs)

def centre(doc, text, bold=False, size=None, space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def toc_line(doc, chapter, title, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent * 0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{chapter}  {title}")
    run.font.size = Pt(10)
    if indent == 0:
        run.bold = True

# ── CHAPTER VIII ───────────────────────────────────────────────────────────────

def build_chapter_8(doc):
    h1(doc, "CHAPTER VIII — CONCLUSIONS")
    body(doc, (
        "This chapter synthesises the key lessons learned during the development of the Kuwait "
        "B2B Hub through three analytical lenses: Robert Martin’s seven design smells "
        "as applied to the codebase, a SWOT analysis of the platform’s strategic position, "
        "and seven concrete recommendations for future work."
    ))

    # §8.1
    h2(doc, "8.1  Seven Design Smells")
    body(doc, (
        "Robert C. Martin’s seven design smells (Agile Software Development, 2002) provide a "
        "principled vocabulary for diagnosing structural weaknesses in object-oriented systems. "
        "Each smell is examined below against evidence drawn directly from the Kuwait B2B Hub "
        "codebase, together with the resolution taken or planned."
    ))

    # Smell 1
    h3(doc, "Smell 1 — Rigidity")
    body(doc, (
        "Problem Statement: Rigidity describes a design so tightly interwoven that a change to one "
        "module forces a cascade of modifications throughout the system. In the Kuwait B2B Hub, "
        "the most prominent instance of rigidity manifests in the readDB() / writeDB() data access "
        "pattern shared across all 25 API route handlers. Every handler in src/app/api/ imports "
        "these two functions from src/lib/db.ts and calls them directly, creating a fan-out "
        "dependency that spans the entire backend. Any structural change to the Database "
        "interface — for example, renaming the orders collection or adding a mandatory "
        "non-nullable field to the User interface — requires simultaneous edits across every "
        "route file."
    ))
    body(doc, (
        "Concrete Example: The Product interface in db.ts was extended with the "
        "bulk_pricing_tiers: BulkPricingTier[] field in CR-006. This required modifications "
        "to the POST /api/orders handler (unit price resolution logic), the POST /api/products "
        "handler (insert path), the PATCH /api/products/[id] handler (update path), and the "
        "UI form in the product creation page — four separate files changed for one data "
        "model extension."
    ))
    body(doc, (
        "Resolution: The team mitigated rigidity by encapsulating all storage operations behind "
        "the three utility functions (readDB(), writeDB(), generateId()), ensuring that the "
        "underlying storage implementation can be replaced without altering route-level business "
        "logic. The PostgreSQL migration planned for Year 1 (SC-01) will complete this abstraction "
        "by introducing a repository layer that isolates SQL query syntax from the application "
        "domain entirely, reducing the blast radius of future schema changes to a single module."
    ))

    # Smell 2
    h3(doc, "Smell 2 — Fragility")
    body(doc, (
        "Problem Statement: A fragile system breaks in unexpected places when a change is made "
        "elsewhere. Fragility typically indicates hidden coupling between modules that share state "
        "through implicit channels rather than explicit interfaces. The Kuwait B2B Hub exhibits "
        "fragility at the boundary between the authentication system and the real-time notification "
        "infrastructure."
    ))
    body(doc, (
        "Concrete Example: The RealtimeProvider.tsx component subscribes to Supabase Realtime "
        "channels using the user session established by getCurrentUser(). If getCurrentUser() "
        "returns null during session expiry or a brief Supabase Auth service interruption, the "
        "subscription attempt silently fails and the client receives no error feedback. The "
        "notification bell becomes unresponsive without any visible error state, and the failure "
        "manifests several component layers removed from the authentication guard that caused it. "
        "This is a textbook fragility: a change to the authentication system (session expiry "
        "behaviour) breaks the notification system without any compile-time or runtime warning."
    ))
    body(doc, (
        "Resolution: The team addressed this fragility by adding explicit null guards in "
        "useRealtimeNotifications.ts and configuring the Zustand notification store to degrade "
        "gracefully to an empty unread count rather than propagating an uncaught exception. "
        "Long-term, the introduction of a formal circuit-breaker pattern around all Supabase "
        "Realtime subscriptions, combined with visible UI indicators for connection status, "
        "will eliminate the silent failure mode."
    ))

    # Smell 3
    h3(doc, "Smell 3 — Immobility")
    body(doc, (
        "Problem Statement: A design is immobile when useful components cannot be extracted "
        "and reused in other contexts because they are too entangled with the surrounding code. "
        "Immobility is particularly costly when it affects high-value domain logic that would "
        "benefit multiple consumers."
    ))
    body(doc, (
        "Concrete Example: The Trust Engine — the subsystem that calculates response_rate, "
        "completion_rate, avg_fulfillment_days, and assigns GCC-calibrated badges for BrandTrust "
        "and BuyerTrust — represents the platform’s most distinctive domain logic and "
        "core competitive differentiation. However, the scoring calculations were developed inline "
        "with the admin dashboard pages rather than isolated in a dedicated service module. This "
        "coupling prevents the Trust Engine from being independently unit-tested, reused in a "
        "future mobile API, or exposed as a standalone analytics endpoint without first "
        "disentangling it from the presentation layer."
    ))
    body(doc, (
        "Resolution: The contrast with src/lib/currencies.ts is instructive: the currency "
        "configuration module is a pure, context-free data module that is successfully consumed "
        "by the CurrencySelector component, the exchange-rate API handler, and all price "
        "formatting utilities with zero coupling. The planned resolution is to extract the Trust "
        "Engine scoring functions into a dedicated src/lib/trustEngine.ts service module, "
        "exposing pure functions with typed inputs and outputs that can be tested in isolation "
        "and consumed by any future consumer without modification."
    ))

    # Smell 4
    h3(doc, "Smell 4 — Viscosity")
    body(doc, (
        "Problem Statement: Viscosity occurs when preserving the designed architectural intent "
        "requires more effort than taking a shortcut, tempting developers to introduce hacks "
        "that gradually degrade the design. Viscosity has two forms: design viscosity (the "
        "right approach is harder than the wrong one) and environment viscosity (the development "
        "environment is slow and discourages discipline)."
    ))
    body(doc, (
        "Concrete Example: During CR-003 (Arabic RTL support, version 3.2.0.0 to 3.3.0.0), the "
        "team encountered severe design viscosity in the internationalisation layer. Adding a "
        "localised string to a component the correct way — by adding a key to both "
        "messages/en.json and messages/ar.json and using the useTranslations() hook with the "
        "next-intl API — required navigating and editing four separate files per string "
        "addition. The easier path was to hard-code English strings directly in the component JSX, "
        "which earlier sprints had normalised. This viscosity caused an estimated 40% of string "
        "literals across 24 components to bypass the translation system, necessitating the "
        "large-scale migration effort in CR-003."
    ))
    body(doc, (
        "Resolution: CR-003 addressed the immediate debt by completing the full migration to "
        "next-intl across all 24 affected components. Three Python audit scripts "
        "(find_missing_keys.py, audit_translations.py, find_all_missing_keys.py) were introduced "
        "to detect translation drift during future development, reducing the environment friction "
        "of maintaining bilingual parity. A pre-commit linting rule prohibiting bare string "
        "literals in JSX is recommended as the permanent viscosity barrier."
    ))

    # Smell 5
    h3(doc, "Smell 5 — Needless Complexity")
    body(doc, (
        "Problem Statement: Needless complexity arises when infrastructure or abstractions are "
        "introduced speculatively before there is a concrete requirement, increasing cognitive "
        "load without delivering immediate value. The canonical cause is over-engineering driven "
        "by anticipation of requirements that may never materialise."
    ))
    body(doc, (
        "Concrete Example: The CURRENCIES registry in src/lib/currencies.ts defines six GCC "
        "currencies — KWD, SAR, AED, QAR, BHD, and OMR — each with a full "
        "CurrencyConfig interface encompassing vatRate, ISO 4217 code, decimal precision, "
        "locale string, phone prefix, date format, and emoji flag. At the time of submission, "
        "the prototype processes all orders exclusively in Kuwaiti Dinar. Live exchange rate "
        "integration, multi-currency invoicing, and GCC market routing are all planned "
        "enhancements not present in the submitted build."
    ))
    body(doc, (
        "Resolution: The team’s position — supported by the GCC expansion roadmap in "
        "Section 6.3 (SC-07) — is that this represents intentional forward-looking "
        "architecture rather than speculative complexity. The CURRENCIES registry is purely "
        "declarative with zero runtime cost when GCC currencies are inactive. Building "
        "multi-currency configuration into the data layer from the outset avoids the far larger "
        "refactoring effort that would be required at expansion time. This design decision is "
        "classified as a managed architectural trade-off: the complexity is real but the "
        "justification is documented and the future activation path is clear."
    ))

    # Smell 6
    h3(doc, "Smell 6 — Needless Repetition")
    body(doc, (
        "Problem Statement: Needless repetition occurs when structurally identical code appears "
        "in multiple locations because the team has not applied a suitable abstraction. Unlike "
        "simple copy-paste duplication, the most insidious form of this smell involves patterns "
        "that are conceptually identical but expressed slightly differently in each location, "
        "making them invisible to naive text search."
    ))
    body(doc, (
        "Concrete Example: The RBAC authentication guard sequence opens every API route handler "
        "in the codebase. The pattern — await getCurrentUser() followed by a null check "
        "returning HTTP 401 Unauthorized — appears verbatim in all 25 route files. Role "
        "checks such as verifying user.role === 'admin' or user.role === 'buyer' appear in "
        "further variants throughout the same files. While the getCurrentUser() call itself is "
        "a correct application of the DRY principle, the null-check guard construction and the "
        "role-assertion pattern are duplicated across all 25 handlers, representing an "
        "abstraction opportunity that was not taken."
    ))
    body(doc, (
        "Resolution: A higher-order function wrapper — for example, withAuth(handler, "
        "requiredRole) that returns a wrapped handler pre-armed with the guard logic — or "
        "a Next.js Middleware extension would eliminate this repetition while making the access "
        "control policy more declarative and centrally enforceable. A single change to the "
        "wrapper function would propagate to all 25 routes simultaneously, dramatically reducing "
        "the surface area for security-relevant modifications. This refactoring is targeted for "
        "the Year 1 API layer hardening sprint."
    ))

    # Smell 7
    h3(doc, "Smell 7 — Opacity")
    body(doc, (
        "Problem Statement: Opacity describes code or design structures whose intent and "
        "constraints are unclear to a reader who was not involved in their original creation. "
        "Opacity is particularly dangerous when it conceals non-obvious invariants or "
        "preconditions that, if violated, produce subtle and hard-to-reproduce failures."
    ))
    body(doc, (
        "Concrete Example: The JSON prototype database architecture is the most opaque construct "
        "in the codebase. The readDB() → db.entity.push(item) → writeDB(db) cycle in "
        "every API route handler does not self-document the fact that it provides no ACID "
        "guarantees, no concurrent write protection, and no rollback capability. A developer "
        "onboarding to the project and adding a new feature is unlikely to recognise that two "
        "simultaneous calls to readDB() will read the same snapshot, and that the second "
        "writeDB() will silently overwrite the first — the root cause of R-01, the "
        "highest-exposure risk in the Risk Assessment Matrix (Exposure Score: 25)."
    ))
    body(doc, (
        "Resolution: Two actions address this opacity. First, the TypeScript interfaces in "
        "src/lib/db.ts make all entity shapes fully explicit and statically checked, ensuring "
        "data structure intent is always visible. Second, a prominent warning comment added "
        "directly to the readDB() and writeDB() function signatures — documenting the "
        "concurrency limitation, the overwite risk, and the planned PostgreSQL migration path "
        "— would serve as a zero-cost safety net for any developer who reads the file. "
        "Long-term, the PostgreSQL migration (SC-01) eliminates the opacity by replacing the "
        "implicit concurrency model with ACID-compliant transactions whose guarantees are "
        "formally specified and universally understood."
    ))

    # §8.2 SWOT
    h2(doc, "8.2  SWOT Analysis")
    body(doc, (
        "The following SWOT analysis evaluates the Kuwait B2B Hub’s strategic position "
        "as a prototype-stage platform at the time of final submission. Each quadrant contains "
        "four items grounded in evidence from the technical and domain analysis conducted "
        "throughout this report."
    ))
    make_table(doc,
        ["Quadrant", "Item", "Evidence / Justification"],
        [
            ["STRENGTH", "TypeScript strict-mode codebase (93 files)",
             "Compile-time type safety eliminates entire classes of runtime errors. All 10 data entity interfaces are formally typed in db.ts, enforced across API routes, components, and hooks."],
            ["STRENGTH", "Bilingual Arabic / English with full RTL layout",
             "next-intl v4.9 with complete messages/ar.json and messages/en.json. RTL layout switching is server-rendered, ensuring first-load correctness for Arabic users."],
            ["STRENGTH", "Differentiated Trust Engine unique to Kuwait/GCC",
             "BrandTrust and BuyerTrust scoring (response_rate, completion_rate, avg_fulfillment_days, GCC-calibrated badges) is custom-built domain IP with no direct equivalent in existing GCC B2B platforms."],
            ["STRENGTH", "Clean RBAC model enforced at every API boundary",
             "Three-tier role model (admin, brand_owner, buyer) enforced via getCurrentUser() guard in all 25 API route handlers. Role-specific data filtering prevents cross-tenant data leakage."],
            ["WEAKNESS", "JSON flat-file database with no transaction isolation",
             "R-01: Concurrent writes can corrupt stock counts and financial records. readDB()/writeDB() cycle is not atomic. Production deployment is not viable without the PostgreSQL migration (SC-01)."],
            ["WEAKNESS", "Zero automated test coverage across 11.35 KLOC",
             "R-03: No test runner, no test files, no coverage gate. Every deployment is a manual regression risk. The Trust Engine and RBAC guards have never been tested programmatically."],
            ["WEAKNESS", "Plain-text password storage in prototype db.json",
             "R-02: Passwords stored without cryptographic hashing. Any read access to db.json exposes all user credentials. Must be resolved before any real user credentials are collected."],
            ["WEAKNESS", "No in-platform payment processing",
             "Orders are confirmed bilaterally via WhatsApp. The platform generates zero financial transactions and cannot produce accounting-compliant records. Commercial viability depends on SC-03."],
            ["OPPORTUNITY", "GCC regional expansion to 5 additional markets",
             "SC-07: currencies.ts already defines KWD, SAR, AED, QAR, BHD, OMR with full VAT rates and locale strings. The latent multi-currency architecture requires activation, not rebuilding."],
            ["OPPORTUNITY", "KNET payment integration to unlock in-platform GMV",
             "SC-03: Kuwait’s national payment network, supplemented by PayTabs/Stripe for GCC cross-border. Converts the platform from a discovery tool to a revenue-generating commerce platform."],
            ["OPPORTUNITY", "Mobile PWA conversion targeting 78% mobile-first users",
             "SC-04: CITRA Kuwait (2024) reports 78% of internet access via mobile. BottomNav.tsx and MobileTopBar.tsx indicate mobile intent; push notification delivery and offline resilience complete the transition."],
            ["OPPORTUNITY", "AI-powered product recommendations to increase AOV",
             "SC-05: Collaborative filtering or AWS Personalize integration. B2B buyers in GCC cite supplier product discovery as the primary purchasing friction. Recommendations directly increase average order value."],
            ["THREAT", "Established GCC competitors with larger user bases",
             "Tradeling (UAE, est. 2019) reported ~$200M GMV by year 4. Sary (KSA) has secured $75M+ in funding. Both have multi-year head starts and established brand relationships in the region."],
            ["THREAT", "Kuwait Personal Data Protection Law (PDPL) compliance",
             "R-07: Law No. 20 of 2014 mandates consent, data retention controls, and breach notification. The platform currently collects PII (name, email, company, WhatsApp) without a privacy policy or consent mechanism."],
            ["THREAT", "Supabase and Vercel vendor lock-in (R-08)",
             "100% of authentication and production hosting are delegated to two external providers. Pricing changes, service discontinuation, or sustained outages translate directly to system unavailability."],
            ["THREAT", "KNET payment compliance requirements",
             "R-11: PCI-DSS certification, a certified integration partner, and Central Bank of Kuwait approval are legally required before processing live payments. Non-compliance constitutes a criminal offence under Kuwait Commercial Law."],
        ],
        col_widths=[0.9, 2.0, 3.6]
    )

    # §8.3 Future Work
    h2(doc, "8.3  Future Work and Recommendations")
    body(doc, (
        "The following seven future work items are recommended in priority order. Each item "
        "addresses a documented gap in the current system and is grounded in the risk register, "
        "system evolution analysis, and codebase evidence presented throughout this report."
    ))

    h3(doc, "FW-01 — PostgreSQL Full Migration (SC-01, Critical)")
    body(doc, (
        "The replacement of src/data/db.json with Supabase-managed PostgreSQL is the single "
        "highest-priority action following submission. The current JSON store provides no "
        "transaction isolation, no concurrent write protection, and no rollback capability. "
        "Under concurrent load — which Vercel’s serverless execution model makes "
        "structurally inevitable — stock overselling is mathematically certain (R-01, "
        "Exposure: 25). The migration path is already partially scaffolded: the Supabase client "
        "factory in src/utils/supabase/server.ts is production-ready, and the ten TypeScript "
        "interfaces in src/lib/db.ts translate directly to PostgreSQL table schemas. A phased "
        "migration with write-shadowing is recommended: run both systems in parallel for two "
        "weeks, verifying consistency before cutting over with a tested rollback path."
    ))

    h3(doc, "FW-02 — Automated Test Suite with Vitest (SC-02, Critical)")
    body(doc, (
        "Introducing Vitest with React Testing Library is the second critical action, and its "
        "cost grows with every month it is deferred. Lehman’s Second Law (Increasing "
        "Complexity) predicts that the 11.35 KLOC codebase will continue to grow, and each "
        "new feature added without corresponding tests increases the regression surface area. "
        "The highest-priority test targets are getCurrentUser() (called by all 25 routes), "
        "the POST /api/orders bulk pricing resolution logic (financial correctness), and the "
        "Trust Engine scoring functions (core platform differentiator). A CI pipeline gate "
        "of ≥60% line coverage should block merges below this threshold to prevent "
        "regression accumulation from the outset."
    ))

    h3(doc, "FW-03 — KNET Payment Gateway Integration (SC-03, Critical)")
    body(doc, (
        "In-platform payment processing is a commercial imperative rather than a feature "
        "enhancement. A B2B marketplace that routes payment confirmation to WhatsApp is a "
        "discovery tool, not a commerce platform — it generates zero revenue, cannot "
        "produce accounting-compliant transaction records, and cannot enforce the financial "
        "accountability that buyers and brand owners require for high-value wholesale "
        "transactions. Integration should proceed through a KNET-certified gateway aggregator "
        "(MyFatoorah or Tap Payments) that provides a pre-certified integration layer, "
        "ensuring the platform never touches raw card data. The order status FSM should be "
        "extended with payment_pending and payment_confirmed states before implementation "
        "to maintain the formal state machine specification documented in Section 3.5."
    ))

    h3(doc, "FW-04 — Mobile-First Progressive Web App (SC-04, High)")
    body(doc, (
        "CITRA Kuwait (2024) reports that over 78% of Kuwaiti internet users access services "
        "via mobile devices. The existing BottomNav.tsx and MobileTopBar.tsx components "
        "demonstrate that mobile was an explicit design intent from the outset, but the current "
        "build lacks the three capabilities that define a production-grade mobile experience: "
        "offline resilience via a service worker cache, native push notification delivery "
        "(bypassing the in-app polling model), and an installable app manifest. The PWA "
        "conversion is achievable within a single sprint using the next-pwa library and Vercel’s "
        "built-in service worker support, with no architectural changes to the existing "
        "Next.js 16 App Router structure."
    ))

    h3(doc, "FW-05 — Arabic NLP Full-Text Search (SC-06, High)")
    body(doc, (
        "Product data in the current marketplace is predominantly entered in English by brand "
        "owners, creating a searchability wall for Arabic-first buyers who cannot locate "
        "products using Arabic keywords. The existing filter-based search provides no "
        "linguistic bridging. Replacing it with Arabic-aware full-text search — using "
        "PostgreSQL tsvector with an Arabic stemming dictionary, or Algolia with Arabic locale "
        "configuration — removes this barrier and directly expands the addressable buyer "
        "base. This recommendation is a prerequisite for the GCC regional expansion (FW-06) "
        "because Arabic-first markets (Saudi Arabia, Qatar, Bahrain) will have even lower "
        "tolerance for English-only search than the bilingual Kuwaiti market."
    ))

    h3(doc, "FW-06 — GCC Regional Expansion (SC-07, High)")
    body(doc, (
        "The architectural groundwork for GCC expansion is already present in the codebase. "
        "The CURRENCIES registry in src/lib/currencies.ts fully specifies KWD, SAR, AED, "
        "QAR, BHD, and OMR with their respective VAT rates, ISO 4217 codes, locale strings, "
        "and phone prefixes. Activating this latent capability requires implementing "
        "country-level VAT calculation in the order flow, currency-native invoice generation, "
        "localised Know Your Customer (KYC) verification workflows aligned with each "
        "jurisdiction’s commercial registration requirements, and Saudi ZATCA e-invoicing "
        "compliance for the KSA market. This expansion directly addresses Lehman’s Sixth "
        "Law (Continuing Growth): without geographic expansion, the platform’s addressable "
        "market remains capped at Kuwait’s relatively small B2B ecosystem."
    ))

    h3(doc, "FW-07 — AI-Powered Product Recommendation Engine (SC-05, High)")
    body(doc, (
        "B2B buyers in GCC markets consistently identify supplier product discovery as the "
        "primary friction in the wholesale purchasing workflow. A recommendation engine "
        "— implementing collaborative filtering based on order history and category "
        "affinity, or integrating AWS Personalize for managed model training — directly "
        "addresses this friction by surfacing relevant products that buyers would not "
        "otherwise encounter. The expected business outcome is a measurable increase in "
        "average order value and repeat purchase frequency, both of which are tracked "
        "by the existing BrandTrust and BuyerTrust metrics and therefore immediately "
        "visible in the platform’s own analytics. The recommendation engine should "
        "be implemented after the PostgreSQL migration (FW-01) to leverage proper relational "
        "query capabilities for co-purchase analysis."
    ))

# ── CHAPTER IX ─────────────────────────────────────────────────────────────────

def build_chapter_9(doc):
    page_break(doc)
    h1(doc, "CHAPTER IX — APPENDICES")

    # Appendix A
    h2(doc, "Appendix A — System Diagram Placeholders")
    body(doc, (
        "The following ten figures are referenced throughout this report. Each diagram is to "
        "be inserted at the location indicated. All diagrams were produced during the design "
        "and implementation phases; final rendered versions are to be attached as image files "
        "in the submitted Word document."
    ))
    make_table(doc,
        ["Figure No.", "Title", "Referenced In", "Status"],
        [
            ["Figure 1", "Sequence Diagram — User Registration & Login",
             "§3.1.1", "[INSERT IMAGE]"],
            ["Figure 2", "Sequence Diagram — Order Placement & Review Lifecycle",
             "§3.1.2", "[INSERT IMAGE]"],
            ["Figure 3", "Sequence Diagram — Brand Verification",
             "§3.1.3", "[INSERT IMAGE]"],
            ["Figure 4", "DFD Level 0 — Context Diagram",
             "§3.2.1", "[INSERT IMAGE]"],
            ["Figure 5", "DFD Level 1 — Sub-Process Decomposition",
             "§3.2.2", "[INSERT IMAGE]"],
            ["Figure 6", "Activity Diagram — Order Placement",
             "§3.4", "[INSERT IMAGE]"],
            ["Figure 7", "Finite State Machine — Order Lifecycle",
             "§3.5.1", "[INSERT IMAGE]"],
            ["Figure 8", "Finite State Machine — Brand Verification",
             "§3.5.2", "[INSERT IMAGE]"],
            ["Figure 9", "Class Diagram — Core Entity Model (10 Entities)",
             "§3.7", "[INSERT IMAGE]"],
            ["Figure 10", "Physical Deployment Diagram — Vercel / Supabase Architecture",
             "§4.6", "[INSERT IMAGE]"],
        ],
        col_widths=[0.8, 2.6, 1.0, 2.1]
    )

    body(doc, "Diagram source files and rendering notes are detailed below.")
    bullet(doc, "Figures 1–3 (Sequence Diagrams): Mermaid sequenceDiagram source "
                "available in Final_Report_Additions.md §Figures 5–7. Render via "
                "mermaid.live or VS Code Mermaid Preview extension.")
    bullet(doc, "Figure 4–5 (DFD): Produce using draw.io or Lucidchart. Context "
                "diagram shows the Kuwait B2B Hub system boundary with three external actors "
                "(Admin, Brand Owner, Business Buyer) and two external services "
                "(Supabase, Vercel CDN).")
    bullet(doc, "Figure 6 (Activity Diagram): Mermaid flowchart TD source available in "
                "Final_Report_Additions.md §Figure 9.")
    bullet(doc, "Figures 7–8 (FSMs): Render as state-transition diagrams with "
                "states and labelled transitions as specified in the FSM tables in "
                "§3.5.1 and §3.5.2.")
    bullet(doc, "Figure 9 (Class Diagram): Mermaid classDiagram source available in "
                "Final_Report_Additions.md §Figure 8. Ten entity classes with "
                "full relationship notation.")
    bullet(doc, "Figure 10 (Deployment Diagram): Render as UML deployment diagram showing "
                "Developer Machine → Git Repository → Vercel CI/CD → "
                "Vercel Edge Network → Supabase Cloud.")

    # Appendix B
    h2(doc, "Appendix B — Critical Source File Extracts")
    body(doc, (
        "The following extracts present the import declarations, interface definitions, and "
        "primary function signatures of the four most architecturally significant source files "
        "in the Kuwait B2B Hub codebase. All extracts were taken verbatim from the live "
        "repository with no modifications applied. Generated files, type declaration files "
        "(.d.ts), and node_modules were excluded consistent with the SLOC methodology "
        "described in §5.1.1."
    ))

    # B.1
    h3(doc, "B.1 — src/lib/auth.ts")
    body(doc, (
        "Centralised authentication resolver. Wraps the Supabase session lookup and enriches "
        "the raw auth token with the full application-level User profile from the profiles "
        "table. All protected API routes call getCurrentUser() as their first guard before "
        "executing any business logic. This function is the single entry point for all "
        "authenticated operations and the enforcement point for the RBAC model."
    ))
    mono(doc, """import { createClient } from '@/utils/supabase/server';
import { User } from './db';

export async function getCurrentUser(): Promise<User | null> {
  const supabase = await createClient();
  const { data: { user }, error } = await supabase.auth.getUser();

  if (error || !user) return null;

  const { data: profile } = await supabase
    .from('profiles')
    .select('*')
    .eq('id', user.id)
    .single();

  if (!profile) return null;

  return profile as User;
}""", fs=8)

    # B.2
    h3(doc, "B.2 — src/app/api/orders/route.ts")
    body(doc, (
        "Next.js App Router API route handler for the /api/orders endpoint. Demonstrates the "
        "Role-Based Access Control (RBAC) pattern applied consistently across all protected "
        "endpoints: buyers see only their own orders, brand owners see only orders belonging "
        "to their brand, and administrators receive the unfiltered dataset. Unauthenticated "
        "requests are rejected with an HTTP 401 response before any database access occurs."
    ))
    mono(doc, """import { NextRequest, NextResponse } from 'next/server';
import { readDB, writeDB, generateId } from '@/lib/db';
import { getCurrentUser } from '@/lib/auth';
import { createNotification } from '@/lib/notifications';

// GET /api/orders — role-filtered order list
export async function GET() {
  const user = await getCurrentUser();
  if (!user) return NextResponse.json({ error: 'Unauthorized' }, { status: 401 });

  const db = readDB();
  let orders = db.orders;

  if (user.role === 'buyer') {
    orders = orders.filter(o => o.buyer_id === user.id);
  } else if (user.role === 'brand_owner') {
    const brand = db.brands.find(b => b.owner_id === user.id);
    orders = brand ? orders.filter(o => o.brand_id === brand.id) : [];
  }
  // admin sees all""", fs=8)

    # B.3
    h3(doc, "B.3 — src/lib/currencies.ts")
    body(doc, (
        "Typed currency configuration module. The CurrencyConfig interface enforces a uniform "
        "schema for all supported Gulf Cooperation Council (GCC) currencies, including VAT "
        "rates, ISO 4217 codes, decimal precision, and locale strings. This interface serves "
        "as a single source of truth consumed by the CurrencySelector component, the live "
        "exchange-rate API handler, and all price formatting utilities throughout the application."
    ))
    mono(doc, """// ─── Currency Configuration ─────────────────────────────────────────────────────────────────────────────

export interface CurrencyConfig {
  code: string;          // ISO 4217
  symbol: string;        // Display symbol
  name: string;          // Full name
  decimals: number;      // Decimal places
  country: string;       // Country name
  flag: string;          // Emoji flag
  locale: string;        // Intl locale
  vatRate: number;       // VAT rate (0.0 – 1.0)
  phonePrefix: string;   // International dialing code
  dateFormat: string;    // Display hint
}

export const CURRENCIES: Record<string, CurrencyConfig> = {
  KWD: {
    code: 'KWD', symbol: 'KD',  name: 'Kuwaiti Dinar',  decimals: 3,
    country: 'Kuwait',  flag: '\U0001f1f0\U0001f1fc', locale: 'en-KW',
    vatRate: 0,  phonePrefix: '+965', dateFormat: 'DD/MM/YYYY',""", fs=8)

    # B.4
    h3(doc, "B.4 — src/utils/supabase/server.ts")
    body(doc, (
        "Server-side Supabase client factory using the @supabase/ssr package. Reads and "
        "writes session tokens exclusively through Next.js cookies(), ensuring authentication "
        "state is scoped to the server request context and never exposed to client-side "
        "JavaScript. This factory is imported by src/lib/auth.ts and all API route handlers "
        "that require an authenticated database connection, establishing it as the foundational "
        "security primitive of the entire backend layer."
    ))
    mono(doc, """import { createServerClient } from '@supabase/ssr'
import { cookies } from 'next/headers'

export async function createClient() {
  const cookieStore = await cookies()

  return createServerClient(
    process.env.NEXT_PUBLIC_SUPABASE_URL!,
    process.env.NEXT_PUBLIC_SUPABASE_ANON_KEY!,
    {
      cookies: {
        getAll() {
          return cookieStore.getAll()
        },
        setAll(cookiesToSet) {
          try {
            cookiesToSet.forEach(({ name, value, options }) =>
              cookieStore.set(name, value, options)
            )
          } catch {
            // Server Component context — cookie writes are ignored
          }
        },
      },
    }
  )
}""", fs=8)

# ── CHAPTER X ──────────────────────────────────────────────────────────────────

def build_chapter_10(doc):
    page_break(doc)
    h1(doc, "CHAPTER X — REFERENCES")
    body(doc, "All references are formatted in IEEE citation style.")
    doc.add_paragraph()

    refs = [
        "[01]  A. R. Rababaah, “Class Notes CSIS 330: Software Engineering,” "
        "American University of Kuwait, Spring 2026.",

        "[02]  I. Sommerville, Software Engineering, 10th ed. Harlow, UK: Pearson Education, 2016.",

        "[03]  R. C. Martin, Agile Software Development: Principles, Patterns, and Practices. "
        "Upper Saddle River, NJ: Prentice Hall, 2002.",

        "[04]  M. Fowler, Refactoring: Improving the Design of Existing Code. "
        "Reading, MA: Addison-Wesley, 1999.",

        "[05]  Vercel, Inc., “Next.js Documentation,” 2024. [Online]. "
        "Available: https://nextjs.org/docs",

        "[06]  Supabase, Inc., “Supabase Documentation,” 2024. [Online]. "
        "Available: https://supabase.com/docs",

        "[07]  B. W. Boehm, Software Engineering Economics. "
        "Englewood Cliffs, NJ: Prentice Hall, 1981.",

        "[08]  International Organization for Standardization, "
        "ISO 31000:2018 — Risk Management: Guidelines. Geneva, Switzerland: ISO, 2018.",

        "[09]  OWASP Foundation, “OWASP Top Ten — 2021,” 2021. [Online]. "
        "Available: https://owasp.org/www-project-top-ten/",

        "[10]  Communications and Information Technology Regulatory Authority (CITRA), "
        "“Kuwait Digital Economy Report 2024,” Kuwait City, Kuwait: CITRA, 2024.",

        "[11]  M. M. Lehman, “Programs, Life Cycles, and Laws of Software Evolution,” "
        "Proceedings of the IEEE, vol. 68, no. 9, pp. 1060–1076, Sep. 1980.",

        "[12]  Vercel, Inc., “Vercel Documentation,” 2024. [Online]. "
        "Available: https://vercel.com/docs",

        "[13]  Microsoft Corporation, “TypeScript Handbook,” 2024. [Online]. "
        "Available: https://www.typescriptlang.org/docs/handbook/",

        "[14]  R. S. Pressman and B. R. Maxim, Software Engineering: A Practitioner’s "
        "Approach, 8th ed. New York, NY: McGraw-Hill Education, 2014.",

        "[15]  World Wide Web Consortium (W3C), “Web Content Accessibility Guidelines "
        "(WCAG) 2.1,” 2018. [Online]. Available: https://www.w3.org/TR/WCAG21/",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(ref).font.size = Pt(10)

# ── FRONT MATTER ───────────────────────────────────────────────────────────────

def build_front_matter(doc):
    """Build cover page, ToC, List of Tables, List of Figures into doc."""

    # ── Cover Page ──────────────────────────────────────────────────────────────
    for _ in range(6):
        doc.add_paragraph()

    centre(doc, "College of Engineering and Applied Sciences (CEAS)", bold=False, size=13)
    centre(doc, "American University of Kuwait", bold=False, size=12)
    doc.add_paragraph()
    centre(doc, "CSIS-330 — Software Engineering", bold=True, size=14)
    centre(doc, "Spring 2026", bold=False, size=12)
    doc.add_paragraph()

    for _ in range(3):
        doc.add_paragraph()

    centre(doc, "Kuwait B2B Hub", bold=True, size=28)
    centre(doc, "Software Engineering Final Project Report", bold=False, size=16)

    for _ in range(4):
        doc.add_paragraph()

    centre(doc, "Course Instructor:", bold=True, size=11)
    centre(doc, "Dr. Aaron Rasheed Rababaah", bold=False, size=12)
    doc.add_paragraph()
    centre(doc, "Project Advisor:", bold=True, size=11)
    centre(doc, "To Be Confirmed (Industry Professional)", bold=False, size=12)

    for _ in range(3):
        doc.add_paragraph()

    centre(doc, "Development Team", bold=True, size=12)
    for member in [
        "Bader Alnoumas         S00059026   (Team Lead)",
        "Abdullah Abduljaleel   S00056040",
        "Abdullah Subhi         S00054535",
        "Salah Abdulfattah      S00052772",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(member)
        run.font.size = Pt(11)
        run.font.name = "Courier New"

    doc.add_paragraph()
    centre(doc, "Team Size: 4 Members", bold=False, size=11)

    for _ in range(3):
        doc.add_paragraph()

    centre(doc, "Date of Submission: 07 May 2026", bold=False, size=12)
    centre(doc, "Version: 3.5.2.1 (Final)", bold=False, size=11)

    doc.add_page_break()

    # ── Table of Contents ───────────────────────────────────────────────────────
    h1(doc, "TABLE OF CONTENTS")

    toc_entries = [
        (0, "Cover Page", ""),
        (0, "Table of Contents", ""),
        (0, "List of Tables", ""),
        (0, "List of Figures", ""),
        (0, "CHAPTER I — INTRODUCTION", ""),
        (1, "1.1  Background", ""),
        (1, "1.2  System Overview", ""),
        (1, "1.3  Document Conventions", ""),
        (1, "1.4  Intended Readership", ""),
        (1, "1.5  Product Scope", ""),
        (1, "1.6  Team Workload Distribution", ""),
        (0, "CHAPTER II — SOFTWARE REQUIREMENTS SPECIFICATION (SRS)", ""),
        (1, "2.1  Feasibility Study", ""),
        (2, "2.1.1  Technical Feasibility", ""),
        (2, "2.1.2  Economic Feasibility", ""),
        (2, "2.1.3  Legal Feasibility", ""),
        (2, "2.1.4  Operational Feasibility", ""),
        (2, "2.1.5  Schedule Feasibility", ""),
        (1, "2.2  User Requirements", ""),
        (1, "2.3  System Requirements", ""),
        (1, "2.4  Non-Functional Requirements", ""),
        (1, "2.5  Functional Requirements Specification (FRS)", ""),
        (1, "2.6  Context Model", ""),
        (1, "2.7  Use Cases", ""),
        (1, "2.8  Requirements Traceability Matrix (RTM)", ""),
        (0, "CHAPTER III — SYSTEM DESIGN", ""),
        (1, "3.1  Sequence Diagrams", ""),
        (1, "3.2  Data Flow Model", ""),
        (1, "3.3  Architectural Pattern", ""),
        (1, "3.4  Activity / Process Model", ""),
        (1, "3.5  Finite State Machine Specifications", ""),
        (2, "3.5.1  Order Lifecycle FSM", ""),
        (2, "3.5.2  Brand Verification FSM", ""),
        (1, "3.6  Component / Unit Design", ""),
        (1, "3.7  Class Diagram — Core Entity Model", ""),
        (1, "3.8  Prototype Screenshots", ""),
        (0, "CHAPTER IV — SYSTEM IMPLEMENTATION", ""),
        (1, "4.1  Make, Buy or Lease (MBL) Analysis", ""),
        (1, "4.2  Software Reuse Analysis", ""),
        (1, "4.3  Coding Artefacts & Project Structure", ""),
        (1, "4.4  Configuration Management", ""),
        (2, "4.4.1  Versioning Scheme", ""),
        (2, "4.4.2  Configuration Items", ""),
        (2, "4.4.3  Change Request Log", ""),
        (2, "4.4.4  Version Baseline Summary", ""),
        (1, "4.5  Development & Deployment Specification", ""),
        (1, "4.6  Physical Deployment Diagram", ""),
        (0, "CHAPTER V — VERIFICATION AND VALIDATION", ""),
        (1, "5.1  Development Testing", ""),
        (2, "5.1.1  Source Lines of Code (SLOC) & Automated Testing Assessment", ""),
        (1, "5.2  System Inspections", ""),
        (2, "5.2.1  SRS Inspection (Sprint 2)", ""),
        (2, "5.2.2  Architecture Inspection (Sprint 3)", ""),
        (2, "5.2.3  Design Models Inspection (Sprint 4)", ""),
        (2, "5.2.4  Code Inspection (Sprint 6)", ""),
        (1, "5.3  System Testing", ""),
        (2, "5.3.4  System Test Cases (20 Test Cases)", ""),
        (1, "5.4  Release Testing", ""),
        (2, "5.4.1  Black-Box Testing", ""),
        (2, "5.4.2  White-Box Testing", ""),
        (2, "5.4.3  Operational Profile Testing", ""),
        (2, "5.4.4  Performance Testing", ""),
        (2, "5.4.5  Use-Case Based Testing", ""),
        (1, "5.5  User Testing", ""),
        (2, "5.5.1  Alpha Testing", ""),
        (2, "5.5.2  Beta Testing", ""),
        (2, "5.5.3  Acceptance Testing", ""),
        (0, "CHAPTER VI — SYSTEM EVOLUTION", ""),
        (1, "6.1  Overview", ""),
        (1, "6.2  Lehman’s Laws Applied to the Kuwait B2B Hub", ""),
        (1, "6.3  Anticipated System Changes (SC-01 to SC-13)", ""),
        (1, "6.4  System Quality and Value: Ten-Year Prediction", ""),
        (1, "6.5  Narrative Justification", ""),
        (1, "6.6  Summary", ""),
        (0, "CHAPTER VII — PLANNING AND MANAGEMENT", ""),
        (1, "7.1  Milestones & Deliverables", ""),
        (1, "7.2  Gantt Chart", ""),
        (1, "7.3  Staff Allocation by Sprint", ""),
        (1, "7.4  Activity Network", ""),
        (1, "7.5  CoCoMo Cost Estimation", ""),
        (1, "7.6  Risk Assessment Matrix (RAM)", ""),
        (2, "7.6.1  Methodology", ""),
        (2, "7.6.2  Scoring Rubrics", ""),
        (2, "7.6.3  Risk Assessment Matrix", ""),
        (2, "7.6.4  Risk Heat Map", ""),
        (2, "7.6.5  Critical Risk Narratives", ""),
        (2, "7.6.6  Residual Risk Summary", ""),
        (1, "7.7  Root-Cause Analysis — Fishbone Diagram", ""),
        (0, "CHAPTER VIII — CONCLUSIONS", ""),
        (1, "8.1  Seven Design Smells", ""),
        (1, "8.2  SWOT Analysis", ""),
        (1, "8.3  Future Work and Recommendations", ""),
        (0, "CHAPTER IX — APPENDICES", ""),
        (1, "Appendix A — System Diagram Placeholders", ""),
        (1, "Appendix B — Critical Source File Extracts", ""),
        (0, "CHAPTER X — REFERENCES", ""),
    ]

    for indent, title, _ in toc_entries:
        toc_line(doc, "", title, indent=indent)

    doc.add_page_break()

    # ── List of Tables ───────────────────────────────────────────────────────────
    h1(doc, "LIST OF TABLES")
    lot_entries = [
        ("Table 1", "Glossary of Terms", "Chapter I"),
        ("Table 2", "Acronyms", "Chapter I"),
        ("Table 3", "Team Workload Distribution", "Chapter I"),
        ("Table 4", "Feasibility Assessment Summary", "Chapter II"),
        ("Table 5", "User Requirements (UR-01 to UR-17)", "Chapter II"),
        ("Table 6", "System Requirements (SR-01 to SR-20)", "Chapter II"),
        ("Table 7", "Non-Functional Requirements (NFR-01 to NFR-10)", "Chapter II"),
        ("Table 8", "FRS — Subsystem 1: Authentication & User Management", "Chapter II"),
        ("Table 9", "FRS — Subsystem 2: Brand Management", "Chapter II"),
        ("Table 10", "FRS — Subsystem 3: Product & Catalogue Management", "Chapter II"),
        ("Table 11", "FRS — Subsystem 4: Order Management", "Chapter II"),
        ("Table 12", "FRS — Subsystem 5: Trust, Reviews & Notifications", "Chapter II"),
        ("Table 13", "Use Case Specifications (UC-01, UC-03, UC-07, UC-12, UC-14)", "Chapter II"),
        ("Table 14", "Requirements Traceability Matrix (RTM)", "Chapter II"),
        ("Table 15", "Architectural Layer Mapping to MVC", "Chapter III"),
        ("Table 16", "Order Lifecycle FSM — State Transitions (T-01 to T-11)", "Chapter III"),
        ("Table 17", "Brand Verification FSM — State Transitions (BV-01 to BV-05)", "Chapter III"),
        ("Table 18", "Component / Unit Design Summary", "Chapter III"),
        ("Table 19", "Make, Buy or Lease (MBL) Analysis (15 Components)", "Chapter IV"),
        ("Table 20", "MBL Decision Summary by Category", "Chapter IV"),
        ("Table 21", "Software Reuse Analysis (12 Components)", "Chapter IV"),
        ("Table 22", "Source Lines of Code (SLOC) by Category", "Chapter V"),
        ("Table 23", "Automated Testing Assessment", "Chapter V"),
        ("Table 24", "Versioning Scheme — Phase Definitions", "Chapter IV"),
        ("Table 25", "Configuration Items (CI-01 to CI-08)", "Chapter IV"),
        ("Table 26", "Change Request Log (CR-001 to CR-010)", "Chapter IV"),
        ("Table 27", "Version Baseline Summary", "Chapter IV"),
        ("Table 28", "Hardware Requirements", "Chapter IV"),
        ("Table 29", "Operating System Support", "Chapter IV"),
        ("Table 30", "Runtime & Package Manager Requirements", "Chapter IV"),
        ("Table 31", "Core Framework & Language Dependencies", "Chapter IV"),
        ("Table 32", "TypeScript Compiler Options", "Chapter IV"),
        ("Table 33", "Cloud Services & Backend Infrastructure", "Chapter IV"),
        ("Table 34", "Key Runtime Dependencies", "Chapter IV"),
        ("Table 35", "IDE & Developer Tooling", "Chapter IV"),
        ("Table 36", "Testing & Quality Assurance Methods", "Chapter IV"),
        ("Table 37", "Deployment Environment Summary", "Chapter IV"),
        ("Table 38", "SRS Inspection Findings (INS1)", "Chapter V"),
        ("Table 39", "Architecture Inspection Findings (INS2)", "Chapter V"),
        ("Table 40", "Design Models Inspection Findings (INS3)", "Chapter V"),
        ("Table 41", "Code Inspection Findings (INS4, 5 Findings)", "Chapter V"),
        ("Table 42", "System Test Cases (TC-001 to TC-020)", "Chapter V"),
        ("Table 43", "Black-Box Test Conditions", "Chapter V"),
        ("Table 44", "White-Box Coverage Assessment", "Chapter V"),
        ("Table 45", "Operational Profile Testing — Transaction Frequencies", "Chapter V"),
        ("Table 46", "Performance Testing Results (6 Endpoints)", "Chapter V"),
        ("Table 47", "Use-Case Based Testing (UC-01 to UC-16)", "Chapter V"),
        ("Table 48", "Alpha Testing — Tester Sessions (4 Testers)", "Chapter V"),
        ("Table 49", "Beta Testing — User Ratings (6 Participants)", "Chapter V"),
        ("Table 50", "Acceptance Testing — UR Compliance (17 URs)", "Chapter V"),
        ("Table 51", "Lehman’s Laws Applied to Kuwait B2B Hub", "Chapter VI"),
        ("Table 52", "Anticipated System Changes (SC-01 to SC-13)", "Chapter VI"),
        ("Table 53", "Ten-Year System Quality and Value Prediction", "Chapter VI"),
        ("Table 54", "Milestones & Deliverables (M-01 to M-13)", "Chapter VII"),
        ("Table 55", "Staff Allocation by Sprint (8 Sprints × 4 Members)", "Chapter VII"),
        ("Table 56", "Activity Network (A1–A15) with Critical Path", "Chapter VII"),
        ("Table 57", "COCOMO Input Parameters", "Chapter VII"),
        ("Table 58", "COCOMO Results Summary", "Chapter VII"),
        ("Table 59", "Likelihood Scoring Rubric", "Chapter VII"),
        ("Table 60", "Impact Scoring Rubric", "Chapter VII"),
        ("Table 61", "Risk Assessment Matrix (R-01 to R-12)", "Chapter VII"),
        ("Table 62", "Residual Risk Summary", "Chapter VII"),
        ("Table 63", "Root-Cause Analysis — Fishbone Contributing Factors", "Chapter VII"),
        ("Table 64", "SWOT Analysis", "Chapter VIII"),
        ("Table 65", "Appendix A — System Diagram Index", "Chapter IX"),
    ]
    make_table(doc,
        ["Table No.", "Title", "Chapter"],
        lot_entries,
        col_widths=[0.8, 4.2, 1.5]
    )

    doc.add_page_break()

    # ── List of Figures ─────────────────────────────────────────────────────────
    h1(doc, "LIST OF FIGURES")
    lof_entries = [
        ("Figure 1", "Sequence Diagram — User Registration & Login", "§3.1.1"),
        ("Figure 2", "Sequence Diagram — Order Placement & Review Lifecycle", "§3.1.2"),
        ("Figure 3", "Sequence Diagram — Brand Verification", "§3.1.3"),
        ("Figure 4", "Data Flow Diagram Level 0 — Context Diagram", "§3.2.1"),
        ("Figure 5", "Data Flow Diagram Level 1 — Sub-Process Decomposition", "§3.2.2"),
        ("Figure 6", "Activity Diagram — Order Placement Process", "§3.4"),
        ("Figure 7", "Finite State Machine — Order Lifecycle (11 Transitions)", "§3.5.1"),
        ("Figure 8", "Finite State Machine — Brand Verification (5 Transitions)", "§3.5.2"),
        ("Figure 9", "Class Diagram — Core Entity Model (10 Entities)", "§3.7"),
        ("Figure 10", "Physical Deployment Diagram — Vercel / Supabase Architecture", "§4.6"),
        ("Figure 11", "Kuwait B2B Hub — Final Project Folder Structure", "§4.3"),
        ("Figure 12", "Gantt Chart — Project Schedule (8 Sprints)", "§7.2"),
        ("Figure 13", "Risk Heat Map — Likelihood × Impact Grid", "§7.6.4"),
        ("Figure 14", "Fishbone Diagram — Root-Cause Analysis", "§7.7"),
    ]
    make_table(doc,
        ["Figure No.", "Caption", "Section"],
        lof_entries,
        col_widths=[0.8, 4.2, 1.5]
    )

    doc.add_page_break()


# ── ASSEMBLY ───────────────────────────────────────────────────────────────────

def main():
    # Step 1: Open v3 and append the three new chapters
    print("Opening Final_Report_v3.docx ...")
    main_doc = Document(INPUT_FILE)

    print("Appending Chapter VIII ...")
    build_chapter_8(main_doc)

    print("Appending Chapter IX ...")
    build_chapter_9(main_doc)

    print("Appending Chapter X ...")
    build_chapter_10(main_doc)

    # Step 2: Create the final document starting with front matter
    print("Building front matter ...")
    final_doc = Document()

    # Remove the single empty paragraph that a fresh Document() starts with
    for p in list(final_doc.paragraphs):
        el = p._element
        el.getparent().remove(el)

    build_front_matter(final_doc)

    # Step 3: Copy all body elements from main_doc into final_doc
    print("Assembling final document ...")
    main_body = main_doc.element.body
    final_body = final_doc.element.body

    # final_body ends with sectPr; insert before it
    final_sectPr = final_body[-1]

    for el in main_body:
        tag = el.tag
        if tag.endswith("}sectPr"):
            continue  # skip source sectPr; keep final_doc's own sectPr
        final_body.insert(list(final_body).index(final_sectPr), copy.deepcopy(el))

    # Step 4: Save
    print(f"Saving to {OUTPUT_FILE} ...")
    final_doc.save(OUTPUT_FILE)

    # Verify
    import os
    from docx import Document as D
    vdoc = D(OUTPUT_FILE)
    size_kb = os.path.getsize(OUTPUT_FILE) // 1024
    tables = len(vdoc.tables)
    paras = len(vdoc.paragraphs)
    print(f"\nVerification: {paras} paragraphs, {tables} tables, {size_kb} KB")
    print("\nALL PHASES COMPLETE — Final report ready")
    print(f"Output: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
