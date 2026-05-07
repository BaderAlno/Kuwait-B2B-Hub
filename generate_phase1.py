#!/usr/bin/env python3
"""
Kuwait B2B Hub — Final Report Generator (Phase 1)
Generates Chapters I and II with full python-docx formatting.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Paths ───────────────────────────────────────────────────────────────────

OUTPUT_DIR = "/Users/baderalnoumas/Desktop/B2B/Final_Report_Build"
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "Final_Report_v1.docx")
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ─── Document Setup ──────────────────────────────────────────────────────────

doc = Document()

# A4 page with academic margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.25)
    section.right_margin  = Inches(1.25)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def page_break(doc):
    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def body(doc, text):
    doc.add_paragraph(text, style="Normal")


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def make_table(doc, headers, rows, col_widths=None):
    """Create a 'Table Grid' table with bold headers."""
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = "Table Grid"

    # Header
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True

    # Body rows
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)

    # Optional column widths
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])

    doc.add_paragraph()  # breathing room after table
    return table


def uc_spec(doc, rows):
    """Render a use-case spec as bold-label: value paragraphs."""
    for label, value in rows:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)
    doc.add_paragraph()


# ─── TITLE PAGE ──────────────────────────────────────────────────────────────

tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run("Kuwait B2B Hub")
r.bold = True
r.font.size = Pt(26)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sp.add_run("Final Project Report")
r2.bold = True
r2.font.size = Pt(20)

doc.add_paragraph()

for line in [
    "Course: Software Engineering",
    "Institution: American University of the Middle East (AUM)",
    "Supervisor: Dr. Rababaah",
    "Date: May 2026",
    "",
    "Platform: B2B Wholesale Marketplace for Kuwait & GCC",
    "",
    "Team Members",
    "Bader Alnoumas — S00059026",
    "Abdullah Abduljaleel — S00056040",
    "Abdullah Subhi — S00054535",
    "Salah Abdulfattah — S00052772",
]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER I — INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════

h1(doc, "CHAPTER I — INTRODUCTION")

# ── §1.1 Background ──────────────────────────────────────────────────────────

h2(doc, "1.1  Background")

body(doc,
    "Kuwait's wholesale and business-to-business (B2B) trading sector has historically been "
    "characterised by an informal, relationship-driven model in which procurement decisions are "
    "negotiated via personal networks, telephone calls, and increasingly through consumer "
    "messaging applications such as WhatsApp. While this model functioned adequately in a smaller, "
    "more localised economy, it has become a significant barrier to efficiency as Kuwait's commercial "
    "sector expands in line with the National Vision 2035 and the Gulf Cooperation Council's (GCC) "
    "broader digital economy agenda. The absence of a structured digital marketplace means that buyers "
    "cannot efficiently discover new suppliers, verify their credentials, compare pricing transparently, "
    "or track the fulfilment status of placed orders — all capabilities that are taken for granted in "
    "B2B platforms serving more mature markets. Wholesale distributors and brand owners, for their part, "
    "lack the tooling to manage catalogue data, enforce minimum order quantities, or build a verifiable "
    "digital reputation that can be shared with prospective buyers without manual reference checks."
)

body(doc,
    "The Kuwait B2B Hub was conceived as a direct response to this structural gap. The platform "
    "provides a managed, trust-centric digital environment in which Kuwaiti and GCC wholesale brands "
    "can register, list their product catalogues with structured bulk-pricing tiers, and transact with "
    "verified business buyers under a regulated order lifecycle. By centralising brand discovery, "
    "product browsing, order management, and reputation building within a single bilingual (Arabic and "
    "English) platform, the Kuwait B2B Hub replaces the fragmented, informal channel with a "
    "professional, auditable, and scalable alternative. The platform further addresses the GCC's "
    "multicurrency reality by supporting live display in Kuwaiti Dinar (KWD), Saudi Riyal (SAR), UAE "
    "Dirham (AED), and Bahraini Dinar (BHD), reducing the cognitive friction that arises when buyers "
    "and sellers operate across different currency jurisdictions. A dual-sided trust scoring engine — "
    "computing composite reputation metrics for both brands and buyers from historical order behaviour "
    "— addresses the fundamental credibility problem that informal markets cannot solve: how to "
    "transact confidently with parties one has never met in person."
)

body(doc,
    "From a technical perspective, the Kuwait B2B Hub is built on a modern full-stack architecture "
    "leveraging Next.js 16 with the App Router paradigm, TypeScript for compile-time type safety "
    "across all 93 source files, Supabase for authentication and production database management, and "
    "Vercel for cloud deployment with automatic CI/CD. The system was designed from the outset with a "
    "separation-of-concerns principle: the presentation layer is decoupled from business logic via a "
    "RESTful API layer comprising 25 serverless route handlers, and data access is abstracted through "
    "a centralised module that permits seamless migration from the MVP's JSON prototype store to a "
    "production-grade PostgreSQL database without modifying the API contracts. The codebase totals "
    "11.35 KLOC across 93 source files, with the largest concentration in reusable UI components "
    "(33.5%) and buyer/brand-facing page routes (40.1% combined). This document constitutes the Final "
    "Project Report for the Kuwait B2B Hub, submitted in partial fulfilment of the Software "
    "Engineering course requirements at the American University of the Middle East."
)

# ── §1.2 System Overview ─────────────────────────────────────────────────────

h2(doc, "1.2  System Overview")

body(doc,
    "The Kuwait B2B Hub is a digital wholesale marketplace platform targeting the Kuwaiti and broader "
    "GCC business community. The following subsections outline the system's primary goals, architectural "
    "structure, user roles, and core technology choices."
)

h3(doc, "Goals")

for goal in [
    "Professionalise B2B wholesale trading in Kuwait by replacing informal WhatsApp-based procurement "
    "with a structured, digital workflow that provides full transaction audit trails.",
    "Provide verified brand discovery, enabling buyers to identify and engage with credentialed "
    "Kuwaiti suppliers confidently through administrative verification and tiered badges.",
    "Enforce structured order management with a full status lifecycle (pending → approved → completed "
    "/ rejected) and in-app notifications at every transition.",
    "Build trust through a dual-sided reputation system covering both brand trustworthiness "
    "(response rate, completion rate, fulfilment speed) and buyer reliability (completion rate, "
    "cancellation rate).",
    "Support bilingual operation — Arabic RTL and English LTR — to serve both local and expatriate "
    "business communities throughout Kuwait and the GCC.",
    "Enable GCC multi-market expansion by supporting multiple currencies, VAT configurations per "
    "country, and market-specific brand visibility settings.",
]:
    bullet(doc, goal)

h3(doc, "System Architecture")

body(doc, "The platform follows a four-layer architectural model:")

make_table(doc,
    ["Layer", "Description"],
    [
        ("Presentation Layer",
         "React/TSX pages and reusable components; CSS Modules for styling; Lucide React for "
         "iconography; Recharts for analytics visualisations on dashboards."),
        ("Application / API Layer",
         "Next.js 16 App Router route handlers providing a RESTful JSON API under /api/*. "
         "Authentication middleware enforces RBAC at the network boundary before business logic executes."),
        ("Data Access Layer",
         "A centralised db.ts module abstracts all reads and writes behind readDB(), writeDB(), and "
         "generateId(). The MVP uses a JSON flat-file store; the production target is Supabase "
         "PostgreSQL, with migration scaffolded and ready."),
        ("External Services Layer",
         "Supabase (JWT authentication + managed database), WhatsApp Business links, Vercel hosting "
         "with global CDN, and a configurable live-exchange-rate API for GCC currency conversions."),
    ],
    col_widths=[1.8, 4.7]
)

h3(doc, "User Roles")

body(doc,
    "The system defines three distinct actor roles, each with a segregated permission set and "
    "dedicated UI portal:"
)

make_table(doc,
    ["Role", "Description"],
    [
        ("Admin",
         "Platform administrators responsible for brand verification, user management, order monitoring, "
         "and review moderation. Full platform-wide data visibility. Accounts are created internally."),
        ("Brand Owner",
         "Kuwaiti wholesale suppliers who create brand profiles, list products with bulk-pricing tiers, "
         "process buyer orders, and respond to reviews via a dedicated brand portal. "
         "Subscription: KD 29 / month with a 14-day free trial."),
        ("Business Buyer",
         "Retail and wholesale procurement professionals who browse verified brands, view product "
         "catalogues, submit bulk orders, track fulfilment, and leave reviews. Free to use."),
    ],
    col_widths=[1.5, 5.0]
)

h3(doc, "Technology Stack")

make_table(doc,
    ["Component", "Technology"],
    [
        ("Frontend Framework",  "Next.js 16.2.2 with App Router; TypeScript 5 (strict mode)"),
        ("UI & Styling",        "CSS Modules; Lucide React icons; Recharts / Chart.js analytics"),
        ("State Management",    "Zustand 5.0.12 (notification store); React Context (currency)"),
        ("Internationalisation","next-intl 4.9.0 — English & Arabic with full RTL support"),
        ("Authentication",      "Supabase Auth with SSR session management via @supabase/ssr ^0.10.0"),
        ("Database (MVP)",      "JSON flat-file (src/data/db.json) via custom db.ts abstraction"),
        ("Database (Prod.)",    "Supabase managed PostgreSQL — migration path scaffolded"),
        ("Data Import",         "xlsx 0.18.5 — Excel/CSV bulk product catalogue import"),
        ("Hosting & CI/CD",     "Vercel — serverless deployment, global CDN, automatic Git pipeline"),
    ],
    col_widths=[2.0, 4.5]
)

# ── §1.3 Document Conventions ────────────────────────────────────────────────

h2(doc, "1.3  Document Conventions")

body(doc,
    "This section defines the terms, acronyms, abbreviations, and symbols used throughout this report. "
    "Readers unfamiliar with any entry should consult the definitions below before proceeding to the "
    "technical chapters."
)

h3(doc, "Glossary of Terms")

make_table(doc,
    ["Term", "Definition"],
    [
        ("B2B",
         "Business-to-Business — commercial transactions conducted between companies rather than "
         "between a company and an individual consumer."),
        ("Brand",
         "A registered wholesale supplier or manufacturer on the Kuwait B2B Hub platform, subject "
         "to administrative verification before appearing on the marketplace."),
        ("Brand Trust Score",
         "A composite metric calculated from a brand's order response rate, completion rate, and "
         "average fulfilment time, used to signal reliability to prospective buyers."),
        ("Bulk Pricing Tier",
         "A quantity-based price schedule where unit price decreases as order quantity increases, "
         "defined per product as an array of min_qty / max_qty / price bands."),
        ("Buyer Trust Score",
         "A metric reflecting a buyer's commercial reliability, computed from order completion rate "
         "and cancellation rate and displayed to brand owners evaluating order requests."),
        ("GCC",
         "Gulf Cooperation Council — the regional intergovernmental union comprising Bahrain, Kuwait, "
         "Oman, Qatar, Saudi Arabia, and the United Arab Emirates."),
        ("Marketplace",
         "The public-facing section of the platform (/marketplace) where buyers discover and browse "
         "all administratively approved brand profiles and product catalogues."),
        ("MOQ",
         "Minimum Order Quantity — the smallest number of units a buyer may include in a single order "
         "line item for a given product, enforced at order submission by the API."),
        ("Platform Admin",
         "An internal operator account with elevated privileges to manage users, approve or reject "
         "brands, moderate reviews, and access platform-wide analytics."),
        ("RTM",
         "Requirements Traceability Matrix — a document that links each requirement to its design "
         "artefact, implementation file, and test evidence, confirming end-to-end coverage."),
        ("Trust Engine",
         "The subsystem within the API layer responsible for computing and persisting trust scores "
         "and earning badges for both brands and buyers based on historical order behaviour."),
        ("Verification Tier",
         "An administrative quality classification assigned to approved brands: Premium (highest), "
         "Verified, or New — displayed as a visual badge throughout the platform."),
    ],
    col_widths=[1.8, 4.7]
)

h3(doc, "Acronyms")

make_table(doc,
    ["Acronym", "Expansion"],
    [
        ("API",   "Application Programming Interface"),
        ("CRUD",  "Create, Read, Update, Delete"),
        ("CSS",   "Cascading Style Sheets"),
        ("FR",    "Functional Requirement"),
        ("FRS",   "Functional Requirements Specification"),
        ("GCC",   "Gulf Cooperation Council"),
        ("HTTP",  "HyperText Transfer Protocol"),
        ("i18n",  "Internationalisation (18 letters between 'i' and 'n')"),
        ("JWT",   "JSON Web Token"),
        ("KLOC",  "Thousand Lines of Code"),
        ("KWD",   "Kuwaiti Dinar (ISO 4217 currency code)"),
        ("MBL",   "Make-Buy-Lease (procurement decision framework)"),
        ("MOQ",   "Minimum Order Quantity"),
        ("MVP",   "Minimum Viable Product"),
        ("NFR",   "Non-Functional Requirement"),
        ("RBAC",  "Role-Based Access Control"),
        ("REST",  "Representational State Transfer"),
        ("RTL",   "Right-To-Left (text direction for Arabic script)"),
        ("RTM",   "Requirements Traceability Matrix"),
        ("SLOC",  "Source Lines of Code"),
        ("SRS",   "Software Requirements Specification"),
        ("SSR",   "Server-Side Rendering"),
        ("TSX",   "TypeScript JSX — TypeScript files containing React JSX syntax"),
        ("UC",    "Use Case"),
        ("UI",    "User Interface"),
        ("UML",   "Unified Modelling Language"),
    ],
    col_widths=[1.2, 5.3]
)

h3(doc, "Abbreviations")

make_table(doc,
    ["Abbreviation", "Meaning"],
    [
        ("AED",    "UAE Dirham — currency of the United Arab Emirates"),
        ("AUM",    "American University of the Middle East"),
        ("BHD",    "Bahraini Dinar — currency of Bahrain"),
        ("CI/CD",  "Continuous Integration / Continuous Delivery"),
        ("E2E",    "End-to-End (testing methodology)"),
        ("LTR",    "Left-To-Right text direction (English layout)"),
        ("PR",     "Pull Request — a code review submission in Git"),
        ("QAR",    "Qatari Riyal — currency of Qatar"),
        ("RTL",    "Right-To-Left text direction (Arabic layout)"),
        ("SAR",    "Saudi Riyal — currency of Saudi Arabia"),
        ("v",      "Version (as in v1.0.0 or v3.5.2.1)"),
        ("VCS",    "Version Control System (Git in this project)"),
    ],
    col_widths=[1.5, 5.0]
)

h3(doc, "Symbols")

make_table(doc,
    ["Symbol", "Meaning in This Report"],
    [
        ("→",    "State transition or data flow (e.g., pending → approved → completed)"),
        ("≥",    "Greater than or equal to (MOQ and threshold expressions)"),
        ("<",    "Strictly less than (performance target thresholds)"),
        ("%",    "Percentage — completion metrics, chapter weight distribution, code coverage"),
        ("*",    "Wildcard in API route path patterns (e.g., /api/*)"),
        ("[ ]",  "Array notation — list-typed field (e.g., bulk_pricing_tiers[])"),
        ("~",    "Approximately equal to (estimation figures)"),
        ("KD",   "Display symbol for Kuwaiti Dinar (ISO code: KWD)"),
        ("/",    "File system and URL path separator"),
        ("§",    "Section reference (e.g., §2.1 = Section 2.1 of this report)"),
    ],
    col_widths=[0.9, 5.6]
)

# ── §1.4 Intended Readership ─────────────────────────────────────────────────

h2(doc, "1.4  Intended Readership")

body(doc,
    "This report is written to serve multiple audiences. The table below identifies the five primary "
    "readership groups and the sections of the report most pertinent to each."
)

make_table(doc,
    ["Readership Group", "Role", "Key Sections"],
    [
        ("Platform Administrators",
         "Operational staff responsible for running the live Kuwait B2B Hub platform post-launch.",
         "§1.2 (roles & features), §2.2–2.4 (requirements), §4 (implementation), §7 (project management)"),
        ("Brand Owners",
         "Wholesale suppliers and manufacturers considering or actively using the platform.",
         "§1.2 (overview), §2.5 (FRS — brand & product subsystems), §4.4 (make/buy analysis)"),
        ("Business Buyers",
         "Retail and wholesale procurement professionals evaluating the platform for sourcing.",
         "§1.2 (overview), §2.7 use cases UC-03, UC-05, UC-06, §5 (system evaluation)"),
        ("Dr. Rababaah (Supervisor)",
         "Academic supervisor assessing the project against Software Engineering course criteria.",
         "All chapters. Particular focus: §2 (SRS), §3 (design), §4 (implementation), §5 (evaluation), "
         "§7 (project management, COCOMO, risk)."),
        ("Future Developers",
         "Engineers who will extend, maintain, or deploy the platform beyond the academic submission.",
         "§3 (design), §4 (implementation & deployment spec), §4.5 (configuration management), "
         "§6 (system evolution), Appendix B (source file extracts)."),
    ],
    col_widths=[1.7, 1.8, 3.0]
)

# ── §1.5 Product Scope ───────────────────────────────────────────────────────

h2(doc, "1.5  Product Scope")

h3(doc, "Functional Scope")

body(doc,
    "The Kuwait B2B Hub delivers a comprehensive set of functional capabilities organised across five "
    "primary subsystems. The Authentication and User Management subsystem provides secure account "
    "registration with role selection (buyer or brand owner), login and logout with HTTP-only "
    "session cookies persisting for seven days, and role-based route protection enforced at both the "
    "Next.js middleware layer and within each API route handler. Admin accounts are created internally "
    "by the platform operator."
)

body(doc,
    "The Brand Management subsystem enables brand owners to create and maintain brand profiles, submit "
    "listings for administrative verification, achieve one of three verification tiers (Premium, "
    "Verified, New), configure WhatsApp contact details with automatic click-count tracking, and "
    "set business hours and automated reply messages. The Product and Catalogue Management subsystem "
    "allows brand owners to create, edit, and delete product listings with full attributes including "
    "images, stock levels, minimum order quantities, and structured bulk-pricing tiers. Brands may "
    "bulk-import entire product catalogues from Microsoft Excel (.xlsx) files, reducing manual "
    "catalogue population effort significantly."
)

body(doc,
    "The Order Management subsystem supports the full order lifecycle: buyers submit requests with "
    "MOQ and stock validation enforced at submission time; brand owners review and approve or reject "
    "requests; orders transition through a defined status sequence (pending → approved → completed / "
    "rejected); both parties receive targeted in-app notifications at each transition. The Trust, "
    "Review, and Notification subsystem provides 1-to-5 star review submission with optional "
    "anonymity, brand replies to reviews, admin moderation tools, and a Trust Engine that computes "
    "composite trust scores for both brands (response rate, completion rate, average fulfilment days) "
    "and buyers (total orders, completion rate, cancellation rate) from historical order data."
)

h3(doc, "Non-Functional Scope")

body(doc,
    "Beyond its functional features, the Kuwait B2B Hub is governed by eight non-functional "
    "requirements that define the quality attributes of the system. Performance targets require API "
    "response times under 300 milliseconds for standard read operations under normal load. The security "
    "posture mandates HTTP-only, SameSite-protected session cookies; HTTPS enforcement via Vercel's "
    "edge network; and a mandatory transition from the MVP's plain-text credential storage to "
    "Supabase's bcrypt-based authentication before any public deployment. Usability requirements "
    "mandate that all primary workflows are operable on mobile viewports as small as 375 pixels wide, "
    "supported by a dedicated bottom navigation bar and mobile-optimised top bar. Internationalisation "
    "requirements specify that every user-facing string is externalised to locale-specific message "
    "files, enabling complete Arabic RTL and English LTR rendering without any hard-coded string "
    "literals in component files. Scalability requirements ensure that the data access abstraction "
    "layer can be migrated from the JSON prototype store to Supabase PostgreSQL without modifying any "
    "API route handler — a key architectural decision that protects the application layer from "
    "infrastructure churn during the production transition."
)

# ── §1.6 Team Workload ───────────────────────────────────────────────────────

h2(doc, "1.6  Team Workload Distribution")

body(doc,
    "The following table documents each team member's authoring responsibilities across the ten "
    "chapters of the Final Report. Chapter weights reflect the relative academic and technical "
    "significance of each section as a proportion of total report effort. A weight of 25% for Chapter "
    "IV (Implementation) reflects its status as the largest technical deliverable, encompassing "
    "codebase documentation, architecture evidence, make/buy analysis, configuration management, and "
    "deployment specification."
)

make_table(doc,
    ["Chapter", "Title", "Weight", "Primary Lead", "Support / Contributors"],
    [
        ("Ch I",   "Introduction",          "7%",  "Bader Alnoumas",        "Abdullah Abduljaleel"),
        ("Ch II",  "SRS",                   "7%",  "Salah Abdulfattah",     "Abdullah Subhi"),
        ("Ch III", "System Design",         "20%", "Abdullah Abduljaleel",  "All members"),
        ("Ch IV",  "Implementation",        "25%", "Abdullah Subhi",        "All members"),
        ("Ch V",   "System Evaluation",     "15%", "Bader Alnoumas",        "Salah Abdulfattah"),
        ("Ch VI",  "System Evolution",      "10%", "Salah Abdulfattah",     "Abdullah Abduljaleel"),
        ("Ch VII", "Project Management",    "7%",  "Bader Alnoumas",        "Abdullah Abduljaleel"),
        ("Ch VIII","Conclusion",            "5%",  "Abdullah Subhi",        "Bader Alnoumas"),
        ("Ch IX",  "References",            "2%",  "All members",           "—"),
        ("Ch X",   "Appendices",            "2%",  "All members",           "—"),
    ],
    col_widths=[0.7, 1.7, 0.7, 1.7, 1.7]
)

body(doc, "Team Member Student IDs:")

make_table(doc,
    ["Member", "Student ID", "Primary Chapter Responsibilities"],
    [
        ("Bader Alnoumas",      "S00059026", "Lead: Ch I, Ch V, Ch VII"),
        ("Abdullah Abduljaleel","S00056040", "Lead: Ch III; Support: Ch I, Ch VI, Ch VII"),
        ("Abdullah Subhi",      "S00054535", "Lead: Ch IV, Ch VIII; Support: Ch II, Ch VI"),
        ("Salah Abdulfattah",   "S00052772", "Lead: Ch II, Ch VI; Support: Ch V, Ch VII"),
    ],
    col_widths=[1.8, 1.3, 3.4]
)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
# CHAPTER II — SOFTWARE REQUIREMENTS SPECIFICATION
# ════════════════════════════════════════════════════════════════════════════

h1(doc, "CHAPTER II — SOFTWARE REQUIREMENTS SPECIFICATION (SRS)")

body(doc,
    "This chapter presents the complete Software Requirements Specification for the Kuwait B2B Hub. "
    "It documents the feasibility analysis across five dimensions, all 17 user requirements, all "
    "20 system requirements, all 8 non-functional requirements, the full Functional Requirements "
    "Specification across the five platform subsystems, the context model with external entity "
    "definitions, the complete use-case catalogue with five detailed specifications, the preliminary "
    "system architecture, and the Requirements Traceability Matrix."
)

# ── §2.1 Feasibility Study ───────────────────────────────────────────────────

h2(doc, "2.1  Feasibility Study")

body(doc,
    "The feasibility of the Kuwait B2B Hub was assessed across five standard dimensions prior to "
    "committing to full-scale development. This analysis informed both the technical choices made "
    "during implementation and the scope boundaries established in the requirements phase."
)

h3(doc, "2.1.1  Technical Feasibility")

body(doc,
    "Technical feasibility was evaluated by determining whether the required functionality could be "
    "delivered within the team's collective skill set, the available open-source tooling, and the "
    "time constraints of a single academic semester. The verdict was affirmative on all three "
    "dimensions. The Next.js 16 App Router provides a mature, well-documented full-stack framework "
    "that unifies server-side rendering, static generation, API route handling, and middleware "
    "pipeline management into a single coherent platform. TypeScript's compile-time type system "
    "enforces data integrity contracts across all 93 source files, dramatically reducing the risk of "
    "runtime type errors in a complex multi-role application with 25 API endpoints. Supabase offers "
    "a managed PostgreSQL instance with built-in JWT authentication, real-time subscriptions, and a "
    "JavaScript SDK that abstracts database administration complexity. Vercel's first-party Next.js "
    "hosting provides zero-configuration deployment, automatic TLS, and a global content delivery "
    "network. All selected technologies are open-source or available on free tiers, confirming that "
    "no technical licensing barrier exists. The team's proficiency in TypeScript and React, combined "
    "with the comprehensive documentation of the chosen stack, supports a finding of "
    "High Technical Feasibility."
)

h3(doc, "2.1.2  Economic Feasibility")

body(doc,
    "The Kuwait B2B Hub prototype was developed at zero direct monetary cost. All frameworks — "
    "Next.js, React, TypeScript, Zustand, next-intl, Recharts — are MIT-licensed open-source "
    "software with no runtime royalties. The Supabase free tier provides a managed PostgreSQL "
    "instance, JWT authentication, and real-time infrastructure at a scale appropriate for "
    "prototype testing. Vercel's Hobby tier provides production-grade serverless hosting and CI/CD "
    "for Next.js at no cost. No proprietary tools or paid SaaS subscriptions were required for "
    "development or deployment during the academic phase. The COCOMO estimation documented in §7.5 "
    "projects that equivalent commercial development effort would total approximately 30.5 "
    "person-months — representing a significant capital cost that was entirely eliminated through "
    "academic labour. Commercial viability rests on the proposed brand subscription model of "
    "KD 29 per month, which at projected adoption rates would generate sufficient recurring revenue "
    "to cover cloud infrastructure costs from as few as 20 paying brand subscribers. Economic "
    "feasibility is assessed as High for the prototype phase and Moderate-to-High for the commercial "
    "version, subject to successful KNET payment gateway integration."
)

h3(doc, "2.1.3  Legal Feasibility")

body(doc,
    "Legal feasibility for the Kuwait B2B Hub involves compliance with several regulatory frameworks. "
    "Kuwait's Personal Data Protection Law (Law No. 20 of 2014) governs the collection, storage, "
    "and processing of personal data including names, email addresses, WhatsApp numbers, and company "
    "identifiers — all collected at platform registration. The prototype does not currently implement "
    "explicit consent mechanisms, a published Privacy Policy, or documented data retention policies; "
    "these are identified as near-term evolution requirements in §6.3. Payment processing via the "
    "Kuwait national payment network (KNET) is subject to Central Bank of Kuwait approval and "
    "PCI-DSS compliance obligations; the current prototype defers all monetary transaction processing "
    "and collects no card data. All open-source licences governing the chosen frameworks (MIT, "
    "Apache 2.0) permit commercial use without royalty obligations. Legal feasibility is assessed as "
    "Moderate — achievable with specific compliance investments — rather than blocked."
)

h3(doc, "2.1.4  Operational Feasibility")

body(doc,
    "Operational feasibility addresses whether the platform can be adopted and used effectively by "
    "its intended user base in the Kuwaiti market context. The Kuwait B2B Hub addresses operational "
    "viability on multiple dimensions. Full bilingual support — Arabic RTL and English LTR — ensures "
    "that native Arabic-speaking brand owners and buyers can operate the platform in their primary "
    "language without linguistic barriers. The mobile-first responsive design directly acknowledges "
    "that over 78% of Kuwaiti internet users access digital services via mobile devices (CITRA, 2024), "
    "and the dedicated BottomNav and MobileTopBar components target this usage pattern. The brand "
    "verification workflow and trust scoring system address the core operational challenge of building "
    "credibility in a market where informal personal relationships have historically substituted for "
    "institutional trust infrastructure. The self-service model — in which brands and buyers register "
    "and operate without requiring platform staff intervention for standard transactions — ensures that "
    "the platform can scale without proportionate staffing growth. Operational feasibility is assessed "
    "as High, contingent on reaching the critical mass of verified brands required to generate a "
    "compelling discovery experience for buyers."
)

h3(doc, "2.1.5  Schedule Feasibility")

body(doc,
    "The Kuwait B2B Hub was developed within a single academic semester, from project inception in "
    "January 2026 to final submission in June 2026 — approximately five to six months of part-time "
    "collaborative development by a four-person team. The Agile sprint model (eight sprints, each "
    "delivering a vertical feature slice) proved well-suited to this compressed timeline. Core "
    "business workflows — authentication, marketplace, product management, order lifecycle, and trust "
    "scoring — were delivered within Sprints 1 through 6. Sprint 7 addressed localisation and quality "
    "hardening; Sprint 8 targets production readiness and academic deliverable completion. The COCOMO "
    "estimation in §7.5 projects a recommended development timeline of approximately 9.2 months for "
    "a team of 3–4 persons; the academic project was completed in roughly 65% of that estimated "
    "commercial timeline, an acceleration that is explicable by the reduced MVP scope and the "
    "elimination of business overhead. Schedule feasibility is assessed as Achieved for the prototype, "
    "with the three Critical post-submission evolution items (PostgreSQL migration, automated test "
    "suite, KNET payment integration) requiring an additional 6–9 months of dedicated commercial "
    "development effort."
)

# ── §2.2 User Requirements ───────────────────────────────────────────────────

h2(doc, "2.2  User Requirements")

body(doc,
    "User Requirements (URs) express what the system must allow users to accomplish, stated in "
    "user-facing language without prescribing implementation detail. The following 17 user "
    "requirements were derived from stakeholder analysis, the platform's business model, and GCC "
    "B2B market research."
)

make_table(doc,
    ["UR ID", "Actor(s)", "Requirement Statement", "Priority"],
    [
        ("UR-01", "Buyer / Brand Owner / Admin",
         "A user shall be able to register a new account on the platform by providing their full name, "
         "email address, password, company name, and selecting a role (Buyer or Brand Owner). "
         "Admin accounts are provisioned internally.",
         "High"),
        ("UR-02", "Buyer / Brand Owner / Admin",
         "A registered user shall be able to log in using their email and password, and log out, "
         "with their session persisting for up to seven days without requiring re-authentication.",
         "High"),
        ("UR-03", "Admin",
         "A platform administrator shall be able to review all pending brand registration "
         "applications, approve or reject each, and assign a verification tier (Premium, Verified, "
         "or New) to approved brands.",
         "High"),
        ("UR-04", "Brand Owner",
         "A brand owner shall be able to create, view, edit, and delete product listings within "
         "their brand catalogue, specifying name, description, price, stock level, minimum order "
         "quantity, and product image.",
         "High"),
        ("UR-05", "Brand Owner",
         "A brand owner shall be able to define multiple bulk-pricing tiers for any product, each "
         "specifying a quantity range (min_qty to max_qty) and the corresponding discounted unit "
         "price for orders within that range.",
         "High"),
        ("UR-06", "Buyer",
         "A buyer shall be able to browse all verified brands on the marketplace, applying filters "
         "for category, verification tier, minimum order quantity, and star rating, and shall be "
         "able to view each brand's full profile including products and trust metrics.",
         "High"),
        ("UR-07", "Buyer",
         "A buyer shall be able to view all product prices displayed in their selected currency "
         "(KWD, SAR, AED, or BHD), with amounts automatically converted using the platform's "
         "configurable exchange rates.",
         "Medium"),
        ("UR-08", "Buyer",
         "A buyer shall be able to submit a bulk order request for one or more products from a "
         "single brand. The system shall enforce minimum order quantity and available stock "
         "constraints at the time of submission.",
         "High"),
        ("UR-09", "Brand Owner",
         "A brand owner shall be able to view all incoming order requests, inspect buyer identity "
         "and order details, and approve or reject each order, updating its status accordingly.",
         "High"),
        ("UR-10", "Buyer / Brand Owner",
         "Both buyers and brand owners shall be able to view the current status of all their "
         "respective orders and monitor all status transitions (pending, approved, completed, "
         "rejected) over the full order lifecycle.",
         "High"),
        ("UR-11", "Buyer",
         "A buyer shall be able to submit a star rating (1–5 stars) and written review for any "
         "brand they have interacted with, with the option to submit the review anonymously.",
         "Medium"),
        ("UR-12", "Brand Owner",
         "A brand owner shall be able to post a single written reply to any review submitted for "
         "their brand, visible to all platform visitors on the brand's public profile.",
         "Medium"),
        ("UR-13", "Admin",
         "A platform administrator shall be able to view all submitted reviews, flag reviews "
         "for investigation, and permanently remove reviews that violate the platform's "
         "community guidelines.",
         "Medium"),
        ("UR-14", "All Users",
         "All authenticated users shall receive in-app notifications for events relevant to their "
         "role — including new orders, order status changes, brand approval decisions, and new "
         "reviews — with the ability to mark individual notifications as read.",
         "High"),
        ("UR-15", "All Users",
         "The platform interface shall be fully available in both English (LTR layout) and Arabic "
         "(RTL layout), with users able to toggle between languages at any time without losing "
         "current page state.",
         "High"),
        ("UR-16", "Brand Owner",
         "A brand owner shall be able to import a product catalogue from a Microsoft Excel (.xlsx) "
         "file, with the system parsing the file and creating product records from its contents "
         "in a single bulk operation.",
         "Low"),
        ("UR-17", "Buyer",
         "A buyer shall be able to contact a brand owner directly via WhatsApp by clicking a "
         "platform-provided contact link, and the system shall record each such click event per "
         "brand for analytics purposes.",
         "Low"),
    ],
    col_widths=[0.7, 1.3, 3.8, 0.7]
)

# ── §2.3 System Requirements ─────────────────────────────────────────────────

h2(doc, "2.3  System Requirements")

body(doc,
    "System Requirements (SRs) specify the observable behaviour that the software must exhibit in "
    "response to inputs, events, or internal states. The following 20 system requirements are "
    "derived from the user requirements above and the non-functional constraints of the Kuwait B2B "
    "operating context."
)

make_table(doc,
    ["SR ID", "Requirement Statement", "Priority"],
    [
        ("SR-01",
         "The system shall enforce role-based access control (RBAC) for all page routes and API "
         "endpoints, rejecting unauthorised requests with HTTP 401 (unauthenticated) or "
         "HTTP 403 (authenticated but insufficient role) before executing any business logic.",
         "High"),
        ("SR-02",
         "The system shall manage user sessions exclusively via HTTP-only, SameSite=Lax cookies "
         "with a maximum age of seven days. Session tokens shall never be stored in browser "
         "localStorage or sessionStorage.",
         "High"),
        ("SR-03",
         "The system shall hash all user passwords using a cryptographic one-way function "
         "(bcrypt or argon2id, cost factor ≥ 12) before persisting them to the database. "
         "Plain-text passwords shall never be stored or logged.",
         "Critical"),
        ("SR-04",
         "The system shall generate globally unique identifiers for all entity records (users, "
         "brands, products, orders, reviews, notifications, messages) using a deterministic "
         "algorithm that prevents ID collisions under concurrent creation.",
         "High"),
        ("SR-05",
         "The system shall validate that each order line item's requested quantity meets or exceeds "
         "the product's configured minimum order quantity (MOQ) at submission time, returning a "
         "descriptive HTTP 400 error if the constraint is violated.",
         "High"),
        ("SR-06",
         "The system shall validate that each order line item's requested quantity does not exceed "
         "the product's current available stock level at submission time, returning a "
         "descriptive HTTP 400 error if the constraint is violated.",
         "High"),
        ("SR-07",
         "The system shall resolve the applicable unit price for each order line item by iterating "
         "the product's bulk-pricing tiers, matching the order quantity to the appropriate "
         "min_qty / max_qty range, and applying that tier's unit price. If no tier matches, "
         "the base product price shall be used.",
         "High"),
        ("SR-08",
         "The system shall compute and persist a Brand Trust Score for each approved brand, derived "
         "from the brand's order response rate, completion rate, average fulfilment days, and "
         "badge threshold criteria.",
         "Medium"),
        ("SR-09",
         "The system shall compute and persist a Buyer Trust Score for each buyer, derived from the "
         "buyer's total completed order count, order completion rate, and cancellation rate.",
         "Medium"),
        ("SR-10",
         "The system shall generate in-app notifications for the appropriate recipient(s) upon each "
         "of the following events: new user registration (→ admins), new brand registration "
         "(→ admins), brand approval / rejection (→ brand owner), new order (→ brand owner), "
         "order status change (→ buyer), new review (→ brand owner), new review reply (→ buyer).",
         "High"),
        ("SR-11",
         "The system shall increment a WhatsApp click counter on the brand's database record each "
         "time a user activates the platform's WhatsApp contact link for that brand.",
         "Low"),
        ("SR-12",
         "The system shall render all user-facing text from locale-specific message files "
         "(en.json and ar.json), with Arabic strings displayed in right-to-left layout when the "
         "Arabic locale is active.",
         "High"),
        ("SR-13",
         "The system shall display all product prices in the buyer's currently selected currency "
         "by applying conversion rates from the centralised CURRENCIES configuration registry "
         "in src/lib/currencies.ts.",
         "Medium"),
        ("SR-14",
         "The system shall restrict all brand verification actions (status update, tier assignment) "
         "exclusively to users with the administrator role, returning HTTP 403 for any attempt by "
         "a non-admin user.",
         "High"),
        ("SR-15",
         "The system shall parse .xlsx files submitted via the bulk catalogue import endpoint, "
         "extracting product field values (name, description, price, MOQ, stock, image URL) from "
         "a defined column layout and creating corresponding product records in a single operation.",
         "Low"),
        ("SR-16",
         "The system shall provide an administrator dashboard displaying real-time platform "
         "statistics: total registered user count, total approved brand count, total order volume, "
         "pending brand approval queue length, and active review moderation queue.",
         "Medium"),
        ("SR-17",
         "The system shall record a timestamped audit entry for each brand verification decision, "
         "capturing the previous status, the new status, the acting administrator's ID, and the "
         "ISO 8601 timestamp of the change.",
         "Medium"),
        ("SR-18",
         "The system shall support assignment of one of three verification tiers — Premium, "
         "Verified, or New — to each approved brand, and shall display the assigned tier as a "
         "visual badge on all brand-facing UI elements.",
         "Medium"),
        ("SR-19",
         "The system shall accept review submissions with an anonymity flag that, when true, "
         "suppresses the reviewer's name and company from all public-facing display of that review.",
         "Medium"),
        ("SR-20",
         "The system shall store a single brand reply per review and display it alongside the "
         "original review in the brand's public profile, attributed to the brand rather than "
         "to an individual named user.",
         "Medium"),
    ],
    col_widths=[0.7, 5.2, 0.7]
)

# ── §2.4 Non-Functional Requirements ────────────────────────────────────────

h2(doc, "2.4  Non-Functional Requirements")

body(doc,
    "Non-Functional Requirements (NFRs) define quality attributes that constrain how the system "
    "performs its functions. The following eight NFRs govern the Kuwait B2B Hub across the dimensions "
    "of performance, security, usability, reliability, maintainability, scalability, "
    "internationalisation, and compliance."
)

make_table(doc,
    ["NFR ID", "Category", "Requirement Statement", "Priority"],
    [
        ("NFR-01", "Performance",
         "The median API response time for all standard read operations shall be under 300 ms under "
         "a load of up to 50 concurrent users. The Marketplace page shall achieve a Largest "
         "Contentful Paint (LCP) of under 3 seconds on a 4G mobile connection.",
         "High"),
        ("NFR-02", "Security",
         "All session cookies shall carry HttpOnly=true, SameSite=Lax, and Secure=true flags. "
         "All network traffic shall be encrypted over HTTPS. User passwords shall be stored using "
         "bcrypt or argon2id with a cost factor of ≥ 12. No credential data shall be written to "
         "any client-accessible storage mechanism.",
         "Critical"),
        ("NFR-03", "Usability",
         "All primary user workflows (registration, marketplace browsing, order submission, order "
         "tracking, notification viewing) shall be fully operable on a mobile viewport of 375 px "
         "width without horizontal scrolling. All interactive elements shall have a touch target "
         "of at least 44 × 44 px.",
         "High"),
        ("NFR-04", "Reliability",
         "The production system shall target a monthly uptime of ≥ 99.5%, leveraging Vercel's "
         "global edge network SLA and Supabase's managed infrastructure availability guarantees. "
         "Unplanned downtime shall not exceed 3.6 hours in any single calendar month.",
         "High"),
        ("NFR-05", "Maintainability",
         "The codebase shall pass ESLint (eslint-config-next) with zero errors on every "
         "production build. TypeScript strict mode shall remain enabled. All API route handlers "
         "shall invoke the getCurrentUser() guard as their first operation before executing "
         "any business logic.",
         "High"),
        ("NFR-06", "Scalability",
         "The data access layer shall be fully encapsulated behind the readDB() / writeDB() "
         "interface, such that migration from the JSON prototype store to Supabase PostgreSQL "
         "requires no modifications to any API route handler implementation.",
         "High"),
        ("NFR-07", "Internationalisation",
         "All user-facing string literals shall be externalised to locale-specific message files "
         "(en.json and ar.json). No hard-coded English or Arabic strings shall appear in .tsx "
         "component files. The platform shall render correctly in Arabic RTL layout on all "
         "supported pages.",
         "High"),
        ("NFR-08", "Compliance",
         "The platform shall not process, transmit, or store raw payment card data under any "
         "circumstances. All in-platform payment flows shall be delegated to a PCI-DSS-certified "
         "gateway aggregator. User data collection shall be limited to the minimum necessary for "
         "platform operation.",
         "Critical"),
    ],
    col_widths=[0.8, 1.2, 4.1, 0.8]
)

# ── §2.5 FRS ─────────────────────────────────────────────────────────────────

h2(doc, "2.5  Functional Requirements Specification (FRS)")

body(doc,
    "The following section presents the complete Functional Requirements Specification across the "
    "five platform subsystems. Each of the 25 functions is described with its input, processing "
    "logic, output, and priority classification."
)

# 2.5.1
h3(doc, "2.5.1  Subsystem 1 — Authentication & User Management")

make_table(doc,
    ["ID", "Function", "Input", "Process", "Output", "Priority"],
    [
        ("F-001", "User Registration",
         "POST /api/auth/register: {name, email, password, role, company_name}",
         "Validate all fields present; check email uniqueness (409 if duplicate); create user "
         "record; if brand_owner → create pending brand + notify admins; set "
         "verification_status = approved (buyer) or pending (brand_owner); create session cookies.",
         "201 user object + two session cookies (b2b_user_id httpOnly, b2b_user_role); "
         "admin notification(s) created.",
         "High"),
        ("F-002", "User Login",
         "POST /api/auth/login: {email, password}",
         "Validate fields present; find user by email; compare credential; if match → set "
         "session cookie pair with 7-day maxAge.",
         "200 user object + two session cookies; or 401 Invalid credentials.",
         "High"),
        ("F-003", "User Logout",
         "POST /api/auth/logout (authenticated session required)",
         "Expire b2b_user_id and b2b_user_role cookies by setting maxAge=0 on NextResponse.",
         "200 success; both cookies cleared in client browser.",
         "High"),
        ("F-004", "Session Validation",
         "GET /api/auth/me (cookie-authenticated)",
         "Read b2b_user_id cookie; locate user profile in database; return profile or "
         "HTTP 401 if not found.",
         "200 user profile object; or 401 Unauthorized.",
         "High"),
        ("F-005", "User Management (Admin)",
         "GET /api/admin/users; PATCH /api/admin/users/:id: {role?, verification_status?}",
         "RBAC guard confirms admin role; return full user list (GET) or apply field updates "
         "(PATCH) to specified user record.",
         "200 user list or updated user object; 403 Forbidden if non-admin.",
         "Medium"),
    ],
    col_widths=[0.6, 1.1, 1.4, 2.0, 1.4, 0.7]
)

# 2.5.2
h3(doc, "2.5.2  Subsystem 2 — Brand Management")

make_table(doc,
    ["ID", "Function", "Input", "Process", "Output", "Priority"],
    [
        ("F-006", "Brand Profile Creation",
         "Triggered automatically on brand_owner registration (F-001)",
         "Create brand record: status='pending', owner_id=userId, logo_url from UI-Avatars API. "
         "Persist to database.",
         "Brand record created; admin notified of new brand awaiting approval.",
         "High"),
        ("F-007", "Brand Profile Update",
         "PATCH /api/brands/:id: {brand_name?, description?, logo_url?, whatsapp_number?, "
         "business_hours?, auto_reply_message?} (brand_owner)",
         "Validate brand.owner_id = current user; update specified fields; persist.",
         "200 updated brand object; 403 if not owner.",
         "High"),
        ("F-008", "Brand Verification",
         "PATCH /api/admin/brands/:id: {status, verification_tier?} (admin only)",
         "Validate admin role; validate status ∈ {approved, rejected}; update brand.status and "
         "owner.verification_status atomically; optionally set tier; notify brand owner and "
         "(on approval) all buyers.",
         "200 updated brand; notifications dispatched; 400/403/404 on error.",
         "High"),
        ("F-009", "WhatsApp Click Tracking",
         "POST /api/brands/whatsapp-click: {brand_id}",
         "Increment brand.whatsapp_clicks counter in database.",
         "200 success with updated click count.",
         "Low"),
        ("F-010", "Brand Discovery",
         "GET /api/brands: {search?, category?, minRating?, verification_tier?}",
         "Return all brands with status='approved'; apply optional filter parameters; sort by "
         "rating or relevance.",
         "200 array of approved brand objects matching filters.",
         "High"),
    ],
    col_widths=[0.6, 1.1, 1.4, 2.0, 1.4, 0.7]
)

# 2.5.3
h3(doc, "2.5.3  Subsystem 3 — Product & Catalogue Management")

make_table(doc,
    ["ID", "Function", "Input", "Process", "Output", "Priority"],
    [
        ("F-011", "Product Creation",
         "POST /api/products: {brand_id, name, description, price, moq, stock, image_url, "
         "bulk_pricing_tiers[]} (brand_owner)",
         "Validate brand ownership; create product record with all specified fields; persist "
         "bulk_pricing_tiers[] as embedded array.",
         "201 created product object.",
         "High"),
        ("F-012", "Product Edit & Delete",
         "PATCH/DELETE /api/products/:id (brand_owner)",
         "Validate brand ownership for the product's parent brand; update specified fields "
         "(PATCH) or remove record (DELETE).",
         "200 updated product; 204 on delete; 403/404 on error.",
         "High"),
        ("F-013", "Bulk Pricing Tier Management",
         "Included in F-011 / F-012: "
         "bulk_pricing_tiers: [{min_qty, max_qty, price}]",
         "Validate tier ranges are logically valid (non-overlapping); replace existing tier "
         "array with the submitted one on each product update.",
         "Updated product with new bulk_pricing_tiers[] array persisted.",
         "High"),
        ("F-014", "Product Browsing",
         "GET /api/products?brand_id=X or GET /api/products/:id",
         "Return all products for a brand (filtered by brand_id) or a single product by ID, "
         "including bulk_pricing_tiers.",
         "200 array of products or single product object.",
         "High"),
        ("F-015", "Excel Catalogue Import",
         "POST /api/products/bulk: multipart .xlsx file upload (brand_owner)",
         "Parse .xlsx using xlsx library; extract product rows; validate each row against "
         "required fields; create product records in batch; report outcome.",
         "200 import summary: {created: N, errors: [{row, reason}]}.",
         "Low"),
    ],
    col_widths=[0.6, 1.1, 1.4, 2.0, 1.4, 0.7]
)

# 2.5.4
h3(doc, "2.5.4  Subsystem 4 — Order Management")

make_table(doc,
    ["ID", "Function", "Input", "Process", "Output", "Priority"],
    [
        ("F-016", "Order Submission",
         "POST /api/orders: {brand_id, items: [{product_id, quantity}]} (buyer)",
         "Validate buyer role; locate approved brand; for each item: check product belongs to "
         "brand, validate qty ≥ MOQ (SR-05), validate qty ≤ stock (SR-06), resolve bulk-pricing "
         "unit_price (SR-07); compute total_amount; INSERT order + order_items; notify brand owner.",
         "201 created order with items; brand owner notification dispatched.",
         "High"),
        ("F-017", "Order Status Update",
         "PATCH /api/orders/:id: {status: approved|rejected|completed} "
         "(brand_owner or admin)",
         "Validate role; brand_owner: verify order.brand_id matches owner's brand; update "
         "order.status; create buyer notification based on new status value.",
         "200 updated order; buyer notification; 403/404 on error.",
         "High"),
        ("F-018", "Order Listing",
         "GET /api/orders (authenticated)",
         "Apply role-based data filter: buyers → own orders only; brand_owners → orders for "
         "their brand only; admins → all orders.",
         "200 array of order objects including items.",
         "High"),
        ("F-019", "Order Detail View",
         "GET /api/orders/:id (authenticated)",
         "Return single order with full item detail; apply role-based ownership check before "
         "returning data.",
         "200 order object with items; 403 if not authorised to view.",
         "High"),
        ("F-020", "Admin Order Analytics",
         "GET /api/admin/orders (admin only)",
         "Validate admin role; aggregate platform-wide order statistics: total count, total "
         "revenue, status distribution, recent order list.",
         "200 analytics payload; 403 if non-admin.",
         "Medium"),
    ],
    col_widths=[0.6, 1.1, 1.4, 2.0, 1.4, 0.7]
)

# 2.5.5
h3(doc, "2.5.5  Subsystem 5 — Trust, Reviews & Notifications")

make_table(doc,
    ["ID", "Function", "Input", "Process", "Output", "Priority"],
    [
        ("F-021", "Review Submission",
         "POST /api/reviews: {brand_id, order_id?, rating, content, anonymous} (buyer)",
         "Validate buyer role; validate rating ∈ [1, 5]; create review with status='active' and "
         "anonymous flag; persist; notify brand owner.",
         "201 created review object; brand owner notification.",
         "Medium"),
        ("F-022", "Brand Reply to Review",
         "PATCH /api/reviews/:id/reply: {brand_reply} (brand_owner)",
         "Validate brand ownership (review.brand_id = current user's brand); update "
         "review.brand_reply field; persist; notify review author.",
         "200 updated review with brand_reply text.",
         "Medium"),
        ("F-023", "Admin Review Moderation",
         "PATCH /api/admin/reviews/:id: {status?, flagged?} (admin)",
         "Validate admin role; update review.status (active / removed) and/or review.flagged "
         "boolean as specified.",
         "200 updated review object; 403 if non-admin.",
         "Medium"),
        ("F-024", "Trust Score Computation",
         "GET /api/trust/:id or triggered on order status change",
         "Calculate BrandTrust: response_rate = responded / total; completion_rate = completed "
         "/ total; avg_fulfillment_days from order timestamps; assign badge tier by threshold. "
         "Calculate BuyerTrust: completion_rate, cancellation_rate from order history.",
         "200 BrandTrust or BuyerTrust object; persisted to database.",
         "Medium"),
        ("F-025", "Notification Management",
         "GET /api/notifications; PATCH /api/notifications/:id: {read: true}",
         "GET: return all notifications for current user ordered by created_at DESC. "
         "PATCH: mark specified notification record as read.",
         "200 notification array (GET); 200 updated notification (PATCH).",
         "High"),
    ],
    col_widths=[0.6, 1.1, 1.4, 2.0, 1.4, 0.7]
)

# ── §2.6 Context Model ───────────────────────────────────────────────────────

h2(doc, "2.6  Context Model")

body(doc,
    "The Context Model defines the boundary of the Kuwait B2B Hub system and identifies the external "
    "entities that interact with it across that boundary. The system boundary encompasses all "
    "Next.js App Router page routes, the 25 serverless API route handlers, the Zustand notification "
    "store, the currency configuration module, and the data access abstraction layer. Everything "
    "outside this boundary is an external entity that communicates with the system via defined "
    "interfaces."
)

body(doc,
    "The system interfaces with four external entity categories, each described in the table below. "
    "Data flows across the boundary include: user input (HTTP requests from browsers), platform "
    "events (notifications dispatched to the notification service), authentication tokens "
    "(JWT sessions exchanged with Supabase Auth), financial exchange data (currency rates from "
    "an external API), and deployment and serving (all static and dynamic assets served via "
    "Vercel's edge network)."
)

make_table(doc,
    ["External Entity", "Role & Interaction", "Interface"],
    [
        ("Business Buyers",
         "Human actors who register accounts, browse brand catalogues, submit bulk orders, "
         "track order status, submit reviews, and view notifications. They interact exclusively "
         "via the browser-rendered Next.js UI.",
         "HTTPS (browser → Vercel edge → Next.js SSR + API routes)"),
        ("Brand Owners",
         "Human actors who register brand profiles, manage product catalogues, process incoming "
         "orders, reply to reviews, and view brand analytics dashboards. Interact via the "
         "/brand/* portal routes.",
         "HTTPS (browser → Vercel edge → Next.js SSR + API routes)"),
        ("Admin Users",
         "Internal platform operators who approve/reject brands, moderate reviews, manage user "
         "accounts, and access platform-wide analytics. Confined to the /admin/* route namespace.",
         "HTTPS (browser → Vercel edge → /admin/* routes)"),
        ("External Services",
         "Four cloud platforms the system depends on: (1) Supabase — provides JWT session "
         "validation via supabase.auth.getUser() and managed PostgreSQL storage; "
         "(2) Vercel — serverless execution environment and global CDN; "
         "(3) WhatsApp Business — receives buyer contact intent via platform-generated deep links; "
         "(4) Exchange Rate API — supplies live GCC currency rates consumed by CurrencyContext.",
         "HTTPS/WebSocket (Supabase), Git webhook (Vercel CI/CD), URI deep-link (WhatsApp), "
         "REST (Exchange Rate API)"),
    ],
    col_widths=[1.4, 3.3, 1.8]
)

# ── §2.7 Use Cases ───────────────────────────────────────────────────────────

h2(doc, "2.7  Use Cases")

body(doc,
    "The table below summarises all 16 use cases identified for the Kuwait B2B Hub. Detailed "
    "specifications for five representative use cases (UC-01, UC-03, UC-07, UC-12, UC-14) are "
    "provided in §2.7.1–2.7.5."
)

make_table(doc,
    ["UC ID", "Use Case Name", "Primary Actor", "Brief Description", "Priority"],
    [
        ("UC-01", "Register Account",            "Buyer / Brand Owner",
         "User creates a new platform account with role selection and receives a session.",          "High"),
        ("UC-02", "Login & Logout",               "All Users",
         "Authenticated user establishes or terminates a persistent session.",                       "High"),
        ("UC-03", "Browse & Search Brands",       "Buyer",
         "Buyer discovers verified brands on the marketplace using search and filter controls.",     "High"),
        ("UC-04", "View Product Catalogue",       "Buyer",
         "Buyer views a brand's product listing with pricing and bulk-pricing tiers.",              "High"),
        ("UC-05", "Submit Bulk Order",            "Buyer",
         "Buyer submits an order request; system enforces MOQ, stock, and pricing rules.",           "High"),
        ("UC-06", "Track Order Status",           "Buyer / Brand Owner",
         "User monitors the full order lifecycle from pending through to completion.",               "High"),
        ("UC-07", "Verify Brand (Admin)",         "Admin",
         "Admin reviews a pending brand application, approves or rejects, and sets tier.",           "High"),
        ("UC-08", "Approve / Reject Order",       "Brand Owner",
         "Brand owner reviews an incoming order and updates its status.",                            "High"),
        ("UC-09", "Manage Product Listings",      "Brand Owner",
         "Brand owner creates, edits, or deletes product entries in their catalogue.",               "High"),
        ("UC-10", "Bulk Import Catalogue",        "Brand Owner",
         "Brand owner uploads an Excel file; system creates product records in batch.",              "Low"),
        ("UC-11", "Submit Review",                "Buyer",
         "Buyer submits a star rating and written review for a brand, optionally anonymously.",      "Medium"),
        ("UC-12", "Reply to Review",              "Brand Owner",
         "Brand owner posts a written response to a review on their brand profile.",                 "Medium"),
        ("UC-13", "Moderate Review",              "Admin",
         "Admin removes or flags a review that violates platform community guidelines.",             "Medium"),
        ("UC-14", "Manage Notifications",         "All Users",
         "User opens the notification panel and marks notifications as read.",                       "High"),
        ("UC-15", "Switch Language & Currency",   "All Users",
         "User toggles the interface language (EN/AR) and selects a display currency.",             "Medium"),
        ("UC-16", "View Trust Scores & Badges",   "Buyer / Brand Owner",
         "User inspects trust metrics and earned reputation badges for brands and buyers.",          "Medium"),
    ],
    col_widths=[0.7, 1.6, 1.3, 2.6, 0.7]
)

# ── §2.7.1 UC-01 ─────────────────────────────────────────────────────────────

h3(doc, "2.7.1  Use Case Specification: UC-01 — Register Account")

uc_spec(doc, [
    ("Use Case ID",      "UC-01"),
    ("Use Case Name",    "Register Account"),
    ("Primary Actor",    "Buyer / Brand Owner"),
    ("Secondary Actors", "Platform Notification Service; Admin Users (notified on completion)"),
    ("Trigger",          "Unauthenticated user navigates to /register and submits the registration form"),
    ("Pre-conditions",   "No existing account uses the provided email address. The platform is accessible."),
    ("Main Flow",
     "1. User selects a role: Buyer or Brand Owner.\n"
     "2. User enters full name, email address, password, and company name.\n"
     "3. System validates all five fields are present and non-empty (HTTP 400 if not).\n"
     "4. System checks that no existing user record has the provided email (HTTP 409 if duplicate).\n"
     "5. System creates a user record with the submitted details.\n"
     "6. If role = brand_owner: system creates a brand record with status='pending' and "
        "owner_id = new user's ID.\n"
     "7. System sets verification_status = 'approved' for buyers; 'pending' for brand owners.\n"
     "8. System sets two session cookies on the response (b2b_user_id httpOnly, b2b_user_role) "
        "with a 7-day maxAge.\n"
     "9. System dispatches admin notification: 'New User Registration'.\n"
     "10. If brand_owner: system dispatches second admin notification: 'New Brand Awaiting Approval'."),
    ("Alternative Flow A", "Step 4: Email already registered → HTTP 409 Conflict; user redirected to login page."),
    ("Alternative Flow B", "Step 3: Any required field empty → HTTP 400 Bad Request; form field error displayed."),
    ("Post-conditions",
     "New user record exists in the database. If brand_owner: a pending brand record also exists. "
     "At least one admin notification created. User is logged in with an active 7-day session."),
    ("Priority",         "High"),
    ("Related Requirements", "SR-01, SR-02, SR-03, SR-04, SR-10; UR-01"),
])

# ── §2.7.2 UC-03 ─────────────────────────────────────────────────────────────

h3(doc, "2.7.2  Use Case Specification: UC-03 — Browse & Search Brands")

uc_spec(doc, [
    ("Use Case ID",      "UC-03"),
    ("Use Case Name",    "Browse & Search Brands"),
    ("Primary Actor",    "Buyer"),
    ("Secondary Actors", "None"),
    ("Trigger",          "Buyer navigates to /marketplace"),
    ("Pre-conditions",   "At least one brand with status='approved' exists. Buyer is authenticated."),
    ("Main Flow",
     "1. Buyer accesses the Marketplace page via the navigation bar.\n"
     "2. System calls GET /api/brands and retrieves all brands with status='approved'.\n"
     "3. System renders brand cards, each showing brand name, logo, verification tier badge, "
        "average star rating, and minimum order quantity.\n"
     "4. Buyer optionally enters a keyword in the search field; system filters brands by name "
        "and description in real time.\n"
     "5. Buyer optionally selects one or more filters: category, verification tier, "
        "minimum rating, minimum MOQ.\n"
     "6. System updates the displayed brand list instantly as each filter is applied.\n"
     "7. Buyer clicks a brand card to navigate to that brand's full profile page (/brands/:id)."),
    ("Alternative Flow A", "Step 2: No approved brands exist → system displays an empty-state message with a prompt."),
    ("Alternative Flow B",
     "Step 5: Applied filters yield zero matching brands → system displays 'No results found' "
     "with a 'Clear filters' affordance."),
    ("Post-conditions",  "Buyer is viewing the filtered brand marketplace. No database state is modified."),
    ("Priority",         "High"),
    ("Related Requirements", "SR-01, SR-13, SR-18; UR-06, UR-07"),
])

# ── §2.7.3 UC-07 ─────────────────────────────────────────────────────────────

h3(doc, "2.7.3  Use Case Specification: UC-07 — Verify Brand (Admin)")

uc_spec(doc, [
    ("Use Case ID",      "UC-07"),
    ("Use Case Name",    "Verify Brand (Admin)"),
    ("Primary Actor",    "Admin"),
    ("Secondary Actors", "Brand Owner (notified); All active Buyers (notified on approval)"),
    ("Trigger",          "Admin navigates to the Admin Dashboard and opens the Pending Brands queue"),
    ("Pre-conditions",
     "At least one brand with status='pending' exists. Admin is authenticated with the admin role."),
    ("Main Flow",
     "1. Admin navigates to Admin Dashboard → Brands → Pending queue.\n"
     "2. System displays all brands with status='pending', sorted by registration date.\n"
     "3. Admin selects a brand to review; system displays brand details and owner information.\n"
     "4. Admin optionally assigns a verification tier (Premium, Verified, or New) from a dropdown.\n"
     "5. Admin clicks 'Approve'.\n"
     "6. System sends PATCH /api/admin/brands/:id {status: 'approved', verification_tier?}.\n"
     "7. System validates admin role (HTTP 403 if not admin).\n"
     "8. System updates brand.status = 'approved' and brand_owner.verification_status = 'approved' "
        "atomically in the database.\n"
     "9. System creates an 'account_approved' notification for the brand owner.\n"
     "10. System creates a 'new_brand' notification for all active buyer accounts.\n"
     "11. Brand becomes visible on the marketplace."),
    ("Alternative Flow A",
     "Step 5: Admin clicks 'Reject' → system sends status='rejected'; updates both records; "
     "creates 'account_rejected' notification for brand owner; skips buyer notifications."),
    ("Alternative Flow B",
     "Step 7: status value not in {approved, rejected} → system returns HTTP 400 Bad Request."),
    ("Alternative Flow C",
     "Step 7: brand ID not found → system returns HTTP 404 Not Found."),
    ("Post-conditions",
     "Brand and brand owner verification statuses updated. Notifications dispatched. "
     "Approved brand visible on marketplace; rejected brand hidden."),
    ("Priority",         "High"),
    ("Related Requirements", "SR-01, SR-10, SR-14, SR-17, SR-18; UR-03"),
])

# ── §2.7.4 UC-12 ─────────────────────────────────────────────────────────────

h3(doc, "2.7.4  Use Case Specification: UC-12 — Reply to Review")

uc_spec(doc, [
    ("Use Case ID",      "UC-12"),
    ("Use Case Name",    "Reply to Review"),
    ("Primary Actor",    "Brand Owner"),
    ("Secondary Actors", "Original Reviewer / Buyer (notified)"),
    ("Trigger",          "Brand owner views a review on their brand profile and chooses to respond"),
    ("Pre-conditions",
     "A review with status='active' exists for the brand. Brand owner is authenticated. "
     "The review does not already have a brand reply (or the owner wishes to update an existing one)."),
    ("Main Flow",
     "1. Brand owner navigates to their Brand Dashboard → Reviews section.\n"
     "2. System displays all reviews for the brand, with rating, content, date, and reviewer "
        "identity (suppressed if anonymous).\n"
     "3. Brand owner locates the review to respond to and clicks 'Reply'.\n"
     "4. Brand owner enters reply text in the reply input field.\n"
     "5. Brand owner submits via PATCH /api/reviews/:id/reply {brand_reply: 'text'}.\n"
     "6. System validates that the review's brand_id matches the authenticated user's brand "
        "(HTTP 403 if not).\n"
     "7. System updates review.brand_reply with the submitted text and persists the change.\n"
     "8. System creates a notification for the original reviewer: 'Brand replied to your review'."),
    ("Alternative Flow A",
     "Step 6: Review does not belong to this brand owner → HTTP 403 Forbidden; no change persisted."),
    ("Alternative Flow B",
     "Step 3: Review already has a brand_reply → existing reply is overwritten with new text "
     "(most recent reply is displayed)."),
    ("Post-conditions",
     "review.brand_reply updated in the database. Reviewer notification created. "
     "Reply visible on the brand's public profile page."),
    ("Priority",         "Medium"),
    ("Related Requirements", "SR-01, SR-10, SR-19, SR-20; UR-12"),
])

# ── §2.7.5 UC-14 ─────────────────────────────────────────────────────────────

h3(doc, "2.7.5  Use Case Specification: UC-14 — Manage Notifications")

uc_spec(doc, [
    ("Use Case ID",      "UC-14"),
    ("Use Case Name",    "Manage Notifications"),
    ("Primary Actor",    "All Users (Buyer / Brand Owner / Admin)"),
    ("Secondary Actors", "None"),
    ("Trigger",          "User clicks the notification bell icon in the top navigation bar"),
    ("Pre-conditions",   "User is authenticated. One or more notifications exist for the current user."),
    ("Main Flow",
     "1. User clicks the notification bell icon (NotificationBell.tsx).\n"
     "2. System calls GET /api/notifications and fetches all notifications for the current user, "
        "ordered by created_at descending.\n"
     "3. System renders the notification panel, displaying title, body, timestamp, and "
        "read/unread indicator for each item.\n"
     "4. System displays the unread count as a badge on the bell icon.\n"
     "5. User clicks a specific notification item.\n"
     "6. System sends PATCH /api/notifications/:id {read: true}.\n"
     "7. System updates the notification record and decrements the unread count badge.\n"
     "8. If the notification carries an action_url, the system navigates the user to that page."),
    ("Alternative Flow A",
     "Step 2: No notifications exist → panel displays 'No notifications yet' empty-state message."),
    ("Alternative Flow B",
     "User activates 'Mark all as read' control → system sends batch PATCH requests for all "
     "unread notifications; badge count resets to zero."),
    ("Post-conditions",
     "Selected notification(s) marked as read in the database. Unread count badge updated in UI. "
     "If action_url present: user navigated to the relevant platform page."),
    ("Priority",         "High"),
    ("Related Requirements", "SR-01, SR-10; F-025; UR-14"),
])

# ── §2.7.6 Preliminary System Architecture ───────────────────────────────────

h3(doc, "2.7.6  Preliminary System Architecture")

body(doc,
    "The Kuwait B2B Hub is structured as a four-layer architecture that enforces a strict separation "
    "of concerns between interface rendering, API-mediated business logic, data access abstraction, "
    "and external service integration. The architecture was designed to permit independent evolution "
    "of each layer — in particular, to allow the data access layer to migrate from the JSON "
    "prototype store to Supabase PostgreSQL without any changes to the API route handlers above it."
)

make_table(doc,
    ["Layer", "Responsibilities", "Key Components", "Runtime"],
    [
        ("L1 — Presentation",
         "Renders SSR HTML pages via Next.js App Router; provides client-side interactivity "
         "via React hydration; manages notification and currency state client-side.",
         "24 UI components (src/components/*); 15 page routes (src/app/**/*page.tsx); "
         "Zustand notificationStore; CurrencyContext",
         "Browser + Next.js SSR (Vercel)"),
        ("L2 — Application / API",
         "Processes all business logic; enforces RBAC via getCurrentUser() guard pattern; "
         "validates inputs; orchestrates data reads/writes and notification dispatch.",
         "25 API route handlers (src/app/api/**/*route.ts); src/lib/auth.ts; "
         "src/lib/notifications.ts; src/middleware.ts",
         "Vercel Serverless Functions"),
        ("L3 — Data Access",
         "Abstracts all database operations behind a uniform interface; isolates API handlers "
         "from the underlying storage technology.",
         "src/lib/db.ts: readDB(), writeDB(), generateId(); "
         "src/data/db.json (MVP) → Supabase PostgreSQL (production)",
         "Node.js file I/O (MVP) / Supabase SDK (production)"),
        ("L4 — External Services",
         "Provides authentication (Supabase Auth), hosting and CDN (Vercel), "
         "buyer-to-brand communication (WhatsApp), and live currency rates (Exchange Rate API).",
         "src/utils/supabase/server.ts; src/lib/currencies.ts; "
         "Vercel edge network; WhatsApp Business deep-links",
         "Supabase Cloud, Vercel Edge, WhatsApp, External REST API"),
    ],
    col_widths=[1.2, 2.2, 2.0, 1.1]
)

# ── §2.8 RTM ─────────────────────────────────────────────────────────────────

h2(doc, "2.8  Requirements Traceability Matrix (RTM)")

body(doc,
    "The following RTM maps each of the 20 system requirements to its originating source, the "
    "design artefact that realises it, the primary implementation file(s), the current "
    "implementation status, an associated test case identifier, and the test result. This matrix "
    "provides full traceability from requirement statement to delivery evidence."
)

make_table(doc,
    ["Req ID", "Requirement Summary", "Priority", "Source", "Design Artefact",
     "Implementation File", "Status", "TC ID", "Test Result"],
    [
        ("SR-01", "RBAC on all routes/endpoints",        "High",     "UC-01–16",    "RBAC guard pattern",           "src/middleware.ts; all route.ts files",         "Complete",       "TC-01", "Pass"),
        ("SR-02", "HTTP-only session cookies",           "High",     "UC-01, UC-02","Cookie spec (httpOnly, Lax)",   "api/auth/login/route.ts; register/route.ts",    "Complete",       "TC-02", "Pass"),
        ("SR-03", "Password hashing (bcrypt/argon2)",    "Critical", "UC-01, UC-02","Auth architecture",            "api/auth/register/route.ts",                   "Partial",        "TC-03", "Blocked"),
        ("SR-04", "Unique ID generation",                "High",     "UC-01–16",    "generateId() utility",         "src/lib/db.ts",                                "Complete",       "TC-04", "Pass"),
        ("SR-05", "MOQ validation on order submit",      "High",     "UC-05",       "Order submission flow",        "api/orders/route.ts (POST)",                   "Complete",       "TC-05", "Pass"),
        ("SR-06", "Stock validation on order submit",    "High",     "UC-05",       "Order submission flow",        "api/orders/route.ts (POST)",                   "Complete",       "TC-06", "Pass"),
        ("SR-07", "Bulk pricing tier resolution",        "High",     "UC-05",       "Pricing engine",               "api/orders/route.ts (POST)",                   "Complete",       "TC-07", "Pass"),
        ("SR-08", "Brand Trust Score computation",       "Medium",   "UC-16",       "BrandTrust schema",            "api/trust/[id]/route.ts",                      "Partial (70%)",  "TC-08", "Manual"),
        ("SR-09", "Buyer Trust Score computation",       "Medium",   "UC-16",       "BuyerTrust schema",            "api/trust/[id]/route.ts",                      "Partial (70%)",  "TC-09", "Manual"),
        ("SR-10", "In-app notifications for all events", "High",     "UC-14",       "Notification service",         "src/lib/notifications.ts",                     "Complete",       "TC-10", "Pass"),
        ("SR-11", "WhatsApp click tracking",             "Low",      "UC-15",       "Click counter endpoint",       "api/brands/whatsapp-click/route.ts",            "Complete",       "TC-11", "Pass"),
        ("SR-12", "Arabic RTL localisation",             "High",     "UC-15",       "next-intl i18n config",        "src/i18n/request.ts; messages/ar.json",         "Partial (80%)",  "TC-12", "Partial"),
        ("SR-13", "Multi-currency price display",        "Medium",   "UC-15",       "CurrencyContext + registry",   "src/lib/currencies.ts; CurrencyContext.tsx",    "Complete",       "TC-13", "Pass"),
        ("SR-14", "Admin-only brand verification",       "High",     "UC-07",       "Admin RBAC guard",             "api/admin/brands/[id]/route.ts",               "Complete",       "TC-14", "Pass"),
        ("SR-15", "Excel catalogue bulk import",         "Low",      "UC-10",       "xlsx parsing endpoint",        "api/products/bulk/route.ts",                   "Complete",       "TC-15", "Pass"),
        ("SR-16", "Admin analytics dashboard",           "Medium",   "UC-07",       "Dashboard page & charts",      "src/app/admin/dashboard/page.tsx",             "Partial (75%)",  "TC-16", "Partial"),
        ("SR-17", "Brand verification audit trail",      "Medium",   "UC-07",       "Status update with timestamp", "api/admin/brands/[id]/route.ts",               "Partial",        "TC-17", "Not started"),
        ("SR-18", "Verification tier badges",            "Medium",   "UC-03, UC-07","Badge UI components",          "BrandCard.tsx; VerifiedBadge.tsx",             "Complete",       "TC-18", "Pass"),
        ("SR-19", "Anonymous review submission",         "Medium",   "UC-11",       "Review schema — anon flag",    "api/reviews/route.ts (POST)",                  "Complete",       "TC-19", "Pass"),
        ("SR-20", "Brand reply to review",               "Medium",   "UC-12",       "Review reply endpoint",        "api/reviews/[id]/reply/route.ts",              "Complete",       "TC-20", "Pass"),
    ],
    col_widths=[0.6, 1.5, 0.7, 0.8, 1.2, 1.5, 0.8, 0.5, 0.9]
)

# ─── SAVE ─────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_FILE)
print("PHASE 1 COMPLETE")
print(f"Output saved to: {OUTPUT_FILE}")
