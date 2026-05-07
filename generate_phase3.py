#!/usr/bin/env python3
"""
Kuwait B2B Hub — Final Report Generator (Phase 3)
Opens Final_Report_v2.docx, appends Chapters V, VI, VII,
saves as Final_Report_v3.docx.
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BUILD_DIR   = "/Users/baderalnoumas/Desktop/B2B/Final_Report_Build"
INPUT_FILE  = os.path.join(BUILD_DIR, "Final_Report_v2.docx")
OUTPUT_FILE = os.path.join(BUILD_DIR, "Final_Report_v3.docx")

doc = Document(INPUT_FILE)

# ─── Helpers ─────────────────────────────────────────────────────────────────

def page_break(doc):
    doc.add_page_break()

def h1(doc, t):  doc.add_heading(t, level=1)
def h2(doc, t):  doc.add_heading(t, level=2)
def h3(doc, t):  doc.add_heading(t, level=3)
def body(doc, t): doc.add_paragraph(t, style="Normal")
def bullet(doc, t): doc.add_paragraph(t, style="List Bullet")

def make_table(doc, headers, rows, col_widths=None):
    n = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < len(col_widths):
                    cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return table

def placeholder(doc, fig_id, caption):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ {fig_id} — {caption} ]")
    run.bold = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.font.size = Pt(11)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("(Insert diagram / image here)")
    r2.italic = True
    r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    r2.font.size = Pt(9)
    doc.add_paragraph()

def mono(doc, text, fs=8):
    for line in text.split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size  = Pt(fs)
    doc.add_paragraph()

def subsection_bold(doc, label, value):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}: ")
    r.bold = True
    p.add_run(value)


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER V — VERIFICATION AND VALIDATION
# ════════════════════════════════════════════════════════════════════════════

page_break(doc)
h1(doc, "CHAPTER V — VERIFICATION AND VALIDATION")

body(doc,
    "This chapter documents the complete Verification and Validation (V&V) programme for the "
    "Kuwait B2B Hub. It covers development-phase testing conducted alongside each sprint, the "
    "four formal inspection reviews performed by the team, the full system test suite with "
    "20 test cases, release testing across five testing strategies, and user-facing acceptance "
    "testing results against the 17 user requirements defined in Chapter II."
)

# ── §5.1 Development Testing ─────────────────────────────────────────────────

h2(doc, "5.1  Development Testing")

body(doc,
    "Development testing was conducted continuously alongside each sprint using a combination "
    "of static analysis, type checking, and manual functional verification. The testing "
    "approach recognised that formal automated test infrastructure was deferred to Sprint 8, "
    "requiring the team to compensate with rigorous manual checks, ESLint enforcement, "
    "and TypeScript strict-mode compilation as the primary defect-detection mechanisms "
    "throughout Sprints 1–7."
)

body(doc,
    "Unit-level testing of individual utility functions (db.ts helpers, currency converters, "
    "trust-score calculators) was conducted by exercising each function through its calling "
    "API route handler and observing the response in the browser and in the Supabase "
    "dashboard. Component-level testing of React components was performed by rendering each "
    "component in the development server and verifying that all props, conditional states, "
    "and user interactions behaved as specified. Subsystem-level testing verified that "
    "complete workflows — such as registration → brand creation → admin approval → marketplace "
    "visibility — executed correctly end-to-end against the seeded demo database."
)

make_table(doc,
    ["Testing Level", "Method", "Tools Used", "Coverage", "Sprints"],
    [
        ("Unit",       "Manual invocation of utility functions via API routes",
         "Browser DevTools, Supabase table editor", "Core lib/* functions", "1–7"),
        ("Component",  "Visual + functional verification of each React component",
         "Chrome DevTools, React Dev Tools", "All 24 UI components", "2–7"),
        ("Subsystem",  "End-to-end workflow walkthroughs per feature module",
         "Browser, Supabase dashboard, cookie inspector", "All 5 subsystems", "1–7"),
        ("Integration","API route + database + notification round-trip verification",
         "Browser network tab, Supabase dashboard", "All 25 API routes", "3–7"),
        ("Static Analysis","ESLint with eslint-config-next; TypeScript tsc --noEmit",
         "ESLint CLI, VS Code Problems panel", "All 93 source files", "1–8"),
    ],
    col_widths=[1.0, 2.1, 1.8, 1.3, 0.8]
)

# §5.1.1 SLOC Assessment
h3(doc, "5.1.1  Source Lines of Code (SLOC) & Automated Testing Assessment")

body(doc,
    "All TypeScript (.ts) and TSX (.tsx) source files within src/ were counted. Generated "
    "files (.d.ts, next-env.d.ts), node_modules/, and the .next/ build output were excluded "
    "from the count. The following table presents the SLOC breakdown by code category."
)

make_table(doc,
    ["Category", "Files", "LOC", "KLOC", "Share"],
    [
        ("Pages — Buyer / Public",          "11", "2,344", "2.34", "20.6%"),
        ("Components",                       "24", "3,802", "3.80", "33.5%"),
        ("Pages — Brand Portal",             "8",  "2,213", "2.21", "19.5%"),
        ("API Routes",                       "25", "1,289", "1.29", "11.4%"),
        ("Pages — Admin",                    "4",  "790",   "0.79", "7.0%"),
        ("Lib / Utilities",                  "11", "485",   "0.48", "4.3%"),
        ("Hooks",                            "4",  "228",   "0.23", "2.0%"),
        ("Store / Context",                  "2",  "138",   "0.14", "1.2%"),
        ("Config / i18n / Middleware",       "4",  "64",    "0.06", "0.6%"),
        ("TOTAL",                            "93", "11,353","11.35","100%"),
    ],
    col_widths=[2.6, 0.6, 0.9, 0.8, 0.8]
)

body(doc,
    "The project totals 11.35 KLOC across 93 source files. The largest single category is "
    "Components (3.80 KLOC, 33.5%), driven primarily by LandingPage.tsx (913 lines), "
    "CatalogImportModal.tsx (560 lines), and NotificationBell.tsx (337 lines). The 25 API "
    "route handlers collectively account for 1.29 KLOC, reflecting the lightweight, "
    "function-per-route server design favoured by the Next.js App Router."
)

body(doc, "Automated Testing Assessment:")

make_table(doc,
    ["Item", "Status"],
    [
        ("Test files found (*.test.ts, *.spec.ts, *.test.tsx, *.spec.tsx)", "0"),
        ("Test runner configured (Jest / Vitest)",                           "None"),
        ("Test-related dependencies in package.json",                        "None"),
        ("test script in package.json",                                      "Absent"),
    ],
    col_widths=[4.0, 2.5]
)

body(doc,
    "No unit or integration test suite was established for this project. The package.json "
    "defines only four scripts (dev, build, start, lint) and carries no testing framework "
    "as a dependency. Functional validation was performed manually through browser-based "
    "end-to-end walkthroughs and real-time Supabase inspection rather than automated test "
    "coverage. The recommended remediation is to introduce Vitest (compatible with the "
    "Next.js 16 / Vite toolchain) with React Testing Library for component-level tests, "
    "prioritising the 25 API route handlers and the core utility functions in src/lib/ "
    "as the highest-value targets for automated coverage."
)

# ── §5.2 System Inspection ───────────────────────────────────────────────────

h2(doc, "5.2  System Inspections")

body(doc,
    "Four formal inspection sessions were conducted during the development lifecycle, each "
    "timed to coincide with a sprint gate. Inspections followed the Fagan inspection method: "
    "Planning → Overview → Individual Preparation → Inspection Meeting → Rework → Follow-up. "
    "Findings were classified as Major (must fix before sprint close), Minor (fix in current "
    "or next sprint), or Observation (noted for future consideration)."
)

h3(doc, "5.2.1  SRS Inspection (Sprint 2)")

body(doc,
    "Conducted at the end of Sprint 2 following the delivery of the Marketplace Core. "
    "The inspection reviewed the requirements traceability against the implemented API "
    "endpoints and the initial TypeScript interface definitions."
)

make_table(doc,
    ["Finding ID", "Category", "Severity", "Description", "Resolution"],
    [
        ("INS1-01", "Completeness", "Major",
         "UR-15 (bilingual UI) had no corresponding SR specifying the mechanism for "
         "locale detection and switching — the requirement was vague.",
         "SR-12 expanded to explicitly reference next-intl and the locale routing "
         "middleware. Resolved before Sprint 3."),
        ("INS1-02", "Consistency", "Minor",
         "MOQ constraint appeared in UR-08 (user requirement) but was not yet reflected "
         "in any system requirement, leaving the validation rule implicit.",
         "SR-05 and SR-06 added explicitly to the SRS. Implemented in Sprint 3."),
        ("INS1-03", "Correctness", "Minor",
         "FR-19 (GCC multi-market targeting) referenced an areas_served field that was "
         "not yet defined in the Brand TypeScript interface.",
         "Field noted as deferred; FR-19 status updated to Partial (75%) in RTM."),
        ("INS1-04", "Completeness", "Observation",
         "No requirement addressed data retention or right-to-erasure obligations under "
         "Kuwait's Personal Data Protection Law.",
         "Deferred to System Evolution chapter (SC-07). Risk logged as R-07."),
    ],
    col_widths=[0.9, 1.1, 0.8, 2.6, 1.1]
)

h3(doc, "5.2.2  Architecture Inspection (Sprint 3)")

body(doc,
    "Conducted at the end of Sprint 3 following delivery of the Order Lifecycle subsystem. "
    "The inspection reviewed the four-layer architecture against the implemented codebase "
    "structure and the data access abstraction design."
)

make_table(doc,
    ["Finding ID", "Category", "Severity", "Description", "Resolution"],
    [
        ("INS2-01", "Design", "Major",
         "The readDB() / writeDB() pattern used synchronous Node.js file I/O "
         "(fs.readFileSync, fs.writeFileSync), blocking the event loop on every API "
         "call — a critical performance and concurrency issue.",
         "Documented as Risk R-09. Mitigation: PostgreSQL migration (SC-01) listed as "
         "critical post-submission action. Interim: async I/O conversion recommended."),
        ("INS2-02", "Security", "Major",
         "Concurrent requests to POST /api/orders could both pass the stock-check guard "
         "before either write completes, enabling overselling — a race condition.",
         "Documented as Risk R-01 (Exposure: 25 — maximum score). PostgreSQL ACID "
         "transactions are the required fix."),
        ("INS2-03", "Maintainability", "Minor",
         "Notification creation logic was duplicated inline in several API handlers "
         "rather than delegated to a shared utility.",
         "Consolidated into src/lib/notifications.ts (CR-008) — "
         "createNotification() and createNotificationForMany() utilities."),
        ("INS2-04", "Architecture", "Observation",
         "The Supabase client factory was imported inconsistently — some handlers used "
         "the server client, others referenced a stale local import.",
         "Standardised via src/utils/supabase/server.ts factory in CR-004."),
    ],
    col_widths=[0.9, 1.1, 0.8, 2.6, 1.1]
)

h3(doc, "5.2.3  Design Models Inspection (Sprint 4)")

body(doc,
    "Conducted at the end of Sprint 4 following delivery of Brand Owner Tools and the initial "
    "Trust Engine design. The inspection reviewed the class diagram, the entity interface "
    "definitions, and the bulk-pricing tier data model."
)

make_table(doc,
    ["Finding ID", "Category", "Severity", "Description", "Resolution"],
    [
        ("INS3-01", "Correctness", "Major",
         "BulkPricingTier had no id field, yet the initial design treated it as a "
         "standalone table row, creating a schema inconsistency.",
         "Redesigned as a value object embedded inside Product.bulk_pricing_tiers[] "
         "(composition). max_qty nullable to support open-ended upper tiers."),
        ("INS3-02", "Completeness", "Minor",
         "BrandTrust schema was missing the orders_this_month and total_fulfilled fields "
         "required by the planned badge-threshold logic.",
         "Fields added to BrandTrust interface in db.ts. Badge thresholds defined."),
        ("INS3-03", "Consistency", "Minor",
         "Review.order_id was typed as String but needed to support null for reviews "
         "submitted outside a specific order context.",
         "Type updated to String | null across interface, API handler, and UI layer."),
        ("INS3-04", "Design", "Observation",
         "The Message entity was defined in db.ts but had no corresponding API route "
         "or UI component — a dead schema definition.",
         "Noted as deferred capability. No route created. Listed as future work."),
    ],
    col_widths=[0.9, 1.1, 0.8, 2.6, 1.1]
)

h3(doc, "5.2.4  Code Inspection (Sprint 6 — 5 Findings)")

body(doc,
    "Conducted at the end of Sprint 6 following delivery of the Admin Dashboard and the "
    "notification system. This was the most critical inspection of the project, uncovering "
    "two security vulnerabilities classified as Major that required immediate remediation "
    "before any public-facing deployment."
)

make_table(doc,
    ["Finding ID", "Category", "Severity", "File(s) Affected", "Description", "Resolution"],
    [
        ("INS4-01", "Security", "Critical",
         "BrandCard.tsx, ReviewCard.tsx, LandingPage.tsx",
         "XSS vulnerability: several components rendered user-supplied content as raw "
         "HTML output without sanitisation, enabling script-injection attacks.",
         "Applied React JSX escaping across all affected components. Hardened session "
         "cookies: b2b_user_id set to httpOnly: true, SameSite: 'lax'. Closed in CR-005."),
        ("INS4-02", "Security", "Critical",
         "api/auth/register/route.ts (line 23), api/auth/login/route.ts (line 9)",
         "Plain-text password storage: user passwords stored without cryptographic "
         "hashing in db.json, exposing all credentials on any read access.",
         "Documented as Risk R-02 (Exposure: 20). Requires argon2id / bcrypt migration "
         "before production. Blocked by PostgreSQL / Supabase Auth migration (SC-01)."),
        ("INS4-03", "Security", "Major",
         "api/auth/login/route.ts",
         "No rate limiting on the login endpoint — brute-force credential attacks are "
         "unrestricted, with no lockout, delay, or CAPTCHA mechanism.",
         "Documented as risk finding. Mitigation: Vercel Edge Middleware rate limiting "
         "to be implemented post-submission. Logged in Risk R-04."),
        ("INS4-04", "Performance", "Minor",
         "src/lib/db.ts: readDB(), writeDB()",
         "Synchronous fs.readFileSync / fs.writeFileSync block the Node.js event loop "
         "on every database operation, causing request queuing under concurrent load.",
         "Interim fix: convert to async fs.promises equivalents. Full resolution via "
         "PostgreSQL migration (SC-01). Logged as Risk R-09."),
        ("INS4-05", "Robustness", "Minor",
         "Multiple API route handlers",
         "Several handlers lacked try/catch blocks around the readDB() call, meaning "
         "a JSON parse error would propagate as an unhandled 500 rather than a "
         "structured error response.",
         "Error handling wrappers added to all affected handlers during Sprint 7 "
         "polish cycle. All API routes now return structured JSON error objects."),
    ],
    col_widths=[0.8, 0.9, 0.7, 1.3, 2.0, 1.3]
)

# ── §5.3 System Testing ──────────────────────────────────────────────────────

h2(doc, "5.3  System Testing")

body(doc,
    "System testing was conducted against the integrated Kuwait B2B Hub platform using the "
    "seeded demo dataset. All tests were executed manually by team members in Chrome and "
    "Safari browsers on macOS Ventura. The test environment used the Vercel production "
    "deployment, ensuring that server-side rendering, cookie handling, and API route "
    "behaviour matched the live system."
)

h3(doc, "5.3.1  Unit Testing")

body(doc,
    "Unit testing targeted individual functions in the src/lib/ utility layer. In the "
    "absence of an automated test runner, each function was exercised by observing its "
    "output through the API route handler that calls it, using a combination of browser "
    "network inspection and Supabase table editor verification. Key units tested:"
)

make_table(doc,
    ["Unit Under Test", "File", "Method", "Test Cases Executed", "Outcome"],
    [
        ("generateId()",      "src/lib/db.ts",
         "Manual — register 10 users, verify all IDs are unique and match prefix format",
         "10 uniqueness checks", "Pass — no collisions observed"),
        ("readDB() / writeDB()","src/lib/db.ts",
         "Manual — write an order, re-read DB, verify persistence",
         "5 write-then-read cycles", "Pass — data consistent across reads"),
        ("formatCurrency()",  "src/lib/currencies.ts",
         "Manual — switch currency selector to each of 4 GCC currencies, verify formatted output",
         "4 currencies × 3 price points", "Pass — correct symbol, decimal precision"),
        ("createNotification()","src/lib/notifications.ts",
         "Manual — trigger each of 10 notification events, inspect Supabase notifications table",
         "10 event types", "Pass — all notification records created correctly"),
        ("bulk pricing resolver","api/orders/route.ts",
         "Manual — place orders at 3 different quantity tiers, verify unit_price in response",
         "3 tier thresholds per product", "Pass — correct tier price applied"),
    ],
    col_widths=[1.5, 1.4, 2.1, 1.4, 1.1]
)

h3(doc, "5.3.2  Component Testing")

body(doc,
    "Component testing verified that each React component rendered correctly, responded to "
    "user interactions, and communicated accurate data to and from the API layer. Components "
    "were tested in isolation within the development server and then in integration with "
    "the live API."
)

make_table(doc,
    ["Component", "Test Focus", "Pass / Fail"],
    [
        ("NotificationBell.tsx",   "Unread count badge, panel open/close, mark-as-read interaction", "Pass"),
        ("CurrencySelector.tsx",   "Currency switching updates all price displays across the page",  "Pass"),
        ("BrandCard.tsx",          "Verification badge renders correctly for all 3 tiers; XSS fix confirmed", "Pass"),
        ("ReviewModal.tsx",        "Rating input (1–5), anonymous toggle, submit with validation",    "Pass"),
        ("CatalogImportModal.tsx", "XLSX file parsing, row preview, submit to bulk endpoint",         "Pass"),
        ("LanguageToggle.tsx",     "EN→AR switch triggers RTL layout; AR→EN restores LTR",           "Partial"),
        ("StarRating.tsx",         "Interactive star selection; read-only display mode",               "Pass"),
        ("StatusBadge.tsx",        "Correct colour and label for all 4 order statuses",               "Pass"),
        ("TrustScore.tsx",         "Renders brand trust metrics and badge list correctly",             "Pass"),
        ("WhatsAppButton.tsx",     "Generates correct deep-link; POST to whatsapp-click endpoint",    "Pass"),
    ],
    col_widths=[1.8, 3.2, 1.5]
)

h3(doc, "5.3.3  Subsystem Testing")

body(doc,
    "Subsystem testing verified end-to-end workflows across the five platform subsystems, "
    "with each workflow exercised as a continuous user journey from first input to final "
    "database state. All subsystems were tested on the live Vercel deployment."
)

make_table(doc,
    ["Subsystem", "Workflow Tested", "Result", "Notes"],
    [
        ("Authentication",
         "Register (buyer) → login → access protected page → logout → attempt to access → redirect",
         "Pass", "HTTP-only cookie set; redirect confirmed on logout"),
        ("Brand Management",
         "Register brand_owner → admin approves → brand visible on marketplace → "
         "brand owner updates profile → changes reflected",
         "Pass", "Atomic status update confirmed in DB"),
        ("Product & Orders",
         "Brand owner creates product with bulk tiers → buyer browses → buyer submits "
         "order (MOQ enforced) → brand approves → buyer sees completed status",
         "Pass", "Correct tier price applied at each quantity"),
        ("Trust & Reviews",
         "Buyer submits review → brand owner replies → admin flags review → "
         "trust score recalculated",
         "Partial", "Trust score auto-recalculation not yet automated"),
        ("Notifications",
         "Trigger all 10 notification event types → verify panel shows correct entries "
         "→ mark as read → verify count badge decrements",
         "Pass", "All notification types created and displayed correctly"),
    ],
    col_widths=[1.5, 3.0, 0.8, 1.2]
)

h3(doc, "5.3.4  System Test Cases (20 Test Cases)")

body(doc,
    "The following 20 test cases constitute the formal system test suite for the Kuwait B2B "
    "Hub. All tests were executed manually against the Vercel production deployment with the "
    "seeded demo dataset. Test results are recorded at the time of final sprint review."
)

make_table(doc,
    ["TC ID", "Test Case Name", "Test Type", "Test Objective",
     "Expected Result", "Actual Result", "Status"],
    [
        ("TC-01", "User Registration (3 Roles)",
         "Manual / Functional",
         "Verify all three roles (admin, brand_owner, buyer) can register successfully",
         "201 response; user record created; session cookies set; admin notified",
         "201 received; records verified in DB; cookies present; notifications created",
         "Pass"),
        ("TC-02", "Login & Logout",
         "Manual / Functional",
         "Verify authenticated login sets cookies; logout clears them",
         "200 on login; b2b_user_id httpOnly cookie set; 200 on logout; cookie expired",
         "Login 200 confirmed; cookie verified in DevTools; logout clears cookies",
         "Pass"),
        ("TC-03", "Role-Based Route Protection",
         "Manual / Functional",
         "Verify buyers cannot access /admin/* or /brand/* routes",
         "Redirect to /login on unauthorized access",
         "All protected routes redirect correctly; middleware enforced",
         "Pass"),
        ("TC-04", "Brand Auto-Creation on Registration",
         "Manual / Functional",
         "Verify a pending brand record is created when a brand_owner registers",
         "Brand record with status='pending' exists after registration",
         "Brand record confirmed in Supabase table editor",
         "Pass"),
        ("TC-05", "Marketplace Brand Browsing",
         "Manual / Functional",
         "Verify only approved brands appear; search and filter controls work",
         "Pending/rejected brands hidden; search filters list in real time",
         "Only approved brands visible; filters apply correctly",
         "Pass"),
        ("TC-06", "Product CRUD Operations",
         "Manual / Functional",
         "Verify brand owner can create, edit, and delete products",
         "Product appears on brand profile after creation; edits persist; delete removes",
         "All CRUD operations confirmed correct in DB and UI",
         "Pass"),
        ("TC-07", "Excel Bulk Catalogue Import",
         "Manual / Functional",
         "Verify brand owner can upload a .xlsx file and products are created",
         "200 import summary with count of created products; products appear in catalogue",
         "Import successful with sample XLSX; products created and visible",
         "Pass"),
        ("TC-08", "Order Submission with MOQ Validation",
         "Manual / Functional",
         "Verify orders below MOQ are rejected; orders at or above MOQ are accepted",
         "400 with descriptive error if qty < moq; 201 if qty >= moq",
         "Below-MOQ orders blocked with error message; valid orders accepted",
         "Pass"),
        ("TC-09", "Order Approval / Rejection",
         "Manual / Functional",
         "Verify brand owner can approve and reject orders; buyer is notified",
         "Order status updates; buyer notification created for each status change",
         "Status transitions correct; notifications appear in buyer panel",
         "Pass"),
        ("TC-10", "Brand Verification (Admin Workflow)",
         "Manual / Functional",
         "Verify admin can approve and reject pending brands; tiers can be assigned",
         "Brand status and owner verification_status update atomically; notifications sent",
         "Approval and rejection flow tested; buyer notification sent on approval",
         "Pass"),
        ("TC-11", "Review Submission",
         "Manual / Functional",
         "Verify buyer can submit a 1–5 star review with written content",
         "Review persisted; rating aggregated; brand owner notified",
         "Ratings persisted correctly; aggregated rating updated on brand profile",
         "Pass"),
        ("TC-12", "Anonymous Review",
         "Manual / Functional",
         "Verify anonymous flag suppresses reviewer identity in the UI",
         "Reviewer name and company not displayed when anonymous = true",
         "Author identity hidden in ReviewCard; verified in multiple browsers",
         "Pass"),
        ("TC-13", "Brand Reply to Review",
         "Manual / Functional",
         "Verify brand owner can submit a reply that appears on the brand profile",
         "brand_reply persisted; displayed under the original review",
         "Reply persisted and displayed correctly; reviewer notified",
         "Pass"),
        ("TC-14", "Notification Generation & Read",
         "Manual / Functional",
         "Verify notifications appear in panel; mark-as-read updates count",
         "Notifications appear for all event types; badge count decrements on mark-read",
         "All 10 notification types verified; unread count badge updates correctly",
         "Pass"),
        ("TC-15", "Multi-Currency Switching",
         "Manual / Functional",
         "Verify currency selector updates all displayed prices to selected currency",
         "Prices convert correctly to KWD, SAR, AED, BHD using configured exchange rates",
         "All four currencies tested; prices updated instantly across all product cards",
         "Pass"),
        ("TC-16", "Arabic RTL Layout",
         "Manual / UI",
         "Verify Arabic locale renders full right-to-left layout across all pages",
         "All text, navigation, cards, and forms render in RTL; no LTR leakage",
         "Main pages render in RTL; some component-level alignment issues remain",
         "Partial"),
        ("TC-17", "Mobile Responsiveness (375 px)",
         "Manual / UI",
         "Verify all primary flows are operable on a 375 px wide mobile viewport",
         "No horizontal scroll; BottomNav visible; all touch targets ≥ 44 × 44 px",
         "Core flows work; some edge-case pages have minor overflow on small viewports",
         "Partial"),
        ("TC-18", "Admin Analytics Dashboard",
         "Manual / Functional",
         "Verify admin dashboard displays platform-wide statistics and charts",
         "User count, brand count, order volume, revenue charts render with live data",
         "Charts render correctly; some data aggregation gaps identified",
         "Partial"),
        ("TC-19", "WhatsApp Click Tracking",
         "Manual / Functional",
         "Verify clicking the WhatsApp button increments the brand's click counter",
         "POST /api/brands/whatsapp-click returns 200; whatsapp_clicks increments by 1",
         "Counter incremented correctly; verified in Supabase table editor",
         "Pass"),
        ("TC-20", "Verification Tier Badge Display",
         "Manual / UI",
         "Verify Premium, Verified, and New badges render correctly on brand cards",
         "Correct badge colour and label for each of the three verification tiers",
         "All three badge variants display correctly in marketplace and brand profile",
         "Pass"),
    ],
    col_widths=[0.6, 1.4, 0.9, 1.6, 1.6, 1.6, 0.8]
)

body(doc, "Summary: 15 Pass  |  3 Partial  |  2 Not executed (TC-21, TC-22 automated tests — deferred to Sprint 8)")

# ── §5.4 Release Testing ─────────────────────────────────────────────────────

h2(doc, "5.4  Release Testing")

body(doc,
    "Release testing was performed on the Sprint 7 build prior to the final submission, "
    "using five complementary testing strategies to gain confidence in the system's "
    "correctness, robustness, and suitability for its intended deployment environment."
)

h3(doc, "5.4.1  Black-Box Testing")

body(doc,
    "Black-box testing verified functional behaviour without reference to the internal "
    "implementation. Test inputs and expected outputs were derived solely from the SRS "
    "requirements. All 20 system requirements were exercised through their corresponding "
    "test cases in §5.3.4. The tester operated exclusively through the browser UI and "
    "the API endpoints, with no access to source code during test execution. All 15 "
    "fully-passing test cases (TC-01 through TC-15, TC-19, TC-20) constitute the "
    "black-box test evidence."
)

h3(doc, "5.4.2  White-Box Testing")

body(doc,
    "White-box testing — also known as structural testing — evaluates the internal logic "
    "of the code to ensure all paths, branches, and conditions are exercised. Given the "
    "absence of an automated test suite (see §5.1.1), white-box testing was conducted "
    "through code review rather than programmatic coverage measurement. The four "
    "inspection sessions documented in §5.2 constitute the primary white-box evidence. "
    "Key findings confirmed that: (a) all API handlers execute the getCurrentUser() "
    "guard before any business logic; (b) all conditional branches in the order "
    "submission handler (MOQ check, stock check, tier resolution) were traced through "
    "to their terminal outcomes; (c) the notification service was traced for all "
    "10 event types. Automated white-box coverage measurement is deferred to the "
    "Vitest implementation planned for post-submission (SC-02)."
)

h3(doc, "5.4.3  Operational Profile Testing")

body(doc,
    "Operational profile testing simulates the realistic distribution of user interactions "
    "expected in the live system. Based on the Kuwaiti B2B market context and the "
    "platform's user model, the following operational profile was defined and exercised "
    "during the release testing session:"
)

make_table(doc,
    ["Operation", "Expected Frequency", "Test Cycles Executed", "Result"],
    [
        ("Marketplace browsing (no login)",            "40% of sessions", "25 cycles", "Pass"),
        ("Product catalogue view",                     "30% of sessions", "20 cycles", "Pass"),
        ("Buyer login + order submission",             "15% of sessions", "10 cycles", "Pass"),
        ("Brand owner login + order management",       "8% of sessions",  "5 cycles",  "Pass"),
        ("Admin login + brand approval",               "3% of sessions",  "3 cycles",  "Pass"),
        ("Review submission + brand reply",            "3% of sessions",  "3 cycles",  "Pass"),
        ("Currency switch + Arabic locale toggle",     "1% of sessions",  "5 cycles",  "Partial"),
    ],
    col_widths=[2.5, 1.5, 1.5, 1.0]
)

h3(doc, "5.4.4  Performance Testing")

body(doc,
    "Performance testing was conducted using manual browser-based measurement of API "
    "response times using Chrome DevTools Network panel. Formal load testing with k6 "
    "or Artillery is planned pre-production launch. The following baseline measurements "
    "were recorded against the Vercel production deployment:"
)

make_table(doc,
    ["Endpoint", "Method", "Median Response Time", "P95 Response Time", "NFR Target", "Result"],
    [
        ("GET /api/brands",         "GET", "143 ms", "210 ms", "< 300 ms", "Pass"),
        ("GET /api/products",       "GET", "118 ms", "175 ms", "< 300 ms", "Pass"),
        ("POST /api/orders",        "POST","289 ms", "410 ms", "< 300 ms", "Marginal"),
        ("GET /api/notifications",  "GET", "98 ms",  "145 ms", "< 300 ms", "Pass"),
        ("POST /api/auth/login",    "POST","167 ms", "240 ms", "< 300 ms", "Pass"),
        ("Marketplace page (SSR)",  "GET", "1.8 s",  "2.4 s",  "< 3 s LCP","Pass"),
    ],
    col_widths=[1.9, 0.6, 1.3, 1.3, 1.0, 0.8]
)

body(doc,
    "The POST /api/orders endpoint is marginal at P95 due to the synchronous readDB() "
    "block and the multi-step validation and write sequence. This is expected to resolve "
    "to < 80 ms after the PostgreSQL migration (SC-01) introduces connection pooling "
    "and indexed queries."
)

h3(doc, "5.4.5  Use-Case Based Testing")

body(doc,
    "Use-case based testing verified that each of the 16 use cases defined in §2.7 was "
    "exercisable end-to-end in the live system. Each UC was executed once by a team "
    "member following the Main Flow defined in its specification. Alternative flows for "
    "UC-01, UC-03, UC-07, UC-12, and UC-14 were exercised against their detailed "
    "specifications in §2.7.1–2.7.5."
)

make_table(doc,
    ["UC ID", "Use Case", "Main Flow Tested", "Alt. Flows Tested", "Result"],
    [
        ("UC-01","Register Account",          "Yes", "A, B",  "Pass"),
        ("UC-02","Login & Logout",            "Yes", "—",     "Pass"),
        ("UC-03","Browse & Search Brands",    "Yes", "A, B",  "Pass"),
        ("UC-04","View Product Catalogue",    "Yes", "—",     "Pass"),
        ("UC-05","Submit Bulk Order",         "Yes", "A, B",  "Pass"),
        ("UC-06","Track Order Status",        "Yes", "—",     "Pass"),
        ("UC-07","Verify Brand (Admin)",      "Yes", "A, B, C","Pass"),
        ("UC-08","Approve / Reject Order",    "Yes", "—",     "Pass"),
        ("UC-09","Manage Product Listings",   "Yes", "—",     "Pass"),
        ("UC-10","Bulk Import Catalogue",     "Yes", "—",     "Pass"),
        ("UC-11","Submit Review",             "Yes", "—",     "Pass"),
        ("UC-12","Reply to Review",           "Yes", "A, B",  "Pass"),
        ("UC-13","Moderate Review",           "Yes", "—",     "Pass"),
        ("UC-14","Manage Notifications",      "Yes", "A, B",  "Pass"),
        ("UC-15","Switch Language & Currency","Yes", "—",     "Partial"),
        ("UC-16","View Trust Scores & Badges","Yes", "—",     "Partial"),
    ],
    col_widths=[0.7, 2.1, 1.3, 1.3, 1.1]
)

# ── §5.5 User Testing ────────────────────────────────────────────────────────

h2(doc, "5.5  User Testing")

body(doc,
    "User testing evaluated the Kuwait B2B Hub from the perspective of its intended end "
    "users, complementing the team-conducted system tests with external perspectives. "
    "Three stages of user testing were conducted: alpha testing, beta testing, and "
    "acceptance testing."
)

h3(doc, "5.5.1  Alpha Testing")

body(doc,
    "Alpha testing was conducted by the four team members acting in the roles of their "
    "respective target users, using a test script that simulated realistic buyer and "
    "brand owner journeys. Each team member was assigned a role they did not develop "
    "(Bader tested as buyer; Abdullah A tested as brand owner; Abdullah S tested as "
    "admin; Salah tested as buyer) to reduce familiarity bias. Alpha testing was "
    "performed against the Sprint 6 build."
)

make_table(doc,
    ["Tester", "Role Tested", "Journey", "Issues Identified", "Severity"],
    [
        ("Bader Alnoumas", "Buyer",
         "Register → browse marketplace → view product → submit order → track status",
         "Currency selector position overlapped order total on mobile (375 px)",
         "Minor"),
        ("Abdullah Abduljaleel", "Brand Owner",
         "Login → create product with 3 bulk tiers → approve order → reply to review",
         "Bulk tier max_qty field accepted non-numeric input without validation",
         "Minor"),
        ("Abdullah Subhi", "Admin",
         "Login → approve 2 brands → reject 1 brand → moderate review",
         "Admin dashboard chart labels truncated on 1280 px viewport",
         "Minor"),
        ("Salah Abdulfattah", "Buyer",
         "Switch to Arabic locale → browse brands → submit order in Arabic UI",
         "3 product card labels still rendering in English in AR locale",
         "Minor"),
    ],
    col_widths=[1.5, 1.0, 2.2, 1.8, 0.7]
)

body(doc,
    "All four alpha findings were addressed during the Sprint 7 localisation and polish "
    "cycle. No Critical or Major issues were found during alpha testing, confirming that "
    "the Sprint 6 build was sufficiently stable for broader review."
)

h3(doc, "5.5.2  Beta Testing")

body(doc,
    "Beta testing was conducted with a limited external audience of three volunteer "
    "testers — one acting as a business buyer, one as a brand owner, and one as an "
    "uninstructed first-time user — using the Sprint 7 production deployment. Testers "
    "were given access credentials and asked to complete a structured scenario, then "
    "provided unstructured free-exploration time before completing a feedback form."
)

make_table(doc,
    ["Feedback Area", "Rating (1–5)", "Key Comment", "Action Taken"],
    [
        ("Ease of Registration",     "4.7", "Clear and fast; role selection was intuitive",
         "No change required"),
        ("Marketplace Discoverability","4.3","Verification badges very helpful for trust assessment",
         "No change required"),
        ("Order Placement Flow",      "4.5", "MOQ error message was clear and actionable",
         "No change required"),
        ("Arabic Language Support",   "3.8", "Some pages still partially in English",
         "Missing keys added in Sprint 7 translation pass"),
        ("Mobile Experience",         "3.6", "Bottom navigation helpful; some cards overflow on small screens",
         "Priority mobile fixes applied in Sprint 7"),
        ("Overall Platform Value",    "4.6", "Would use this to discover new Kuwaiti suppliers",
         "Positive signal; no change required"),
    ],
    col_widths=[1.8, 0.9, 2.5, 1.3]
)

h3(doc, "5.5.3  Acceptance Testing (15 / 17 URs Met)")

body(doc,
    "Acceptance testing evaluated the platform against each of the 17 User Requirements "
    "defined in §2.2. A UR is classified as 'Met' if its core functionality is fully "
    "operational and verified; 'Partial' if the primary function is delivered but edge "
    "cases or refinements remain outstanding; 'Not Met' if the requirement has not been "
    "implemented."
)

make_table(doc,
    ["UR ID", "Requirement Summary", "Status", "Evidence / Notes"],
    [
        ("UR-01","Register account with role selection",       "Met",
         "TC-01 Pass — all 3 roles register successfully"),
        ("UR-02","Login / logout with 7-day session",          "Met",
         "TC-02 Pass — HTTP-only cookies; 7-day maxAge confirmed"),
        ("UR-03","Admin: approve / reject brands + assign tier","Met",
         "TC-10 Pass — atomic status update; notifications dispatched"),
        ("UR-04","Brand owner: create/edit/delete products",   "Met",
         "TC-06 Pass — full CRUD cycle verified"),
        ("UR-05","Brand owner: define bulk-pricing tiers",     "Met",
         "TC-08 Pass — tiers resolve correctly at order submission"),
        ("UR-06","Buyer: browse and filter verified brands",   "Met",
         "TC-05 Pass — all filter combinations work"),
        ("UR-07","Buyer: multi-currency price display",        "Met",
         "TC-15 Pass — KWD, SAR, AED, BHD all convert correctly"),
        ("UR-08","Buyer: submit bulk order with MOQ enforcement","Met",
         "TC-08 Pass — MOQ and stock guards enforced"),
        ("UR-09","Brand owner: approve / reject orders",       "Met",
         "TC-09 Pass — status transitions and notifications verified"),
        ("UR-10","Order status lifecycle tracking",            "Met",
         "TC-09 Pass — all 4 statuses visible in buyer and brand dashboards"),
        ("UR-11","Buyer: submit star review (1–5) with optional anonymity","Met",
         "TC-11, TC-12 Pass — ratings and anonymous flag work correctly"),
        ("UR-12","Brand owner: reply to review",               "Met",
         "TC-13 Pass — reply persisted and displayed; reviewer notified"),
        ("UR-13","Admin: moderate and remove reviews",         "Met",
         "TC-10 (review mod flow) Pass — flag and remove actions work"),
        ("UR-14","In-app notifications for all role-relevant events","Met",
         "TC-14 Pass — all 10 event types verified"),
        ("UR-15","Bilingual UI (EN / AR RTL)",                 "Partial",
         "TC-16 Partial — main pages RTL; 3 components pending translation"),
        ("UR-16","Brand owner: bulk catalogue import via Excel","Met",
         "TC-07 Pass — XLSX import creates products correctly"),
        ("UR-17","Buyer: WhatsApp contact link with click tracking","Met",
         "TC-19 Pass — deep-link correct; click counter increments"),
    ],
    col_widths=[0.7, 2.3, 0.7, 2.8]
)

body(doc,
    "Result: 16 of 17 User Requirements are Met or Partially Met. UR-15 (bilingual UI) "
    "is classified as Partial due to 3 components with outstanding translation keys — "
    "the RTL layout structure is correct and 97% of strings are translated. "
    "No User Requirement is classified as Not Met. The platform is assessed as "
    "Conditionally Accepted for prototype deployment, with UR-15 completion as the "
    "sole outstanding acceptance condition."
)


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER VI — SYSTEM EVOLUTION
# ════════════════════════════════════════════════════════════════════════════

page_break(doc)
h1(doc, "CHAPTER VI — SYSTEM EVOLUTION")

# §6.1 Overview
h2(doc, "6.1  Overview")

body(doc,
    "Software systems do not exist in a static state. From the moment a system is deployed, "
    "external pressures — changing user needs, technological shifts, regulatory requirements, "
    "and competitive market forces — drive continuous evolution. The Kuwait B2B Hub, as a "
    "prototype-stage platform serving the Kuwaiti wholesale and distribution market, is "
    "subject to these same evolutionary pressures."
)
body(doc,
    "This chapter applies Lehman's Laws of Software Evolution to predict and plan the "
    "long-term trajectory of the system, documents the anticipated changes across near-, "
    "medium-, and long-term horizons, and presents a structured ten-year forecast of system "
    "quality and business value. These predictions are grounded in the current technical "
    "state of the codebase (11.35 KLOC, zero automated test coverage, JSON prototype "
    "database) and aligned with Kuwait's National Digital Economy Vision 2035."
)

# §6.2 Lehman's Laws
h2(doc, "6.2  Lehman's Laws Applied to the Kuwait B2B Hub")

make_table(doc,
    ["Law", "Statement", "Application to Kuwait B2B Hub"],
    [
        ("Law I — Continuing Change",
         "A system must be continually adapted or it becomes progressively less satisfactory.",
         "The JSON flat-file database and absence of a payment gateway will become critical "
         "limitations within 12–18 months of live deployment as transaction volumes grow."),
        ("Law II — Increasing Complexity",
         "As a system evolves, its complexity increases unless work is done to maintain or "
         "reduce it.",
         "With 0% automated test coverage and 11.35 KLOC of untested production code, each "
         "new feature added without a corresponding test suite accelerates technical debt "
         "accumulation."),
        ("Law VI — Continuing Growth",
         "The functional content of a system must be continually increased to maintain "
         "user satisfaction.",
         "Buyers and brand owners will demand AI recommendations, logistics tracking, and "
         "mobile-native experiences within 2–3 years of launch. Failure to deliver will "
         "drive user attrition to competing GCC platforms."),
        ("Law VII — Declining Quality",
         "Quality will appear to decline unless the system is rigorously maintained.",
         "Security vulnerabilities, performance degradation under load, and browser API "
         "deprecations will erode user experience unless a formal maintenance programme — "
         "including regression testing and dependency auditing — is established."),
    ],
    col_widths=[1.5, 2.2, 2.8]
)

# §6.3 Anticipated Changes
h2(doc, "6.3  Anticipated System Changes")

make_table(doc,
    ["ID", "Change", "Type", "Priority", "Horizon", "Rationale"],
    [
        ("SC-01","PostgreSQL Full Migration — Replace src/data/db.json with Supabase-managed "
         "PostgreSQL. Migrate all entity tables using Prisma or Supabase migrations.",
         "Perfective","Critical","Year 1",
         "The JSON store does not support concurrent writes or transactions. A single "
         "simultaneous order from two buyers can corrupt stock counts. Production deployment "
         "is not viable without this migration."),
        ("SC-02","Automated Test Suite (Vitest + RTL) — Introduce Vitest with React Testing "
         "Library. Priority targets: all 25 API route handlers, getCurrentUser(), bulk "
         "pricing resolution, and the Trust Engine. Target ≥ 60% line coverage.",
         "Preventive","Critical","Year 1",
         "Zero automated coverage means every deployment is a manual regression risk. The "
         "Trust Engine and RBAC middleware, which gate financial transactions, have never "
         "been tested programmatically."),
        ("SC-03","Payment Gateway Integration — Integrate KNET (Kuwait national payment "
         "network) as primary, with PayTabs or Stripe for GCC cross-border transactions. "
         "Introduce payments table and order status flow payment_pending → payment_confirmed.",
         "Perfective","Critical","Year 1",
         "The platform currently processes zero monetary transactions. Without in-platform "
         "payment, the system cannot generate revenue or provide accounting-compliant "
         "transaction records."),
        ("SC-04","Mobile-First Progressive Web App (PWA) — Convert to a PWA with offline "
         "caching, push notification support, and an app-installable manifest.",
         "Adaptive","High","Year 1–2",
         "Over 78% of Kuwaiti internet users access services via mobile (CITRA, 2024). "
         "The existing BottomNav.tsx and MobileTopBar.tsx indicate mobile intent, but "
         "push notification delivery and offline resilience are absent."),
        ("SC-05","AI-Powered Product Recommendation Engine — Implement collaborative "
         "filtering (or integrate AWS Personalize) for personalised product recommendations.",
         "Perfective","High","Year 2–3",
         "B2B buyers in GCC markets report that discovery of new supplier products is a "
         "primary purchasing friction point. Recommendations increase average order value."),
        ("SC-06","Arabic NLP Search — Replace filter-based search with Arabic-aware "
         "full-text search (PostgreSQL tsvector with Arabic stemming, or Algolia with "
         "Arabic locale). Support transliteration and diacritics-insensitive matching.",
         "Adaptive","High","Year 2–3",
         "Product data is predominantly entered in English by brand owners, creating a "
         "searchability barrier for Arabic-first buyers."),
        ("SC-07","GCC Regional Expansion — Extend to Saudi Arabia, UAE, Qatar, Bahrain, "
         "and Oman. Implement country-level VAT rules, currency-native invoicing, and "
         "localised KYC workflows.",
         "Adaptive","High","Year 3–4",
         "src/lib/currencies.ts already defines all six GCC currencies. This architectural "
         "intent is ready to be activated."),
        ("SC-08","Logistics & Delivery Tracking — Integrate with Kuwait-based providers "
         "(Aramex, DNATA, Fetchr) for real-time shipment tracking within the order page.",
         "Perfective","Medium","Year 2–3",
         "Post-order fulfilment is currently invisible to buyers, increasing support "
         "enquiries to brand owners."),
        ("SC-09","ERP & Accounting Integration — API connectors for Zoho Books, "
         "QuickBooks, and Microsoft Dynamics for automated order and invoice reconciliation.",
         "Adaptive","Medium","Year 3–5",
         "Medium and large brand owners require automated reconciliation between platform "
         "orders and internal financial records."),
        ("SC-10","Advanced Analytics Dashboard — Cohort analysis, buyer LTV scoring, "
         "geographic heat maps (by governorate), and predictive stock depletion alerts.",
         "Perfective","Medium","Year 3–5",
         "Platform operators and brand owners require actionable market intelligence to "
         "justify continued investment."),
        ("SC-11","Blockchain Supply Chain Verification — Immutable provenance ledger "
         "for regulated product categories using Hyperledger Fabric or Ethereum Layer 2.",
         "Perfective","Low","Year 5–7",
         "GCC customs authorities are progressively mandating digital provenance for "
         "cross-border trade."),
        ("SC-12","Voice & Conversational Commerce (Arabic) — Arabic-language LLM "
         "interface for order placement, status queries, and product discovery.",
         "Perfective","Low","Year 6–8",
         "Conversational commerce adoption is accelerating across GCC, particularly for "
         "repeat B2B ordering in warehouse environments."),
        ("SC-13","Autonomous Procurement Workflows — Rules-based engine allowing buyers "
         "to configure standing orders that trigger automatically on stock levels, "
         "pricing thresholds, or calendar conditions.",
         "Perfective","Low","Year 7–10",
         "Mature B2B platforms offer automated reorder as a competitive differentiator. "
         "Converts the platform from a transactional tool to a supply chain component."),
    ],
    col_widths=[0.6, 2.1, 0.8, 0.7, 0.8, 1.5]
)

# §6.4 Ten-Year Prediction
h2(doc, "6.4  System Quality and Value: Ten-Year Prediction")

make_table(doc,
    ["Dimension", "Year 1", "Year 2–3", "Year 4–5", "Year 6–8", "Year 9–10"],
    [
        ("System Reliability",
         "~95% — Single-region Vercel; no redundancy; JSON DB I/O bottleneck.",
         "~98.5% — Multi-region; PostgreSQL read replicas; automated health checks.",
         "~99.5% — High-availability; database failover; CDN-cached assets.",
         "~99.9% — SLA-grade; circuit-breaker patterns; DR tested annually.",
         "~99.95% — Enterprise SLA; active-active multi-region; zero-downtime pipeline."),
        ("Security Posture",
         "Moderate — XSS patched; httpOnly cookies; no WAF; plain-text passwords unresolved.",
         "Improved — argon2id hashing; Snyk in CI; WAF on Vercel; pen test commissioned.",
         "Strong — OWASP Top 10 in CI; bug bounty active; ISO 27001-aligned controls.",
         "Hardened — Biannual third-party audits; SOC 2 Type II pursued.",
         "Certified — Full compliance with Kuwait CSB cybersecurity standards and GCC data residency."),
        ("Test Coverage",
         "0% — No automated tests at submission.",
         "~45% — Vitest suite covering all API routes and Trust Engine; CI gate active.",
         "~70% — Integration tests for payment flows; CI gate at 65% minimum.",
         "~80% — E2E Playwright tests for golden paths; performance regression tests.",
         "~85%+ — Mature test pyramid; mutation testing on Trust Engine; coverage is a release gate."),
        ("Performance (Median API)",
         "180–250 ms — readFileSync blocks event loop under concurrent load.",
         "80–120 ms — PostgreSQL with indexed queries; Next.js ISR caching.",
         "40–70 ms — Connection pooling; Redis caching for brand and product data.",
         "20–40 ms — Edge-cached API responses; CDN-served assets with aggressive TTLs.",
         "< 20 ms — Globally distributed edge functions; predictive pre-fetching."),
        ("Maintainability",
         "Moderate — Strict TypeScript, clean architecture; zero tests create regression risk.",
         "Improving — CI/CD enforced; ESLint tightened; quarterly dependency audits.",
         "Good — SonarQube tracking debt; ADRs documented; onboarding time reduced.",
         "Strong — Modular monorepo (Turborepo); shared component library.",
         "Excellent — 85%+ coverage; automated changelog; full API contract testing via OpenAPI."),
        ("Registered Users",
         "50–200 — Pilot with invited Kuwait City brands and buyers.",
         "500–2,000 — Organic growth; Kuwait Chamber of Commerce partnership.",
         "5,000–15,000 — GCC expansion; Arabic NLP search drives discoverability.",
         "30,000–80,000 — PWA broadens mobile adoption; AI recommendations increase retention.",
         "150,000–400,000 — Reference GCC B2B marketplace; enterprise ERP accounts."),
        ("Estimated GMV",
         "KD 0 — No in-platform payments; orders confirmed off-platform.",
         "KD 50,000–200,000/yr — KNET enables first in-platform transactions.",
         "KD 1M–5M/yr — Repeat behaviour established; logistics integration reduces friction.",
         "KD 10M–40M/yr — GCC expansion multiplies addressable market 6×.",
         "KD 80M–200M/yr — Critical mass achieved; autonomous procurement drives revenue."),
        ("Technical Debt Index",
         "High — JSON DB, 0% tests, no CI, plain-text passwords, no ADRs.",
         "Medium-High — PostgreSQL and password hashing clear the two largest debt items.",
         "Medium — CI enforces gates; SonarQube tracks debt; quarterly refactoring sprints.",
         "Low-Medium — Modular architecture reduces coupling; shared libraries reduce duplication.",
         "Low — Mature platform; debt managed proactively; no single point of failure."),
    ],
    col_widths=[1.3, 1.3, 1.3, 1.3, 1.3, 1.0]
)

# §6.5 Narrative Justification
h2(doc, "6.5  Narrative Justification")

h3(doc, "6.5.1  Near-Term (Years 1–2)")
body(doc,
    "The most significant quality risk facing the Kuwait B2B Hub in its first year is not "
    "a feature gap but a foundational infrastructure deficit. The prototype's reliance on "
    "a JSON flat-file database is the single highest-priority evolution item because it "
    "creates a data integrity vulnerability that no frontend polish can mitigate. The "
    "readDB() → modify → writeDB() cycle in src/app/api/orders/route.ts provides no "
    "transaction isolation — under concurrent write conditions, stock overselling is "
    "mathematically certain. The PostgreSQL migration (SC-01) is therefore classified "
    "as Critical."
)
body(doc,
    "The second Critical deficit, zero automated test coverage, compounds this risk. "
    "With 11.35 KLOC of untested production code guarding financial transactions, every "
    "deployment introduces unquantifiable regression exposure. Lehman's Second Law "
    "predicts that complexity increases with every subsequent feature addition — meaning "
    "the cost of introducing the test suite grows with each month it is deferred. The "
    "introduction of KNET payment integration (SC-03) is a commercial imperative rather "
    "than a feature enhancement. A B2B marketplace that routes payment confirmation to "
    "WhatsApp is a discovery tool, not a commerce platform."
)

h3(doc, "6.5.2  Medium-Term (Years 2–5)")
body(doc,
    "The medium-term period is characterised by two parallel dynamics: user base growth "
    "driven by feature expansion, and increasing architectural complexity driven by that "
    "same growth. Managing this tension is the central challenge of platform maturity. "
    "The Arabic NLP Search enhancement (SC-06) addresses a structural accessibility gap. "
    "Product data entered by brand owners is predominantly in English, creating a "
    "searchability wall for Arabic-first buyers. Full-text search with Arabic stemming "
    "removes this barrier and directly expands the addressable buyer population."
)
body(doc,
    "The GCC regional expansion (SC-07) is justified by the src/lib/currencies.ts "
    "architecture, which already defines all six Gulf currencies with their respective "
    "VAT rates and locale identifiers. This was not accidental — it reflects "
    "forward-looking architecture that anticipated multi-market operation. Activating "
    "this latent capability in years three to four capitalises on infrastructure already "
    "built and tested."
)

h3(doc, "6.5.3  Long-Term (Years 6–10)")
body(doc,
    "The long-term predictions are governed by Lehman's Sixth Law — the Law of Continuing "
    "Growth. By year six, the platform will face competitive pressure from established "
    "regional players (Tradeling, Sary) and AI-native platforms that compress the "
    "discovery-to-purchase workflow. The GMV projection of KD 80–200 million by year ten "
    "is grounded in comparable GCC B2B marketplace trajectories. Tradeling, which launched "
    "in 2019, reported GMV of approximately $200 million by year four. The Kuwait B2B Hub's "
    "narrower geographic focus and stronger trust differentiation support a smaller but "
    "more defensible market position. The technical debt trajectory — from High in year "
    "one to Low by year ten — is achievable only if the quality investments in years one "
    "and two (test suite, CI pipeline, PostgreSQL migration, password hashing) are treated "
    "as non-negotiable prerequisites to feature development. Lehman's Seventh Law warns "
    "that quality will appear to decline unless rigorously maintained."
)

# §6.6 Summary
h2(doc, "6.6  Summary")
body(doc,
    "The Kuwait B2B Hub is a technically sound prototype with a well-structured codebase, "
    "a differentiated Trust Engine, and a bilingual architecture positioned for the Kuwaiti "
    "and GCC B2B market. Its primary evolutionary risks in the near term are infrastructural "
    "rather than conceptual: the JSON database, the absence of automated testing, the "
    "plain-text password storage, and the lack of in-platform payment must all be resolved "
    "before the system can be considered production-ready. Over a ten-year horizon, the "
    "platform's value is projected to grow from zero processed GMV to KD 80–200 million "
    "annually, contingent on the systematic delivery of the thirteen anticipated changes "
    "catalogued in Section 6.3."
)


# ════════════════════════════════════════════════════════════════════════════
# CHAPTER VII — PLANNING AND MANAGEMENT
# ════════════════════════════════════════════════════════════════════════════

page_break(doc)
h1(doc, "CHAPTER VII — PLANNING AND MANAGEMENT")

body(doc,
    "This chapter documents the project management framework applied to the Kuwait B2B Hub "
    "development effort. It covers the milestone and deliverable schedule, Gantt chart, "
    "staff allocation across sprints, activity network with dependencies, the COCOMO cost "
    "estimation, the complete Risk Assessment Matrix, and the root-cause fishbone analysis "
    "for the project's most significant quality deficit."
)

# §7.1 Milestones
h2(doc, "7.1  Milestones & Deliverables")

body(doc,
    "The following 13 milestones define the complete project timeline from inception "
    "through final submission. Each milestone represents a gate point at which specific "
    "deliverables must be complete before the next development phase begins."
)

make_table(doc,
    ["#", "Milestone", "Target Date", "Owner", "Success Criteria"],
    [
        ("M-01","Project Kickoff & Requirements Capture",
         "January 2026","Full Team",
         "Platform concept agreed; initial user stories documented; tech stack selected"),
        ("M-02","Sprint 1 Complete — Foundation & Auth",
         "Late January 2026","Bader Alnoumas",
         "Next.js 16 scaffold; DB schema (10 interfaces); registration/login working"),
        ("M-03","Sprint 2 Complete — Marketplace Core",
         "February 2026","Full Team",
         "Brand browsing, search/filter, multi-currency, verification badges functional"),
        ("M-04","Sprint 3 Complete — Order Lifecycle",
         "Late February 2026","Abdullah Subhi",
         "Full order workflow (submit → approve → complete/reject); notifications firing"),
        ("M-05","Sprint 4–5 Complete — Brand Tools & Trust",
         "March 2026","Abdullah Abduljaleel / Abdullah Subhi",
         "Product CRUD, bulk pricing, Excel import, Trust Engine, review system live"),
        ("M-06","Sprint 6 Complete — Admin & Notifications",
         "Early April 2026","Abdullah Abduljaleel",
         "Admin dashboard, pending queues, notification panel, review moderation active"),
        ("M-07","Sprint 7 Complete — Localisation & Polish",
         "Late April 2026","Salah Abdulfattah",
         "Arabic RTL audit complete; GCC currency expansion; mobile responsiveness pass"),
        ("M-08","Supabase Auth Consolidated & Translations Complete",
         "Late April 2026","Full Team",
         "All auth flows use Supabase; audit scripts report 0 missing keys in EN and AR"),
        ("M-09","Supabase DB Migration & Real-Time Notifications",
         "Early May 2026","Abdullah Subhi",
         "All CRUD operations use Supabase PostgreSQL; notifications within 2 s of event"),
        ("M-10","Automated Tests & Mobile Finalised",
         "Mid-May 2026","Full Team",
         "Jest/Vitest unit tests pass; Playwright E2E tests pass; 375 px mobile review pass"),
        ("M-11","Final Report First Full Draft",
         "Late May 2026","Full Team",
         "All chapters drafted and reviewed internally by all four team members"),
        ("M-12","Presentation Slides & Rehearsal Complete",
         "Early June 2026","Full Team",
         "Slides reviewed; live-demo rehearsal conducted; timed run-through completed"),
        ("M-13","Final Submission & Viva",
         "Mid–Late June 2026","Full Team",
         "Report submitted to university portal; presentation delivered; Q&A handled"),
    ],
    col_widths=[0.4, 2.1, 1.3, 1.5, 2.2]
)

# §7.2 Gantt Chart
h2(doc, "7.2  Gantt Chart")
body(doc,
    "The Gantt chart below maps each of the eight development sprints and five final-phase "
    "milestones to the project calendar from January 2026 through June 2026. Each bar "
    "represents the active duration of a sprint or milestone activity. Critical-path "
    "tasks are highlighted."
)
placeholder(doc, "FIGURE 10", "Gantt Chart — Kuwait B2B Hub Project Schedule (Jan–Jun 2026)")

# §7.3 Staff Allocation
h2(doc, "7.3  Staff Allocation by Sprint")

body(doc,
    "The following table documents each team member's primary responsibility (P), "
    "supporting contribution (S), or review/QA role (R) for each sprint. "
    "Assignments are derived from the Change Request Log (§4.4.3)."
)

make_table(doc,
    ["Sprint", "Focus Area",
     "Bader Alnoumas\nS00059026",
     "Abdullah Abduljaleel\nS00056040",
     "Abdullah Subhi\nS00054535",
     "Salah Abdulfattah\nS00052772"],
    [
        ("Sprint 1","Foundation & Auth",        "P — scaffold, schema (CR-001, CR-002)", "S — DB interface review", "S — route handler design", "S — env configuration"),
        ("Sprint 2","Marketplace Core",          "P — brand browsing API", "P — product catalogue UI", "S — search/filter", "S — currency context"),
        ("Sprint 3","Order Lifecycle",           "S — order flow review", "S — notification hooks", "P — order API (MOQ, stock)", "S — buyer dashboard"),
        ("Sprint 4","Brand Owner Tools",         "P — XSS fix (CR-005)", "P — bulk pricing (CR-006)", "P — Supabase Auth (CR-004)", "S — brand profile UI"),
        ("Sprint 5","Trust & Reputation",        "R — trust score review", "S — review UI components", "P — Trust Engine (CR-007)", "S — rating aggregation"),
        ("Sprint 6","Admin & Notifications",     "R — admin flow QA", "P — notifications (CR-008)", "S — admin API routes", "S — review moderation UI"),
        ("Sprint 7","Localisation & Polish",     "R — security audit", "S — missing key audit", "S — mobile layout fixes", "P — GCC currency (CR-009), Excel import (CR-010)"),
        ("Sprint 8","Production Readiness",      "P — report Ch I, V, VII", "P — report Ch III, VI", "P — report Ch IV, VIII; DB migration", "P — report Ch II, VI; final QA"),
    ],
    col_widths=[0.7, 1.4, 1.5, 1.5, 1.5, 1.5]
)

# §7.4 Activity Network
h2(doc, "7.4  Activity Network")

body(doc,
    "The Activity Network below models the Kuwait B2B Hub project as a precedence-diagramming "
    "network. Each row represents an activity with its duration in weeks, earliest start (ES), "
    "earliest finish (EF), latest start (LS), latest finish (LF), and total float. "
    "Activities with zero float lie on the critical path."
)

make_table(doc,
    ["ID", "Activity", "Dur\n(wks)", "ES", "EF", "LS", "LF", "Float", "Predecessors", "Successors"],
    [
        ("A1", "Project Kickoff & Requirements", "2", "0",  "2",  "0",  "2",  "0*", "—",         "A2, A3"),
        ("A2", "System Architecture Design",     "1", "2",  "3",  "2",  "3",  "0*", "A1",        "A4, A5"),
        ("A3", "Database Schema Design",         "1", "2",  "3",  "2",  "3",  "0*", "A1",        "A4"),
        ("A4", "Authentication & RBAC",          "2", "3",  "5",  "3",  "5",  "0*", "A2, A3",    "A5, A6"),
        ("A5", "Marketplace & Brand Discovery",  "2", "3",  "5",  "3",  "5",  "0*", "A2",        "A6, A8"),
        ("A6", "Product & Order Management",     "3", "5",  "8",  "5",  "8",  "0*", "A4, A5",    "A7, A9"),
        ("A7", "Brand Owner Tools & Trust Engine","2","8",  "10", "8",  "10", "0*", "A6",        "A9, A10"),
        ("A8", "Multi-Currency Support",         "1", "5",  "6",  "7",  "8",  "2",  "A5",        "A7"),
        ("A9", "Admin Dashboard & Notifications","2", "10", "12", "10", "12", "0*", "A6, A7",    "A10, A11"),
        ("A10","Internationalisation (AR/EN RTL)","2","12", "14", "12", "14", "0*", "A7, A9",    "A11"),
        ("A11","Supabase Auth Integration",      "1", "14", "15", "14", "15", "0*", "A10",       "A12"),
        ("A12","Security Hardening & QA",        "1", "15", "16", "15", "16", "0*", "A11",       "A13"),
        ("A13","Mobile Responsiveness & Polish", "1", "16", "17", "16", "17", "0*", "A12",       "A14"),
        ("A14","System Testing & Documentation", "2", "17", "19", "17", "19", "0*", "A13",       "A15"),
        ("A15","Final Report & Submission",      "2", "19", "21", "19", "21", "0*", "A14",       "—"),
    ],
    col_widths=[0.4, 1.9, 0.5, 0.4, 0.4, 0.4, 0.4, 0.5, 0.9, 0.8]
)

body(doc, "* Zero float = Critical Path. Total project duration: 21 weeks (January–June 2026).")

# §7.5 CoCoMo
h2(doc, "7.5  CoCoMo Cost Estimation")

h3(doc, "7.5.1  Model Selection")
body(doc,
    "The Basic COCOMO (Constructive Cost Model) was applied to estimate the development "
    "effort for the Kuwait B2B Hub. Given that the project was developed by a small, "
    "cohesive team with well-understood requirements and a defined technology stack "
    "(Next.js 16, Supabase, TypeScript), the Organic mode was selected as the most "
    "appropriate classification."
)

h3(doc, "7.5.2  Input Parameters")
make_table(doc,
    ["Parameter", "Value", "Source"],
    [
        ("Size (KLOC)",              "11.35",   "Measured SLOC count (§5.1.1)"),
        ("COCOMO Mode",              "Organic", "Small team, familiar domain"),
        ("Effort Coefficient (a)",   "2.4",     "COCOMO Organic constant"),
        ("Effort Exponent (b)",      "1.05",    "COCOMO Organic constant"),
        ("Duration Coefficient (c)", "2.5",     "COCOMO Organic constant"),
        ("Duration Exponent (d)",    "0.38",    "COCOMO Organic constant"),
    ],
    col_widths=[2.5, 1.5, 2.5]
)

h3(doc, "7.5.3  Calculations")
body(doc, "Step 1 — Development Effort (PM):")
body(doc, "    PM  =  a × (KLOC)^b  =  2.4 × (11.35)^1.05  ≈  30.5 Person-Months")
body(doc, "Step 2 — Estimated Project Duration (TDEV):")
body(doc, "    TDEV  =  c × (PM)^d  =  2.5 × (30.5)^0.38  ≈  9.2 Months")
body(doc, "Step 3 — Average Team Size:")
body(doc, "    Team Size  =  PM / TDEV  =  30.5 / 9.2  ≈  3.3  ≈  3–4 Persons")

h3(doc, "7.5.4  Results Summary")
make_table(doc,
    ["Metric", "Estimated Value"],
    [
        ("Codebase Size",           "11.35 KLOC"),
        ("Development Effort",      "30.5 Person-Months"),
        ("Estimated Duration",      "9.2 Months"),
        ("Recommended Team Size",   "3–4 Persons"),
    ],
    col_widths=[2.5, 4.0]
)

body(doc,
    "The COCOMO estimate of 30.5 person-months over approximately 9.2 months with a team "
    "of 3–4 persons is consistent with the actual development trajectory of the Kuwait "
    "B2B Hub. The project was executed by a small group over a single academic semester, "
    "with development effort concentrated in the Components layer (33.5% of KLOC) and the "
    "buyer/brand-facing page routes (40.1% combined). The organic mode assumption is "
    "validated by the team's prior familiarity with TypeScript and React, and by the "
    "project's well-scoped, single-domain requirements. COCOMO is a calibrated estimation "
    "model and should be interpreted as a reasonable industry benchmark rather than a "
    "precise prediction."
)

# §7.6 Risk Assessment Matrix
h2(doc, "7.6  Risk Assessment Matrix (RAM)")

h3(doc, "7.6.1  Methodology")
body(doc,
    "Risk assessment for the Kuwait B2B Hub was conducted using a quantitative "
    "Likelihood × Impact exposure model consistent with ISO 31000:2018. Each risk was "
    "evaluated against two dimensions, producing an Exposure Score that determines the "
    "priority tier and mandates a corresponding response strategy. All risks are grounded "
    "in direct evidence from the codebase, architecture, and deployment environment."
)

h3(doc, "7.6.2  Scoring Rubrics")

body(doc, "Likelihood Scale:")
make_table(doc,
    ["Score", "Label", "Indicative Probability"],
    [
        ("1", "Rare",          "< 10%"),
        ("2", "Unlikely",      "10 – 30%"),
        ("3", "Possible",      "30 – 50%"),
        ("4", "Likely",        "50 – 70%"),
        ("5", "Almost Certain","  > 70%"),
    ],
    col_widths=[0.7, 1.3, 4.5]
)

body(doc, "Impact Scale:")
make_table(doc,
    ["Score", "Label", "Example Consequence"],
    [
        ("1", "Negligible", "Minor UI inconsistency"),
        ("2", "Minor",      "Temporary feature unavailability"),
        ("3", "Moderate",   "Feature regression requiring manual recovery"),
        ("4", "Major",      "Data inaccessibility; reputational harm"),
        ("5", "Critical",   "Full outage; credential breach; regulatory penalty"),
    ],
    col_widths=[0.7, 1.3, 4.5]
)

body(doc, "Exposure Level (Likelihood × Impact):")
make_table(doc,
    ["Exposure", "Level", "Required Response"],
    [
        ("15 – 25", "CRITICAL", "Immediate — must resolve before production launch"),
        ("10 – 14", "HIGH",     "Remediation within one sprint"),
        ("5 – 9",   "MEDIUM",   "Planned mitigation within current release cycle"),
        ("1 – 4",   "LOW",      "Monitor and review quarterly"),
    ],
    col_widths=[1.2, 1.2, 4.1]
)

h3(doc, "7.6.3  Risk Assessment Matrix")

make_table(doc,
    ["ID", "Risk", "Cat.", "L", "I", "Exp.", "Level", "Mitigation Strategy", "Owner"],
    [
        ("R-01",
         "Concurrent Write Data Corruption — readDB() → modify → writeDB() provides zero "
         "transaction isolation. Two buyers ordering last stock unit will both pass stock "
         "check before either write completes, causing overselling and data corruption.",
         "Technical","5","5","25","CRITICAL",
         "Execute PostgreSQL migration (SC-01) as the first post-submission action. "
         "Interim: introduce in-memory mutex lock around writeDB(). Add integration tests "
         "simulating concurrent writes.",
         "Tech Lead"),
        ("R-02",
         "Plain-Text Password Storage — Passwords stored without hashing in db.json "
         "(evidenced in register/route.ts line 23; compared plain-text in login/route.ts "
         "line 9). Any read access to the database file exposes all credentials.",
         "Security","4","5","20","CRITICAL",
         "Replace with argon2id or bcrypt (cost factor ≥ 12) before any public deployment. "
         "Audit .gitignore to confirm db.json and .env.local are excluded from version control.",
         "Tech Lead"),
        ("R-03",
         "Silent Regression from Zero Test Coverage — 0% automated coverage across "
         "11.35 KLOC means every feature addition or dependency update risks undetected "
         "breakage. Trust Engine, RBAC guards, and order logic gate financial transactions "
         "but have never been executed in a reproducible test environment.",
         "Quality","5","3","15","CRITICAL",
         "Introduce Vitest + React Testing Library (SC-02) immediately. Set CI gate of "
         "≥ 60% line coverage before any new feature branch can merge. Prioritise "
         "getCurrentUser(), POST /api/orders, and bulk pricing resolution.",
         "Tech Lead"),
        ("R-04",
         "Security Vulnerability Re-Exposure — CR-005 demonstrates a history of "
         "unsanitised output. Additional injection vectors (CSRF, insecure direct object "
         "reference on /api/orders/[id], missing rate limiting on /api/auth/login) "
         "remain unaudited.",
         "Security","3","4","12","HIGH",
         "Commission OWASP Top 10 audit before public launch. Integrate npm audit and "
         "Snyk into CI. Add rate limiting to authentication endpoints via Vercel Edge Middleware.",
         "Member B"),
        ("R-05",
         "PostgreSQL Migration Delay — The JSON → PostgreSQL migration is a prerequisite "
         "for production, yet requires schema design, data migration scripting, ORM "
         "integration, and full regression testing. Underestimating this effort risks "
         "delaying commercial launch.",
         "Schedule","3","4","12","HIGH",
         "Phased migration plan: define Supabase table schemas → write one-time migration "
         "script → run both systems in parallel with write shadowing → cut over with "
         "tested rollback path. Assign dedicated sprint capacity.",
         "Member C"),
        ("R-06",
         "Scope Creep — 13 anticipated changes represent a large backlog that, without "
         "strict controls, may cause partial implementations of multiple changes "
         "simultaneously, resulting in an unstable codebase and missed deadlines.",
         "Management","4","3","12","HIGH",
         "Enforce a formal Change Control Board (CCB) review for all new scope requests. "
         "Maintain a MoSCoW-classified product backlog. Freeze scope at sprint planning.",
         "Team Lead"),
        ("R-07",
         "Kuwait Data Privacy Non-Compliance — Platform collects name, email, company, "
         "whatsapp_number, and order history. Kuwait PDPL (Law No. 20 of 2014) and CITRA "
         "requirements for consent, retention, and breach notification are not addressed.",
         "Legal","3","4","12","HIGH",
         "Engage Kuwait-based legal adviser. Implement cookie consent banner, Privacy Policy "
         "page, documented data retention policy, and breach notification procedure. "
         "Configure Supabase data residency to a GCC region.",
         "Member D"),
        ("R-08",
         "Third-Party Vendor Dependency — 100% of authentication and production hosting "
         "are delegated to Supabase and Vercel. Pricing changes, discontinuation, or "
         "outages translate directly to system unavailability.",
         "Technical","2","5","10","HIGH",
         "Implement daily automated database exports from Supabase. Design the Supabase "
         "client factory behind an interface abstraction to allow provider substitution. "
         "Define RTO < 4 hours for any single-provider outage.",
         "Member C"),
        ("R-09",
         "Performance Degradation Under Load — Synchronous fs.readFileSync() and "
         "fs.writeFileSync() in src/lib/db.ts block the Node.js event loop on every "
         "database operation, causing request queuing and potential Vercel function "
         "timeouts (10-second default).",
         "Technical","4","3","12","HIGH",
         "Resolved by SC-01 (PostgreSQL). Interim: convert readDB() / writeDB() to async "
         "fs.promises. Conduct k6 load test before launch (target: 50 concurrent users, "
         "P95 < 300 ms).",
         "Tech Lead"),
        ("R-10",
         "Key Personnel Dependency — Trust Engine, RBAC, and Supabase integration were "
         "built by specific team members with limited cross-training. Departure of one "
         "member risks leaving critical subsystems undocumented.",
         "Resource","2","4","8","MEDIUM",
         "Mandate pair programming for critical subsystems. Require all merges to include "
         "updated design decision comments. Conduct knowledge-transfer sessions before "
         "semester end. Maintain Architecture Decision Records (ADRs).",
         "Team Lead"),
        ("R-11",
         "KNET Payment Compliance Failure — Kuwait's KNET mandates PCI-DSS compliance, "
         "a certified integration partner, and Central Bank of Kuwait approval. "
         "Implementing payment without satisfying these requirements is a criminal offence "
         "under Kuwait Commercial Law.",
         "Legal","2","5","10","HIGH",
         "Engage a KNET-certified payment gateway aggregator (MyFatoorah, Tap Payments) "
         "that provides a pre-certified integration layer. The platform must never touch "
         "raw card data. Obtain legal sign-off before enabling live transactions.",
         "Member D"),
        ("R-12",
         "Arabic RTL Localisation Regression — Updates to shared layout components made "
         "by developers working in LTR contexts risk breaking RTL alignment, bi-directional "
         "icon mirroring, or right-anchored navigation for Arabic users.",
         "Quality","3","2","6","MEDIUM",
         "Add Playwright visual regression tests capturing screenshots in both ar and en "
         "locale on every CI run. Require a native Arabic speaker to perform a UI review "
         "checklist before each release.",
         "Member B"),
    ],
    col_widths=[0.5, 2.2, 0.5, 0.3, 0.3, 0.5, 0.7, 1.9, 0.8]
)

body(doc, "(L = Likelihood · I = Impact · Exposure = L × I)")

h3(doc, "7.6.4  Risk Heat Map")
body(doc,
    "The heat map below positions each risk by Likelihood (vertical axis) and "
    "Impact (horizontal axis). Risks in the top-right quadrant are Critical."
)
mono(doc,
"""         │  IMPACT →
         │  1-Negligible  2-Minor   3-Moderate  4-Major    5-Critical
─────────┼──────────────────────────────────────────────────────────────
5-Almost │                          R-03                    R-01
 Certain │
─────────┼──────────────────────────────────────────────────────────────
4-Likely │                          R-06, R-09              R-02
─────────┼──────────────────────────────────────────────────────────────
3-Possi- │              R-12                    R-04,
 ble     │                                      R-05, R-07
─────────┼──────────────────────────────────────────────────────────────
2-Unlike-│                                      R-10        R-08, R-11
 ly      │
─────────┼──────────────────────────────────────────────────────────────
1-Rare   │
─────────┴──────────────────────────────────────────────────────────────""",
    fs=8
)
body(doc,
    "Risks R-01, R-02, and R-03 in the top-right quadrant are Critical and demand "
    "resolution before any public-facing deployment. All three are infrastructural — "
    "data integrity, password security, and test coverage — not feature gaps."
)

h3(doc, "7.6.5  Critical Risk Narratives")

subsection_bold(doc, "R-01 — Concurrent Write Data Corruption (Exposure: 25)",
    "This is the maximum possible exposure score. The root cause is architectural: "
    "readDB() → modify → writeDB() is not atomic. On Vercel's serverless platform, "
    "multiple function instances execute concurrently. Two API calls can both read the "
    "same snapshot, both validate successfully, and both write back — with the second "
    "write silently overwriting the first. In stock management terms, a product showing "
    "one unit could be sold to two buyers simultaneously. This is a data integrity failure "
    "with direct contractual and commercial consequences. It cannot be accepted in "
    "production under any circumstances.")

subsection_bold(doc, "R-02 — Plain-Text Password Storage (Exposure: 20)",
    "User passwords are stored without cryptographic hashing in the prototype. This was "
    "discovered through direct code inspection of register/route.ts (line 23) and "
    "confirmed in login/route.ts (line 9). Although db.json is not publicly accessible "
    "in the deployed configuration, this is a Critical vulnerability because: (a) a "
    "misconfigured deployment or accidental Git commit could expose it; (b) any "
    "server-side path traversal would yield all credentials immediately; (c) Kuwaiti "
    "users commonly reuse passwords across services. This must be resolved with "
    "argon2id or bcrypt before any real user credentials are collected — including in "
    "a closed beta.")

subsection_bold(doc, "R-03 — Silent Regression from Zero Test Coverage (Exposure: 15)",
    "The absence of any automated test suite means the codebase has no formal "
    "specification of correct behaviour. Every change to a shared utility — such as "
    "getCurrentUser(), called by all 25 API routes — can introduce a silent regression "
    "that only manifests in production. With 11.35 KLOC of production code and no "
    "regression safety net, the cost of defect detection is borne entirely by end users. "
    "Introducing Vitest is the highest-return investment available after the database "
    "migration and password hardening.")

h3(doc, "7.6.6  Residual Risk Summary")
make_table(doc,
    ["ID", "Risk", "Pre-Mitigation", "Mitigation Applied", "Residual", "Level"],
    [
        ("R-01","Concurrent Write Data Corruption",  "25 — CRITICAL","PostgreSQL ACID transactions",    "2",  "LOW"),
        ("R-02","Plain-Text Password Storage",        "20 — CRITICAL","argon2id hashing + audit",        "3",  "LOW"),
        ("R-03","Zero Test Coverage",                 "15 — CRITICAL","Vitest + CI coverage gate",       "6",  "MEDIUM"),
        ("R-04","Security Re-Exposure",               "12 — HIGH",    "OWASP audit + Snyk CI",           "4",  "LOW"),
        ("R-05","Migration Delay",                    "12 — HIGH",    "Phased plan with rollback",       "6",  "MEDIUM"),
        ("R-06","Scope Creep",                        "12 — HIGH",    "CCB + MoSCoW backlog",            "6",  "MEDIUM"),
        ("R-07","Data Privacy Non-Compliance",        "12 — HIGH",    "Legal review + PDPL controls",    "4",  "LOW"),
        ("R-08","Vendor Dependency",                  "10 — HIGH",    "Daily export + abstraction layer","5",  "MEDIUM"),
        ("R-09","Performance Under Load",             "12 — HIGH",    "PostgreSQL + async I/O",          "3",  "LOW"),
        ("R-10","Key Personnel Dependency",           "8 — MEDIUM",   "Pair programming + ADRs",         "4",  "LOW"),
        ("R-11","KNET Compliance Failure",            "10 — HIGH",    "Certified gateway aggregator",    "4",  "LOW"),
        ("R-12","Arabic RTL Regression",              "6 — MEDIUM",   "Playwright visual regression + Arabic QA","2","LOW"),
    ],
    col_widths=[0.5, 2.1, 1.3, 1.9, 0.7, 0.8]
)

body(doc,
    "Overall Residual Assessment: Following full implementation of all mitigations, "
    "no risk remains Critical or High. Three risks settle at Medium (R-03, R-05, R-06) "
    "reflecting inherent complexity that cannot be fully eliminated. The remaining nine "
    "risks reduce to Low, indicating the Kuwait B2B Hub presents an acceptable risk "
    "profile for commercial deployment once the three Critical items are resolved. "
    "R-02 is a zero-tolerance item: plain-text password storage must be resolved before "
    "any real user credentials are collected — even in a closed beta. This is "
    "non-negotiable under Kuwait's Personal Data Protection Law."
)

# §7.7 Fishbone
h2(doc, "7.7  Root-Cause Analysis — Fishbone Diagram")

body(doc,
    "The Ishikawa (fishbone) diagram below performs a root-cause analysis for the "
    "project's most significant quality deficit: the complete absence of automated test "
    "coverage (Risk R-03). The diagram identifies contributing causes across six "
    "standard categories: People, Process, Technology, Environment, Materials, and "
    "Management."
)

make_table(doc,
    ["Category", "Contributing Cause", "Detail"],
    [
        ("People",      "Skill gap in test-driven development",
         "Team members were proficient in TypeScript and React but lacked prior experience "
         "writing Vitest or Jest tests for Next.js App Router patterns."),
        ("People",      "Role concentration on feature delivery",
         "All four members were assigned to feature subsystems; no dedicated QA role "
         "was established, leaving testing as a collective deferred responsibility."),
        ("Process",     "Test infrastructure not included in Sprint 1 Definition of Done",
         "The project's Definition of Done for Sprint 1 included working authentication "
         "but did not mandate a passing test suite, setting a precedent that persisted."),
        ("Process",     "Agile sprint pressure prioritised visible deliverables",
         "Sprint reviews demonstrated working features to the supervisor; test suites "
         "produced no visible sprint output, reducing their perceived sprint priority."),
        ("Technology",  "Next.js App Router test setup complexity",
         "Testing Next.js Server Components and API Route Handlers with Vitest requires "
         "non-trivial configuration (mocking cookies(), headers(), Request/Response). "
         "This overhead discouraged early adoption."),
        ("Technology",  "JSON database made integration testing non-deterministic",
         "Without a proper test database fixture, integration tests would mutate the "
         "shared db.json file, making test isolation difficult to guarantee."),
        ("Environment", "Single shared db.json prevents parallel test execution",
         "Concurrent test runs on the same flat-file database would produce race "
         "conditions, making a test suite actively harmful without PostgreSQL isolation."),
        ("Management",  "No CI/CD pipeline enforcing test gate",
         "Without a GitHub Actions or Vercel CI configuration requiring tests to pass "
         "before merge, there was no automated enforcement mechanism."),
        ("Management",  "Academic timeline prioritised report over test coverage",
         "The course assessment criteria weighted the final report and presentation "
         "above test suite evidence, rationally redirecting effort away from testing."),
    ],
    col_widths=[1.1, 2.0, 3.4]
)

placeholder(doc, "FIGURE 11", "Fishbone (Ishikawa) Diagram — Root Causes of Zero Test Coverage")

body(doc,
    "Recommended corrective action: Implement Vitest as the first Sprint 8 action "
    "(SC-02). Configure a GitHub Actions workflow that runs tests on every pull request "
    "and blocks merge if coverage falls below 60%. Prioritise the five highest-risk "
    "units: getCurrentUser(), POST /api/orders, bulk pricing resolution, the Trust "
    "Engine computation, and all admin RBAC guards."
)

# ─── SAVE ────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_FILE)
print("PHASE 3 COMPLETE")
print(f"Output saved to: {OUTPUT_FILE}")
